import os
import json
import copy
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches
from lxml import etree

NORMATIVIDAD_DIR = "normatividad"
PLANTILLAS_DIR = "plantillas"
PLAN_TEMPLATE = os.path.join(PLANTILLAS_DIR, "Plan de Evaluación EC0301.docx")
PORTADA_TEMPLATE = os.path.join(PLANTILLAS_DIR, "Portada y Subportada.docx")
SERVICIO_TEMPLATE = os.path.join(PLANTILLAS_DIR, "Servicio a Usuarios.docx")
ENCUESTA_TEMPLATE = os.path.join(PLANTILLAS_DIR, "Encuesta de Satisfacción (Caritas).docx")

EVALUATOR_SYSTEM_PROMPT = """Eres un Evaluador Experto del CONOCER certificado en el EC0076. Tu objetivo es auditar y generar documentos formales de evaluación para candidatos del EC0301, cumpliendo estrictamente con los lineamientos de Verificación Interna y Portafolio de Evidencias.

REGLAS INQUEBRANTABLES:
1. Todo dictamen, criterio y contenido que generes DEBE estar sustentado EXCLUSIVAMENTE en los documentos normativos proporcionados como contexto. NO inventes, inferiras ni agregues criterios fuera del marco legal.
2. Debes alinearte a las Reglas Generales del Sistema Nacional de Competencias (SNC), los Manuales de Evaluación y Verificación del EC0076, y el estándar EC0301.
3. Cada actividad de evaluación debe especificar claramente la fase (Gabinete o Campo), la técnica de evaluación, el instrumento y la evidencia esperada.
4. Usa lenguaje formal, técnico y profesional propio del ámbito de certificación CONOCER."""


def get_openai_client():
    # Migrado a Gemini (google-genai). Se conserva el nombre de la función
    # para no tocar el resto de este archivo.
    return genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))


def _read_pdf_text(path):
    try:
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        if parts:
            return "\n".join(parts)
    except Exception:
        pass
    return ""


def load_normatividad_context():
    texts = {}
    if not os.path.isdir(NORMATIVIDAD_DIR):
        return texts
    for fname in sorted(os.listdir(NORMATIVIDAD_DIR)):
        fpath = os.path.join(NORMATIVIDAD_DIR, fname)
        if fname.lower().endswith(".pdf"):
            content = _read_pdf_text(fpath)
            if content:
                texts[fname] = content
    return texts


def _build_normative_context():
    docs = load_normatividad_context()
    if not docs:
        return "[No se encontraron documentos normativos en /normatividad/]"
    parts = []
    for fname, content in docs.items():
        truncated = content[:8000]
        parts.append(f"=== {fname} ===\n{truncated}")
    return "\n\n".join(parts)


def _inject_logo_into_doc(doc, logo_path, width_inches=1.5):
    if not logo_path or not os.path.isfile(logo_path):
        return
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_pic = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    body = doc.element.body
    has_image = False
    for elem in body.iter():
        tag = etree.QName(elem).localname if isinstance(elem.tag, str) else ''
        if tag == 'pic':
            has_image = True
            break
    if has_image:
        return
    try:
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            run = first_para.add_run()
            run.add_picture(logo_path, width=Inches(width_inches))
    except Exception:
        pass


def _set_cell_text(cell, text):
    cell.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = str(text)
    else:
        cell.add_paragraph(str(text))


def fill_plan_template(plan_data, datos_candidato, datos_evaluador, fechas, logo_path=None):
    if not os.path.isfile(PLAN_TEMPLATE):
        raise FileNotFoundError(f"Plantilla no encontrada: {PLAN_TEMPLATE}")

    doc = Document(PLAN_TEMPLATE)
    num_tables = len(doc.tables)

    try:
        t0 = doc.tables[0]
        if len(t0.rows) >= 5 and len(t0.rows[1].cells) >= 2:
            _set_cell_text(t0.rows[1].cells[1], fechas.get("fecha_emision", ""))
            _set_cell_text(t0.rows[3].cells[1], datos_evaluador.get("nombre", ""))
            _set_cell_text(t0.rows[4].cells[1], datos_candidato.get("nombre", ""))
            if datos_evaluador.get("clave_conocer") and datos_evaluador["clave_conocer"] != "N/A":
                _set_cell_text(t0.rows[0].cells[1], datos_evaluador["clave_conocer"])
    except Exception:
        pass

    try:
        diag = plan_data.get("resultado_diagnostico", {})
        if num_tables > 1:
            t1 = doc.tables[1]
            if len(t1.rows) >= 3 and len(t1.rows[1].cells) >= 5:
                sugiere = diag.get("sugiere_capacitacion", "No") or "No"
                procede = diag.get("procede_evaluacion", "Sí") or "Sí"
                if sugiere.lower().startswith("s"):
                    _set_cell_text(t1.rows[1].cells[2], "X")
                else:
                    _set_cell_text(t1.rows[1].cells[4], "X")
                if procede.lower().startswith("s"):
                    _set_cell_text(t1.rows[2].cells[2], "X")
                else:
                    _set_cell_text(t1.rows[2].cells[4], "X")
    except Exception:
        pass

    try:
        if num_tables > 2:
            t2 = doc.tables[2]
            fecha_gabinete = fechas.get("fecha_gabinete", "")
            fecha_campo = fechas.get("fecha_campo", "")
            for ri, row in enumerate(t2.rows):
                if ri == 0:
                    continue
                try:
                    cells = row.cells
                    if len(cells) < 4:
                        continue
                    cell_text = cells[1].text.strip().lower()
                    instrument_text = cells[2].text.strip().lower() if len(cells) > 2 else ""
                    merged_header = all(c.text.strip() == cells[0].text.strip() for c in cells)
                    if merged_header:
                        continue
                    if "verificaré" in cell_text or "proceso" in cell_text:
                        _set_cell_text(cells[3], fecha_campo)
                    elif "conocimiento" in cell_text or "cuestionario" in instrument_text:
                        _set_cell_text(cells[3], fecha_gabinete)
                    elif cells[3].text.strip() == "":
                        _set_cell_text(cells[3], fecha_gabinete)
                except Exception:
                    pass
    except Exception:
        pass

    try:
        if num_tables > 5 and len(doc.tables[5].rows) > 2 and len(doc.tables[5].rows[2].cells) >= 2:
            _set_cell_text(doc.tables[5].rows[2].cells[1], fechas.get("fecha_campo", ""))
    except Exception:
        pass

    try:
        if num_tables > 6 and len(doc.tables[6].rows) > 2 and len(doc.tables[6].rows[2].cells) >= 2:
            fecha_resultados = fechas.get("fecha_resultados", fechas.get("fecha_campo", ""))
            _set_cell_text(doc.tables[6].rows[2].cells[1], fecha_resultados)
    except Exception:
        pass

    try:
        if num_tables > 7 and len(doc.tables[7].rows) > 1 and len(doc.tables[7].rows[1].cells) >= 3:
            _set_cell_text(doc.tables[7].rows[1].cells[0], datos_evaluador.get("nombre", ""))
            _set_cell_text(doc.tables[7].rows[1].cells[2], datos_candidato.get("nombre", ""))
    except Exception:
        pass

    _inject_logo_into_doc(doc, logo_path)

    return doc


def _get_paragraph_text(p_elem):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = []
    for r in p_elem.findall(f'.//{{{ns_w}}}r'):
        for t in r.findall(f'{{{ns_w}}}t'):
            if t.text:
                parts.append(t.text)
    return ''.join(parts)


def _set_paragraph_runs_text(p_elem, new_full_text):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    runs = p_elem.findall(f'.//{{{ns_w}}}r')
    if not runs:
        return
    first_t = runs[0].find(f'{{{ns_w}}}t')
    if first_t is not None:
        first_t.text = new_full_text
        first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for r in runs[1:]:
        for t in r.findall(f'{{{ns_w}}}t'):
            t.text = ''


def _replace_textbox_text(doc, replacements):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    for txbx in body.iter():
        tag = etree.QName(txbx).localname if isinstance(txbx.tag, str) else ''
        if tag == 'txbxContent':
            for p_elem in txbx.findall(f'.//{{{ns_w}}}p'):
                full_text = _get_paragraph_text(p_elem)
                if not full_text.strip():
                    continue
                new_text = full_text
                changed = False
                for old_text, replacement in replacements.items():
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, replacement)
                        changed = True
                if changed:
                    _set_paragraph_runs_text(p_elem, new_text)


def fill_portada_template(datos_candidato, datos_evaluador, logo_path=None):
    if not os.path.isfile(PORTADA_TEMPLATE):
        raise FileNotFoundError(f"Plantilla no encontrada: {PORTADA_TEMPLATE}")

    doc = Document(PORTADA_TEMPLATE)

    nombre_candidato = datos_candidato.get("nombre", "")
    nombre_evaluador = datos_evaluador.get("nombre", "")
    clave_ce = datos_evaluador.get("clave_conocer", "")
    estandar_text = "EC0301: Diseño de cursos de formación del capital humano de manera presencial grupal, sus instrumentos de evaluación y manuales del curso"

    replacements = {
        "Nombre Candidato: ": f"Nombre Candidato: {nombre_candidato}",
        "Candidato: --------------------------------": f"Candidato: {nombre_candidato}",
        "EC          : Nombre del estándar": estandar_text,
        "EC0217: Impartición de cursos de formación de capital humano de manera presencial grupal": estandar_text,
        "Clave del CE/EI. ": f"Clave del CE/EI: {clave_ce}",
        "CE1497-ECE020-10": clave_ce if clave_ce and clave_ce != "N/A" else "CE1497-ECE020-10",
        "Nombre Evaluador: ": f"Nombre Evaluador: {nombre_evaluador}",
        "Evaluador: -----------------------------------": f"Evaluador: {nombre_evaluador}",
    }

    _replace_textbox_text(doc, replacements)

    _inject_logo_into_doc(doc, logo_path)

    return doc


def fill_servicio_template(datos_candidato, fechas, logo_path=None):
    if not os.path.isfile(SERVICIO_TEMPLATE):
        raise FileNotFoundError(f"Plantilla no encontrada: {SERVICIO_TEMPLATE}")

    doc = Document(SERVICIO_TEMPLATE)
    num_tables = len(doc.tables)

    if num_tables < 1:
        return doc

    try:
        t0 = doc.tables[0]
        if len(t0.rows) >= 3 and len(t0.rows[2].cells) >= 2:
            nombre = datos_candidato.get("nombre", "")
            _set_cell_text(t0.rows[2].cells[1], nombre)
    except Exception:
        pass

    try:
        for p in doc.paragraphs:
            if "Fecha de aplicación:" in p.text:
                fecha = fechas.get("fecha_emision", fechas.get("fecha_campo", ""))
                for run in p.runs:
                    if "Fecha de aplicación:" in run.text:
                        run.text = f"Fecha de aplicación: {fecha}"
                        break
                break
    except Exception:
        pass

    for ti in range(2, min(13, num_tables)):
        try:
            table = doc.tables[ti]
            if len(table.rows) >= 1 and len(table.rows[0].cells) >= 2:
                _set_cell_text(table.rows[0].cells[1], "X")
        except Exception:
            pass

    _inject_logo_into_doc(doc, logo_path)

    return doc


def fill_encuesta_template(datos_candidato, fechas, logo_path=None):
    if not os.path.isfile(ENCUESTA_TEMPLATE):
        raise FileNotFoundError(f"Plantilla no encontrada: {ENCUESTA_TEMPLATE}")

    doc = Document(ENCUESTA_TEMPLATE)
    num_tables = len(doc.tables)

    if num_tables < 1:
        return doc

    try:
        t0 = doc.tables[0]
        if len(t0.rows) >= 2 and len(t0.rows[0].cells) >= 2:
            nombre = datos_candidato.get("nombre", "")
            fecha = fechas.get("fecha_emision", fechas.get("fecha_campo", ""))
            _set_cell_text(t0.rows[0].cells[1], nombre)
            _set_cell_text(t0.rows[1].cells[1], fecha)
    except Exception:
        pass

    _inject_logo_into_doc(doc, logo_path)

    return doc


def generate_plan_evaluacion(datos_candidato, datos_evaluador, fechas, logo_path=None):
    client = get_openai_client()
    normative_context = _build_normative_context()

    user_prompt = f"""Con base EXCLUSIVA en los documentos normativos proporcionados, genera los datos para completar un Plan de Evaluación formal del estándar EC0301.

DATOS DEL CANDIDATO:
- Nombre: {datos_candidato.get('nombre', 'N/A')}
- Correo: {datos_candidato.get('correo', 'N/A')}
- WhatsApp: {datos_candidato.get('whatsapp', 'N/A')}

DATOS DEL EVALUADOR:
- Nombre del Evaluador: {datos_evaluador.get('nombre', 'N/A')}
- Centro Evaluador: {datos_evaluador.get('centro', 'N/A')}
- Clave de Registro CONOCER: {datos_evaluador.get('clave_conocer', 'N/A')}

FECHAS:
- Fecha de Gabinete (revisión documental): {fechas.get('fecha_gabinete', 'N/A')}
- Fecha de Campo (evaluación presencial): {fechas.get('fecha_campo', 'N/A')}
- Fecha de Emisión del Plan: {fechas.get('fecha_emision', 'N/A')}

Genera un JSON con esta estructura exacta:
{{
  "actividades_evaluacion": [
    {{
      "fase": "Gabinete o Campo",
      "actividad": "Descripción de la actividad evaluada",
      "tecnica": "Técnica de evaluación",
      "instrumento": "Instrumento aplicado",
      "evidencia": "Tipo de evidencia",
      "fecha": "Fecha programada"
    }}
  ],
  "resultado_diagnostico": {{
    "sugiere_capacitacion": "Sí o No",
    "procede_evaluacion": "Sí o No"
  }},
  "requerimientos_adicionales": "Observaciones sobre materiales o requerimientos"
}}

Responde ÚNICAMENTE con el JSON válido, sin texto adicional ni bloques de código."""

    system_content = f"""{EVALUATOR_SYSTEM_PROMPT}

DOCUMENTOS NORMATIVOS DE REFERENCIA (BASE DE CONOCIMIENTO OBLIGATORIA):
{normative_context}"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=system_content,
            max_output_tokens=4096,
            temperature=0.2,
        ),
    )

    raw = (response.text or "").strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()

    plan_data = json.loads(raw)

    filled_doc = fill_plan_template(plan_data, datos_candidato, datos_evaluador, fechas, logo_path=logo_path)

    return plan_data, filled_doc
