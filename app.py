import os
import re
import io
from sqlalchemy import text
import csv
import json
import stripe
import logging
import smtplib
import traceback
import unicodedata
import threading
from datetime import datetime, timezone, timedelta
from zoneinfo import ZoneInfo
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
from flask import Flask, render_template, request, jsonify, send_from_directory, send_file, abort, redirect, url_for, flash, Response, session
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from ai_helper import chat_with_ai, load_all_reference_docs, extract_course_info_ai
from doc_generator import generate_custom_docx, generate_from_template, generate_pptx_from_slides, list_generated_docs, pop_last_pexels_usage
from evaluator_ai import generate_plan_evaluacion as ai_generate_plan_evaluacion, fill_portada_template, fill_servicio_template, fill_encuesta_template
import analytics_rules
import reengagement

logger = logging.getLogger(__name__)
import time as _time

_ai_active_count = [0]
_ai_queue_waiting = [0]
_ai_queue_lock = threading.Lock()
_ai_condition = threading.Condition(_ai_queue_lock)

def _normalize_topic(text):
    if not text:
        return ''
    text = text.lower().strip()
    text = unicodedata.normalize('NFD', text)
    text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
    return text

app = Flask(__name__)
# WhiteNoise para servir archivos estáticos en producción (Cloud Run)
from whitenoise import WhiteNoise

app.wsgi_app = WhiteNoise(app.wsgi_app, root='static/', prefix='static/')
app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY', 'pertinentia-secret-key-2026')
database_url = os.environ.get('DATABASE_URL', 'sqlite:///pertinentia.db')
if database_url.startswith('postgres://'):
    database_url = database_url.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 700 * 1024 * 1024
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300,
    'pool_size': 5,
    'max_overflow': 15,
}

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Inicia sesión para acceder a la plataforma.'
login_manager.login_message_category = 'info'

ADMIN_EMAIL = 'arturogarciac@pertinentia.com'

_BOOT_TS = str(int(_time.time()))

@app.after_request
def add_cache_headers(response):
    if request.path.startswith('/static/js/'):
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

_PROMO_RE = __import__('re').compile(r'^[a-z0-9_\-]{1,30}$')
_PROMO_TTL_SECONDS = 7 * 24 * 3600

@app.before_request
def _capture_promo_param():
    try:
        raw = (request.args.get('promo') or '').strip().lower()[:30]
        if not raw or not _PROMO_RE.match(raw):
            return
        import time as _tp
        session['promo_code'] = raw
        session['promo_expires'] = int(_tp.time()) + _PROMO_TTL_SECONDS
        session.permanent = True
    except Exception:
        return

def _get_session_promo():
    try:
        import time as _tp
        exp = int(session.get('promo_expires', 0) or 0)
        if exp and _tp.time() < exp:
            return (session.get('promo_code') or '').lower()
    except Exception:
        pass
    return ''

def _is_maestro_campaign_active():
    """Día del Maestro (México, 15 mayo). Campaña activa todo mayo 2026.
    Override con env MAESTRO_CAMPAIGN_FORCE=1 (forzar ON) o MAESTRO_CAMPAIGN_FORCE=0 (forzar OFF)."""
    try:
        force = (os.environ.get('MAESTRO_CAMPAIGN_FORCE') or '').strip()
        if force == '1':
            return True
        if force == '0':
            return False
        from datetime import datetime as _dt
        try:
            now_cdmx = _dt.now(CDMX_TZ)
        except Exception:
            now_cdmx = _dt.utcnow()
        return now_cdmx.year == 2026 and now_cdmx.month == 5
    except Exception:
        return False

def _queue_pixel_event(name, params=None):
    try:
        if not name:
            return
        evs = session.get('_pixel_events') or []
        if not isinstance(evs, list):
            evs = []
        evs.append({'name': str(name)[:60], 'params': (params if isinstance(params, dict) else {})})
        session['_pixel_events'] = evs[-10:]
    except Exception as _epx:
        try:
            logger.warning(f"_queue_pixel_event failed for {name}: {_epx}")
        except Exception:
            pass

@app.context_processor
def inject_cache_bust():
    base = (os.environ.get('PUBLIC_BASE_URL') or 'https://pertinentia.com').rstrip('/')
    try:
        path = request.path if request else '/'
    except Exception:
        path = '/'
    pixel_events = []
    try:
        pixel_events = session.pop('_pixel_events', []) or []
        if not isinstance(pixel_events, list):
            pixel_events = []
    except Exception:
        pixel_events = []
    return {
        'cache_v': _BOOT_TS,
        'public_base_url': base,
        'canonical_url': base + path,
        'ga4_id': os.environ.get('GA4_MEASUREMENT_ID', '').strip(),
        'gsc_verification': os.environ.get('GSC_VERIFICATION_TOKEN', '').strip(),
        'meta_pixel_id': os.environ.get('META_PIXEL_ID', '').strip(),
        'session_promo': _get_session_promo(),
        'maestro_campaign_active': _is_maestro_campaign_active(),
        'pixel_events': pixel_events,
    }

STRIPE_LINKS = {
    'PRO_PROJECT': 'https://buy.stripe.com/9B600idtXdkS0L83w88k80k',
    'PRO_MULTICURSO': 'https://buy.stripe.com/28EeVc3TneoWfG2aYA8k80q',
    'PREMIUM_MONTHLY': 'https://buy.stripe.com/3cI7sKcpTgx4dxU6Ik8k80n',
    'PREMIUM_ANNUAL': 'https://buy.stripe.com/28EaEW9dH0y679wd6I8k80o',
    'PRO_AFFILIATE': 'https://buy.stripe.com/dRm14mdtX2Ge2Tg9Uw8k80p',
    'ALACARTE_E1': 'https://buy.stripe.com/5kQ4gy2Pj1Ca1Pc6Ik8k80g',
    'ALACARTE_E2': 'https://buy.stripe.com/bJefZg9dH0y679w9Uw8k80h',
    'ALACARTE_E3': 'https://buy.stripe.com/00wcN4cpT5SqalI5Eg8k80i',
    'ALACARTE_E4': 'https://buy.stripe.com/4gMfZg9dH2GealI3w88k80j',
    'ALACARTE_E5': 'https://buy.stripe.com/3cIfZg1Lf1Ca3Xkd6I8k80r',
}

TIER_LABELS = {
    'PRO_AFFILIATE': 'PRO Afiliado — $500',
    'PRO_PROJECT': 'PRO 1 Curso — $999',
    'PRO_MULTICURSO': 'PRO 5 Cursos — $2,129',
    'PREMIUM_MONTHLY': 'PREMIUM Mensual — $2,997',
    'PREMIUM_ANNUAL': 'PREMIUM Anual',
    'ALACARTE_E1': 'A la Carta — E1 Carta Descriptiva ($299)',
    'ALACARTE_E2': 'A la Carta — E2 Instrumentos ($399)',
    'ALACARTE_E3': 'A la Carta — E3 Manuales ($399)',
    'ALACARTE_E4': 'A la Carta — E4 Autodiagnóstico ($249)',
    'ALACARTE_E5': 'A la Carta — E5 Evaluación EC0217.01 ($349)',
}

TIER_PRICES = {
    'PRO_AFFILIATE': 500,
    'PRO_PROJECT': 999,
    'PRO_MULTICURSO': 2129,
    'PREMIUM_MONTHLY': 2997,
    'PREMIUM_ANNUAL': 29970,
    'ALACARTE_E1': 299,
    'ALACARTE_E2': 399,
    'ALACARTE_E3': 399,
    'ALACARTE_E4': 249,
    'ALACARTE_E5': 349,
}

def _get_stripe_link_for_tier(tier_key):
    if tier_key == 'PRO_AFFILIATE':
        return Config.get('STRIPE_CHECKOUT_URL_AFFILIATE', '') or STRIPE_LINKS.get('PRO_AFFILIATE', '')
    return STRIPE_LINKS.get(tier_key, '')

STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
stripe.api_key = os.environ.get('STRIPE_SECRET_KEY', '')

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    # NOTA: el correo YA NO es único. Se permiten varias cuentas distintas que
    # comparten correo/WhatsApp (creadas solo por admin). El inicio de sesión
    # desambigua por contraseña. Ver migración DROP CONSTRAINT user_email_key.
    email = db.Column(db.String(150), nullable=False, index=True)
    password_hash = db.Column(db.String(256), nullable=False)
    full_name = db.Column(db.String(150), nullable=False)
    first_name = db.Column(db.String(80), nullable=True)
    apellido_paterno = db.Column(db.String(80), nullable=True)
    apellido_materno = db.Column(db.String(80), nullable=True)
    tier = db.Column(db.String(10), default='FREE', nullable=False)
    whatsapp = db.Column(db.String(20), nullable=True)
    terms_accepted = db.Column(db.Boolean, default=False, nullable=False)
    marketing_consent = db.Column(db.Boolean, default=False, nullable=False)
    marketing_consent_source = db.Column(db.String(20), nullable=True)
    mkt_unsubscribed_at = db.Column(db.DateTime, nullable=True)
    email_bounced_at = db.Column(db.DateTime, nullable=True)
    email_bounce_type = db.Column(db.String(10), nullable=True)
    email_bounce_note = db.Column(db.String(255), nullable=True)
    ultimo_contacto_at = db.Column(db.DateTime, nullable=True)
    notas_contacto = db.Column(db.Text, nullable=True)
    normative_agreement_accepted = db.Column(db.Boolean, default=False, nullable=False)
    chat_usage_count = db.Column(db.Integer, default=0, nullable=False)
    free_downloads_used = db.Column(db.Integer, default=0, nullable=False)
    alacarte_e1 = db.Column(db.Integer, default=0, nullable=False)
    alacarte_e2 = db.Column(db.Integer, default=0, nullable=False)
    alacarte_e3 = db.Column(db.Integer, default=0, nullable=False)
    alacarte_e4 = db.Column(db.Integer, default=0, nullable=False)
    alacarte_e5 = db.Column(db.Integer, default=0, nullable=False)
    e5_preview_used_at = db.Column(db.DateTime, nullable=True)
    ec0217_grants = db.Column(db.Integer, default=0, nullable=False)
    pro_courses_remaining = db.Column(db.Integer, default=0, nullable=False)
    pro_active_course = db.Column(db.String(300), nullable=True)
    is_affiliate = db.Column(db.Boolean, default=False)
    affiliate_terms_accepted = db.Column(db.Boolean, default=False)
    tax_regime = db.Column(db.String(20), default='RESICO')
    referred_by = db.Column(db.Integer, nullable=True)
    tier_usage_stats = db.Column(db.Text, default='{}')
    rfc = db.Column(db.String(20), nullable=True)
    razon_social = db.Column(db.String(250), nullable=True)
    domicilio_fiscal = db.Column(db.String(300), nullable=True)
    nombre_representante_legal = db.Column(db.String(200), nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    active_course_session_id = db.Column(db.Integer, nullable=True, index=True)
    # Timestamp del último "hard reset" de estado del curso disparado server-side
    # (p.ej. tras pago Stripe FREE→PRO). El frontend compara con su localStorage
    # `pertinentia_last_reset_ack`; si difiere, purga caches locales y hace ACK.
    # Nunca afecta créditos.
    needs_state_reset_at = db.Column(db.DateTime, nullable=True)
    utm_source = db.Column(db.String(80), nullable=True, index=True)
    utm_medium = db.Column(db.String(80), nullable=True, index=True)
    utm_campaign = db.Column(db.String(150), nullable=True, index=True)
    utm_term = db.Column(db.String(200), nullable=True, index=True)
    utm_content = db.Column(db.String(250), nullable=True, index=True)
    utm_id = db.Column(db.String(50), nullable=True)
    utm_landing = db.Column(db.String(120), nullable=True)
    utm_captured_at = db.Column(db.DateTime, nullable=True)

    def get_usage_stats(self):
        import json
        try:
            return json.loads(self.tier_usage_stats or '{}')
        except Exception:
            return {}

    def add_token_usage(self, tier, tokens):
        import json
        stats = self.get_usage_stats()
        stats[tier] = stats.get(tier, 0) + tokens
        self.tier_usage_stats = json.dumps(stats)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    @property
    def is_admin(self):
        return self.email == ADMIN_EMAIL

    @property
    def is_pro(self):
        return self.tier in ('PRO', 'PREMIUM')

    @property
    def is_premium(self):
        return self.tier == 'PREMIUM'

    def has_alacarte(self, element_num):
        col = {1: 'alacarte_e1', 2: 'alacarte_e2', 3: 'alacarte_e3', 4: 'alacarte_e4', 5: 'alacarte_e5'}.get(element_num)
        return col and getattr(self, col, 0) > 0

    def use_alacarte(self, element_num):
        col = {1: 'alacarte_e1', 2: 'alacarte_e2', 3: 'alacarte_e3', 4: 'alacarte_e4', 5: 'alacarte_e5'}.get(element_num)
        if not col:
            return False
        result = db.session.execute(
            db.text(f'UPDATE "user" SET {col} = {col} - 1 WHERE id = :uid AND {col} > 0'),
            {"uid": self.id}
        )
        db.session.commit()
        if result.rowcount > 0:
            db.session.refresh(self)
            return True
        return False

    @property
    def ec0217_authorized(self):
        return (self.ec0217_grants or 0) > 0

    def use_ec0217_grant(self):
        """Descuenta 1 evaluación EC0217.01 del saldo autorizado por el admin.
        Atómico: sólo descuenta si hay saldo. Devuelve True si descontó."""
        result = db.session.execute(
            db.text('UPDATE "user" SET ec0217_grants = ec0217_grants - 1 '
                    'WHERE id = :uid AND ec0217_grants > 0'),
            {"uid": self.id}
        )
        db.session.commit()
        if result.rowcount > 0:
            db.session.refresh(self)
            return True
        return False

class Ec0217Ledger(db.Model):
    """Bitácora de auditoría de autorizaciones EC0217.01 otorgadas por el admin
    y de cada servicio (evaluación) consumido. Da claridad ante reclamos:
    quién autorizó, cuándo, cuántas, y día/hora de cada consumo."""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    tipo = db.Column(db.String(12), nullable=False)  # 'grant' | 'consume' | 'revoke'
    cantidad = db.Column(db.Integer, default=0, nullable=False)  # +N grant, -1 consume, set->0 revoke
    saldo_after = db.Column(db.Integer, default=0, nullable=False)
    admin_email = db.Column(db.String(150), nullable=True)
    portafolio_id = db.Column(db.Integer, nullable=True)
    nota = db.Column(db.String(300), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)

class CEProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), unique=True, nullable=False)
    ce_name = db.Column(db.String(200), default='', nullable=False)
    ce_key = db.Column(db.String(100), default='', nullable=False)
    evaluator_name = db.Column(db.String(200), default='', nullable=False)
    logo_path = db.Column(db.String(300), default='', nullable=False)

class Candidato(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ce_user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    nombre_completo = db.Column(db.String(200), nullable=False)
    apellidos = db.Column(db.String(200), default='', nullable=False)
    curp = db.Column(db.String(18), default='', nullable=False)
    estatus_autodiagnostico = db.Column(db.Boolean, default=False, nullable=False)
    estatus_plan = db.Column(db.Boolean, default=False, nullable=False)
    estatus_dictamen = db.Column(db.Boolean, default=False, nullable=False)
    ce_owner = db.relationship('User', backref=db.backref('candidatos', lazy='dynamic'))

class PortafolioEvaluacion(db.Model):
    __tablename__ = 'portafolio_evaluacion'
    id = db.Column(db.Integer, primary_key=True)
    candidato_id = db.Column(db.Integer, db.ForeignKey('candidato.id', ondelete='CASCADE'), nullable=False, index=True)
    ce_user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    estandar = db.Column(db.String(20), nullable=False)
    status = db.Column(db.String(30), default='pendiente', nullable=False)
    resultado_json = db.Column(db.Text, nullable=True)
    error_msg = db.Column(db.Text, nullable=True)
    tokens_prompt = db.Column(db.Integer, default=0, nullable=False)
    tokens_completion = db.Column(db.Integer, default=0, nullable=False)
    cost_usd = db.Column(db.Float, default=0.0, nullable=False)
    cost_mxn = db.Column(db.Float, default=0.0, nullable=False)
    model_used = db.Column(db.String(60), default='', nullable=False)
    portafolio_integrado_pdf = db.Column(db.LargeBinary, nullable=True)
    portafolio_integrado_size = db.Column(db.BigInteger, default=0, nullable=False)
    hash_sha256 = db.Column(db.String(80), default='', nullable=False)
    autorizado_por_evaluador = db.Column(db.Boolean, default=False, nullable=False)
    autorizado_at = db.Column(db.DateTime, nullable=True)
    firma_evaluador_json = db.Column(db.Text, nullable=True)
    video_transcripcion = db.Column(db.Text, nullable=True)
    video_duracion_seg = db.Column(db.Integer, default=0, nullable=False)
    video_idioma = db.Column(db.String(10), default='', nullable=False)
    video_segundos_transcritos = db.Column(db.Integer, default=0, nullable=False)
    video_cost_usd = db.Column(db.Float, default=0.0, nullable=False)
    video_cost_mxn = db.Column(db.Float, default=0.0, nullable=False)
    video_procesado_at = db.Column(db.DateTime, nullable=True)
    video_modelo = db.Column(db.String(40), default='', nullable=False)
    dictamen_json = db.Column(db.Text, nullable=True)
    dictamen_final = db.Column(db.String(40), default='', nullable=False)
    dictamen_pdf = db.Column(db.LargeBinary, nullable=True)
    dictamen_pdf_size = db.Column(db.BigInteger, default=0, nullable=False)
    dictamen_generado_at = db.Column(db.DateTime, nullable=True)
    dictamen_tokens_prompt = db.Column(db.Integer, default=0, nullable=False)
    dictamen_tokens_completion = db.Column(db.Integer, default=0, nullable=False)
    dictamen_cost_usd = db.Column(db.Float, default=0.0, nullable=False)
    dictamen_cost_mxn = db.Column(db.Float, default=0.0, nullable=False)
    e5_unlocked = db.Column(db.Boolean, default=False, nullable=False)
    e5_unlocked_at = db.Column(db.DateTime, nullable=True)
    e5_unlocked_via = db.Column(db.String(20), default='', nullable=False)
    e5_mode = db.Column(db.String(10), default='', nullable=False)
    video_preview_seg = db.Column(db.Integer, default=0, nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now(), index=True)
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())
    candidato = db.relationship('Candidato', backref=db.backref('portafolio_evaluaciones', lazy='dynamic', cascade='all, delete-orphan'))

class PortafolioArchivo(db.Model):
    __tablename__ = 'portafolio_archivo'
    id = db.Column(db.Integer, primary_key=True)
    portafolio_id = db.Column(db.Integer, db.ForeignKey('portafolio_evaluacion.id', ondelete='CASCADE'), nullable=False, index=True)
    tipo = db.Column(db.String(20), nullable=False)
    filename = db.Column(db.String(300), default='', nullable=False)
    contenido = db.Column(db.LargeBinary, nullable=True)
    mime = db.Column(db.String(80), default='', nullable=False)
    size_bytes = db.Column(db.BigInteger, default=0, nullable=False)
    source_url = db.Column(db.String(500), default='', nullable=False)
    categoria = db.Column(db.String(60), default='', nullable=False)
    orden_oficial = db.Column(db.Integer, default=999, nullable=False)
    descripcion = db.Column(db.String(500), default='', nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    portafolio = db.relationship('PortafolioEvaluacion', backref=db.backref('archivos', lazy='dynamic', cascade='all, delete-orphan'))

class PortafolioUploadAttempt(db.Model):
    __tablename__ = 'portafolio_upload_attempt'
    id = db.Column(db.Integer, primary_key=True)
    ce_user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    candidato_id = db.Column(db.Integer, db.ForeignKey('candidato.id', ondelete='SET NULL'), nullable=True)
    estandar = db.Column(db.String(20), default='', nullable=False)
    tipo = db.Column(db.String(20), default='', nullable=False)
    filename = db.Column(db.String(300), default='', nullable=False)
    size_bytes = db.Column(db.BigInteger, default=0, nullable=False)
    success = db.Column(db.Boolean, default=False, nullable=False, index=True)
    error_code = db.Column(db.String(80), default='', nullable=False)
    error_msg = db.Column(db.Text, nullable=True)
    user_agent = db.Column(db.String(300), default='', nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now(), index=True)

class EvaluationProcess(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    autodiagnostico_completado = db.Column(db.Boolean, default=False, nullable=False)
    plan_evaluacion_aprobado = db.Column(db.Boolean, default=False, nullable=False)
    dictamen_competente = db.Column(db.Boolean, default=False, nullable=False)
    user = db.relationship('User', backref=db.backref('evaluation_process', uselist=False))

class AffiliateLead(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sponsor_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    prospect_name = db.Column(db.String(150), nullable=False)
    prospect_email = db.Column(db.String(150), unique=True, nullable=False)
    prospect_whatsapp = db.Column(db.String(20), nullable=True)
    ce_key = db.Column(db.String(50), nullable=True)
    status = db.Column(db.String(20), default='Pendiente', nullable=False)
    payment_link = db.Column(db.String(300), nullable=True)
    cfdi_pdf = db.Column(db.String(255), nullable=True)
    cfdi_xml = db.Column(db.String(255), nullable=True)
    payout_status = db.Column(db.String(50), default='Pendiente CFDI')
    commission_amount = db.Column(db.Float, nullable=True)
    selected_tier = db.Column(db.String(30), default='PRO_AFFILIATE')
    created_via = db.Column(db.String(30), default='manual')
    sponsor = db.relationship('User', backref=db.backref('affiliate_leads', lazy=True))

class AffiliateVideoView(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sponsor_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    ref_id = db.Column(db.String(50), nullable=True)
    ip = db.Column(db.String(64), nullable=True)
    user_agent = db.Column(db.String(300), nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now(), index=True)

class AffiliateCommission(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    lead_id = db.Column(db.Integer, db.ForeignKey('affiliate_lead.id'), nullable=True)
    sponsor_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    amount = db.Column(db.Float, nullable=False)
    commission_amount = db.Column(db.Float, nullable=False)
    payment_type = db.Column(db.String(30), nullable=False)
    stripe_session_id = db.Column(db.String(300), nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    lead = db.relationship('AffiliateLead', backref=db.backref('commissions', lazy=True))
    sponsor = db.relationship('User', foreign_keys=[sponsor_id], backref=db.backref('earned_commissions', lazy='dynamic'))

class StoredFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    filename = db.Column(db.String(300), nullable=False, index=True)
    content = db.Column(db.LargeBinary, nullable=False)
    content_type = db.Column(db.String(100), default='application/octet-stream')
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    file_category = db.Column(db.String(50), default='document')
    course_session_id = db.Column(db.Integer, nullable=True, index=True)
    user = db.relationship('User', backref=db.backref('stored_files', lazy='dynamic'))

class CourseSession(db.Model):
    __tablename__ = 'course_session'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    session_num = db.Column(db.Integer, nullable=False, default=1)
    topic = db.Column(db.String(300), nullable=True)
    master_doc = db.Column(db.Text, nullable=True)
    course_info_json = db.Column(db.Text, nullable=True)
    is_active = db.Column(db.Boolean, default=True, nullable=False, index=True)
    started_at = db.Column(db.DateTime, server_default=db.func.now())
    last_activity_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())
    # Marca explícita "esta sesión es del demo público" para reset automático
    # post-upgrade. Identificación por bandera (NO por coincidencia de título).
    is_demo = db.Column(db.Boolean, default=False, nullable=False, index=True)
    user = db.relationship('User', backref=db.backref('course_sessions', lazy='dynamic'))

class DatosPersonales(db.Model):
    """Etapa 2: Datos personales del candidato/usuario, persistidos del lado
    servidor (no en cookie). Alimentan el Autodiagnóstico (Tabla 1) y, más
    adelante, la Ficha de Registro y el Portafolio de Evidencias (E5).
    Un registro por usuario (upsert)."""
    __tablename__ = 'datos_personales'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, unique=True, index=True)
    nombre_completo = db.Column(db.String(200), default='')
    curp = db.Column(db.String(18), default='')
    domicilio = db.Column(db.String(300), default='')
    ultimo_grado = db.Column(db.String(150), default='')
    tel_casa = db.Column(db.String(30), default='')
    tel_celular = db.Column(db.String(30), default='')
    correo = db.Column(db.String(150), default='')
    fecha_aplicacion = db.Column(db.String(40), default='')
    lugar_nacimiento = db.Column(db.String(150), default='')
    nacionalidad = db.Column(db.String(80), default='')
    genero = db.Column(db.String(40), default='')
    fecha_nacimiento = db.Column(db.String(40), default='')
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())

class StripeProcessedEvent(db.Model):
    __tablename__ = 'stripe_processed_events'
    event_id = db.Column(db.String(255), primary_key=True)
    processed_at = db.Column(db.DateTime, default=datetime.utcnow)

class UserPurchase(db.Model):
    __tablename__ = 'user_purchases'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    payment_type = db.Column(db.String(40), nullable=False)
    amount_mxn = db.Column(db.Float, default=0.0)
    credits_granted = db.Column(db.Integer, default=0)
    stripe_session_id = db.Column(db.String(255), unique=True, nullable=True, index=True)
    source = db.Column(db.String(20), default='stripe_webhook')
    created_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)

class AdCampaign(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.String(500), default='')
    bg_gradient = db.Column(db.String(200), default='from-primary/15 to-primary/5')
    target_url = db.Column(db.String(500), nullable=False)
    status = db.Column(db.String(20), default='Pendiente', nullable=False)
    impressions = db.Column(db.Integer, default=0, nullable=False)
    clicks = db.Column(db.Integer, default=0, nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    advertiser_email = db.Column(db.String(150), nullable=True)
    total_cost = db.Column(db.Float, nullable=True)
    starts_at = db.Column(db.Date, nullable=True)
    ends_at = db.Column(db.Date, nullable=True)
    stripe_session_id = db.Column(db.String(300), nullable=True)
    signature_image = db.Column(db.LargeBinary, nullable=True)
    signing_ip = db.Column(db.String(50), nullable=True)
    advertiser_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    image_url = db.Column(db.String(500), nullable=True)

class AdIssue(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user_email = db.Column(db.String(150), nullable=False)
    issue_type = db.Column(db.String(50), nullable=False)
    description = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(20), default='Abierto', nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    admin_notes = db.Column(db.Text, default='')

class Config(db.Model):
    _cache = {}
    _CACHE_TTL = 60

    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(db.String(100), unique=True, nullable=False, index=True)
    value = db.Column(db.String(500), nullable=False)

    @staticmethod
    def get(key, default=None):
        now = _time.time()
        cached = Config._cache.get(key)
        if cached and (now - cached[1]) < Config._CACHE_TTL:
            return cached[0]
        row = Config.query.filter_by(key=key).first()
        val = row.value if row else default
        Config._cache[key] = (val, now)
        return val

    @staticmethod
    def set(key, value):
        row = Config.query.filter_by(key=key).first()
        if row:
            row.value = str(value)
        else:
            row = Config(key=key, value=str(value))
            db.session.add(row)
        db.session.commit()
        Config._cache[key] = (str(value), _time.time())

class BetaFeedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    session_id = db.Column(db.String(100), nullable=True)
    rating = db.Column(db.Integer, nullable=True)
    comment = db.Column(db.Text, nullable=True)
    page = db.Column(db.String(200), nullable=True)
    element_num = db.Column(db.Integer, nullable=True)
    generated_file = db.Column(db.String(300), nullable=True)
    feedback_type = db.Column(db.String(30), nullable=False, default='widget')
    created_at = db.Column(db.DateTime, server_default=db.func.now())

class BetaMetric(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    session_id = db.Column(db.String(100), nullable=True)
    metric_type = db.Column(db.String(50), nullable=False)
    element_num = db.Column(db.Integer, nullable=True)
    value_int = db.Column(db.Integer, nullable=True)
    value_text = db.Column(db.String(500), nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())

class ChatSpec(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True, index=True)
    session_id = db.Column(db.String(64), nullable=True, index=True)
    element_num = db.Column(db.Integer, nullable=False)
    content = db.Column(db.Text, nullable=False)
    filenames = db.Column(db.String(800), nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())

class ChatHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True, index=True)
    session_id = db.Column(db.String(64), nullable=True, index=True)
    element_num = db.Column(db.Integer, nullable=False)
    messages_json = db.Column(db.Text, nullable=False, default='[]')
    course_topic = db.Column(db.String(300), nullable=True)
    generated_files = db.Column(db.Text, nullable=True)
    course_session_id = db.Column(db.Integer, nullable=True, index=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())

class UserEvent(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    event_category = db.Column(db.String(50), nullable=False)
    event_action = db.Column(db.String(100), nullable=False)
    url = db.Column(db.String(500), nullable=True)
    metadata_json = db.Column(db.Text, nullable=True)
    ip_address = db.Column(db.String(50), nullable=True)
    user_agent = db.Column(db.String(500), nullable=True)
    timestamp = db.Column(db.DateTime, server_default=db.func.now())

class PexelsUsage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True, index=True)
    query = db.Column(db.String(300), nullable=True)
    photo_id = db.Column(db.String(40), nullable=True)
    photographer = db.Column(db.String(200), nullable=True)
    photographer_url = db.Column(db.String(500), nullable=True)
    photo_url = db.Column(db.String(500), nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now(), index=True)

CDMX_TZ = ZoneInfo('America/Mexico_City')

def _fmt_cdmx(dt, fmt='%Y-%m-%d %H:%M'):
    if not dt:
        return ''
    try:
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(CDMX_TZ).strftime(fmt)
    except Exception:
        try:
            return dt.strftime(fmt)
        except Exception:
            return ''

def track_event(category, action, user_id=None, extra_data=None):
    try:
        import json
        ip = None
        ua = None
        try:
            ip = request.headers.get('X-Forwarded-For', request.remote_addr)
            if ip and ',' in ip:
                ip = ip.split(',')[0].strip()
            ua = (request.user_agent.string or '')[:500]
        except Exception:
            pass
        evt = UserEvent(
            user_id=user_id,
            event_category=category,
            event_action=action,
            url=request.url if request else None,
            metadata_json=json.dumps(extra_data, ensure_ascii=False) if extra_data else None,
            ip_address=ip,
            user_agent=ua
        )
        db.session.add(evt)
        db.session.commit()
    except Exception:
        db.session.rollback()

class EmailLog(db.Model):
    __tablename__ = 'email_log'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True, index=True)
    recipient_email = db.Column(db.String(254), nullable=False, index=True)
    email_type = db.Column(db.String(40), nullable=False, index=True)
    subject = db.Column(db.String(300), nullable=True)
    document_filename = db.Column(db.String(300), nullable=True)
    sender_email = db.Column(db.String(254), nullable=True)
    status = db.Column(db.String(20), nullable=False, default='pending', index=True)
    smtp_response = db.Column(db.Text, nullable=True)
    error_message = db.Column(db.Text, nullable=True)
    attempted_at = db.Column(db.DateTime, server_default=db.func.now())
    sent_at = db.Column(db.DateTime, nullable=True)
    retry_count = db.Column(db.Integer, default=0, nullable=False)
    metadata_json = db.Column(db.Text, nullable=True)
    opened_at = db.Column(db.DateTime, nullable=True)
    open_count = db.Column(db.Integer, default=0, nullable=False)


EMAIL_REGEX_STRICT = re.compile(
    r'^[A-Za-z0-9._%+\-]+@[A-Za-z0-9](?:[A-Za-z0-9\-]{0,61}[A-Za-z0-9])?'
    r'(?:\.[A-Za-z0-9](?:[A-Za-z0-9\-]{0,61}[A-Za-z0-9])?)+$'
)
COMMON_TYPO_DOMAINS = {
    'gmial.com': 'gmail.com', 'gmai.com': 'gmail.com', 'gnail.com': 'gmail.com',
    'gmail.co': 'gmail.com', 'gmail.cm': 'gmail.com', 'gmail.con': 'gmail.com',
    'gmaill.com': 'gmail.com', 'gmail.om': 'gmail.com', 'gmali.com': 'gmail.com',
    'hotmial.com': 'hotmail.com', 'hotmai.com': 'hotmail.com', 'hotmial.es': 'hotmail.es',
    'hotmail.co': 'hotmail.com', 'hotmail.cm': 'hotmail.com', 'hotmail.con': 'hotmail.com',
    'hotmaill.com': 'hotmail.com', 'hotnail.com': 'hotmail.com',
    'yahooo.com': 'yahoo.com', 'yaho.com': 'yahoo.com', 'yahho.com': 'yahoo.com',
    'yahoo.co': 'yahoo.com', 'yhoo.com': 'yahoo.com', 'yahoo.com.mc': 'yahoo.com.mx',
    'outloo.com': 'outlook.com', 'outlokk.com': 'outlook.com', 'outlook.co': 'outlook.com',
    'iclud.com': 'icloud.com', 'icloud.co': 'icloud.com',
    'live.co': 'live.com', 'live.cm': 'live.com',
}
DISPOSABLE_DOMAINS = {
    'mailinator.com', 'tempmail.com', 'temp-mail.org', 'guerrillamail.com',
    'guerrillamail.info', 'guerrillamail.biz', 'sharklasers.com',
    '10minutemail.com', '10minutemail.net', 'yopmail.com', 'trashmail.com',
    'throwawaymail.com', 'maildrop.cc', 'getnada.com', 'fakeinbox.com',
    'mintemail.com', 'mohmal.com', 'tempinbox.com', 'mytemp.email',
    'spamgourmet.com', 'dispostable.com', 'mailcatch.com', 'inboxbear.com',
    'emailondeck.com', 'tempr.email', 'fake-mail.net', 'spamex.com',
    'mvrht.net', 'mt2015.com', 'mailnesia.com', 'spam4.me',
    'einrot.com', 'mailtemp.info', 'mintemail.com', 'mailbox.in.ua',
    'mailnull.com', 'tempmailer.com', 'tempemail.net', 'tempymail.com',
}

from collections import OrderedDict
_MX_CACHE = OrderedDict()
_MX_CACHE_TTL_SEC = 3600
_MX_CACHE_MAX = 2000

def _mx_cache_set(domain, ok, reason):
    import time as _t
    if domain in _MX_CACHE:
        try: _MX_CACHE.move_to_end(domain)
        except Exception: pass
    _MX_CACHE[domain] = (_t.time(), ok, reason)
    while len(_MX_CACHE) > _MX_CACHE_MAX:
        try: _MX_CACHE.popitem(last=False)
        except Exception: break

def _check_mx_record(domain):
    """Devuelve (ok: bool, motivo: str). ok=True si dominio acepta correo (MX o A fallback)."""
    import time as _t
    domain = (domain or '').lower().strip()
    if not domain:
        return False, 'dominio vacío'
    cached = _MX_CACHE.get(domain)
    if cached and (_t.time() - cached[0]) < _MX_CACHE_TTL_SEC:
        try: _MX_CACHE.move_to_end(domain)
        except Exception: pass
        return cached[1], cached[2]
    try:
        import dns.resolver
        resolver = dns.resolver.Resolver()
        resolver.lifetime = 4.0
        resolver.timeout = 4.0
        try:
            answers = resolver.resolve(domain, 'MX')
            if len(answers) > 0:
                _mx_cache_set(domain, True, 'MX OK')
                return True, 'MX OK'
        except dns.resolver.NoAnswer:
            try:
                resolver.resolve(domain, 'A')
                _mx_cache_set(domain, True, 'A fallback OK')
                return True, 'A fallback OK'
            except Exception:
                _mx_cache_set(domain, False, 'sin MX ni A')
                return False, 'el dominio no acepta correo (sin MX ni A)'
        except dns.resolver.NXDOMAIN:
            _mx_cache_set(domain, False, 'NXDOMAIN')
            return False, 'el dominio no existe'
        except dns.resolver.NoNameservers:
            return True, 'DNS sin respuesta — permitido'
        except Exception as e:
            return True, f'DNS error tolerado: {e}'
    except Exception as e:
        return True, f'verificación MX no disponible: {e}'
    return True, 'OK'

def validate_email_full(email):
    """Devuelve (ok: bool, mensaje_usuario: str, motivo_log: str).
    A: regex estricta. B: MX. C: dominio desechable + typos comunes.
    Diseñado para ser fail-open ante errores de red (no bloquea por DNS caído)."""
    e = (email or '').strip().lower()
    if not e:
        return False, 'El correo es obligatorio.', 'vacío'
    if len(e) > 254:
        return False, 'El correo es demasiado largo.', 'longitud >254'
    if not EMAIL_REGEX_STRICT.match(e):
        return False, 'El formato del correo no es válido. Revisa que no tenga espacios ni caracteres extraños.', 'regex_fail'
    try:
        local, domain = e.rsplit('@', 1)
    except ValueError:
        return False, 'El correo debe contener un @.', 'sin_arroba'
    if domain in COMMON_TYPO_DOMAINS:
        suggested = COMMON_TYPO_DOMAINS[domain]
        return False, f'¿Quisiste decir "{local}@{suggested}"? Detectamos un error de tipeo común.', f'typo:{domain}->{suggested}'
    if domain in DISPOSABLE_DOMAINS:
        return False, 'No aceptamos correos desechables. Por favor usa tu correo personal o profesional.', f'disposable:{domain}'
    mx_ok, mx_reason = _check_mx_record(domain)
    if not mx_ok:
        return False, f'El dominio "{domain}" no parece recibir correos. Verifica que esté bien escrito.', f'mx_fail:{mx_reason}'
    return True, '', f'ok ({mx_reason})'


def _split_full_name(full_name):
    """Heurística para desglosar nombres mexicanos a {first_name, apellido_paterno, apellido_materno}.
    Convención asumida: 'Nombre(s) ApellidoPaterno ApellidoMaterno'.
    - 1 token  -> solo first_name.
    - 2 tokens -> first_name + apellido_paterno.
    - 3 tokens -> first_name + paterno + materno.
    - 4+ tokens -> primeros (n-2) tokens = first_name; últimos 2 = paterno + materno.
    Es 'best-effort'; el admin puede corregir manualmente desde el panel."""
    if not full_name:
        return {'first_name': '', 'apellido_paterno': '', 'apellido_materno': ''}
    parts = [p for p in str(full_name).strip().split() if p]
    n = len(parts)
    if n == 0:
        return {'first_name': '', 'apellido_paterno': '', 'apellido_materno': ''}
    if n == 1:
        return {'first_name': parts[0][:80], 'apellido_paterno': '', 'apellido_materno': ''}
    if n == 2:
        return {'first_name': parts[0][:80], 'apellido_paterno': parts[1][:80], 'apellido_materno': ''}
    if n == 3:
        return {'first_name': parts[0][:80], 'apellido_paterno': parts[1][:80], 'apellido_materno': parts[2][:80]}
    return {
        'first_name': ' '.join(parts[:-2])[:80],
        'apellido_paterno': parts[-2][:80],
        'apellido_materno': parts[-1][:80],
    }


def _log_email_attempt(email_type, recipient, sender=None, subject=None, user_id=None, document_filename=None, extra=None):
    """Inserta registro 'pending' en email_log usando una conexión aislada (no toca db.session
    para evitar commits accidentales de cambios pendientes en la transacción del request).
    Devuelve id o None si falla. Nunca debe romper el envío."""
    try:
        import json as _json
        params = {
            'user_id': user_id,
            'recipient_email': (recipient or '')[:254],
            'email_type': email_type,
            'subject': (subject or '')[:300] if subject else None,
            'document_filename': (document_filename or '')[:300] if document_filename else None,
            'sender_email': (sender or '')[:254] if sender else None,
            'status': 'pending',
            'metadata_json': _json.dumps(extra, ensure_ascii=False) if extra else None,
        }
        with db.engine.begin() as conn:
            row = conn.execute(db.text(
                "INSERT INTO email_log (user_id, recipient_email, email_type, subject, document_filename, sender_email, status, metadata_json, attempted_at, retry_count) "
                "VALUES (:user_id, :recipient_email, :email_type, :subject, :document_filename, :sender_email, :status, :metadata_json, NOW(), 0) "
                "RETURNING id"
            ), params).fetchone()
        return row[0] if row else None
    except Exception as _e:
        logger.warning(f"EmailLog: no se pudo crear registro pending ({email_type} -> {recipient}): {_e}")
        return None


def _log_email_result(log_id, success, error=None, smtp_response=None):
    """Actualiza registro a 'sent' o 'failed' usando una conexión aislada (no toca db.session).
    No-op silencioso si falla."""
    if not log_id:
        return
    try:
        params = {'log_id': log_id}
        if success:
            params['status'] = 'sent'
            params['error_message'] = None
            params['smtp_response'] = str(smtp_response)[:2000] if smtp_response else None
            sql = ("UPDATE email_log SET status=:status, sent_at=NOW(), error_message=:error_message, "
                   "smtp_response=COALESCE(:smtp_response, smtp_response) WHERE id=:log_id")
        else:
            params['status'] = 'failed'
            params['error_message'] = (str(error) if error else 'Unknown error')[:2000]
            params['smtp_response'] = str(smtp_response)[:2000] if smtp_response else None
            sql = ("UPDATE email_log SET status=:status, error_message=:error_message, "
                   "smtp_response=COALESCE(:smtp_response, smtp_response) WHERE id=:log_id")
        with db.engine.begin() as conn:
            conn.execute(db.text(sql), params)
    except Exception as _e:
        logger.warning(f"EmailLog: no se pudo actualizar registro {log_id}: {_e}")


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

def _get_table_columns(conn, table_name):
    dialect = db.engine.dialect.name
    if dialect == 'sqlite':
        result = conn.execute(db.text(f"PRAGMA table_info({table_name})"))
        return [row[1] for row in result.fetchall()]
    else:
        result = conn.execute(db.text(
            "SELECT column_name FROM information_schema.columns "
            "WHERE table_name = :tname"
        ), {"tname": table_name})
        return [row[0] for row in result.fetchall()]

with app.app_context():
    db.create_all()
    with db.engine.connect() as conn:
        # Permitir varias cuentas distintas con el mismo correo (creadas solo por admin).
        # Elimina el constraint/índice único de correo si aún existe. Idempotente.
        try:
            conn.execute(db.text('ALTER TABLE "user" DROP CONSTRAINT IF EXISTS user_email_key'))
            conn.commit()
        except Exception:
            conn.rollback()
        try:
            conn.execute(db.text('CREATE INDEX IF NOT EXISTS ix_user_email ON "user" (email)'))
            conn.commit()
        except Exception:
            conn.rollback()
        columns = _get_table_columns(conn, 'user')
        if columns and 'chat_usage_count' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN chat_usage_count INTEGER DEFAULT 0 NOT NULL'))
            conn.commit()
        if columns and 'free_downloads_used' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN free_downloads_used INTEGER DEFAULT 0 NOT NULL'))
            conn.commit()
        if columns and 'is_affiliate' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN is_affiliate BOOLEAN DEFAULT FALSE'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN affiliate_terms_accepted BOOLEAN DEFAULT FALSE'))
            conn.execute(db.text("ALTER TABLE \"user\" ADD COLUMN tax_regime VARCHAR(20) DEFAULT 'RESICO'"))
            conn.commit()
        if columns and 'tier_usage_stats' not in columns:
            conn.execute(db.text("ALTER TABLE \"user\" ADD COLUMN tier_usage_stats TEXT DEFAULT '{}'"))
            conn.commit()
        al_columns = _get_table_columns(conn, 'affiliate_lead')
        if al_columns and 'cfdi_pdf' not in al_columns:
            conn.execute(db.text("ALTER TABLE affiliate_lead ADD COLUMN cfdi_pdf VARCHAR(255)"))
            conn.execute(db.text("ALTER TABLE affiliate_lead ADD COLUMN cfdi_xml VARCHAR(255)"))
            conn.execute(db.text("ALTER TABLE affiliate_lead ADD COLUMN payout_status VARCHAR(50) DEFAULT 'Pendiente CFDI'"))
            conn.commit()
        if al_columns and 'commission_amount' not in al_columns:
            conn.execute(db.text("ALTER TABLE affiliate_lead ADD COLUMN commission_amount FLOAT"))
            conn.commit()
        if al_columns and 'selected_tier' not in al_columns:
            conn.execute(db.text("ALTER TABLE affiliate_lead ADD COLUMN selected_tier VARCHAR(30) DEFAULT 'PRO_AFFILIATE'"))
            conn.execute(db.text("ALTER TABLE affiliate_lead ADD COLUMN created_via VARCHAR(30) DEFAULT 'manual'"))
            conn.commit()
        elif al_columns and 'created_via' in al_columns:
                    pass

        ac_columns = set()
        try:
            ac_result = conn.execute(db.text("SELECT column_name FROM information_schema.columns WHERE table_name='affiliate_commission'"))
            ac_columns = {r[0] for r in ac_result}
        except Exception:
            pass
        if ac_columns and 'user_id' in ac_columns:
            try:
                conn.execute(db.text("ALTER TABLE affiliate_commission ALTER COLUMN user_id DROP NOT NULL"))
                conn.commit()
            except Exception:
                conn.rollback()

        if columns and 'alacarte_e1' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN alacarte_e1 INTEGER DEFAULT 0 NOT NULL'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN alacarte_e2 INTEGER DEFAULT 0 NOT NULL'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN alacarte_e3 INTEGER DEFAULT 0 NOT NULL'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN alacarte_e4 INTEGER DEFAULT 0 NOT NULL'))
            conn.commit()

        if columns and 'alacarte_e5' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN alacarte_e5 INTEGER DEFAULT 0 NOT NULL'))
            conn.commit()

        if columns and 'e5_preview_used_at' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN e5_preview_used_at TIMESTAMP NULL'))
            conn.commit()

        if columns and 'ec0217_grants' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN ec0217_grants INTEGER DEFAULT 0 NOT NULL'))
            conn.commit()

        try:
            pe_cols = _get_table_columns(conn, 'portafolio_evaluacion')
        except Exception:
            pe_cols = []
        if pe_cols and 'e5_unlocked' not in pe_cols:
            conn.execute(db.text('ALTER TABLE portafolio_evaluacion ADD COLUMN e5_unlocked BOOLEAN DEFAULT FALSE NOT NULL'))
            conn.execute(db.text('ALTER TABLE portafolio_evaluacion ADD COLUMN e5_unlocked_at TIMESTAMP NULL'))
            conn.execute(db.text("ALTER TABLE portafolio_evaluacion ADD COLUMN e5_unlocked_via VARCHAR(20) DEFAULT '' NOT NULL"))
            conn.commit()
        if pe_cols and 'e5_mode' not in pe_cols:
            conn.execute(db.text("ALTER TABLE portafolio_evaluacion ADD COLUMN e5_mode VARCHAR(10) DEFAULT '' NOT NULL"))
            conn.execute(db.text("ALTER TABLE portafolio_evaluacion ADD COLUMN video_preview_seg INTEGER DEFAULT 0 NOT NULL"))
            conn.commit()

        if columns and 'created_at' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN created_at TIMESTAMP DEFAULT NOW()'))
            conn.commit()
        if columns and 'rfc' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN rfc VARCHAR(20)'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN razon_social VARCHAR(250)'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN domicilio_fiscal VARCHAR(300)'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN nombre_representante_legal VARCHAR(200)'))
            conn.commit()

        if columns and 'pro_courses_remaining' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN pro_courses_remaining INTEGER DEFAULT 0 NOT NULL'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN pro_active_course VARCHAR(300)'))
            conn.commit()

        if columns and 'referred_by' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN referred_by INTEGER'))
            conn.commit()

        try:
            email_log_cols = _get_table_columns(conn, 'email_log')
            for _ecol, _etype in [
                ('opened_at', 'TIMESTAMP'),
                ('open_count', 'INTEGER DEFAULT 0 NOT NULL'),
            ]:
                if email_log_cols and _ecol not in email_log_cols:
                    try:
                        conn.execute(db.text(f'ALTER TABLE email_log ADD COLUMN {_ecol} {_etype}'))
                        conn.commit()
                    except Exception as _e:
                        conn.rollback()
                        logger.warning(f"email_log migration {_ecol} skipped: {_e}")
        except Exception as _e:
            logger.warning(f"email_log column check skipped: {_e}")

        # Migración: needs_state_reset_at (señal de hard-reset al frontend
        # tras upgrade FREE→PRO). Idempotente.
        if columns and 'needs_state_reset_at' not in columns:
            try:
                conn.execute(db.text('ALTER TABLE "user" ADD COLUMN needs_state_reset_at TIMESTAMP NULL'))
                conn.commit()
            except Exception as _e:
                conn.rollback()
                logger.warning(f"needs_state_reset_at migration skipped: {_e}")

        # Migración: course_session.is_demo (marca sesiones demo para purgar)
        try:
            cs_cols = _get_table_columns(conn, 'course_session')
        except Exception:
            cs_cols = []
        if cs_cols and 'is_demo' not in cs_cols:
            try:
                conn.execute(db.text('ALTER TABLE course_session ADD COLUMN is_demo BOOLEAN DEFAULT FALSE NOT NULL'))
                conn.execute(db.text('CREATE INDEX IF NOT EXISTS ix_course_session_is_demo ON course_session (is_demo)'))
                conn.commit()
            except Exception as _e:
                conn.rollback()
                logger.warning(f"course_session.is_demo migration skipped: {_e}")

        if columns and 'mkt_unsubscribed_at' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN mkt_unsubscribed_at TIMESTAMP'))
            conn.commit()

        if columns and 'marketing_consent_source' not in columns:
            try:
                conn.execute(db.text('ALTER TABLE "user" ADD COLUMN IF NOT EXISTS marketing_consent_source VARCHAR(20)'))
                conn.commit()
            except Exception:
                conn.rollback()
            columns.append('marketing_consent_source')

        if columns:
            for _bcol, _btype in [
                ('email_bounced_at', 'TIMESTAMP'),
                ('email_bounce_type', 'VARCHAR(10)'),
                ('email_bounce_note', 'VARCHAR(255)'),
                ('ultimo_contacto_at', 'TIMESTAMP'),
                ('notas_contacto', 'TEXT'),
            ]:
                if _bcol not in columns:
                    conn.execute(db.text(f'ALTER TABLE "user" ADD COLUMN {_bcol} {_btype}'))
                    conn.commit()
                    columns.append(_bcol)

        if columns and 'first_name' not in columns:
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN first_name VARCHAR(80)'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN apellido_paterno VARCHAR(80)'))
            conn.execute(db.text('ALTER TABLE "user" ADD COLUMN apellido_materno VARCHAR(80)'))
            conn.commit()

        try:
            pending_rows = conn.execute(db.text(
                'SELECT id, full_name FROM "user" '
                "WHERE full_name IS NOT NULL AND full_name <> '' "
                'AND (first_name IS NULL OR first_name = \'\')'
            )).fetchall()
            _bf_ok, _bf_fail = 0, 0
            for _row in pending_rows:
                try:
                    _parsed = _split_full_name(_row[1])
                    conn.execute(db.text(
                        'UPDATE "user" SET first_name=:fn, apellido_paterno=:ap, apellido_materno=:am WHERE id=:uid'
                    ), {
                        'fn': _parsed['first_name'] or None,
                        'ap': _parsed['apellido_paterno'] or None,
                        'am': _parsed['apellido_materno'] or None,
                        'uid': _row[0],
                    })
                    conn.commit()
                    _bf_ok += 1
                except Exception as _row_e:
                    conn.rollback()
                    _bf_fail += 1
                    logger.warning(f"Backfill desglose: fila id={_row[0]} falló: {_row_e}")
            if pending_rows:
                logger.info(f"Backfill desglose de nombres: {_bf_ok} ok, {_bf_fail} fallidas (de {len(pending_rows)} pendientes).")
        except Exception as _bf_e:
            try: conn.rollback()
            except Exception: pass
            logger.warning(f"Backfill desglose de nombres no pudo iniciarse: {_bf_e}")

        ue_columns = _get_table_columns(conn, 'user_event')
        if ue_columns and 'ip_address' not in ue_columns:
            conn.execute(db.text("ALTER TABLE user_event ADD COLUMN ip_address VARCHAR(50)"))
            conn.commit()
        if ue_columns and 'user_agent' not in ue_columns:
            conn.execute(db.text("ALTER TABLE user_event ADD COLUMN user_agent VARCHAR(500)"))
            conn.commit()

        u_columns = _get_table_columns(conn, 'user')
        _utm_cols = [
            ('utm_source', 'VARCHAR(80)'),
            ('utm_medium', 'VARCHAR(80)'),
            ('utm_campaign', 'VARCHAR(150)'),
            ('utm_term', 'VARCHAR(200)'),
            ('utm_content', 'VARCHAR(250)'),
            ('utm_id', 'VARCHAR(50)'),
            ('utm_landing', 'VARCHAR(120)'),
            ('utm_captured_at', 'TIMESTAMP'),
        ]
        for _cn, _ct in _utm_cols:
            if u_columns and _cn not in u_columns:
                try:
                    conn.execute(db.text(f'ALTER TABLE "user" ADD COLUMN {_cn} {_ct}'))
                    conn.commit()
                except Exception as _e_utm:
                    try: conn.rollback()
                    except Exception: pass
                    logger.warning(f"ALTER user ADD {_cn} falló: {_e_utm}")
        try:
            conn.execute(db.text('CREATE INDEX IF NOT EXISTS ix_user_utm_content ON "user"(utm_content)'))
            conn.execute(db.text('CREATE INDEX IF NOT EXISTS ix_user_utm_term ON "user"(utm_term)'))
            conn.execute(db.text('CREATE INDEX IF NOT EXISTS ix_user_utm_campaign ON "user"(utm_campaign)'))
            conn.commit()
        except Exception:
            try: conn.rollback()
            except Exception: pass

        ac_columns = _get_table_columns(conn, 'ad_campaign')
        if ac_columns and 'image_url' not in ac_columns:
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN image_url VARCHAR(500)"))
            conn.commit()
        if ac_columns and 'total_cost' not in ac_columns:
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN total_cost FLOAT"))
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN starts_at DATE"))
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN ends_at DATE"))
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN stripe_session_id VARCHAR(300)"))
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN signature_image BYTEA"))
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN signing_ip VARCHAR(50)"))
            conn.execute(db.text("ALTER TABLE ad_campaign ADD COLUMN advertiser_id INTEGER REFERENCES \"user\"(id)"))
            conn.commit()

    with db.engine.connect() as conn:
        try:
            conn.execute(db.text(
                "CREATE UNIQUE INDEX IF NOT EXISTS ux_chat_spec_user_elem "
                "ON chat_spec (user_id, element_num) WHERE user_id IS NOT NULL"))
            conn.execute(db.text(
                "CREATE UNIQUE INDEX IF NOT EXISTS ux_chat_spec_sid_elem "
                "ON chat_spec (session_id, element_num) WHERE session_id IS NOT NULL AND user_id IS NULL"))
            conn.commit()
        except Exception as _e:
            logger.warning(f"chat_spec unique index create failed: {_e}")

    with db.engine.connect() as conn:
        try:
            user_cols_now = _get_table_columns(conn, 'user')
            if user_cols_now and 'active_course_session_id' not in user_cols_now:
                conn.execute(db.text('ALTER TABLE "user" ADD COLUMN active_course_session_id INTEGER'))
                conn.commit()
            ch_cols_now = _get_table_columns(conn, 'chat_history')
            if ch_cols_now and 'course_session_id' not in ch_cols_now:
                conn.execute(db.text('ALTER TABLE chat_history ADD COLUMN course_session_id INTEGER'))
                conn.commit()
            sf_cols_now = _get_table_columns(conn, 'stored_file')
            if sf_cols_now and 'course_session_id' not in sf_cols_now:
                conn.execute(db.text('ALTER TABLE stored_file ADD COLUMN course_session_id INTEGER'))
                conn.commit()
            try:
                conn.execute(db.text("CREATE INDEX IF NOT EXISTS ix_chat_history_user_elem_cs ON chat_history (user_id, element_num, course_session_id)"))
                conn.execute(db.text("CREATE INDEX IF NOT EXISTS ix_stored_file_user_cs ON stored_file (user_id, course_session_id)"))
                conn.execute(db.text("CREATE INDEX IF NOT EXISTS ix_course_session_user_active ON course_session (user_id, is_active)"))
                conn.execute(db.text("CREATE UNIQUE INDEX IF NOT EXISTS uq_course_session_user_num ON course_session (user_id, session_num)"))
                conn.execute(db.text("CREATE UNIQUE INDEX IF NOT EXISTS uq_course_session_user_active_one ON course_session (user_id) WHERE is_active = TRUE"))
                conn.execute(db.text("CREATE UNIQUE INDEX IF NOT EXISTS uq_chat_history_user_elem_cs ON chat_history (user_id, element_num, course_session_id) WHERE user_id IS NOT NULL AND course_session_id IS NOT NULL"))
                conn.commit()
            except Exception as _eidx:
                conn.rollback()
                logger.warning(f"course_session indexes skipped: {_eidx}")
        except Exception as _emig:
            try: conn.rollback()
            except Exception: pass
            logger.warning(f"course_session migration failed: {_emig}")

        try:
            lock_row = conn.execute(db.text("SELECT pg_try_advisory_lock(778899001)")).fetchone()
            got_lock = bool(lock_row and lock_row[0])
            if not got_lock:
                logger.info("course_session backfill: skipped (another worker holds the advisory lock)")
            else:
                try:
                    backfill_users = conn.execute(db.text(
                        'SELECT DISTINCT u.id FROM "user" u '
                        'WHERE u.active_course_session_id IS NULL AND ('
                        '  EXISTS (SELECT 1 FROM chat_history ch WHERE ch.user_id = u.id) '
                        '  OR EXISTS (SELECT 1 FROM stored_file sf WHERE sf.user_id = u.id)'
                        ') AND NOT EXISTS (SELECT 1 FROM course_session cs WHERE cs.user_id = u.id)'
                    )).fetchall()
                    for (uid,) in backfill_users:
                        topic_row = conn.execute(db.text(
                            'SELECT course_topic FROM chat_history WHERE user_id = :uid AND course_topic IS NOT NULL '
                            'ORDER BY updated_at DESC NULLS LAST, id DESC LIMIT 1'
                        ), {'uid': uid}).fetchone()
                        topic_val = (topic_row[0] if topic_row and topic_row[0] else None)
                        ins = conn.execute(db.text(
                            'INSERT INTO course_session (user_id, session_num, topic, is_active) '
                            'SELECT :uid, 1, :topic, TRUE '
                            'WHERE NOT EXISTS (SELECT 1 FROM course_session WHERE user_id = :uid) '
                            'RETURNING id'
                        ), {'uid': uid, 'topic': topic_val})
                        row_ins = ins.fetchone()
                        if not row_ins:
                            existing = conn.execute(db.text(
                                'SELECT id FROM course_session WHERE user_id = :uid ORDER BY id ASC LIMIT 1'
                            ), {'uid': uid}).fetchone()
                            new_cs_id = existing[0] if existing else None
                        else:
                            new_cs_id = row_ins[0]
                        if not new_cs_id:
                            conn.commit()
                            continue
                        conn.execute(db.text('UPDATE "user" SET active_course_session_id = :cs WHERE id = :uid AND active_course_session_id IS NULL'),
                                     {'cs': new_cs_id, 'uid': uid})
                        conn.execute(db.text('UPDATE chat_history SET course_session_id = :cs WHERE user_id = :uid AND course_session_id IS NULL'),
                                     {'cs': new_cs_id, 'uid': uid})
                        conn.execute(db.text('UPDATE stored_file SET course_session_id = :cs WHERE user_id = :uid AND course_session_id IS NULL'),
                                     {'cs': new_cs_id, 'uid': uid})
                        conn.commit()
                    if backfill_users:
                        logger.info(f"course_session backfill: {len(backfill_users)} users provisioned with session_num=1")
                    try:
                        upd_ch = conn.execute(db.text(
                            'UPDATE chat_history SET course_session_id = u.active_course_session_id '
                            'FROM "user" u WHERE chat_history.user_id = u.id '
                            'AND chat_history.course_session_id IS NULL AND u.active_course_session_id IS NOT NULL'
                        ))
                        upd_sf = conn.execute(db.text(
                            'UPDATE stored_file SET course_session_id = u.active_course_session_id '
                            'FROM "user" u WHERE stored_file.user_id = u.id '
                            'AND stored_file.course_session_id IS NULL AND u.active_course_session_id IS NOT NULL'
                        ))
                        conn.commit()
                        logger.info(f"course_session backfill (post-release NULLs): chat_history={upd_ch.rowcount} stored_file={upd_sf.rowcount}")
                    except Exception as _enull:
                        try: conn.rollback()
                        except Exception: pass
                        logger.warning(f"course_session backfill (post-release NULLs) skipped: {_enull}")
                finally:
                    try:
                        conn.execute(db.text("SELECT pg_advisory_unlock(778899001)"))
                        conn.commit()
                    except Exception:
                        pass
        except Exception as _ebf:
            try: conn.rollback()
            except Exception: pass
            logger.warning(f"course_session backfill skipped: {_ebf}")

    if not Config.get('COMISION_PORCENTAJE'):
        Config.set('COMISION_PORCENTAJE', '0.30')
    if not Config.get('COST_PER_DAY_ADS'):
        Config.set('COST_PER_DAY_ADS', '150')
    if not Config.get('MAX_CONCURRENT_ADS'):
        Config.set('MAX_CONCURRENT_ADS', '5')
    if not Config.get('MAX_CONCURRENT_AI'):
        Config.set('MAX_CONCURRENT_AI', '3')

    logger.info("Pre-cargando documentos de referencia EC0301 al arranque...")
    reference_docs_cache = load_all_reference_docs()
    pass

    from datetime import datetime, date
    admin_user = User.query.filter_by(email=ADMIN_EMAIL).first()
    if not admin_user:
        admin_user = User(
            email=ADMIN_EMAIL,
            full_name='Arturo García',
            tier='PREMIUM',
            whatsapp='5591046391',
            terms_accepted=True,
            marketing_consent=False,
            normative_agreement_accepted=True,
            chat_usage_count=0,
            free_downloads_used=0
        )
        admin_user.set_password('gYC43Nz3yqBEkEF')
        db.session.add(admin_user)
        db.session.commit()
        logger.info(f"Admin user created: {ADMIN_EMAIL}")
    else:
        admin_user.set_password('gYC43Nz3yqBEkEF')
        if admin_user.tier != 'PREMIUM':
            admin_user.tier = 'PREMIUM'
        db.session.commit()
        logger.info(f"Admin user password & tier ensured: {ADMIN_EMAIL}")

    martha = User.query.filter_by(email='martha_ville@hotmail.com').first()
    if not martha:
        martha = User(
            email='martha_ville@hotmail.com',
            full_name='MARTHA VILLEGAS',
            tier='PRO',
            whatsapp='312 157 6507',
            terms_accepted=True,
            marketing_consent=False,
            normative_agreement_accepted=False,
            chat_usage_count=2,
            free_downloads_used=0,
            is_affiliate=False,
            affiliate_terms_accepted=False,
            tax_regime='RESICO',
            tier_usage_stats='{"FREE": 4886}'
        )
        martha.password_hash = 'scrypt:32768:8:1$wZPVYMg9CM0NRJ5e$38507aeecdddb2146287d505c83c3d48857c068ad78b19ee4952a49d206565816c3da0c7da221b9c7d24a0c4248332ee9346637b10f23c6d42d96a449aeab7e5'
        db.session.add(martha)
        db.session.commit()
        logger.info("User Martha Villegas migrated to this environment")

    adv_id = admin_user.id if admin_user else None
    seed_ads = [
        {
            'title': 'Crea tu Campaña en pocos clics',
            'description': 'La descripción aparecerá aquí..',
            'bg_gradient': 'from-primary/15 to-primary/5',
            'target_url': 'https://app.pertinentia.com',
            'status': 'Activo',
            'impressions': 0, 'clicks': 0,
            'created_at': datetime(2026, 3, 24, 5, 24, 36),
            'advertiser_email': 'arturogarciac@pertinentia.com',
            'total_cost': 500.0,
            'starts_at': date(2026, 3, 24),
            'ends_at': date(2026, 5, 13),
            'signing_ip': '127.0.0.1',
            'advertiser_id': adv_id,
            'image_url': '/static/uploads/ads/312f5511007646abb5de5f4d9cc71e46.png'
        },
        {
            'title': 'PertinentIA\u00ae',
            'description': 'La App para la CAPACITACI\u00d3N CERTIFICADA en M\u00e9xico',
            'bg_gradient': 'from-primary/15 to-primary/5',
            'target_url': 'https://app.pertinentia.com',
            'status': 'Activo',
            'impressions': 0, 'clicks': 0,
            'created_at': datetime(2026, 3, 24, 18, 43, 47),
            'advertiser_email': 'arturogarciac@pertinentia.com',
            'total_cost': 500.0,
            'starts_at': date(2026, 3, 24),
            'ends_at': date(2026, 5, 13),
            'signing_ip': '127.0.0.1',
            'advertiser_id': adv_id,
            'image_url': '/static/uploads/ads/3619ff9894474bb0b81349fc681863ff.png'
        },
        {
            'title': 'Virtual Labor Inspector',
            'description': 'Auditoría laboral inteligente con IA',
            'bg_gradient': 'from-primary/15 to-primary/5',
            'target_url': 'https://app.pertinentia.com',
            'status': 'Activo',
            'impressions': 0, 'clicks': 0,
            'created_at': datetime(2026, 3, 25, 3, 37, 0),
            'advertiser_email': 'arturogarciac@pertinentia.com',
            'total_cost': 500.0,
            'starts_at': date(2026, 3, 25),
            'ends_at': date(2026, 5, 14),
            'signing_ip': '127.0.0.1',
            'advertiser_id': adv_id,
            'image_url': '/static/uploads/ads/872d2f7a8a274a588072ee70a625eb4a.png'
        },
    ]
    existing_titles = {a.title for a in AdCampaign.query.all()}
    added = 0
    for ad_data in seed_ads:
        if ad_data['title'] not in existing_titles:
            db.session.add(AdCampaign(**ad_data))
            added += 1
    if added:
        db.session.commit()
        logger.info(f'Synced {added} ad campaign(s) to database.')

@app.template_filter('strip_user_prefix')
def strip_user_prefix_filter(filename):
    return re.sub(r'^u\d+_', '', filename)

@app.template_filter('fmt_cdmx')
def fmt_cdmx_filter(dt, fmt='%d/%m/%Y'):
    return _fmt_cdmx(dt, fmt)

def _prefix_user_file(filepath):
    if not filepath or not os.path.isfile(filepath):
        return filepath
    directory = os.path.dirname(filepath)
    basename = os.path.basename(filepath)
    prefixed = f"u{current_user.id}_{basename}"
    new_path = os.path.join(directory, prefixed)
    os.rename(filepath, new_path)
    return new_path

def _sanitize_topic_for_filename(topic):
    if not topic:
        return ""
    clean = re.sub(r'[^\w\s\-]', '', topic, flags=re.UNICODE)
    clean = re.sub(r'\s+', '_', clean.strip())
    clean = clean[:40]
    return clean

def _prefix_anon_file(filepath, course_topic=None):
    if not filepath or not os.path.isfile(filepath):
        return filepath
    directory = os.path.dirname(filepath)
    basename = os.path.basename(filepath)
    topic_slug = _sanitize_topic_for_filename(course_topic)
    if topic_slug:
        name_part, ext = os.path.splitext(basename)
        prefixed = f"{name_part}_{topic_slug}{ext}"
    else:
        prefixed = f"anon_{basename}"
    new_path = os.path.join(directory, prefixed)
    os.rename(filepath, new_path)
    return new_path

def _persist_file_to_db(filepath, user_id, category='document', course_session_id=None):
    if not filepath or not os.path.isfile(filepath):
        return
    filename = os.path.basename(filepath)
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    ct_map = {'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
              'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
              'pdf': 'application/pdf', 'xml': 'application/xml'}
    content_type = ct_map.get(ext, 'application/octet-stream')
    with open(filepath, 'rb') as f:
        file_data = f.read()
    existing = StoredFile.query.filter_by(user_id=user_id, filename=filename).first()
    cs_id_for_file = course_session_id
    if cs_id_for_file is None:
        try:
            _u = User.query.get(user_id)
            if _u is not None:
                cs_id_for_file = _u.active_course_session_id
        except Exception:
            cs_id_for_file = None
    if existing:
        existing.content = file_data
        existing.content_type = content_type
        existing.file_category = category
        if cs_id_for_file and not existing.course_session_id:
            existing.course_session_id = cs_id_for_file
    else:
        sf = StoredFile(user_id=user_id, filename=filename, content=file_data,
                        content_type=content_type, file_category=category,
                        course_session_id=cs_id_for_file)
        db.session.add(sf)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        raise

def _get_file_from_db(filename, user_id=None):
    q = StoredFile.query.filter_by(filename=filename)
    if user_id is not None:
        q = q.filter_by(user_id=user_id)
    return q.first()

def _delete_file_from_db(filename, user_id=None):
    q = StoredFile.query.filter_by(filename=filename)
    if user_id is not None:
        q = q.filter_by(user_id=user_id)
    sf = q.first()
    if sf:
        db.session.delete(sf)
        db.session.commit()
        return True
    return False

def _list_user_files_from_db(user_id, category=None):
    q = StoredFile.query.filter_by(user_id=user_id)
    if category:
        q = q.filter_by(file_category=category)
    return [sf.filename for sf in q.order_by(StoredFile.created_at.desc()).all()]

def _user_generated_docs():
    prefix = f"u{current_user.id}_"
    db_docs = _list_user_files_from_db(current_user.id, category='document')
    disk_docs = [d for d in list_generated_docs() if d.startswith(prefix)]
    all_docs = list(dict.fromkeys(db_docs + disk_docs))
    return all_docs

def get_reference_docs():
    global reference_docs_cache
    if reference_docs_cache is None:
        logger.warning("reference_docs_cache was None — reloading (should not happen after startup)")
        reference_docs_cache = load_all_reference_docs()
    return reference_docs_cache

ELEMENT_CONFIG = {
    1: {
        "page_title": "Carta Descriptiva",
        "page_description": "Elemento 1: Diseña la carta descriptiva de tu curso conforme al EC0301",
        "active_page": "e1",
        "welcome_message": "Bienvenido al Módulo de Carta Descriptiva. Indícame el tema de tu curso y generaré automáticamente la Carta Descriptiva estructurada y el Contrato de Aprendizaje.",
        "welcome_message_ready": "Listo para generar los productos de Carta Descriptiva para tu curso. Selecciona una acción rápida o escribe tu solicitud.",
        "quick_actions": [
            {"label": "Carta Descriptiva Completa", "prompt": "Genera la carta descriptiva completa de mi curso con todos los campos requeridos por el EC0301: información general, objetivo general, objetivos particulares (cognitivo, psicomotor, afectivo), requerimientos, estrategias de evaluación, y las 3 etapas (apertura, desarrollo, cierre) con tiempos, técnicas y materiales."},
            {"label": "Contrato de Aprendizaje", "prompt": "Genera el Contrato de Aprendizaje para mi curso. Incluye una lista clara de Compromisos del Instructor (puntualidad, dominio del tema, retroalimentación, materiales, evaluación justa) y Compromisos del Participante (asistencia, participación activa, cumplimiento de actividades, respeto, entrega de evidencias). Alineados al tema del curso y estructurados para integrarse en la plantilla oficial CONTRATO_DE_APRENDIZAJE_301.docx. No uses tablas, usa texto estructurado."},
            {"label": "Lista de Verificación", "prompt": "Genera la Lista de Verificación de Requerimientos para impartir mi curso, alineada a la Carta Descriptiva ya generada. Devuelve EXACTAMENTE este formato Markdown con cinco secciones, una descripción concreta y específica al tema en cada item (mobiliario, equipos, materiales y consumibles puntuales que el instructor llevará o requerirá en el lugar). NO uses tablas, NO uses asteriscos, solo numeración decimal. Cada descripción debe ser una sola línea breve (máx 18 palabras). Conserva los encabezados literalmente:\n\n## INSTALACIONES (5 items)\n1. Aula con capacidad para al menos N participantes y mesas de trabajo.\n2. ...\n3. ...\n4. ...\n5. ...\n\n## EQUIPO_DE_APOYO (6 items)\n1. ...\n2. ...\n3. ...\n4. ...\n5. ...\n6. ...\n\n## MATERIALES_DIDACTICOS (16 items)\n1. ...\n... hasta 16.\n\n## REQUERIMIENTOS_HUMANOS (2 items)\n1. Instructor titular con perfil afín al tema.\n2. ...\n\n## OTROS_REQUERIMIENTOS (8 items)\n1. ...\n... hasta 8.\n\nNo agregues introducción ni cierre, solo las cinco secciones."},
        ]
    },
    2: {
        "page_title": "IECs (Instrumentos de Evaluación)",
        "page_description": "Elemento 2: Diseña los Instrumentos de Evaluación de la Competencia con las plantillas oficiales",
        "active_page": "e2",
        "welcome_message": "Bienvenido al Módulo de Instrumentos de Evaluación. Para comenzar, indícame el tema o materia de tu curso de capacitación.",
        "welcome_message_ready": "Listo para generar los Instrumentos de Evaluación para tu curso. Selecciona una acción rápida o escribe tu solicitud.",
        "quick_actions": [
            {"label": "Evaluación Diagnóstica", "prompt": "Genera la evaluación diagnóstica (cuestionario) con instrucciones para instructor y participante, y al menos 5 reactivos específicos para mi curso. Valor 0%, solo referencial."},
            {"label": "Guía de Observación", "prompt": "Genera la evaluación formativa como Guía de Observación para evaluar el desempeño del participante. Incluye al menos 5 reactivos con valores y criterios de cumplimiento."},
            {"label": "Lista de Cotejo", "prompt": "Genera la evaluación formativa como Lista de Cotejo para evaluar las características del producto final. Incluye al menos 5 reactivos sobre calidad y especificaciones."},
            {"label": "Evaluación Sumativa", "prompt": "Genera la evaluación sumativa (cuestionario final) con al menos 8 reactivos que cubran conocimientos teóricos y prácticos, con valores asignados."},
            {"label": "Evaluación de Satisfacción", "prompt": "Genera la evaluación de satisfacción/reacción basada en el formato oficial del EC0301."},
            {"label": "Hojas de Respuestas", "prompt": "Genera las hojas de respuestas para todos los instrumentos de evaluación (diagnóstica, formativa y sumativa)."},
        ]
    },
    3: {
        "page_title": "Manuales del Curso",
        "page_description": "Elemento 3: Diseña el Manual del Instructor y del Participante",
        "active_page": "e3",
        "welcome_message": "Bienvenido al Módulo de Manuales. Para comenzar, indícame el tema o materia de tu curso de capacitación.",
        "welcome_message_ready": "Listo para generar los Manuales para tu curso. Selecciona una acción rápida o escribe tu solicitud.",
        "quick_actions": [
            {"label": "Manual del Instructor", "prompt": "Genera el Manual del Instructor completo con: portada, índice, introducción, estructura del taller, modalidad, requerimientos, contenido temático con sugerencias de apoyo, técnicas de desarrollo, formas/criterios/tiempos de evaluación, actividades de refuerzo, conclusión, resumen y fuentes de información."},
            {"label": "Manual del Participante", "prompt": "Genera el Manual del Participante completo con: portada, índice, presentación (bienvenida, recomendaciones, organización), introducción (resumen, beneficios, enfoque didáctico), objetivos, temas desglosados de lo simple a lo complejo con actividades y síntesis, y fuentes en formato APA."},
            {"label": "Presentación Ejecutiva", "tooltip": "Diapositivas visuales y ligeras para proyectar frente al grupo: pocos puntos por lámina, enfoque en el participante y las indicaciones del instructor en las notas del orador.", "prompt": "Genera el guión Slide-by-Slide de la Presentación Ejecutiva del curso: una presentación VISUAL y LIGERA para proyectar e impartir frente al grupo. Usa pocos bullets por diapositiva (máximo 6), redacción corta y orientada al participante. NO incluyas el momento cero (comprobación de recursos), ni la lista de verificación, ni los instrumentos de evaluación detallados, ni los cronometrajes internos del instructor. Mantén: portada, TÉCNICA ROMPE HIELO con nombre específico, objetivos (en tabla markdown), temario, beneficios del curso, reflexión inicial, los temas del desarrollo con su práctica, TÉCNICA ENERGIZANTE con nombre específico, y el cierre (resumen, logro de expectativas, compromisos de aplicación, agradecimiento). TODAS las instrucciones operativas del instructor van en 'Notas para el presentador', no en el cuerpo de la diapositiva. Las tablas en formato markdown: cada fila inicia y termina con '|' y debajo del encabezado una fila separadora '|---|---|'."},
            {"label": "Presentación de Facilitación", "tooltip": "Presentación completa alineada al EC0217.01 con todo el contenido normativo (momento cero, evaluaciones, evidencias, tiempos), con tablas reales y notas del orador para el instructor.", "prompt": "Genera el guión Slide-by-Slide de la Presentación de Facilitación del curso cumpliendo el EC0217.01. El Slide 1 (momento cero) DEBE ser obligatoriamente la Comprobación de la existencia y el funcionamiento de los recursos requeridos para la sesión (Lista de verificación antes del inicio). Luego desglosa los Slides para: Encuadre (presentación del instructor, TÉCNICA ROMPE HIELO con nombre específico e instrucciones, objetivos generales y particulares, temario, reglas y contrato de aprendizaje, evaluación diagnóstica), Desarrollo (temas con técnicas Expositiva/Demostrativa/Diálogo-Discusión, evaluación formativa, DESCANSO de 10 min, TÉCNICA ENERGIZANTE con nombre específico e instrucciones para reactivar al grupo, temas restantes) y Cierre (TÉCNICA GRUPAL DE CIERRE: conclusiones, resumen, logro de expectativas/objetivos, sugerencias de continuidad, compromisos, evaluación de satisfacción, agradecimiento). Las 3 técnicas grupales (Rompe hielo, Energizante, Cierre) son OBLIGATORIAS y deben aparecer como slides independientes con instrucciones completas. Incluye notas para el presentador en cada slide. Las tablas (objetivos, tipos y momentos de evaluación) en formato markdown: cada fila inicia y termina con '|' y debajo del encabezado una fila separadora '|---|---|'."},
        ]
    },
    4: {
        "page_title": "Evaluación EC0301",
        "page_description": "Módulo Evaluador: Revisa tus productos contra la rúbrica del estándar",
        "active_page": "e4",
        "welcome_message": "Selecciona una opción arriba para iniciar tu evaluación. Usa la Opción A si tus documentos fueron generados por la plataforma, o la Opción B para auditar archivos externos.",
        "welcome_message_ready": "Selecciona una opción arriba para iniciar tu evaluación.",
        "quick_actions": []
    }
}

@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    next_target = request.args.get('next', '')
    if next_target and 'admin' in next_target.lower() and request.method == "GET":
        track_event('Seguridad', 'Intento Admin Anónimo', extra_data={'next': next_target})
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        # El correo ya no es único: puede haber varias cuentas distintas que lo
        # comparten. Entramos a la que coincida con la contraseña (la contraseña
        # desambigua). Orden estable por id para usuarios de correo único.
        candidates = User.query.filter_by(email=email).order_by(User.id.asc()).all()
        user = next((u for u in candidates if u.check_password(password)), None)
        if user:
            login_user(user, remember=True)
            track_event('Auth', 'Login', user_id=user.id, extra_data={'tier': user.tier})
            next_page = request.args.get('next')
            if next_page and (not next_page.startswith('/') or next_page.startswith('//')):
                next_page = None
            return redirect(next_page or url_for('home'))
        flash('Correo electrónico o contraseña incorrectos.', 'error')
    return render_template("login.html", title="Iniciar Sesión", active_page="login")

def _get_serializer():
    return URLSafeTimedSerializer(app.config['SECRET_KEY'])

def _send_reset_email(recipient_email, reset_url):
    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_email = os.environ.get('SMTP_EMAIL')
    smtp_password = os.environ.get('SMTP_PASSWORD')
    _u = User.query.filter_by(email=recipient_email).first()
    log_id = _log_email_attempt(
        email_type='password_reset',
        recipient=recipient_email,
        sender=smtp_email,
        subject='Restablecer contraseña - PertinentIA',
        user_id=_u.id if _u else None,
    )
    if not all([smtp_server, smtp_email, smtp_password]):
        logger.error("SMTP credentials not configured")
        _log_email_result(log_id, False, error="SMTP credentials not configured")
        return False
    msg = MIMEMultipart('alternative')
    msg['Subject'] = 'Restablecer contraseña - PertinentIA'
    msg['From'] = smtp_email
    msg['To'] = recipient_email
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:500px;margin:0 auto;padding:20px;">
        <h2 style="color:#2d2d2d;">Restablecer tu contraseña</h2>
        <p>Recibimos una solicitud para restablecer la contraseña de tu cuenta en <strong>PertinentIA</strong>.</p>
        <p>Haz clic en el siguiente botón para crear una nueva contraseña:</p>
        <a href="{reset_url}" style="display:inline-block;padding:12px 24px;background:#FA8072;color:white;text-decoration:none;border-radius:6px;font-weight:600;margin:16px 0;">Restablecer Contraseña</a>
        <p style="color:#6b7280;font-size:0.85em;">Este enlace expira en 1 hora. Si no solicitaste este cambio, puedes ignorar este correo.</p>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
        <p style="color:#9ca3af;font-size:0.75em;">PertinentIA &mdash; Fábrica de Productos EC0301</p>
    </div>
    """
    msg.attach(MIMEText(html_body, 'html'))
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.sendmail(smtp_email, recipient_email, msg.as_string())
        _log_email_result(log_id, True, smtp_response="250 Accepted by Titan")
        return True
    except Exception as e:
        logger.error(f"Error sending reset email: {e}")
        _log_email_result(log_id, False, error=e)
        return False

def _send_tier_change_email(user, prev_tier, new_tier, prev_credits, new_credits, applied_label, reason=''):
    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_email = os.environ.get('SMTP_EMAIL')
    smtp_password = os.environ.get('SMTP_PASSWORD')
    log_id = _log_email_attempt(
        email_type='tier_change',
        recipient=user.email,
        sender=smtp_email,
        subject=f'Cambio de plan: {prev_tier} -> {applied_label or new_tier}',
        user_id=user.id,
        extra={'prev_tier': prev_tier, 'new_tier': new_tier, 'prev_credits': prev_credits, 'new_credits': new_credits, 'applied_label': applied_label, 'reason': reason or ''},
    )
    if not all([smtp_server, smtp_email, smtp_password]):
        logger.error(f"Tier-change email: SMTP no configurado, omitiendo aviso a {user.email}")
        _log_email_result(log_id, False, error="SMTP credentials not configured")
        return False

    label_pretty = {
        'PRO_PROJECT': 'PRO &mdash; 1 Curso',
        'PRO_MULTICURSO': 'PRO &mdash; 5 Cursos',
        'PREMIUM': 'Premium',
        'FREE': 'FREE'
    }.get(applied_label, applied_label)

    is_upgrade = (prev_tier == 'FREE' and new_tier in ('PRO', 'PREMIUM')) or \
                 (prev_tier == 'PRO' and new_tier == 'PREMIUM') or \
                 (prev_tier == 'PRO' and applied_label == 'PRO_MULTICURSO' and (new_credits or 0) > (prev_credits or 0))
    is_downgrade_to_free = (new_tier == 'FREE' and prev_tier != 'FREE')
    is_downgrade_other = (prev_tier == 'PREMIUM' and new_tier == 'PRO')

    if is_downgrade_to_free:
        subject = 'Cambio en tu plan PertinentIA'
        intro = f'<p>Te informamos que tu plan en <strong>PertinentIA</strong> fue cambiado a <strong>FREE</strong>.</p>'
        if (prev_credits or 0) > 0:
            intro += f'<p style="background:#fff7ed;border-left:4px solid #f97316;padding:10px 14px;border-radius:6px;"><strong>Buenas noticias:</strong> tus <strong>{prev_credits} cr&eacute;dito{"s" if prev_credits != 1 else ""} PRO</strong> quedan preservados en reserva. Si vuelves a PRO, los recuperar&aacute;s autom&aacute;ticamente.</p>'
        body_color = '#dc2626'
    elif is_downgrade_other:
        subject = 'Cambio en tu plan PertinentIA'
        intro = f'<p>Tu plan fue ajustado de <strong>Premium</strong> a <strong>{label_pretty}</strong>.</p>'
        body_color = '#d97706'
    elif is_upgrade:
        subject = f'¡Tu plan PertinentIA ahora es {label_pretty.replace("&mdash;","-").replace("&nbsp;"," ")}!'
        if applied_label == 'PRO_MULTICURSO':
            intro = f'<p>&iexcl;Excelentes noticias! Tu cuenta ahora cuenta con <strong>{new_credits} cr&eacute;dito{"s" if new_credits != 1 else ""} PRO</strong> activos para generar hasta {new_credits} curso{"s" if new_credits != 1 else ""} completo{"s" if new_credits != 1 else ""} con todas las funciones PRO.</p>'
        elif applied_label == 'PRO_PROJECT':
            intro = f'<p>&iexcl;Excelentes noticias! Tu cuenta ahora cuenta con <strong>{new_credits} cr&eacute;dito PRO</strong> activo para generar tu pr&oacute;ximo curso completo con todas las funciones PRO.</p>'
        else:
            intro = f'<p>&iexcl;Excelentes noticias! Tu cuenta ha sido elevada al plan <strong>{label_pretty}</strong> con acceso completo a Centro Evaluador IA, gesti&oacute;n de candidatos, marca blanca y todas las funciones premium.</p>'
        body_color = '#16a34a'
    else:
        subject = 'Actualizaci&oacute;n en tu plan PertinentIA'
        intro = f'<p>Tu plan fue actualizado de <strong>{prev_tier}</strong> ({prev_credits or 0} cr&eacute;d.) a <strong>{label_pretty}</strong> ({new_credits or 0} cr&eacute;d.).</p>'
        body_color = '#1e40af'

    reason_html = ''
    if reason:
        reason_html = f'<p style="background:#f3f4f6;border-left:3px solid #6b7280;padding:10px 14px;border-radius:6px;font-size:0.9em;color:#374151;"><strong>Motivo indicado por el equipo:</strong><br>{reason}</p>'

    cortesia_html = ''
    if applied_label in ('PRO_PROJECT', 'PRO_MULTICURSO', 'PREMIUM') and is_upgrade:
        cortesia_html = '<p style="font-size:0.78em;color:#6b7280;font-style:italic;margin-top:14px;">Esta asignaci&oacute;n fue realizada directamente por el equipo de PertinentIA como cortes&iacute;a administrativa. No genera comprobante fiscal de Stripe.</p>'

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = smtp_email
    msg['To'] = user.email
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:0 auto;padding:20px;">
        <h2 style="color:{body_color};margin-bottom:8px;">Hola {user.full_name or ''},</h2>
        {intro}
        {reason_html}
        <p style="margin-top:18px;">Plan anterior: <strong>{prev_tier}</strong>{f' &middot; {prev_credits} cr&eacute;d.' if prev_credits else ''}<br>
        Plan actual: <strong>{label_pretty}</strong>{f' &middot; {new_credits} cr&eacute;d.' if new_credits else ''}</p>
        <a href="https://pertinentia.com/precios" style="display:inline-block;padding:12px 24px;background:#FA8072;color:white;text-decoration:none;border-radius:6px;font-weight:600;margin:16px 0;">Ver mi cuenta</a>
        {cortesia_html}
        <p style="font-size:0.78em;color:#9ca3af;margin-top:18px;">Si tienes alguna duda, responde a este correo o escribe a soporte@pertinentia.com</p>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
        <p style="color:#9ca3af;font-size:0.72em;">PertinentIA &mdash; F&aacute;brica de Productos EC0301 &middot; Aviso transaccional de servicio</p>
    </div>
    """
    msg.attach(MIMEText(html_body, 'html'))
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.sendmail(smtp_email, user.email, msg.as_string())
        logger.info(f"Tier-change email sent to {user.email} ({prev_tier}->{new_tier})")
        _log_email_result(log_id, True, smtp_response="250 Accepted by Titan")
        return True
    except Exception as e:
        logger.error(f"Failed to send tier-change email to {user.email}: {e}")
        _log_email_result(log_id, False, error=e)
        return False

def _send_first_doc_nudge_email(user):
    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_email = os.environ.get('SMTP_EMAIL')
    smtp_password = os.environ.get('SMTP_PASSWORD')
    subject = '¿Te ayudo a generar tu primer documento en PertinentIA?'
    log_id = _log_email_attempt(
        email_type='first_doc_nudge',
        recipient=user.email,
        sender=smtp_email,
        subject=subject,
        user_id=user.id,
    )
    if not all([smtp_server, smtp_email, smtp_password]):
        logger.error(f"first_doc_nudge: SMTP no configurado, omitiendo envío a {user.email}")
        _log_email_result(log_id, False, error="SMTP credentials not configured")
        return False
    import html as _html
    nombre_raw = (user.first_name or user.full_name or '').split(' ')[0] or 'hola'
    nombre = _html.escape(nombre_raw)
    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = smtp_email
    msg['To'] = user.email
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:0 auto;padding:20px;color:#1f2937;">
        <h2 style="color:#1e40af;margin-bottom:8px;">Hola {nombre},</h2>
        <p>Vi que te registraste en <strong>PertinentIA</strong> pero aún no has generado tu primer documento.</p>
        <p>Si te quedaste atorado en algún paso o no sabes por dónde empezar, te ayudo en 2 minutos:</p>
        <ul style="line-height:1.7;">
            <li><strong>Carta Descriptiva</strong> — el documento más usado, ideal para empezar.</li>
            <li>Solo necesitas el nombre del curso y a quién va dirigido.</li>
            <li>Genera, descarga, listo. Sin tarjeta de crédito.</li>
        </ul>
        <a href="https://pertinentia.com/elemento/1" style="display:inline-block;padding:14px 28px;background:#FA8072;color:white;text-decoration:none;border-radius:6px;font-weight:600;margin:16px 0;">Generar mi primer documento</a>
        <p style="margin-top:18px;">¿Necesitas ayuda? Solo responde a este correo y te contesto personalmente.</p>
        <p style="color:#6b7280;font-size:0.85em;margin-top:18px;">Un saludo,<br><strong>Arturo García</strong><br>PertinentIA</p>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
        <p style="color:#9ca3af;font-size:0.72em;">PertinentIA &mdash; Fábrica de Productos EC0301 &middot; Aviso transaccional de servicio</p>
    </div>
    """
    msg.attach(MIMEText(html_body, 'html'))
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.sendmail(smtp_email, user.email, msg.as_string())
        logger.info(f"first_doc_nudge enviado a {user.email}")
        _log_email_result(log_id, True, smtp_response="250 Accepted by Titan")
        return True
    except Exception as e:
        logger.error(f"Error enviando first_doc_nudge a {user.email}: {e}")
        _log_email_result(log_id, False, error=e)
        return False

def _send_document_email(recipient_email, filename, filepath, user_id=None):
    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_email_docs = os.environ.get('SMTP_EMAIL_DOCS') or os.environ.get('SMTP_EMAIL')
    smtp_password_docs = os.environ.get('SMTP_PASSWORD_DOCS') or os.environ.get('SMTP_PASSWORD')
    clean_name = filename.replace('anon_', '').replace('_', ' ')
    if clean_name.startswith('u') and '_' in filename:
        parts = filename.split('_', 1)
        if len(parts) > 1:
            clean_name = parts[1].replace('_', ' ')
    subject_str = f'Tu documento PertinentIA: {clean_name}'
    log_id = _log_email_attempt(
        email_type='document',
        recipient=recipient_email,
        sender=smtp_email_docs,
        subject=subject_str,
        user_id=user_id,
        document_filename=filename,
    )
    if not all([smtp_server, smtp_email_docs, smtp_password_docs]):
        logger.error("SMTP credentials not configured for document email")
        _log_email_result(log_id, False, error="SMTP credentials not configured")
        return False
    smtp_email = smtp_email_docs
    smtp_password = smtp_password_docs
    msg = MIMEMultipart('mixed')
    msg['Subject'] = subject_str
    msg['From'] = smtp_email
    msg['To'] = recipient_email
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:550px;margin:0 auto;padding:20px;">
        <h2 style="color:#2d2d2d;">Tu documento est&aacute; listo</h2>
        <p>Hemos generado exitosamente tu documento:</p>
        <p style="background:#f3f4f6;padding:12px 16px;border-radius:8px;font-weight:600;color:#374151;">{clean_name}</p>
        <p style="color:#6b7280;font-size:0.9em;">Encontrar&aacute;s el archivo adjunto en este correo.</p>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
        <p style="font-size:0.85em;color:#6b7280;">&#127775; <strong>&iquest;Quieres documentos sin marca de agua, plantillas profesionales y descargas ilimitadas?</strong></p>
        <a href="https://pertinentia.com/precios" style="display:inline-block;padding:10px 20px;background:linear-gradient(135deg,#FA8072,#e8705f);color:white;text-decoration:none;border-radius:6px;font-weight:600;font-size:0.9em;">Ver Planes PRO</a>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
        <p style="color:#9ca3af;font-size:0.75em;">PertinentIA &mdash; F&aacute;brica de Productos EC0301</p>
    </div>
    """
    msg.attach(MIMEText(html_body, 'html'))
    try:
        file_data = None
        if os.path.isfile(filepath):
            with open(filepath, 'rb') as f:
                file_data = f.read()
        else:
            sf = StoredFile.query.filter_by(filename=os.path.basename(filepath)).first()
            if sf and sf.content:
                file_data = sf.content
        if file_data:
            from email.mime.base import MIMEBase
            from email import encoders
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(file_data)
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(filepath)}"')
            msg.attach(part)
    except Exception as e:
        logger.warning(f"Could not attach file to email: {e}")
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.sendmail(smtp_email, recipient_email, msg.as_string())
        logger.info(f"Document email sent to {recipient_email}: {filename}")
        _log_email_result(log_id, True, smtp_response="250 Accepted by Titan")
        return True
    except Exception as e:
        logger.error(f"Error sending document email: {e}")
        _log_email_result(log_id, False, error=e)
        return False

@app.route("/api/admin/ce/costos")
@login_required
def admin_ce_costos():
    """Rollup de costos del Centro Evaluador IA (EC0301 + EC0217.01) consolidado."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    desde = request.args.get('desde', '').strip()
    hasta = request.args.get('hasta', '').strip()
    q = PortafolioEvaluacion.query
    try:
        if desde:
            q = q.filter(PortafolioEvaluacion.created_at >= datetime.fromisoformat(desde))
        if hasta:
            q = q.filter(PortafolioEvaluacion.created_at <= datetime.fromisoformat(hasta))
    except ValueError:
        return jsonify({"error": "Formato de fecha inválido (usa YYYY-MM-DD)"}), 400

    evals = q.all()
    by_user = {}
    by_estandar = {'EC0301': {'evaluaciones': 0, 'auditoria_mxn': 0.0, 'video_mxn': 0.0,
                              'dictamen_mxn': 0.0, 'auditoria_usd': 0.0, 'video_usd': 0.0,
                              'dictamen_usd': 0.0, 'tokens_prompt': 0, 'tokens_completion': 0,
                              'video_segundos': 0, 'dictaminadas': 0},
                   'EC0217.01': {'evaluaciones': 0, 'auditoria_mxn': 0.0, 'video_mxn': 0.0,
                                 'dictamen_mxn': 0.0, 'auditoria_usd': 0.0, 'video_usd': 0.0,
                                 'dictamen_usd': 0.0, 'tokens_prompt': 0, 'tokens_completion': 0,
                                 'video_segundos': 0, 'dictaminadas': 0}}
    totales = {'evaluaciones': 0, 'auditoria_mxn': 0.0, 'video_mxn': 0.0,
               'dictamen_mxn': 0.0, 'total_mxn': 0.0, 'total_usd': 0.0,
               'tokens_prompt': 0, 'tokens_completion': 0, 'video_segundos': 0,
               'dictaminadas': 0, 'usuarios_unicos': 0}

    for ev in evals:
        est = ev.estandar if ev.estandar in by_estandar else 'EC0301'
        a_mxn = float(ev.cost_mxn or 0); v_mxn = float(ev.video_cost_mxn or 0); d_mxn = float(ev.dictamen_cost_mxn or 0)
        a_usd = float(ev.cost_usd or 0); v_usd = float(ev.video_cost_usd or 0); d_usd = float(ev.dictamen_cost_usd or 0)
        tp = int((ev.tokens_prompt or 0) + (ev.dictamen_tokens_prompt or 0))
        tc = int((ev.tokens_completion or 0) + (ev.dictamen_tokens_completion or 0))
        vs = int(ev.video_segundos_transcritos or 0)
        is_dict = bool(ev.dictamen_pdf)

        be = by_estandar[est]
        be['evaluaciones'] += 1
        be['auditoria_mxn'] += a_mxn; be['video_mxn'] += v_mxn; be['dictamen_mxn'] += d_mxn
        be['auditoria_usd'] += a_usd; be['video_usd'] += v_usd; be['dictamen_usd'] += d_usd
        be['tokens_prompt'] += tp; be['tokens_completion'] += tc
        be['video_segundos'] += vs
        if is_dict: be['dictaminadas'] += 1

        totales['evaluaciones'] += 1
        totales['auditoria_mxn'] += a_mxn; totales['video_mxn'] += v_mxn; totales['dictamen_mxn'] += d_mxn
        totales['total_mxn'] += a_mxn + v_mxn + d_mxn
        totales['total_usd'] += a_usd + v_usd + d_usd
        totales['tokens_prompt'] += tp; totales['tokens_completion'] += tc
        totales['video_segundos'] += vs
        if is_dict: totales['dictaminadas'] += 1

        uid = ev.ce_user_id
        if uid not in by_user:
            by_user[uid] = {'user_id': uid, 'email': '', 'nombre': '', 'evaluaciones': 0,
                            'auditoria_mxn': 0.0, 'video_mxn': 0.0, 'dictamen_mxn': 0.0,
                            'total_mxn': 0.0, 'tokens_total': 0, 'video_segundos': 0,
                            'dictaminadas': 0}
        u = by_user[uid]
        u['evaluaciones'] += 1
        u['auditoria_mxn'] += a_mxn; u['video_mxn'] += v_mxn; u['dictamen_mxn'] += d_mxn
        u['total_mxn'] += a_mxn + v_mxn + d_mxn
        u['tokens_total'] += tp + tc
        u['video_segundos'] += vs
        if is_dict: u['dictaminadas'] += 1

    if by_user:
        users = User.query.filter(User.id.in_(list(by_user.keys()))).all()
        for u in users:
            if u.id in by_user:
                by_user[u.id]['email'] = u.email or ''
                by_user[u.id]['nombre'] = u.nombre or ''
    totales['usuarios_unicos'] = len(by_user)

    for d in (totales, *by_estandar.values()):
        for k, v in list(d.items()):
            if isinstance(v, float): d[k] = round(v, 4)
    for u in by_user.values():
        for k, v in list(u.items()):
            if isinstance(v, float): u[k] = round(v, 4)

    usuarios_lista = sorted(by_user.values(), key=lambda x: x['total_mxn'], reverse=True)
    return jsonify({
        "totales": totales,
        "por_estandar": by_estandar,
        "por_usuario": usuarios_lista,
        "filtros": {"desde": desde, "hasta": hasta},
        "generado_at": datetime.utcnow().isoformat(),
    })


@app.route("/api/admin/ce/costos.csv")
@login_required
def admin_ce_costos_csv():
    if not current_user.is_admin:
        return ("No autorizado", 403)
    import csv
    from io import StringIO
    q = PortafolioEvaluacion.query.order_by(PortafolioEvaluacion.created_at.desc())
    desde = request.args.get('desde', '').strip()
    hasta = request.args.get('hasta', '').strip()
    try:
        if desde: q = q.filter(PortafolioEvaluacion.created_at >= datetime.fromisoformat(desde))
        if hasta: q = q.filter(PortafolioEvaluacion.created_at <= datetime.fromisoformat(hasta))
    except ValueError:
        return ("Fecha inválida", 400)
    evals = q.all()
    user_cache = {}
    if evals:
        users = User.query.filter(User.id.in_(list({e.ce_user_id for e in evals}))).all()
        user_cache = {u.id: u for u in users}

    sio = StringIO()
    w = csv.writer(sio)
    w.writerow(['id', 'creada_at', 'ce_user_id', 'ce_email', 'estandar', 'status',
                'candidato_id', 'autorizado', 'dictaminado',
                'auditoria_tokens_prompt', 'auditoria_tokens_completion', 'auditoria_cost_usd', 'auditoria_cost_mxn',
                'video_segundos', 'video_idioma', 'video_modelo', 'video_cost_usd', 'video_cost_mxn',
                'dictamen_tokens_prompt', 'dictamen_tokens_completion', 'dictamen_cost_usd', 'dictamen_cost_mxn',
                'dictamen_final', 'total_mxn'])
    for ev in evals:
        u = user_cache.get(ev.ce_user_id)
        total_mxn = float(ev.cost_mxn or 0) + float(ev.video_cost_mxn or 0) + float(ev.dictamen_cost_mxn or 0)
        w.writerow([
            ev.id, ev.created_at.isoformat() if ev.created_at else '',
            ev.ce_user_id, (u.email if u else ''), ev.estandar, ev.status,
            ev.candidato_id, int(bool(ev.autorizado_por_evaluador)), int(bool(ev.dictamen_pdf)),
            ev.tokens_prompt or 0, ev.tokens_completion or 0,
            round(float(ev.cost_usd or 0), 6), round(float(ev.cost_mxn or 0), 4),
            ev.video_segundos_transcritos or 0, ev.video_idioma or '', ev.video_modelo or '',
            round(float(ev.video_cost_usd or 0), 6), round(float(ev.video_cost_mxn or 0), 4),
            ev.dictamen_tokens_prompt or 0, ev.dictamen_tokens_completion or 0,
            round(float(ev.dictamen_cost_usd or 0), 6), round(float(ev.dictamen_cost_mxn or 0), 4),
            ev.dictamen_final or '', round(total_mxn, 4),
        ])
    resp = make_response(sio.getvalue())
    resp.headers['Content-Type'] = 'text/csv; charset=utf-8'
    resp.headers['Content-Disposition'] = f'attachment; filename="ce_costos_{datetime.utcnow().strftime("%Y%m%d_%H%M%S")}.csv"'
    resp.headers['Cache-Control'] = 'private, no-store'
    return resp


@app.route("/admin/ce_costos")
@login_required
def admin_ce_costos_view():
    if not current_user.is_admin:
        return redirect(url_for('login'))
    return render_template('admin_ce_costos.html')


@app.route("/admin/pexels")
@login_required
def admin_pexels_view():
    if not current_user.is_admin:
        return redirect(url_for('login'))
    from sqlalchemy import func as _f
    now = datetime.utcnow()
    hour_ago = now - timedelta(hours=1)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    day_ago = now - timedelta(hours=24)
    week_ago = now - timedelta(days=7)

    def _count(since):
        return db.session.query(_f.count(PexelsUsage.id)).filter(PexelsUsage.created_at >= since).scalar() or 0

    total = db.session.query(_f.count(PexelsUsage.id)).scalar() or 0
    last_hour = _count(hour_ago)
    this_month = _count(month_start)
    last_24h = _count(day_ago)
    last_7d = _count(week_ago)

    HOURLY_LIMIT = 200
    MONTHLY_LIMIT = 25000

    by_day = (db.session.query(_f.date(PexelsUsage.created_at).label('d'), _f.count(PexelsUsage.id))
              .filter(PexelsUsage.created_at >= (now - timedelta(days=30)))
              .group_by('d').order_by('d').all())
    by_day = [(str(d), int(c)) for d, c in by_day]

    top_q = (db.session.query(PexelsUsage.query, _f.count(PexelsUsage.id).label('c'))
             .group_by(PexelsUsage.query).order_by(_f.count(PexelsUsage.id).desc()).limit(15).all())
    top_q = [(q or '(sin término)', int(c)) for q, c in top_q]

    recent = (PexelsUsage.query.order_by(PexelsUsage.id.desc()).limit(20).all())
    recent_rows = [{
        'fecha': _fmt_cdmx(r.created_at),
        'query': r.query or '',
        'photographer': r.photographer or '',
        'photo_url': r.photo_url or '',
        'user_id': r.user_id,
    } for r in recent]

    stats = {
        'has_api_key': bool(os.environ.get('PEXELS_API_KEY', '')),
        'total': total,
        'last_hour': last_hour, 'hourly_limit': HOURLY_LIMIT,
        'hourly_pct': round(100.0 * last_hour / HOURLY_LIMIT, 1) if HOURLY_LIMIT else 0,
        'this_month': this_month, 'monthly_limit': MONTHLY_LIMIT,
        'monthly_pct': round(100.0 * this_month / MONTHLY_LIMIT, 1) if MONTHLY_LIMIT else 0,
        'monthly_remaining': max(0, MONTHLY_LIMIT - this_month),
        'last_24h': last_24h, 'last_7d': last_7d,
        'by_day': by_day, 'top_q': top_q, 'recent': recent_rows,
    }
    return render_template('admin_pexels.html', s=stats)


@app.route("/api/admin/test_email_docs")
@login_required
def test_email_docs():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_email_docs = os.environ.get('SMTP_EMAIL_DOCS') or os.environ.get('SMTP_EMAIL')
    smtp_password_docs = os.environ.get('SMTP_PASSWORD_DOCS') or os.environ.get('SMTP_PASSWORD')
    if not all([smtp_server, smtp_email_docs, smtp_password_docs]):
        return jsonify({"error": "SMTP_EMAIL_DOCS o SMTP_PASSWORD_DOCS no configurados", "vars": {
            "SMTP_SERVER": bool(smtp_server),
            "SMTP_EMAIL_DOCS": bool(os.environ.get('SMTP_EMAIL_DOCS')),
            "SMTP_PASSWORD_DOCS": bool(os.environ.get('SMTP_PASSWORD_DOCS')),
            "SMTP_EMAIL_fallback": bool(os.environ.get('SMTP_EMAIL'))
        }}), 400
    try:
        from email.mime.text import MIMEText as _MT
        msg = _MT(f"<div style='font-family:Arial;padding:20px'><h2>Prueba de Email</h2><p>Este es un correo de prueba enviado desde <strong>{smtp_email_docs}</strong> a las {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}.</p><p>Si recibes esto, la configuraci&oacute;n de SMTP_EMAIL_DOCS funciona correctamente.</p><hr><p style='color:#9ca3af;font-size:0.8em'>PertinentIA - F&aacute;brica de Productos EC0301</p></div>", 'html')
        msg['Subject'] = 'PertinentIA - Prueba de Email (contacto@)'
        msg['From'] = smtp_email_docs
        msg['To'] = current_user.email
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email_docs, smtp_password_docs)
            server.sendmail(smtp_email_docs, current_user.email, msg.as_string())
        return jsonify({"success": True, "from": smtp_email_docs, "to": current_user.email})
    except Exception as e:
        return jsonify({"error": str(e), "from": smtp_email_docs}), 500

def send_error_alert(error, tb_str):
    smtp_server = os.environ.get('SMTP_SERVER')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_email = os.environ.get('SMTP_EMAIL')
    smtp_password = os.environ.get('SMTP_PASSWORD')
    log_id = _log_email_attempt(
        email_type='centinela_alert',
        recipient=smtp_email or 'unknown',
        sender=smtp_email,
        subject='ALERTA Error 500',
        user_id=current_user.id if (current_user and current_user.is_authenticated) else None,
    )
    if not all([smtp_server, smtp_email, smtp_password]):
        logger.error("Centinela: SMTP credentials not configured, cannot send alert")
        _log_email_result(log_id, False, error="SMTP credentials not configured")
        return
    now = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    req_url = request.url if request else 'N/A'
    req_method = request.method if request else 'N/A'
    if current_user and current_user.is_authenticated:
        user_info = f"ID: {current_user.id} | Email: {current_user.email} | Tier: {current_user.tier}"
    else:
        user_info = "Usuario Anónimo"
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:700px;margin:0 auto;padding:20px;">
        <h2 style="color:#dc2626;">⚠️ Error 500 Detectado</h2>
        <table style="width:100%;border-collapse:collapse;margin:16px 0;">
            <tr><td style="padding:8px;border:1px solid #e5e7eb;font-weight:600;width:140px;">Fecha y Hora</td><td style="padding:8px;border:1px solid #e5e7eb;">{now}</td></tr>
            <tr><td style="padding:8px;border:1px solid #e5e7eb;font-weight:600;">URL</td><td style="padding:8px;border:1px solid #e5e7eb;">{req_url}</td></tr>
            <tr><td style="padding:8px;border:1px solid #e5e7eb;font-weight:600;">Método</td><td style="padding:8px;border:1px solid #e5e7eb;">{req_method}</td></tr>
            <tr><td style="padding:8px;border:1px solid #e5e7eb;font-weight:600;">Usuario</td><td style="padding:8px;border:1px solid #e5e7eb;">{user_info}</td></tr>
            <tr><td style="padding:8px;border:1px solid #e5e7eb;font-weight:600;">Error</td><td style="padding:8px;border:1px solid #e5e7eb;color:#dc2626;">{type(error).__name__}: {error}</td></tr>
        </table>
        <h3 style="color:#2d2d2d;">Traceback Completo</h3>
        <pre style="background:#1e1e1e;color:#d4d4d4;padding:16px;border-radius:8px;overflow-x:auto;font-size:0.85em;line-height:1.5;">{tb_str}</pre>
        <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
        <p style="color:#9ca3af;font-size:0.75em;">PertinentIA &mdash; Centinela Silencioso (Sistema de Monitoreo Automático)</p>
    </div>
    """
    msg = MIMEMultipart('alternative')
    msg['Subject'] = '⚠️ ALERTA DE SISTEMA (PertinentIA): Error 500 Detectado'
    msg['From'] = smtp_email
    msg['To'] = smtp_email
    msg.attach(MIMEText(html_body, 'html'))
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.sendmail(smtp_email, smtp_email, msg.as_string())
        logger.info(f"Centinela: alerta enviada por error en {req_url}")
        _log_email_result(log_id, True, smtp_response="250 Accepted by Titan")
    except Exception as mail_err:
        logger.error(f"Centinela: fallo al enviar alerta — {mail_err}")
        _log_email_result(log_id, False, error=mail_err)

@app.errorhandler(Exception)
def handle_global_exception(error):
    from werkzeug.exceptions import HTTPException
    if isinstance(error, HTTPException) and error.code < 500:
        return error
    tb_str = traceback.format_exc()
    logger.error(f"Centinela — Error 500 no controlado: {error}\n{tb_str}")
    try:
        send_error_alert(error, tb_str)
    except Exception:
        pass
    is_api = request.path.startswith('/api/') or request.is_json or \
             request.accept_mimetypes.best == 'application/json' or \
             request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    if is_api:
        return jsonify({"error": "Error interno del servidor. El equipo ha sido notificado."}), 500
    return render_template("500.html", title="Error del Servidor", active_page="error"), 500

@app.route("/forgot_password", methods=["GET", "POST"])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        # El correo puede estar compartido por varias cuentas distintas. El
        # auto-servicio por correo no puede saber a cuál aplica, así que solo
        # se envía cuando hay EXACTAMENTE una cuenta. Las cuentas que comparten
        # correo las administra el admin (reset por id de usuario).
        accounts = User.query.filter_by(email=email).all()
        if len(accounts) == 1:
            user = accounts[0]
            s = _get_serializer()
            token = s.dumps(user.email, salt='password-reset-salt')
            reset_url = url_for('reset_password', token=token, _external=True)
            _send_reset_email(user.email, reset_url)
        elif len(accounts) > 1:
            logger.info(f"forgot_password: correo compartido por {len(accounts)} cuentas, auto-servicio omitido ({email})")
        flash('Si el correo existe en nuestro sistema, recibirás un enlace para restablecer tu contraseña.', 'success')
        return redirect(url_for('forgot_password'))
    return render_template("forgot_password.html", title="Recuperar Contraseña", active_page="login")

@app.route("/reset_password/<token>", methods=["GET", "POST"])
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    s = _get_serializer()
    try:
        email = s.loads(token, salt='password-reset-salt', max_age=3600)
    except SignatureExpired:
        flash('El enlace ha expirado. Solicita uno nuevo.', 'error')
        return redirect(url_for('forgot_password'))
    except BadSignature:
        flash('El enlace es inválido. Solicita uno nuevo.', 'error')
        return redirect(url_for('forgot_password'))
    accounts = User.query.filter_by(email=email).all()
    if not accounts:
        flash('Usuario no encontrado.', 'error')
        return redirect(url_for('forgot_password'))
    if len(accounts) > 1:
        # Correo compartido por varias cuentas: el enlace no puede saber a cuál
        # aplicar. Las gestiona el admin para no cambiar la contraseña equivocada.
        flash('Este correo tiene varias cuentas asociadas. Contacta al administrador para restablecer tu contraseña.', 'error')
        return redirect(url_for('login'))
    user = accounts[0]
    if request.method == "POST":
        password = request.form.get("password", "")
        password2 = request.form.get("password2", "")
        if not password or len(password) < 6:
            flash('La contraseña debe tener al menos 6 caracteres.', 'error')
        elif password != password2:
            flash('Las contraseñas no coinciden.', 'error')
        else:
            user.set_password(password)
            db.session.commit()
            flash('Contraseña actualizada exitosamente. Inicia sesión con tu nueva contraseña.', 'success')
            return redirect(url_for('login'))
    return render_template("reset_password.html", title="Nueva Contraseña", active_page="login", token=token)

@app.route("/registro", methods=["GET", "POST"])
def registro():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    if request.method == "GET":
        try:
            _anon_turns = session.get('anon_chat_count', 0)
            _has_pending = bool(session.get('pending_document'))
            _has_anon_sid = bool(session.get('anon_sid'))
            _ref_param = request.args.get('ref')
            _from_param = (request.args.get('from') or '')[:60]
            track_event('Funnel', 'Pagina Registro Vista', extra_data={
                'anon_turns': _anon_turns,
                'has_pending_doc': _has_pending,
                'has_anon_sid': _has_anon_sid,
                'ref': _ref_param,
                'from': _from_param,
                'referer': (request.headers.get('Referer') or '')[:200],
            })
        except Exception:
            pass
    ref_id = request.args.get('ref', type=int)
    if ref_id and request.method == "GET":
        session['affiliate_ref'] = ref_id
    if request.method == "GET":
        _prefill_email = (request.args.get('email') or '').strip().lower()[:120]
        if _prefill_email and '@' in _prefill_email:
            session['prefill_email'] = _prefill_email
        _coupon_arg = (request.args.get('coupon') or request.args.get('cupon') or '').strip().upper()[:40]
        if _coupon_arg:
            import re as _re_cp
            if _re_cp.fullmatch(r'[A-Z0-9_\-]{3,40}', _coupon_arg):
                session['coupon_ref'] = _coupon_arg
                try:
                    track_event('Funnel', 'Cupon Prefilled', extra_data={'coupon': _coupon_arg, 'source': (request.args.get('from') or 'registro')[:60]})
                except Exception:
                    pass
    _from_arg = (request.args.get('from') or '').strip()[:60]
    if _from_arg and not session.get('seo_origin'):
        import re as _re
        if _re.fullmatch(r'[A-Za-z0-9_\-]{1,60}', _from_arg):
            session['seo_origin'] = _from_arg
    if request.method == "POST":
        first_name_in = request.form.get("first_name", "").strip()
        apellido_paterno_in = request.form.get("apellido_paterno", "").strip()
        apellido_materno_in = request.form.get("apellido_materno", "").strip()
        legacy_full_name = request.form.get("full_name", "").strip()
        legacy_mode = bool(legacy_full_name) and not first_name_in
        if legacy_mode:
            parsed = _split_full_name(legacy_full_name)
            first_name_in = parsed['first_name']
            apellido_paterno_in = parsed['apellido_paterno']
            apellido_materno_in = parsed['apellido_materno']
        first_name_in = first_name_in[:80]
        apellido_paterno_in = apellido_paterno_in[:80]
        apellido_materno_in = apellido_materno_in[:80]
        full_name = " ".join([p for p in [first_name_in, apellido_paterno_in, apellido_materno_in] if p]).strip()[:150]
        email = request.form.get("email", "").strip().lower()
        whatsapp = request.form.get("whatsapp", "").strip()
        password = request.form.get("password", "")
        password2 = request.form.get("password2", "")
        if legacy_mode:
            missing_required = (not full_name or not email or not password or not whatsapp)
            err_msg = 'Todos los campos son obligatorios.'
        else:
            missing_required = (not first_name_in or not apellido_paterno_in or not email or not password or not whatsapp)
            err_msg = 'Nombre, apellido paterno, correo, WhatsApp y contraseña son obligatorios.'
        if missing_required:
            flash(err_msg, 'error')
        elif password != password2:
            flash('Las contraseñas no coinciden.', 'error')
        elif len(password) < 6:
            flash('La contraseña debe tener al menos 6 caracteres.', 'error')
        elif email == ADMIN_EMAIL:
            flash('Este correo está reservado. Contacta al administrador.', 'error')
        elif User.query.filter_by(email=email).first():
            flash('Ya existe una cuenta con este correo electrónico.', 'error')
        else:
            _email_ok, _email_user_msg, _email_log_reason = validate_email_full(email)
            if not _email_ok:
                logger.info(f"Registro rechazado por validación de correo: {email} -> {_email_log_reason}")
                flash(_email_user_msg, 'error')
                return render_template("registro.html",
                    form_email=email, form_first_name=first_name_in,
                    form_apellido_paterno=apellido_paterno_in,
                    form_apellido_materno=apellido_materno_in,
                    form_whatsapp=whatsapp)
            user = User(
                email=email,
                full_name=full_name,
                first_name=first_name_in[:80],
                apellido_paterno=apellido_paterno_in[:80],
                apellido_materno=(apellido_materno_in[:80] if apellido_materno_in else None),
                whatsapp=whatsapp,
                terms_accepted=True,
                marketing_consent=True,
                marketing_consent_source='SIGNUP',
            )
            user.set_password(password)
            try:
                _utm_ft = session.get('utm_first_touch') or {}
                if _utm_ft:
                    user.utm_source = (_utm_ft.get('utm_source') or None) or None
                    user.utm_medium = (_utm_ft.get('utm_medium') or None) or None
                    user.utm_campaign = (_utm_ft.get('utm_campaign') or None) or None
                    user.utm_term = (_utm_ft.get('utm_term') or None) or None
                    user.utm_content = (_utm_ft.get('utm_content') or None) or None
                    user.utm_id = (_utm_ft.get('utm_id') or None) or None
                    user.utm_landing = (_utm_ft.get('utm_landing') or None) or None
                    try:
                        user.utm_captured_at = datetime.fromisoformat(_utm_ft['utm_captured_at']) if _utm_ft.get('utm_captured_at') else datetime.utcnow()
                    except Exception:
                        user.utm_captured_at = datetime.utcnow()
            except Exception:
                pass
            db.session.add(user)
            db.session.flush()
            eval_proc = EvaluationProcess(user_id=user.id)
            db.session.add(eval_proc)
            db.session.commit()
            login_user(user, remember=True)
            try:
                _queue_pixel_event('CompleteRegistration', {'content_name': 'Pertinentia Signup', 'status': 'success'})
                _queue_pixel_event('Lead', {'content_name': 'Pertinentia Signup', 'content_category': 'signup'})
            except Exception:
                pass

            _seo_origin = session.pop('seo_origin', None)
            if _seo_origin:
                try:
                    track_event('Auth', 'Registro SEO', user_id=user.id,
                                extra_data={'seo_origin': _seo_origin})
                except Exception:
                    pass

            anon_sid_at_reg = session.get('anon_sid')
            if anon_sid_at_reg:
                try:
                    transferred = ChatSpec.query.filter_by(session_id=anon_sid_at_reg, user_id=None).update(
                        {'user_id': user.id, 'session_id': None}, synchronize_session=False)
                    db.session.commit()
                    if transferred:
                        logger.info(f"REG_SPECS_TRANSFERRED user_id={user.id} anon_sid={anon_sid_at_reg} rows={transferred}")
                except Exception as e:
                    db.session.rollback()
                    logger.warning(f"REG_SPECS_TRANSFER_FAIL user_id={user.id} error={type(e).__name__}")
                try:
                    moved_hist = _transfer_chat_history_to_user(user.id, anon_sid_at_reg)
                    if moved_hist:
                        logger.info(f"REG_HISTORY_TRANSFERRED user_id={user.id} anon_sid={anon_sid_at_reg} rows={moved_hist}")
                except Exception as e:
                    logger.warning(f"REG_HISTORY_TRANSFER_FAIL user_id={user.id} error={type(e).__name__}")
                try:
                    _ci_anon = session.get('course_info')
                    _topic_anon = session.get('master_doc_topic') or ''
                    track_event('Auth', 'Registro Completado', user_id=user.id, extra_data={
                        'anon_turns': session.get('anon_chat_count', 0),
                        'pending_doc_set': bool(session.get('pending_document')),
                        'had_course_info': bool(_ci_anon),
                        'topic_anon': (_topic_anon or '')[:80],
                    })
                except Exception:
                    pass

            aff_ref = session.pop('affiliate_ref', None)
            if aff_ref:
                try:
                    aff_sponsor = db.session.get(User, int(aff_ref))
                    if aff_sponsor and aff_sponsor.is_affiliate and aff_sponsor.affiliate_terms_accepted:
                        user.referred_by = aff_sponsor.id
                        existing_lead = AffiliateLead.query.filter_by(prospect_email=email).first()
                        if existing_lead and existing_lead.sponsor_id == aff_sponsor.id:
                            if not existing_lead.prospect_whatsapp and whatsapp:
                                existing_lead.prospect_whatsapp = whatsapp
                            if existing_lead.prospect_name != full_name:
                                existing_lead.prospect_name = full_name
                            if existing_lead.created_via == 'manual':
                                existing_lead.created_via = 'manual+autoregistro'
                            db.session.commit()
                            logger.info(f"Affiliate auto-lead: updated existing lead {existing_lead.id} with registration data, sponsor={aff_sponsor.id}")
                        elif not existing_lead:
                            checkout_url = Config.get('STRIPE_CHECKOUT_URL_AFFILIATE', '') or STRIPE_LINKS.get('PRO_AFFILIATE', '')
                            aff_lead = AffiliateLead(
                                sponsor_id=aff_sponsor.id,
                                prospect_name=full_name,
                                prospect_email=email,
                                prospect_whatsapp=whatsapp,
                                payment_link='',
                                created_via='autoregistro'
                            )
                            db.session.add(aff_lead)
                            db.session.flush()
                            if checkout_url:
                                from urllib.parse import quote as url_quote
                                aff_lead.payment_link = (checkout_url
                                    + '?prefilled_email=' + url_quote(email)
                                    + '&client_reference_id=' + f"{aff_sponsor.id}_{aff_lead.id}")
                            else:
                                aff_lead.payment_link = f"{request.host_url.rstrip('/')}/precios"
                            db.session.commit()
                            logger.info(f"Affiliate auto-lead: sponsor={aff_sponsor.id}, lead={aff_lead.id}, email={email}")
                        else:
                            db.session.commit()
                            logger.info(f"Affiliate ref: user {user.id} referred_by={aff_sponsor.id}, lead belongs to different sponsor")
                except Exception:
                    logger.warning(f"Error auto-creating affiliate lead for ref={aff_ref}, email={email}")

            def _track_post_reg(dest_endpoint, reason, **extra):
                try:
                    track_event('Funnel', 'Post-Registro Redirect', user_id=user.id, extra_data={
                        'destino': dest_endpoint,
                        'motivo': reason,
                        'anon_turns': session.get('anon_chat_count', 0),
                        'pending_doc': bool(session.get('pending_document')),
                        **extra,
                    })
                except Exception:
                    pass

            _pending_e5 = (session.pop('pending_e5_url', '') or '').strip()[:500]
            if _pending_e5 and 'youtu' in _pending_e5.lower():
                try:
                    cand_e5 = Candidato.query.filter_by(
                        ce_user_id=user.id, curp='AUTOEVAL_E5'
                    ).first()
                    if not cand_e5:
                        cand_e5 = Candidato(
                            ce_user_id=user.id,
                            nombre_completo=(user.email or 'Auto-evaluación')[:200],
                            apellidos='(Auto-evaluación E5)',
                            curp='AUTOEVAL_E5',
                        )
                        db.session.add(cand_e5); db.session.commit()
                    ev_e5 = PortafolioEvaluacion(
                        candidato_id=cand_e5.id, ce_user_id=user.id,
                        estandar='EC0217.01', status='pendiente_video',
                        e5_mode='preview', e5_unlocked=False,
                    )
                    db.session.add(ev_e5); db.session.commit()
                    db.session.add(PortafolioArchivo(
                        portafolio_id=ev_e5.id, tipo='youtube_url',
                        filename='youtube_link', source_url=_pending_e5, size_bytes=0,
                    ))
                    user.e5_preview_used_at = datetime.utcnow()
                    db.session.commit()
                    try:
                        track_event('autoevaluacion_e5', 'preview_iniciado_postregistro',
                                    user_id=user.id, extra_data={'pid': ev_e5.id})
                    except Exception:
                        pass
                    flash('¡Cuenta creada! Tu pre-dictamen IA gratuito está listo para procesar (primeros 10 min).', 'success')
                    _track_post_reg(f'/evaluar-mi-clase/{ev_e5.id}', 'e5_preview_postreg', pid=ev_e5.id)
                    return redirect(url_for('evaluar_mi_clase_vista', pid=ev_e5.id))
                except Exception as _e_e5:
                    db.session.rollback()
                    logger.warning(f"E5_PREVIEW_POSTREG_FAIL user_id={user.id} err={type(_e_e5).__name__}: {str(_e_e5)[:200]}")

            anon_turns_at_reg = session.get('anon_chat_count', 0)
            try:
                if anon_turns_at_reg > 0 and (user.chat_usage_count or 0) < anon_turns_at_reg:
                    user.chat_usage_count = min(anon_turns_at_reg, 2)
                    db.session.commit()
                    logger.info(f"REG_COUNTER_INHERITED user_id={user.id} anon_turns={anon_turns_at_reg} chat_usage_count={user.chat_usage_count}")
            except Exception as _e_ci:
                try: db.session.rollback()
                except Exception: pass
                logger.warning(f"REG_COUNTER_INHERIT_FAIL user_id={user.id} error={type(_e_ci).__name__}")
            pending_doc = session.get('pending_document')
            if pending_doc:
                try:
                    pending_doc = os.path.basename(pending_doc)
                    if '..' in pending_doc or '/' in pending_doc or '\\' in pending_doc:
                        pending_doc = None
                        raise ValueError("Invalid pending document name")
                    anon_path = os.path.join("generated_docs", pending_doc)
                    if os.path.isfile(anon_path):
                        new_basename = pending_doc.replace("anon_", f"u{user.id}_", 1) if pending_doc.startswith("anon_") else f"u{user.id}_{pending_doc}"
                        new_path = os.path.join("generated_docs", new_basename)
                        os.rename(anon_path, new_path)
                        _persist_file_to_db(new_path, user.id)
                        try:
                            import json as _json_rn
                            _hist_rows = ChatHistory.query.filter_by(user_id=user.id).all()
                            for _hr in _hist_rows:
                                if not _hr.generated_files:
                                    continue
                                try:
                                    _files = _json_rn.loads(_hr.generated_files)
                                    if not isinstance(_files, list):
                                        continue
                                except Exception:
                                    continue
                                _changed = False
                                _new_files = []
                                for _f in _files:
                                    if isinstance(_f, str) and _f == pending_doc:
                                        _new_files.append(new_basename)
                                        _changed = True
                                    elif isinstance(_f, str) and _f.startswith("anon_"):
                                        _new_files.append(_f.replace("anon_", f"u{user.id}_", 1))
                                        _changed = True
                                    else:
                                        _new_files.append(_f)
                                if _changed:
                                    _hr.generated_files = _json_rn.dumps(_new_files, ensure_ascii=False)
                            db.session.commit()
                        except Exception as _e_rn:
                            try: db.session.rollback()
                            except Exception: pass
                            logger.warning(f"REG_HIST_FILES_RENAME_FAIL user_id={user.id} error={type(_e_rn).__name__}")
                        anon_turns = session.get('anon_chat_count', 1)
                        user.chat_usage_count = min(anon_turns, 2)
                        db.session.commit()
                        track_event('Lead', 'Documento Anónimo Vinculado', user_id=user.id, extra_data={'archivo': new_basename})
                        session.pop('pending_document', None)
                        session.pop('anon_chat_count', None)
                        logger.info(f"REG_DOC_LINK_OK user_id={user.id} anon_turns={anon_turns_at_reg} redirect=/elemento/1")
                        flash('¡Cuenta creada! Tu documento está listo para descarga aquí abajo.', 'success')
                        _track_post_reg('/elemento/1', 'doc_link_ok', archivo=new_basename)
                        return redirect(url_for('elemento', num=1))
                    else:
                        logger.warning(f"REG_DOC_FILE_MISSING user_id={user.id} anon_turns={anon_turns_at_reg} pending_doc_set=true redirect=/elemento/1")
                        session.pop('pending_document', None)
                        flash('¡Cuenta creada! Continúa tu conversación aquí para terminar tu Carta Descriptiva.', 'success')
                        _track_post_reg('/elemento/1', 'doc_file_missing')
                        return redirect(url_for('elemento', num=1))
                except Exception as e:
                    logger.warning(f"REG_DOC_LINK_FAIL user_id={user.id} anon_turns={anon_turns_at_reg} error_type={type(e).__name__} redirect=/elemento/1")
                    flash('¡Cuenta creada! Continúa tu conversación aquí para terminar tu Carta Descriptiva.', 'success')
                    _track_post_reg('/elemento/1', 'doc_link_fail', error_type=type(e).__name__)
                    return redirect(url_for('elemento', num=1))
            elif anon_turns_at_reg > 0:
                logger.info(f"REG_RESUME_ELEMENTO user_id={user.id} anon_turns={anon_turns_at_reg} redirect=/elemento/1")
                flash('¡Bienvenido! Continúa donde te quedaste — sigue conversando para terminar tu Carta Descriptiva.', 'success')
                _track_post_reg('/elemento/1', 'resume_elemento')
                return redirect(url_for('elemento', num=1))
            else:
                _intent_elem = session.pop('post_registro_elemento', None)
                if _intent_elem and isinstance(_intent_elem, int) and 1 <= _intent_elem <= 4:
                    logger.info(f"REG_INTENT_ELEMENTO user_id={user.id} elemento={_intent_elem} redirect=/elemento/{_intent_elem}")
                    flash('¡Cuenta creada! Continúa donde te quedaste.', 'success')
                    _track_post_reg(f'/elemento/{_intent_elem}', 'intent_elemento', elemento=_intent_elem)
                    return redirect(url_for('elemento', num=_intent_elem))
                logger.info(f"REG_NO_PENDING_DOC user_id={user.id} anon_turns={anon_turns_at_reg} redirect=/home")

            _signup_promo = _get_session_promo()
            if _signup_promo:
                flash('¡Cuenta creada! Tu promo Día del Maestro queda activa 7 días — pruébalo gratis ahora y aprovéchala cuando quieras.', 'success')
            else:
                flash('Cuenta creada exitosamente. ¡Bienvenido a Pertinentia!', 'success')
            _track_post_reg('/home', 'no_pending_doc', signup_promo=bool(_signup_promo))
            return redirect(url_for('home'))
    return render_template("registro.html", title="Crear Cuenta", active_page="registro")

@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# ============================================================
# SEO — Capa 1 técnica + Cluster 2 landings públicas + Blog
# ============================================================
SEO_LANDINGS = {
    'carta-descriptiva-ec0301-y-ec0217-01-ejemplo': {
        'template': 'seo/carta_descriptiva.html',
        'title': 'Carta Descriptiva EC0301 y EC0217.01 — Ejemplo y Generador IA',
        'priority': 0.9,
    },
    'instrumentos-evaluacion-ec0301-y-ec0217-01-ejemplo': {
        'template': 'seo/instrumentos_evaluacion.html',
        'title': 'Instrumentos de Evaluación EC0301 y EC0217.01 — Ejemplo',
        'priority': 0.9,
    },
    'manuales-curso-ec0301-y-ec0217-01': {
        'template': 'seo/manuales_curso.html',
        'title': 'Manuales del Curso EC0301 y EC0217.01 — Participante e Instructor',
        'priority': 0.9,
    },
    'auto-diagnostico-ec0301-y-ec0217-01': {
        'template': 'seo/auto_diagnostico.html',
        'title': 'Auto-diagnóstico EC0301 y EC0217.01 — Mide tu nivel',
        'priority': 0.9,
    },
    'diseno-instruccional-conocer-para-agentes-capacitadores-externos-stps': {
        'template': 'seo/diseno_instruccional_ace_stps.html',
        'title': 'Diseño Instruccional CONOCER para ACE e instructores STPS',
        'priority': 0.85,
    },
    'alinear-mis-cursos-conocer-sin-certificarme': {
        'template': 'seo/alinear_cursos_sin_certificarme.html',
        'title': 'Alinea tus cursos a la metodología CONOCER sin certificarte',
        'priority': 0.85,
    },
    'red-profesional-capacitacion-certificada-mexico': {
        'template': 'seo/red_profesional_capacitacion.html',
        'title': 'Red Profesional de Capacitación Certificada en México — PertinentIA',
        'priority': 0.85,
    },
}

def _seo_landing_view(slug):
    cfg = SEO_LANDINGS.get(slug)
    if not cfg:
        from flask import abort as _abort
        _abort(404)
    src = (request.args.get('from') or '').strip()[:60] or f'seo_{slug}'
    session['seo_origin'] = src
    _utm_src = (request.args.get('utm_source') or '').strip()[:80]
    _utm_med = (request.args.get('utm_medium') or '').strip()[:80]
    _utm_cmp = (request.args.get('utm_campaign') or '').strip()[:150]
    _utm_trm = (request.args.get('utm_term') or '').strip()[:200]
    _utm_ctn = (request.args.get('utm_content') or '').strip()[:250]
    _utm_id  = (request.args.get('utm_id') or '').strip()[:50]
    if _utm_src or _utm_cmp or _utm_ctn:
        if not session.get('utm_first_touch'):
            session['utm_first_touch'] = {
                'utm_source': _utm_src, 'utm_medium': _utm_med,
                'utm_campaign': _utm_cmp, 'utm_term': _utm_trm,
                'utm_content': _utm_ctn, 'utm_id': _utm_id,
                'utm_landing': slug[:120],
                'utm_captured_at': datetime.utcnow().isoformat(),
            }
        session['utm_last_touch'] = {
            'utm_source': _utm_src, 'utm_medium': _utm_med,
            'utm_campaign': _utm_cmp, 'utm_term': _utm_trm,
            'utm_content': _utm_ctn, 'utm_id': _utm_id,
            'utm_landing': slug[:120],
        }
    try:
        track_event('SEO Landing', 'View',
                    user_id=current_user.id if current_user.is_authenticated else None,
                    extra_data={'slug': slug, 'utm_source': _utm_src,
                                'utm_medium': _utm_med, 'utm_campaign': _utm_cmp,
                                'utm_term': _utm_trm, 'utm_content': _utm_ctn,
                                'utm_id': _utm_id,
                                'referrer': (request.referrer or '')[:200]})
    except Exception:
        pass
    return render_template(cfg['template'], title=cfg['title'], active_page='seo', ads=[])

@app.route('/carta-descriptiva-ec0301-y-ec0217-01-ejemplo')
def seo_carta_descriptiva():
    return _seo_landing_view('carta-descriptiva-ec0301-y-ec0217-01-ejemplo')

@app.route('/instrumentos-evaluacion-ec0301-y-ec0217-01-ejemplo')
def seo_iecs():
    return _seo_landing_view('instrumentos-evaluacion-ec0301-y-ec0217-01-ejemplo')

@app.route('/manuales-curso-ec0301-y-ec0217-01')
def seo_manuales():
    return _seo_landing_view('manuales-curso-ec0301-y-ec0217-01')

@app.route('/auto-diagnostico-ec0301-y-ec0217-01')
def seo_autodiagnostico():
    return _seo_landing_view('auto-diagnostico-ec0301-y-ec0217-01')

@app.route('/diseno-instruccional-conocer-para-agentes-capacitadores-externos-stps')
def seo_ace_stps():
    return _seo_landing_view('diseno-instruccional-conocer-para-agentes-capacitadores-externos-stps')

@app.route('/alinear-mis-cursos-conocer-sin-certificarme')
def seo_alinear_sin_certificarme():
    return _seo_landing_view('alinear-mis-cursos-conocer-sin-certificarme')

@app.route('/red-profesional-capacitacion-certificada-mexico')
def seo_red_profesional():
    return _seo_landing_view('red-profesional-capacitacion-certificada-mexico')

@app.route('/blog')
def blog_index():
    try:
        from seo_blog_posts import list_posts
        posts = list_posts()
    except Exception:
        posts = []
    try:
        track_event('SEO Blog', 'Index View',
                    user_id=current_user.id if current_user.is_authenticated else None)
    except Exception:
        pass
    return render_template('blog_index.html', title='Blog EC0301 y EC0217.01',
                           active_page='blog', ads=[], posts=posts)

@app.route('/blog/<slug>')
def blog_post(slug):
    from flask import abort as _abort
    try:
        from seo_blog_posts import get_post_by_slug
        post = get_post_by_slug(slug)
    except Exception:
        post = None
    if not post:
        _abort(404)
    try:
        track_event('SEO Blog', 'Post View',
                    user_id=current_user.id if current_user.is_authenticated else None,
                    extra_data={'slug': slug})
    except Exception:
        pass
    return render_template('blog_post.html', title=post['title'],
                           active_page='blog', ads=[], post=post)

@app.route('/robots.txt')
def robots_txt():
    base = (os.environ.get('PUBLIC_BASE_URL') or 'https://app.pertinentia.com/sitemap.xml').rstrip('/')
    body = (
        "User-agent: *\n"
        "Allow: /\n"
        "Disallow: /admin-crm\n"
        "Disallow: /api/\n"
        "Disallow: /elemento/\n"
        "Disallow: /documentos\n"
        "Disallow: /perfil\n"
        "Disallow: /dashboard-ce\n"
        "Disallow: /reset-password\n"
        "Disallow: /forgot-password\n"
        "Disallow: /e/o/\n"
        f"\nSitemap: {base}/sitemap.xml\n"
    )
    from flask import Response as _Resp
    return _Resp(body, mimetype='text/plain')

@app.route('/sitemap.xml')
def sitemap_xml():
    base = (os.environ.get('PUBLIC_BASE_URL') or 'https://app.pertinentia.com/sitemap.xml').rstrip('/')
    from datetime import datetime as _dt
    today = _dt.utcnow().strftime('%Y-%m-%d')
    urls = [
        (base + '/', '1.0', today, 'daily'),
        (base + '/precios', '0.8', today, 'weekly'),
        (base + '/certificado', '0.8', today, 'weekly'),
        (base + '/affiliate', '0.7', today, 'monthly'),
        (base + '/blog', '0.8', today, 'weekly'),
    ]
    for slug, cfg in SEO_LANDINGS.items():
        urls.append((base + '/' + slug, str(cfg['priority']), today, 'weekly'))
    try:
        from seo_blog_posts import list_posts
        for p in list_posts():
            urls.append((base + '/blog/' + p['slug'], '0.7',
                         p['published_at'].strftime('%Y-%m-%d'), 'monthly'))
    except Exception:
        pass
    items = '\n'.join(
        f'  <url><loc>{u}</loc><lastmod>{lm}</lastmod><changefreq>{cf}</changefreq><priority>{pr}</priority></url>'
        for (u, pr, lm, cf) in urls
    )
    body = ('<?xml version="1.0" encoding="UTF-8"?>\n'
            '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
            f'{items}\n</urlset>\n')
    from flask import Response as _Resp
    return _Resp(body, mimetype='application/xml')

@app.route('/favicon.ico')
def favicon_ico():
    from flask import send_from_directory as _sfd
    try:
        return _sfd('static', 'favicon.png', mimetype='image/png')
    except Exception:
        from flask import Response as _Resp
        return _Resp('', status=204)

@app.route('/api/admin/seo/metrics')
@login_required
def admin_seo_metrics():
    if not current_user.is_admin:
        return jsonify({'error': 'forbidden'}), 403
    from datetime import datetime as _dt, timedelta as _td
    now = _dt.utcnow()
    windows = {'24h': now - _td(hours=24), '7d': now - _td(days=7), '30d': now - _td(days=30)}
    out = {'windows': {}, 'top_landings': [], 'top_blog': [], 'conversions': {}, 'health': {}}
    try:
        for label, since in windows.items():
            views = UserEvent.query.filter(
                UserEvent.event_category == 'SEO Landing',
                UserEvent.event_action == 'View',
                UserEvent.timestamp >= since
            ).count()
            blog_views = UserEvent.query.filter(
                UserEvent.event_category == 'SEO Blog',
                UserEvent.timestamp >= since
            ).count()
            out['windows'][label] = {'landing_views': views, 'blog_views': blog_views}
        # top landings 30d
        rows = db.session.query(
            UserEvent.metadata_json, db.func.count(UserEvent.id)
        ).filter(
            UserEvent.event_category == 'SEO Landing',
            UserEvent.event_action == 'View',
            UserEvent.timestamp >= windows['30d']
        ).group_by(UserEvent.metadata_json).all()
        import json as _json
        agg = {}
        for meta, c in rows:
            slug = ''
            try:
                slug = (_json.loads(meta or '{}') or {}).get('slug', '') or ''
            except Exception:
                pass
            if not slug:
                continue
            agg[slug] = agg.get(slug, 0) + int(c)
        out['top_landings'] = sorted(
            [{'slug': k, 'views': v} for k, v in agg.items()],
            key=lambda x: x['views'], reverse=True
        )
        # conversions: registros con seo_origin
        conv = UserEvent.query.filter(
            UserEvent.event_category == 'Auth',
            UserEvent.event_action == 'Registro SEO',
            UserEvent.timestamp >= windows['30d']
        ).count()
        out['conversions'] = {
            '30d_registros_desde_seo': conv,
            '30d_landing_views': out['windows']['30d']['landing_views'],
            'tasa_conversion_pct': round((conv / out['windows']['30d']['landing_views'] * 100), 2)
                if out['windows']['30d']['landing_views'] > 0 else 0.0,
        }
    except Exception as _e:
        out['error'] = type(_e).__name__
    out['health'] = {
        'ga4_configured': bool(os.environ.get('GA4_MEASUREMENT_ID', '').strip()),
        'gsc_configured': bool(os.environ.get('GSC_VERIFICATION_TOKEN', '').strip()),
        'public_base_url': (os.environ.get('PUBLIC_BASE_URL') or 'https://app.pertinentia.com/sitemap.xml').rstrip('/'),
        'sitemap_url': (os.environ.get('PUBLIC_BASE_URL') or 'https://app.pertinentia.com/sitemap.xml').rstrip('/') + '/sitemap.xml',
        'robots_url': (os.environ.get('PUBLIC_BASE_URL') or 'https://app.pertinentia.com/sitemap.xml').rstrip('/') + '/robots.txt',
        'public_landings_count': len(SEO_LANDINGS),
    }
    return jsonify(out)

@app.route("/")
def home():
    from datetime import datetime
    if current_user.is_authenticated:
        gen_docs = _user_generated_docs()
        generated_count = len(gen_docs)
    else:
        generated_count = 0
    today = datetime.utcnow().date()
    _hide_ads = False
    _hide_reason = ''
    if current_user.is_authenticated:
        _tier = getattr(current_user, 'tier', 'FREE')
        if _tier in ('PRO', 'PREMIUM'):
            _hide_ads = True
            _hide_reason = f'tier_{_tier}'
        elif getattr(current_user, 'created_at', None):
            try:
                _age = datetime.utcnow() - current_user.created_at
                if _age < timedelta(hours=48):
                    _hide_ads = True
                    _hide_reason = f'recent_user_{_age.total_seconds()/3600:.1f}h'
            except Exception:
                pass
    if _hide_ads:
        ads_data = []
        logger.info(f"ADS_HIDDEN user_id={current_user.id} reason={_hide_reason}")
    else:
        active_ads = AdCampaign.query.filter(
            AdCampaign.status == 'Activo',
            AdCampaign.starts_at <= today,
            AdCampaign.ends_at >= today
        ).all()
        ads_data = [{
            "id": ad.id,
            "title": ad.title,
            "description": ad.description,
            "image_url": ad.image_url,
            "target_url": ad.target_url
        } for ad in active_ads]
    return render_template("home.html",
                           title="Inicio",
                           active_page="home",
                           generated_count=generated_count,
                           ads=ads_data)

@app.route("/elemento/<int:num>")
def elemento(num):
    from datetime import datetime
    if num < 1 or num > 4:
        return "Elemento no encontrado", 404
    _coupon_arg_el = (request.args.get('coupon') or request.args.get('cupon') or '').strip().upper()[:40]
    if _coupon_arg_el:
        import re as _re_cp_el
        if _re_cp_el.fullmatch(r'[A-Z0-9_\-]{3,40}', _coupon_arg_el):
            session['coupon_ref'] = _coupon_arg_el
    _email_arg_el = (request.args.get('email') or '').strip().lower()[:120]
    if _email_arg_el and '@' in _email_arg_el:
        session['prefill_email'] = _email_arg_el
    config = ELEMENT_CONFIG[num]
    if current_user.is_authenticated:
        gen_docs = _user_generated_docs()
        eval_proc = EvaluationProcess.query.filter_by(user_id=current_user.id).first()
    else:
        gen_docs = []
        eval_proc = None
        try:
            session['post_registro_elemento'] = int(num)
        except Exception:
            pass
    today = datetime.utcnow().date()
    active_ads = AdCampaign.query.filter(
        AdCampaign.status == 'Activo',
        AdCampaign.starts_at <= today,
        AdCampaign.ends_at >= today
    ).all()
    ads_data = [{
        "id": ad.id,
        "title": ad.title,
        "description": ad.description,
        "image_url": ad.image_url,
        "target_url": ad.target_url
    } for ad in active_ads]
    free_dl_used = current_user.free_downloads_used if current_user.is_authenticated else 0
    placeholders = {
        1: 'Ej: Repostería Básica, Soldadura Industrial, Primeros Auxilios...',
        2: 'Ej: Evaluación Diagnóstica para curso de Repostería...',
        3: 'Ej: Manual del Instructor para curso de Soldadura...',
        4: ''
    }
    saved_topic = ''
    # El nombre de respaldo del badge "Curso activo" DEBE provenir ÚNICAMENTE del
    # CURSO ACTIVO (su topic persistido o sus propios documentos), nunca de la lista
    # global de documentos del usuario. De lo contrario, al pulsar "Iniciar Nuevo
    # Curso" el badge (y la siguiente generación) resucitaban el nombre del curso
    # anterior a partir de una Carta Descriptiva ya cerrada.
    if current_user.is_authenticated and current_user.active_course_session_id:
        _act_id = current_user.active_course_session_id
        try:
            _act_cs = db.session.get(CourseSession, _act_id)
        except Exception:
            _act_cs = None
        if _act_cs is not None:
            _DIRTY_PREFIXES = ('genera', 'crea', 'diseña', 'disena', 'haz')
            try:
                _topic_clean = _sanitize_course_name_from_message(_act_cs.topic or '') or (_act_cs.topic or '')
            except Exception:
                _topic_clean = (_act_cs.topic or '')
            _topic_clean = (_topic_clean or '').strip()
            if (not _topic_clean) or _topic_clean.lower().startswith(_DIRTY_PREFIXES):
                try:
                    _real = _extract_course_name_from_cs_docs(current_user.id, _act_id)
                    if _real:
                        _topic_clean = _real.strip()
                except Exception:
                    pass
            if _topic_clean and not _topic_clean.lower().startswith(_DIRTY_PREFIXES):
                saved_topic = _topic_clean
    has_alacarte = False
    alacarte_remaining = 0
    if current_user.is_authenticated and current_user.tier == 'FREE':
        has_alacarte = current_user.has_alacarte(num)
        alacarte_remaining = getattr(current_user, f'alacarte_e{num}', 0) if has_alacarte else 0
    if current_user.is_authenticated:
        _no_prior_gen = (len(gen_docs) == 0 and (current_user.chat_usage_count or 0) == 0)
        _tier_ok = (current_user.tier == 'FREE')
    else:
        _no_prior_gen = (session.get('anon_chat_count', 0) == 0)
        _tier_ok = True
    show_demo_precargado = (num == 1 and _no_prior_gen and _tier_ok)
    try:
        _funnel_uid = current_user.id if current_user.is_authenticated else None
        _funnel_tier = current_user.tier if current_user.is_authenticated else 'ANON'
        track_event('Funnel', 'Vista Elemento', user_id=_funnel_uid, extra_data={
            'modulo': num, 'tier': _funnel_tier, 'has_docs': len(gen_docs) > 0,
            'demo_visible': bool(show_demo_precargado)
        })
    except Exception:
        pass
    try:
        if current_user.is_authenticated:
            _get_or_create_active_course_session(current_user)
            try:
                _anon_sid_pending = session.get('anon_sid')
                if _anon_sid_pending and current_user.active_course_session_id:
                    _moved = _transfer_chat_history_to_user(current_user.id, _anon_sid_pending)
                    if _moved:
                        try: session.pop('anon_sid', None)
                        except Exception: pass
                        logger.info(f"render: post-auth retry transferred {_moved} anon rows for user {current_user.id}")
            except Exception as _etr:
                logger.warning(f"render: post-auth transfer retry failed: {_etr}")
            try:
                _cs_now = CourseSession.query.get(current_user.active_course_session_id) if current_user.active_course_session_id else None
                if _cs_now is not None:
                    if _cs_now.master_doc and not session.get('master_doc'):
                        session['master_doc'] = _cs_now.master_doc
                    if _cs_now.topic and not session.get('master_doc_topic'):
                        session['master_doc_topic'] = _cs_now.topic
                    if _cs_now.course_info_json and not session.get('course_info'):
                        try:
                            import json as _jrend
                            _ci_rend = _jrend.loads(_cs_now.course_info_json)
                            if isinstance(_ci_rend, dict):
                                session['course_info'] = _ci_rend
                        except Exception:
                            pass
            except Exception:
                pass
    except Exception:
        pass
    try:
        prior_hist_payload = _load_full_chat_history(num)
    except Exception:
        prior_hist_payload = {'messages': [], 'course_topic': None, 'generated_files': []}
    if not saved_topic and prior_hist_payload.get('course_topic'):
        saved_topic = prior_hist_payload['course_topic']
    course_session_meta = {'active_id': None, 'session_num': None, 'total': 0, 'topic': ''}
    if current_user.is_authenticated:
        try:
            _all_cs = CourseSession.query.filter_by(user_id=current_user.id).count()
            _act = CourseSession.query.get(current_user.active_course_session_id) if current_user.active_course_session_id else None
            course_session_meta = {
                'active_id': current_user.active_course_session_id,
                'session_num': (_act.session_num if _act else None),
                'total': _all_cs,
                'topic': (_act.topic if _act else '') or '',
            }
        except Exception:
            pass
    quick_actions_final = list(config["quick_actions"])
    if num == 2 and session.get('contexto_institucional') == 'uam':
        quick_actions_final.append({"label": "Evaluación Mediadora", "prompt": "Genera la Evaluación Mediadora del curso: evalúa la calidad de las interacciones docente-participante, el tipo de retroalimentación proporcionada durante el proceso, y la mediación entre pares. Especialmente relevante en modalidad en línea o híbrida. Incluye al menos 5 reactivos con valores y criterios de cumplimiento."})
    return render_template("elemento.html",
                           title=config["page_title"],
                           active_page=config["active_page"],
                           page_title=config["page_title"],
                           page_description=config["page_description"],
                           welcome_message=config["welcome_message"],
                           welcome_message_ready=config.get("welcome_message_ready", ""),
                           quick_actions=quick_actions_final,
                           element_num=num,
                           generated_docs=gen_docs,
                           eval_process=eval_proc,
                           free_downloads_used=free_dl_used,
                           module_placeholder=placeholders.get(num, ''),
                           ads=ads_data,
                           saved_course_topic=saved_topic,
                           has_alacarte=has_alacarte,
                           alacarte_remaining=alacarte_remaining,
                           show_demo_precargado=show_demo_precargado,
                           prior_chat_history=prior_hist_payload.get('messages', []),
                           prior_generated_files=prior_hist_payload.get('generated_files', []),
                           course_session_meta=course_session_meta,
                           stripe_links=STRIPE_LINKS)

@app.route("/documentos")
@login_required
def documentos():
    gen_docs = _user_generated_docs()
    return render_template("documentos.html",
                           title="Mis Documentos",
                           active_page="docs",
                           gen_docs=gen_docs,
                           free_downloads_used=current_user.free_downloads_used)

@app.route("/certificado")
@login_required
def certificado():
    return render_template("certificado.html",
                           title="Licencia Centro Evaluador IA",
                           active_page="cert")

@app.route("/precios")
def precios():
    is_anon = not current_user.is_authenticated
    if is_anon:
        track_event('Navegacion', 'Vista Precios Anonimo',
                    extra_data={'promo': request.args.get('promo', ''),
                                'utm_source': request.args.get('utm_source', ''),
                                'utm_content': request.args.get('utm_content', '')})
    else:
        track_event('Navegacion', 'Vista Precios', user_id=current_user.id, extra_data={'tier': current_user.tier})
    is_referred = False
    affiliate_pro_link = ''
    if not is_anon and current_user.referred_by:
        sponsor = db.session.get(User, current_user.referred_by)
        if sponsor and sponsor.is_affiliate and sponsor.affiliate_terms_accepted:
            is_referred = True
            affiliate_pro_link = _get_stripe_link_for_tier('PRO_AFFILIATE')
    return render_template("precios.html",
                           title="Planes y Precios",
                           active_page="precios",
                           stripe_links=STRIPE_LINKS,
                           is_referred=is_referred,
                           affiliate_pro_link=affiliate_pro_link)

@app.route("/publicitar")
@login_required
def publicitar():
    my_ads = AdCampaign.query.filter_by(advertiser_id=current_user.id).order_by(AdCampaign.id.desc()).all()
    cost_per_day = float(Config.get('COST_PER_DAY_ADS', '150'))
    if request.args.get('paid') == '1':
        try:
            from datetime import datetime as _dtp, timedelta as _tdp
            _cutoff = _dtp.utcnow() - _tdp(hours=2)
            recent_ad = AdCampaign.query.filter(
                AdCampaign.advertiser_id == current_user.id,
                AdCampaign.status.in_(['Activa', 'Pagada', 'Aprobada', 'En Revisión']),
                AdCampaign.created_at >= _cutoff
            ).order_by(AdCampaign.id.desc()).first() if hasattr(AdCampaign, 'created_at') else None
            if recent_ad is None:
                recent_ad = my_ads[0] if my_ads else None
            _amount = float(getattr(recent_ad, 'total_cost', 0) or 0) if recent_ad else 0
            _queue_pixel_event('Purchase', {
                'value': _amount,
                'currency': 'MXN',
                'content_name': 'Ad Campaign',
                'content_category': 'advertisement',
                'content_ids': [str(recent_ad.id)] if recent_ad else []
            })
        except Exception as _epp:
            logger.warning(f"publicitar pixel Purchase queue failed: {_epp}")
    return render_template("publicitar.html",
                           title="Portal de Anunciantes",
                           active_page="publicitar",
                           my_ads=my_ads,
                           cost_per_day=cost_per_day)

@app.route("/r/<int:ref_id>")
def affiliate_landing(ref_id):
    sponsor = db.session.get(User, ref_id)
    if not sponsor or not sponsor.is_affiliate or not sponsor.affiliate_terms_accepted:
        return redirect(url_for('home'))
    if current_user.is_authenticated:
        if not current_user.referred_by:
            current_user.referred_by = sponsor.id
            db.session.commit()
    else:
        session['affiliate_ref'] = ref_id
    return render_template("affiliate_landing.html",
                           title="PertinentIA EC0301 — La Plataforma de la Capacitación Certificada",
                           sponsor=sponsor,
                           ref_id=ref_id,
                           sponsor_id=sponsor.id)

@app.route("/api/affiliate-video-played", methods=['POST'])
def affiliate_video_played():
    try:
        data = request.get_json(silent=True) or {}
        sid = data.get('sponsor_id')
        rid = (data.get('ref_id') or '')[:50]
        if not sid:
            return jsonify({'ok': False}), 400
        try:
            sid = int(sid)
        except (TypeError, ValueError):
            return jsonify({'ok': False}), 400
        sponsor = User.query.get(sid)
        if not sponsor or not sponsor.is_affiliate:
            return jsonify({'ok': False}), 404
        ip = (request.headers.get('X-Forwarded-For', request.remote_addr) or '')[:64]
        ua = (request.headers.get('User-Agent', '') or '')[:300]
        view = AffiliateVideoView(sponsor_id=sid, ref_id=rid, ip=ip, user_agent=ua)
        db.session.add(view)
        db.session.commit()
        return jsonify({'ok': True})
    except Exception as e:
        db.session.rollback()
        app.logger.warning(f"affiliate_video_played error: {e}")
        return jsonify({'ok': False}), 500

@app.route("/afiliados")
@login_required
def afiliados():
    leads = AffiliateLead.query.filter_by(sponsor_id=current_user.id).order_by(AffiliateLead.id.desc()).all()
    stripe_checkout_url = Config.get('STRIPE_CHECKOUT_URL_AFFILIATE', '') or STRIPE_LINKS.get('PRO_AFFILIATE', '')
    comision_pct = Config.get('COMISION_PORCENTAJE', '0.30')
    video_views_total = AffiliateVideoView.query.filter_by(sponsor_id=current_user.id).count()
    from datetime import datetime, timedelta
    since_7d = datetime.utcnow() - timedelta(days=7)
    video_views_7d = AffiliateVideoView.query.filter(
        AffiliateVideoView.sponsor_id == current_user.id,
        AffiliateVideoView.created_at >= since_7d
    ).count()
    return render_template("afiliados.html",
                           title="Panel de Afiliados",
                           active_page="afiliados",
                           leads=leads,
                           stripe_links=STRIPE_LINKS,
                           stripe_checkout_url=stripe_checkout_url,
                           comision_pct=comision_pct,
                           tier_labels=TIER_LABELS,
                           tier_prices=TIER_PRICES,
                           video_views_total=video_views_total,
                           video_views_7d=video_views_7d)

@app.route("/api/afiliados/accept_terms", methods=["POST"])
@login_required
def accept_affiliate_terms():
    current_user.affiliate_terms_accepted = True
    current_user.is_affiliate = True
    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/afiliados/update_regime", methods=["POST"])
@login_required
def update_affiliate_regime():
    data = request.json
    regime = data.get("regime", "RESICO")
    if regime not in ("RESICO", "PFAE"):
        return jsonify({"error": "Régimen no válido"}), 400
    current_user.tax_regime = regime
    db.session.commit()
    return jsonify({"success": True, "regime": regime})

@app.route("/api/afiliados/register", methods=["POST"])
@login_required
def register_affiliate_lead():
    if not current_user.affiliate_terms_accepted:
        return jsonify({"error": "Debes aceptar los términos del programa de afiliados."}), 403
    data = request.json
    name = (data.get("prospect_name") or "").strip()
    email = (data.get("prospect_email") or "").strip().lower()
    whatsapp = (data.get("prospect_whatsapp") or "").strip()
    ce_key = (data.get("ce_key") or "").strip()
    selected_tier = (data.get("selected_tier") or "PRO_AFFILIATE").strip()
    valid_tiers = list(STRIPE_LINKS.keys())
    if selected_tier not in valid_tiers:
        selected_tier = 'PRO_AFFILIATE'
    if not name or not email:
        return jsonify({"error": "Nombre y correo del prospecto son obligatorios."}), 400
    existing_lead = AffiliateLead.query.filter_by(prospect_email=email).first()
    if existing_lead:
        if existing_lead.sponsor_id == current_user.id:
            return jsonify({"error": "Este prospecto ya está en tu cartera.", "existing": True, "lead_id": existing_lead.id}), 400
        else:
            return jsonify({"error": "Este correo ya está asignado a otro afiliado."}), 400
    existing_user = User.query.filter_by(email=email).first()
    if existing_user:
        return jsonify({"error": "Este correo ya pertenece a un usuario registrado en Pertinentia."}), 400
    from urllib.parse import quote
    stripe_link = _get_stripe_link_for_tier(selected_tier)
    lead = AffiliateLead(
        sponsor_id=current_user.id,
        prospect_name=name,
        prospect_email=email,
        prospect_whatsapp=whatsapp,
        ce_key=ce_key,
        payment_link='',
        selected_tier=selected_tier,
        created_via='manual'
    )
    db.session.add(lead)
    try:
        db.session.flush()
        if stripe_link:
            payment_link = (stripe_link
                            + '?prefilled_email=' + quote(email)
                            + '&client_reference_id=' + f"{current_user.id}_{lead.id}")
        else:
            payment_link = f"{request.host_url.rstrip('/')}/registro?ref={current_user.id}"
        lead.payment_link = payment_link
        db.session.commit()
    except Exception:
        db.session.rollback()
        return jsonify({"error": "Error al registrar prospecto. Intenta de nuevo."}), 400
    try:
        _tier_amount_map = {
            'PRO_AFFILIATE': 500, 'PRO_PROJECT': 999, 'PRO_MULTICURSO': 2129,
            'PREMIUM_MONTHLY': 2997, 'PREMIUM_ANNUAL': 29970,
            'ALACARTE_E1': 299, 'ALACARTE_E2': 399, 'ALACARTE_E3': 399, 'ALACARTE_E4': 249,
        }
        _queue_pixel_event('Lead', {
            'content_name': selected_tier,
            'content_category': 'affiliate_prospect',
            'value': _tier_amount_map.get(selected_tier, 0),
            'currency': 'MXN'
        })
    except Exception:
        pass
    return jsonify({
        "success": True,
        "message": f"Prospecto '{name}' registrado exitosamente.",
        "payment_link": payment_link,
        "lead_id": lead.id,
        "selected_tier": selected_tier,
        "pixel_event": {"name": "Lead", "content_name": selected_tier}
    })

@app.route("/api/afiliados/delete/<int:lead_id>", methods=["POST", "DELETE"])
@login_required
def delete_affiliate_lead(lead_id):
    lead = db.session.get(AffiliateLead, lead_id)
    if not lead:
        return jsonify({"error": "Prospecto no encontrado."}), 404
    if lead.sponsor_id != current_user.id:
        return jsonify({"error": "No autorizado."}), 403
    db.session.delete(lead)
    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/afiliados/change_tier/<int:lead_id>", methods=["POST"])
@login_required
def change_lead_tier(lead_id):
    if not current_user.affiliate_terms_accepted:
        return jsonify({"error": "No autorizado."}), 403
    lead = db.session.get(AffiliateLead, lead_id)
    if not lead or lead.sponsor_id != current_user.id:
        return jsonify({"error": "Prospecto no encontrado."}), 404
    if lead.status == 'Pagado':
        return jsonify({"error": "No se puede cambiar el plan de un prospecto que ya pagó."}), 400
    data = request.json
    new_tier = (data.get("selected_tier") or "").strip()
    if new_tier not in STRIPE_LINKS:
        return jsonify({"error": "Plan no válido."}), 400
    from urllib.parse import quote
    stripe_link = _get_stripe_link_for_tier(new_tier)
    if stripe_link:
        payment_link = (stripe_link
                        + '?prefilled_email=' + quote(lead.prospect_email)
                        + '&client_reference_id=' + f"{current_user.id}_{lead.id}")
    else:
        payment_link = lead.payment_link
    lead.selected_tier = new_tier
    lead.payment_link = payment_link
    db.session.commit()
    return jsonify({
        "success": True,
        "payment_link": payment_link,
        "selected_tier": new_tier,
        "tier_label": TIER_LABELS.get(new_tier, new_tier)
    })

REGIME_SAT_MAP = {'RESICO': ['625', '626'], 'PFAE': ['612']}

def _audit_cfdi_xml(xml_bytes, lead, sponsor):
    from lxml import etree
    try:
        parser = etree.XMLParser(resolve_entities=False, no_network=True)
        tree = etree.fromstring(xml_bytes, parser=parser)
    except etree.XMLSyntaxError as e:
        return f"El archivo XML no es un CFDI válido: {e}"

    ns_strip = lambda tag: tag.split('}')[-1] if '}' in tag else tag

    comprobante = tree
    if ns_strip(comprobante.tag) != 'Comprobante':
        return "El XML no contiene un nodo Comprobante raíz."

    emisor = None
    for child in comprobante:
        if ns_strip(child.tag) == 'Emisor':
            emisor = child
            break
    if emisor is None:
        return "El XML no contiene un nodo Emisor."

    xml_regimen = emisor.get('RegimenFiscal', '')
    expected_regimes = REGIME_SAT_MAP.get(sponsor.tax_regime, [])
    if xml_regimen not in expected_regimes:
        return (f"Régimen fiscal del CFDI ({xml_regimen}) no coincide con tu régimen registrado "
                f"({sponsor.tax_regime} → códigos SAT esperados: {', '.join(expected_regimes)}).")

    xml_total_str = comprobante.get('Total', '0')
    try:
        xml_total = float(xml_total_str)
    except ValueError:
        return f"El Total del CFDI no es un número válido: {xml_total_str}"

    base = lead.commission_amount
    if base is None or base <= 0:
        return "La comisión del lead no ha sido calculada. El pago aún no fue procesado por el webhook."

    ret_isr_xml = 0.0
    ret_iva_xml = 0.0
    for elem in comprobante.iter():
        tag = ns_strip(elem.tag)
        if tag == 'Retencion':
            impuesto = elem.get('Impuesto', '')
            importe_str = elem.get('Importe', '0')
            try:
                importe = float(importe_str)
            except ValueError:
                importe = 0.0
            if impuesto == '001':
                ret_isr_xml += importe
            elif impuesto == '002':
                ret_iva_xml += importe

    if sponsor.tax_regime == 'RESICO':
        ret_isr_esperada = round(base * 0.0125, 2)
    elif sponsor.tax_regime == 'PFAE':
        ret_isr_esperada = round(base * 0.10, 2)
    else:
        ret_isr_esperada = 0.0

    ret_iva_esperada = round(base * 0.106667, 2)

    tolerancia = 0.10
    errores = []

    if abs(ret_isr_xml - ret_isr_esperada) > tolerancia:
        errores.append(
            f"Retención ISR: CFDI ${ret_isr_xml:.2f} vs esperado ${ret_isr_esperada:.2f} "
            f"(base ${base:.2f} × {'1.25%' if sponsor.tax_regime == 'RESICO' else '10%'})")

    if abs(ret_iva_xml - ret_iva_esperada) > tolerancia:
        errores.append(
            f"Retención IVA: CFDI ${ret_iva_xml:.2f} vs esperado ${ret_iva_esperada:.2f} "
            f"(base ${base:.2f} × 10.6667%)")

    if errores:
        return "Audit-Gate fiscal rechazado:\n• " + "\n• ".join(errores)

    return None

@app.route("/api/afiliados/upload_cfdi/<int:lead_id>", methods=["POST"])
@login_required
def upload_cfdi(lead_id):
    lead = db.session.get(AffiliateLead, lead_id)
    if not lead:
        return jsonify({"error": "Prospecto no encontrado."}), 404
    if lead.sponsor_id != current_user.id:
        return jsonify({"error": "No autorizado."}), 403
    if lead.status != 'Pagado':
        return jsonify({"error": "Solo se puede subir CFDI para prospectos con estatus Pagado."}), 400
    pdf_file = request.files.get('cfdi_pdf')
    xml_file = request.files.get('cfdi_xml')
    if not pdf_file or not xml_file:
        return jsonify({"error": "Debes subir ambos archivos: PDF y XML del CFDI."}), 400
    pdf_ext = pdf_file.filename.rsplit('.', 1)[-1].lower() if '.' in pdf_file.filename else ''
    xml_ext = xml_file.filename.rsplit('.', 1)[-1].lower() if '.' in xml_file.filename else ''
    if pdf_ext != 'pdf':
        return jsonify({"error": "El archivo PDF debe tener extensión .pdf"}), 400
    if xml_ext != 'xml':
        return jsonify({"error": "El archivo XML debe tener extensión .xml"}), 400

    xml_bytes = xml_file.read()
    xml_file.seek(0)

    audit_error = _audit_cfdi_xml(xml_bytes, lead, current_user)
    if audit_error:
        return jsonify({"error": audit_error}), 400

    upload_dir = os.path.join("cfdi_uploads")
    os.makedirs(upload_dir, exist_ok=True)
    prefix = f"aff{current_user.id}_lead{lead_id}"
    pdf_name = f"{prefix}_cfdi.pdf"
    xml_name = f"{prefix}_cfdi.xml"
    pdf_path = os.path.join(upload_dir, pdf_name)
    xml_path = os.path.join(upload_dir, xml_name)
    pdf_file.save(pdf_path)
    xml_file.save(xml_path)
    _persist_file_to_db(pdf_path, current_user.id, category='cfdi')
    _persist_file_to_db(xml_path, current_user.id, category='cfdi')
    lead.cfdi_pdf = pdf_name
    lead.cfdi_xml = xml_name
    lead.payout_status = 'Facturado'
    db.session.commit()
    return jsonify({"success": True, "message": "CFDI validado y facturado correctamente."})

@app.route("/api/admin/update_payout/<int:lead_id>", methods=["POST"])
@login_required
def update_payout_status(lead_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    lead = db.session.get(AffiliateLead, lead_id)
    if not lead:
        return jsonify({"error": "Lead no encontrado"}), 404
    data = request.get_json(silent=True)
    new_status = (data.get("payout_status") or "").strip() if data else ""
    if new_status not in ("Pendiente CFDI", "CFDI en Revisión", "Facturado", "Liquidado"):
        return jsonify({"error": "Estatus inválido"}), 400
    lead.payout_status = new_status
    db.session.commit()
    return jsonify({"success": True, "payout_status": lead.payout_status})

@app.route("/api/admin/download_cfdi/<int:lead_id>/<string:file_type>")
@login_required
def download_cfdi(lead_id, file_type):
    if not current_user.is_admin:
        abort(403)
    lead = db.session.get(AffiliateLead, lead_id)
    if not lead:
        abort(404)
    cfdi_name = None
    if file_type == 'pdf' and lead.cfdi_pdf:
        cfdi_name = lead.cfdi_pdf
    elif file_type == 'xml' and lead.cfdi_xml:
        cfdi_name = lead.cfdi_xml
    if not cfdi_name:
        abort(404)
    disk_path = os.path.join("cfdi_uploads", cfdi_name)
    if os.path.isfile(disk_path):
        return send_from_directory("cfdi_uploads", cfdi_name, as_attachment=True)
    sf = _get_file_from_db(cfdi_name)
    if sf:
        return send_file(io.BytesIO(sf.content), download_name=cfdi_name,
                         as_attachment=True, mimetype=sf.content_type)
    abort(404)

def _record_purchase(user, payment_type, amount_total, stripe_session_id, credits_granted=0, source='stripe_webhook'):
    if not user:
        return
    try:
        if stripe_session_id:
            existing = UserPurchase.query.filter_by(stripe_session_id=stripe_session_id).first()
            if existing:
                logger.info(f"UserPurchase already recorded for session={stripe_session_id}, skipping")
                return
        purchase = UserPurchase(
            user_id=user.id,
            payment_type=payment_type,
            amount_mxn=round((amount_total or 0) / 100.0, 2),
            credits_granted=credits_granted or 0,
            stripe_session_id=stripe_session_id,
            source=source
        )
        db.session.add(purchase)
        db.session.commit()
        logger.info(f"UserPurchase recorded: user={user.id} type={payment_type} amount=${purchase.amount_mxn} credits={credits_granted} session={stripe_session_id}")
    except Exception as e:
        logger.error(f"Failed to record UserPurchase for user {user.id}: {e}")
        db.session.rollback()

def _record_affiliate_commission(user, amount_total, payment_type, stripe_session_id):
    if not user or not user.referred_by:
        return
    sponsor_id = user.referred_by
    sponsor = db.session.get(User, sponsor_id)
    if not sponsor or not sponsor.is_affiliate:
        return
    pago_mxn = amount_total / 100.0
    comision_pct = float(Config.get('COMISION_PORCENTAJE', '0.30'))
    commission_amount = round(pago_mxn * comision_pct, 2)
    if commission_amount <= 0:
        return
    lead = AffiliateLead.query.filter_by(sponsor_id=sponsor_id, prospect_email=user.email).first()
    if not lead:
        lead = AffiliateLead(
            sponsor_id=sponsor_id,
            prospect_name=user.full_name,
            prospect_email=user.email,
            prospect_whatsapp=user.whatsapp or '',
            payment_link='',
            selected_tier=payment_type,
            created_via='atribucion_automatica',
            status='Pagado',
            commission_amount=commission_amount
        )
        db.session.add(lead)
        db.session.flush()
    else:
        lead.status = 'Pagado'
    if stripe_session_id:
        existing_ac = AffiliateCommission.query.filter_by(
            stripe_session_id=stripe_session_id, sponsor_id=sponsor_id, payment_type=payment_type
        ).first()
        if existing_ac:
            logger.info(f"Affiliate commission already exists for session={stripe_session_id}, skipping")
            return
    ac = AffiliateCommission(
        lead_id=lead.id,
        sponsor_id=sponsor_id,
        user_id=user.id,
        amount=pago_mxn,
        commission_amount=commission_amount,
        payment_type=payment_type,
        stripe_session_id=stripe_session_id
    )
    db.session.add(ac)
    db.session.flush()
    db.session.execute(
        db.text('UPDATE affiliate_lead SET commission_amount = (SELECT COALESCE(SUM(commission_amount), 0) FROM affiliate_commission WHERE lead_id = :lid) WHERE id = :lid'),
        {"lid": lead.id}
    )
    db.session.commit()
    logger.info(f"Affiliate commission recorded: sponsor={sponsor_id}, user={user.id}, type={payment_type}, amount={pago_mxn}, commission={commission_amount}")

def _dispatch_welcome_email_async(user, prev_tier, new_tier, prev_credits, new_credits, applied_label, reason=''):
    user_id_snapshot = user.id
    email_snapshot = user.email
    def _run():
        try:
            with app.app_context():
                u = db.session.get(User, user_id_snapshot)
                if not u:
                    logger.warning(f"Welcome email: user {user_id_snapshot} not found in thread context")
                    return
                _send_tier_change_email(u, prev_tier, new_tier, prev_credits, new_credits, applied_label, reason=reason or '')
        except Exception as _e:
            logger.error(f"Welcome email thread crashed for {email_snapshot}: {_e}", exc_info=True)
    try:
        threading.Thread(target=_run, daemon=True).start()
        logger.info(f"Welcome email dispatched (async) to {email_snapshot} for {applied_label}")
    except Exception as _e:
        logger.warning(f"Welcome email dispatch failed for {email_snapshot}: {_e}")

@app.route("/api/webhooks/stripe", methods=["POST"])
def stripe_webhook():
    payload = request.get_data()
    sig_header = request.headers.get("Stripe-Signature", "")

    webhook_secret = Config.get('STRIPE_WEBHOOK_SECRET') or STRIPE_WEBHOOK_SECRET
    if not webhook_secret:
        logger.error("STRIPE_WEBHOOK_SECRET not configured (neither DB nor env)")
        return jsonify({"error": "Webhook not configured"}), 500

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except ValueError:
        logger.warning("Stripe webhook: invalid payload")
        return jsonify({"error": "Invalid payload"}), 400
    except stripe.error.SignatureVerificationError:
        logger.warning("Stripe webhook: invalid signature")
        return jsonify({"error": "Invalid signature"}), 400

    event_id = event.get("id", "")
    if event_id:
        existing = db.session.execute(
            db.text("SELECT 1 FROM stripe_processed_events WHERE event_id = :eid"),
            {"eid": event_id}
        ).fetchone()
        if existing:
            logger.info(f"Stripe webhook: duplicate event {event_id}, skipping")
            return jsonify({"received": True}), 200
        try:
            db.session.execute(
                db.text("INSERT INTO stripe_processed_events (event_id) VALUES (:eid)"),
                {"eid": event_id}
            )
            db.session.commit()
        except Exception:
            db.session.rollback()
            logger.info(f"Stripe webhook: race duplicate event {event_id}, skipping")
            return jsonify({"received": True}), 200

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        client_ref = (session.get("client_reference_id") or "").strip()
        amount_total = session.get("amount_total", 0)
        customer_email = (session.get("customer_details", {}) or {}).get("email", "")
        metadata = session.get("metadata") or {}

        logger.info(f"Stripe checkout.session.completed: client_ref={client_ref}, amount={amount_total}, email={customer_email}, metadata={metadata}")

        if metadata.get("type") == "advertisement":
            ad_id_str = metadata.get("ad_id")
            if ad_id_str and ad_id_str.isdigit():
                ad = db.session.get(AdCampaign, int(ad_id_str))
                if ad:
                    ad.status = 'Pendiente de Revisión'
                    ad.stripe_session_id = session.get("id", ad.stripe_session_id)
                    db.session.commit()
                    logger.info(f"Stripe webhook: ad campaign {ad.id} marked as 'Pendiente de Revisión'")
                else:
                    logger.warning(f"Stripe webhook: ad campaign {ad_id_str} not found")
            return jsonify({"received": True}), 200

        elif metadata.get("type") == "pro_multicurso":
            if client_ref.isdigit():
                user_id = int(client_ref)
                user = db.session.get(User, user_id)
                if user:
                    prev_tier_wh = user.tier
                    prev_credits_wh = user.pro_courses_remaining or 0
                    user.tier = 'PRO'
                    user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 5
                    user.pro_active_course = None
                    db.session.commit()
                    logger.info(f"Stripe webhook: user {user_id} ({user.email}) PRO multicurso (5 credits), remaining={user.pro_courses_remaining}")
                    _record_purchase(user, 'PRO_MULTICURSO', amount_total, session.get("id"), credits_granted=5)
                    _record_affiliate_commission(user, amount_total, 'PRO_MULTICURSO', session.get("id"))
                    _dispatch_welcome_email_async(user, prev_tier_wh, 'PRO', prev_credits_wh, user.pro_courses_remaining or 0, 'PRO_MULTICURSO', reason=f'Compra Stripe (session {session.get("id")})')
                    # Hard-reset del estado de curso activo tras upgrade.
                    # Purga master_doc/course_info_json para evitar contaminación
                    # transversal del demo/borrador previo. NO toca créditos.
                    try:
                        reset_active_course_state(user.id, reason='stripe_upgrade_pro_multicurso')
                    except Exception as _re:
                        logger.warning(f'[stripe_webhook] reset_active_course_state failed user={user.id}: {_re}')
                else:
                    logger.warning(f"Stripe webhook: multicurso - user {client_ref} not found")
            return jsonify({"received": True}), 200

        elif metadata.get("type") == "alacarte" or metadata.get("alacarte"):
            element_str = metadata.get("element", "") or metadata.get("alacarte", "")
            if client_ref.isdigit() and element_str.isdigit():
                user_id = int(client_ref)
                element_num = int(element_str)
                user = db.session.get(User, user_id)
                if user and element_num in (1, 2, 3, 4, 5):
                    col = f'alacarte_e{element_num}'
                    setattr(user, col, getattr(user, col, 0) + 1)
                    db.session.commit()
                    logger.info(f"Stripe webhook: user {user_id} ({user.email}) purchased alacarte E{element_num}")
                    _record_purchase(user, f'ALACARTE_E{element_num}', amount_total, session.get("id"), credits_granted=1)
                    _record_affiliate_commission(user, amount_total, f'ALACARTE_E{element_num}', session.get("id"))
                else:
                    logger.warning(f"Stripe webhook: alacarte - user {client_ref} not found or invalid element {element_str}")
            return jsonify({"received": True}), 200

        elif "_" in client_ref:
            parts = client_ref.split("_", 1)
            try:
                sponsor_id = int(parts[0])
                lead_id = int(parts[1])
            except (ValueError, IndexError):
                logger.error(f"Stripe webhook: invalid affiliate client_reference_id format: {client_ref}")
                return jsonify({"received": True}), 200

            lead = db.session.get(AffiliateLead, lead_id)
            if lead and lead.sponsor_id == sponsor_id:
                lead.status = 'Pagado'
                pago_mxn = amount_total / 100.0
                comision_pct = float(Config.get('COMISION_PORCENTAJE', '0.30'))
                commission_this = round(pago_mxn * comision_pct, 2)
                stripe_sid = session.get("id")
                existing_ac = AffiliateCommission.query.filter_by(
                    stripe_session_id=stripe_sid, sponsor_id=sponsor_id
                ).first() if stripe_sid else None
                if not existing_ac and commission_this > 0:
                    # Correo no único: si varias cuentas comparten el correo del
                    # prospecto, no adivinamos a cuál enlazar el user_id (FK informativo,
                    # nullable). La comisión del sponsor se registra igual.
                    _pu_matches = User.query.filter_by(email=lead.prospect_email).all()
                    prospect_user = _pu_matches[0] if len(_pu_matches) == 1 else None
                    ac = AffiliateCommission(
                        lead_id=lead.id,
                        sponsor_id=sponsor_id,
                        user_id=prospect_user.id if prospect_user else None,
                        amount=pago_mxn,
                        commission_amount=commission_this,
                        payment_type=lead.selected_tier or 'PRO_AFFILIATE',
                        stripe_session_id=stripe_sid
                    )
                    db.session.add(ac)
                total_comisiones = db.session.query(db.func.coalesce(db.func.sum(AffiliateCommission.commission_amount), 0)).filter_by(lead_id=lead.id).scalar()
                lead.commission_amount = round(float(total_comisiones) + (commission_this if not existing_ac else 0), 2)
                db.session.commit()
                logger.info(f"Stripe webhook: affiliate lead {lead_id} marked as Pagado (sponsor {sponsor_id}), commission={lead.commission_amount}")
            else:
                logger.warning(f"Stripe webhook: lead {lead_id} not found or sponsor mismatch (expected {sponsor_id})")

        elif client_ref.isdigit():
            payment_link_url = session.get("payment_link", "") or ""
            _alacarte_link_ids = {v.split('/')[-1] for k, v in STRIPE_LINKS.items() if k.startswith('ALACARTE_')}
            if payment_link_url and any(lid in str(payment_link_url) for lid in _alacarte_link_ids):
                logger.warning(f"Stripe webhook: alacarte payment detected without metadata, skipping tier upgrade for user {client_ref}")
            else:
                user_id = int(client_ref)
                user = db.session.get(User, user_id)
                if user:
                    _prev_tier_main = user.tier
                    _prev_credits_main = user.pro_courses_remaining or 0
                    meta_type = metadata.get("type", "")
                    if meta_type == 'premium' or amount_total >= 299700:
                        user.tier = 'PREMIUM'
                        _ptype = 'PREMIUM'
                    elif meta_type == 'pro_project' or amount_total >= 200000:
                        if meta_type == 'pro_project':
                            user.tier = 'PRO'
                            user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 1
                            user.pro_active_course = None
                            _ptype = 'PRO_PROJECT'
                        else:
                            user.tier = 'PRO'
                            user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 5
                            user.pro_active_course = None
                            _ptype = 'PRO_MULTICURSO'
                    else:
                        user.tier = 'PRO'
                        user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 1
                        user.pro_active_course = None
                        _ptype = 'PRO_PROJECT'
                    db.session.commit()
                    logger.info(f"Stripe webhook: user {user_id} ({user.email}) upgraded to {user.tier}, pro_courses_remaining={user.pro_courses_remaining}")
                    _credits_granted = 5 if _ptype == 'PRO_MULTICURSO' else (1 if _ptype == 'PRO_PROJECT' else 0)
                    _record_purchase(user, _ptype, amount_total, session.get("id"), credits_granted=_credits_granted)
                    _record_affiliate_commission(user, amount_total, _ptype, session.get("id"))
                    _dispatch_welcome_email_async(user, _prev_tier_main, user.tier, _prev_credits_main, user.pro_courses_remaining or 0, _ptype, reason=f'Compra Stripe (session {session.get("id")})')
                    # Hard-reset estado de curso post-upgrade (purga demo/borrador,
                    # blindando créditos). Idempotente y tolerante a JSON vacío.
                    try:
                        reset_active_course_state(user.id, reason=f'stripe_upgrade_{_ptype.lower()}')
                    except Exception as _re:
                        logger.warning(f'[stripe_webhook] reset_active_course_state failed user={user.id}: {_re}')
                else:
                    logger.warning(f"Stripe webhook: user {user_id} not found")

        elif customer_email:
            _ce = customer_email.lower().strip()
            _matches = User.query.filter_by(email=_ce).all()
            if len(_matches) > 1:
                # Correo compartido por varias cuentas distintas: no podemos saber
                # con certeza a cuál abonar. NO adivinamos para no subir de tier la
                # cuenta equivocada. El admin asigna manualmente este pago.
                logger.warning(f"Stripe webhook: email fallback con {len(_matches)} cuentas para {_ce} (amount={amount_total}, metadata={metadata}) — NO aplicado, requiere asignación admin")
                user = None
            else:
                user = _matches[0] if _matches else None
            if user:
                meta_type = (metadata.get("type") or "").lower()
                if meta_type == 'premium' or amount_total >= 299700:
                    user.tier = 'PREMIUM'
                    _ptype = 'PREMIUM'
                elif meta_type == 'pro_multicurso' or amount_total >= 200000:
                    user.tier = 'PRO'
                    user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 5
                    user.pro_active_course = None
                    _ptype = 'PRO_MULTICURSO'
                elif meta_type == 'pro_project':
                    user.tier = 'PRO'
                    user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 1
                    user.pro_active_course = None
                    _ptype = 'PRO_PROJECT'
                else:
                    user.tier = 'PRO'
                    user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 1
                    user.pro_active_course = None
                    _ptype = 'PRO_PROJECT'
                db.session.commit()
                logger.info(f"Stripe webhook: user {user.id} ({user.email}) upgraded to {user.tier} via email fallback (meta_type='{meta_type}', amount={amount_total}), pro_courses_remaining={user.pro_courses_remaining}")
                _credits_granted = 5 if _ptype == 'PRO_MULTICURSO' else (1 if _ptype == 'PRO_PROJECT' else 0)
                _record_purchase(user, _ptype, amount_total, session.get("id"), credits_granted=_credits_granted, source='stripe_webhook_email_fallback')
                _record_affiliate_commission(user, amount_total, _ptype, session.get("id"))
                _dispatch_welcome_email_async(user, 'FREE', user.tier, 0, user.pro_courses_remaining or 0, _ptype, reason=f'Pago Stripe reconciliado por email ({customer_email})')
                # Hard-reset estado de curso (email-fallback path). NO toca créditos.
                try:
                    reset_active_course_state(user.id, reason=f'stripe_upgrade_email_fallback_{_ptype.lower()}')
                except Exception as _re:
                    logger.warning(f'[stripe_webhook] reset_active_course_state failed user={user.id}: {_re}')
            else:
                logger.warning(f"Stripe webhook: email fallback - no user found for {customer_email} (amount={amount_total}, metadata={metadata})")

    return jsonify({"received": True}), 200

@app.route("/api/admin/compras/<int:user_id>")
@login_required
def admin_compras(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    purchases = UserPurchase.query.filter_by(user_id=user_id).order_by(UserPurchase.created_at.desc()).all()
    label_map = {
        'PRO_PROJECT': 'PRO 1 Curso',
        'PRO_MULTICURSO': 'PRO 5 Cursos',
        'PREMIUM': 'Premium',
        'ALACARTE_E1': 'A la carta E1 (Carta+Contrato)',
        'ALACARTE_E2': 'A la carta E2 (IECs)',
        'ALACARTE_E3': 'A la carta E3 (Manuales)',
        'ALACARTE_E4': 'A la carta E4 (Auto-Diagnóstico)'
    }
    result = []
    total = 0.0
    for p in purchases:
        result.append({
            "id": p.id,
            "payment_type": p.payment_type,
            "label": label_map.get(p.payment_type, p.payment_type),
            "amount_mxn": p.amount_mxn,
            "credits_granted": p.credits_granted,
            "stripe_session_id": (p.stripe_session_id or '')[:40],
            "source": p.source,
            "created_at": _fmt_cdmx(p.created_at)
        })
        total += float(p.amount_mxn or 0)
    return jsonify({"purchases": result, "count": len(result), "total_mxn": round(total, 2)})

@app.route("/api/admin/referidos/<int:user_id>")
@login_required
def admin_referidos(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    leads = AffiliateLead.query.filter_by(sponsor_id=user_id).all()
    result = []
    for l in leads:
        result.append({
            "id": l.id,
            "prospect_name": l.prospect_name,
            "prospect_email": l.prospect_email,
            "status": l.status,
            "payout_status": l.payout_status or 'Pendiente CFDI'
        })
    return jsonify({"success": True, "leads": result})

@app.route("/api/admin/actividad/<int:user_id>")
@login_required
def admin_actividad(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    events = UserEvent.query.filter_by(user_id=user_id).order_by(UserEvent.timestamp.desc()).limit(50).all()
    result = []
    for e in events:
        result.append({
            "category": e.event_category,
            "action": e.event_action,
            "timestamp": _fmt_cdmx(e.timestamp),
            "ip": e.ip_address or '',
            "metadata": e.metadata_json or ''
        })
    return jsonify({"success": True, "events": result})

@app.route("/api/state_flags", methods=["GET"])
def api_state_flags():
    """Devuelve banderas de estado server-side que el frontend debe respetar.

    Hoy: needs_state_reset_at — timestamp del último hard-reset disparado
    server-side (p.ej. post-pago Stripe). El cliente lo compara con su
    `localStorage.pertinentia_last_reset_ack`; si difiere, purga sus caches
    locales y POSTea a /api/state_flags/ack con el stamp que vio.

    Fix architect #2: como side-effect, si el stamp del server difiere del que
    la Flask session ya reconoció (`session['reset_ack']`), purgamos también
    las claves del curso en la Flask session — el script local sólo borra
    localStorage/sessionStorage del navegador, pero el render server-side
    también lee de `session['master_doc']` etc. y puede reinyectar el demo.
    """
    if not current_user.is_authenticated:
        return jsonify({'authenticated': False, 'reset_stamp': None})
    stamp = current_user.needs_state_reset_at.isoformat() if current_user.needs_state_reset_at else None
    if stamp and session.get('reset_ack') != stamp:
        # Purga server-side de claves del curso en Flask session (mismo set que
        # el endpoint /api/reset_course manual). Evita contaminación del render.
        for _k in ('course_info', 'master_doc', 'master_doc_topic',
                   'pending_action', 'pending_document', 'pending_e5_url',
                   'course_info_paso0_shown', 'subnorm_pending',
                   'cocreation_skip_session'):
            session.pop(_k, None)
        # Defensa adicional: purga cualquier otra clave con prefijo "pending_"
        # introducida en el futuro sin tener que actualizar este endpoint.
        for _k in [k for k in list(session.keys()) if isinstance(k, str) and k.startswith('pending_')]:
            session.pop(_k, None)
        for _e in (1, 2, 3, 4):
            session.pop(f'cocreation_state_e{_e}', None)
        session['reset_ack'] = stamp
        session.modified = True
    return jsonify({'authenticated': True, 'reset_stamp': stamp,
                    'active_course_session_id': current_user.active_course_session_id})


@app.route("/api/demo/mark_active", methods=["POST"])
def api_demo_mark_active():
    """Marca la CourseSession activa del usuario como is_demo=True.

    Disparado por el botón "Generar demo" en el frontend ANTES de enviar la
    primera prompt. Permite identificar después la sesión como demo por
    BANDERA (no por coincidencia de texto del topic), tal como exige la spec
    de reset_active_course_state.

    Anónimos: no-op (no hay CourseSession persistente).
    """
    if not current_user.is_authenticated:
        return jsonify({'ok': True, 'reason': 'anonymous_noop'})
    try:
        cs = None
        if current_user.active_course_session_id:
            cs = db.session.get(CourseSession, current_user.active_course_session_id)
        if cs is None:
            cs = CourseSession.query.filter_by(
                user_id=current_user.id, is_active=True
            ).order_by(CourseSession.session_num.desc()).first()
        if cs is None:
            return jsonify({'ok': False, 'reason': 'no_active_session'}), 404
        # Guarda anti-misetiquetado (architect): SÓLO marcar demo si la sesión
        # aún está limpia — sin documentos generados, sin master_doc, sin
        # course_info_json y sin topic real. Si ya tiene contenido real, esta
        # marca borraría una sesión válida al próximo reset; mejor rechazar.
        try:
            _doc_count = StoredFile.query.filter_by(
                user_id=current_user.id, course_session_id=cs.id,
                file_category='document'
            ).count()
        except Exception:
            _doc_count = 0
        _has_real_content = bool(
            _doc_count or cs.master_doc or cs.course_info_json or
            (cs.topic and cs.topic.strip())
        )
        if _has_real_content and not cs.is_demo:
            logger.warning(f'[api_demo_mark_active] refused user={current_user.id} cs={cs.id} (already has real content)')
            return jsonify({'ok': False, 'reason': 'session_has_real_content'}), 409
        if not cs.is_demo:
            cs.is_demo = True
            db.session.commit()
        return jsonify({'ok': True, 'cs_id': cs.id, 'is_demo': True})
    except Exception as _e:
        try: db.session.rollback()
        except Exception: pass
        logger.warning(f'[api_demo_mark_active] failed user={current_user.id}: {_e}')
        return jsonify({'ok': False, 'reason': str(_e)}), 500


@app.route("/api/state_flags/ack", methods=["POST"])
@login_required
def api_state_flags_ack():
    """ACK del cliente: confirma que ya purgó su localStorage hasta `stamp`.
    NO borra needs_state_reset_at del DB (el cliente persiste el stamp acked
    en su propio localStorage) — esto permite que múltiples pestañas/devices
    purguen independientemente sin pisarse entre sí.
    """
    payload = request.get_json(silent=True) or {}
    ack_stamp = (payload.get('stamp') or '').strip()
    try:
        track_event('Session', 'hard_reset_ack', user_id=current_user.id,
                    extra_data={'ack_stamp': ack_stamp[:40]})
    except Exception:
        pass
    return jsonify({'ok': True})


@app.route("/api/reset_course", methods=["POST"])
def reset_course():
    """Endpoint público para el botón 'Iniciar Nuevo Curso'.

    Delega en reset_active_course_state() para autenticados (la fuente de verdad
    de purga + blindaje de créditos). Para anónimos sólo limpia session keys.
    """
    new_cs_id = None
    new_cs_num = None
    purge_summary = None
    if current_user.is_authenticated:
        try:
            _persist_course_session_state(current_user)
        except Exception:
            pass
    keys = ('course_info', 'master_doc', 'master_doc_topic', 'pending_action',
            'pending_document', 'pending_e5_url',
            'course_info_paso0_shown', 'subnorm_pending',
            'cocreation_skip_session', 'course_logo_path')
    for k in keys:
        session.pop(k, None)
    for k in [k for k in list(session.keys()) if isinstance(k, str) and k.startswith('pending_')]:
        session.pop(k, None)
    for _e in (1, 2, 3, 4):
        session.pop(f'cocreation_state_e{_e}', None)
    session.modified = True
    if current_user.is_authenticated:
        # Delegamos al helper central: cierra activa, purga JSON, crea nueva,
        # blinda créditos, estampa señal frontend. Idempotente.
        purge_summary = reset_active_course_state(current_user.id, reason='manual_button')
        if purge_summary and purge_summary.get('ok'):
            new_cs_id = purge_summary.get('new_cs_id')
            new_cs_num = purge_summary.get('new_cs_num')
        return jsonify({"success": True, "released_course_slot": True,
                        "new_course_session_id": new_cs_id,
                        "new_course_session_num": new_cs_num,
                        "credit_snapshot_match": purge_summary.get('credit_snapshot_match') if purge_summary else True,
                        "purged_demo": purge_summary.get('purged_demo') if purge_summary else False})
    return jsonify({"success": True, "released_course_slot": False,
                    "new_course_session_id": None, "new_course_session_num": None})


@app.route("/api/reset_course__legacy_unused", methods=["POST"])
def reset_course__legacy_unused():
    """Implementación previa preservada como referencia. NO enrutada en uso.
    Mantener por si necesitamos rollback rápido de la nueva lógica.
    """
    new_cs_id = None
    new_cs_num = None
    if current_user.is_authenticated:
        try:
            _persist_course_session_state(current_user)
        except Exception:
            pass
    keys = ('course_info', 'master_doc', 'master_doc_topic', 'pending_action',
            'pending_document', 'course_info_paso0_shown', 'subnorm_pending',
            'cocreation_skip_session')
    for k in keys:
        session.pop(k, None)
    for _e in (1, 2, 3, 4):
        session.pop(f'cocreation_state_e{_e}', None)
    session.modified = True
    released = False
    if current_user.is_authenticated:
        try:
            if (current_user.pro_active_course or '').strip() and (current_user.pro_courses_remaining or 0) > 0:
                current_user.pro_active_course = None
                db.session.commit()
                released = True
        except Exception:
            db.session.rollback()
        try:
            reused_cs = None
            try:
                _lock_user_row(current_user.id)
                try:
                    db.session.refresh(current_user)
                except Exception:
                    pass
                candidates = CourseSession.query.filter_by(user_id=current_user.id).order_by(CourseSession.session_num.asc()).all()
                for _cs in candidates:
                    if _cs.id == current_user.active_course_session_id:
                        continue
                    _docs = StoredFile.query.filter_by(user_id=current_user.id, course_session_id=_cs.id, file_category='document').count()
                    if _docs > 0:
                        continue
                    _msgs = ChatHistory.query.filter_by(user_id=current_user.id, course_session_id=_cs.id).count()
                    if _msgs > 0:
                        continue
                    reused_cs = _cs
                    break
                if reused_cs is not None:
                    CourseSession.query.filter_by(user_id=current_user.id, is_active=True).update({'is_active': False})
                    reused_cs.is_active = True
                    reused_cs.last_activity_at = datetime.utcnow()
                    current_user.active_course_session_id = reused_cs.id
                    db.session.commit()
                    new_cs_id = reused_cs.id
                    new_cs_num = reused_cs.session_num
            except Exception:
                try: db.session.rollback()
                except Exception: pass
                reused_cs = None
            if reused_cs is None:
                new_cs = _create_new_course_session(current_user)
                if new_cs is not None:
                    new_cs_id = new_cs.id
                    new_cs_num = new_cs.session_num
        except Exception:
            pass
    return jsonify({"success": True, "released_course_slot": released,
                    "new_course_session_id": new_cs_id, "new_course_session_num": new_cs_num})


@app.route("/admin/realign_user_courses", methods=["POST"])
@login_required
def admin_realign_user_courses():
    if not current_user.is_admin:
        return jsonify({"success": False, "error": "No autorizado"}), 403
    raw_uid = None
    try:
        if request.is_json:
            raw_uid = (request.get_json(silent=True) or {}).get('user_id')
        if not raw_uid:
            raw_uid = request.form.get('user_id') or request.args.get('user_id')
        target_user_id = int(raw_uid)
    except Exception:
        return jsonify({"success": False, "error": "user_id requerido (int)"}), 400
    if target_user_id != 3:
        return jsonify({"success": False, "error": "Solo soporta user_id=3 (one-shot)"}), 400
    report = {"created": [], "updated": [], "deleted": [], "skipped": []}
    try:
        try:
            _lock_user_row(target_user_id)
        except Exception:
            pass
        target = User.query.get(target_user_id)
        if target is None:
            return jsonify({"success": False, "error": "Usuario no encontrado"}), 404
        sf56 = StoredFile.query.get(56)
        sf61 = StoredFile.query.get(61)
        sf62 = StoredFile.query.get(62)
        sf63 = StoredFile.query.get(63)
        for sf in (sf56, sf61, sf62, sf63):
            if sf is None or sf.user_id != target_user_id:
                return jsonify({"success": False, "error": f"StoredFile inesperado o user mismatch", "report": report}), 409
        cs1 = CourseSession.query.filter_by(user_id=target_user_id, session_num=1).first()
        if cs1 is None:
            return jsonify({"success": False, "error": "CS#1 no encontrado", "report": report}), 409
        cs_oswaldo = CourseSession.query.filter_by(user_id=target_user_id, topic="Teoría del Delito es de Derecho").first()
        if cs_oswaldo is None:
            cs_oswaldo = CourseSession(user_id=target_user_id,
                                       session_num=(db.session.query(db.func.max(CourseSession.session_num)).filter_by(user_id=target_user_id).scalar() or 0) + 1,
                                       is_active=False,
                                       topic="Teoría del Delito es de Derecho",
                                       started_at=datetime.utcnow(),
                                       last_activity_at=datetime.utcnow())
            db.session.add(cs_oswaldo)
            db.session.flush()
            report["created"].append({"cs_id": cs_oswaldo.id, "session_num": cs_oswaldo.session_num, "topic": cs_oswaldo.topic})
        else:
            report["skipped"].append({"reason": "cs_oswaldo_existed", "cs_id": cs_oswaldo.id})
        cs_blanca = CourseSession.query.filter_by(user_id=target_user_id, topic="Diseño de máster prompts para la automatización docente").first()
        if cs_blanca is None:
            cs_blanca = CourseSession(user_id=target_user_id,
                                      session_num=(db.session.query(db.func.max(CourseSession.session_num)).filter_by(user_id=target_user_id).scalar() or 0) + 1,
                                      is_active=False,
                                      topic="Diseño de máster prompts para la automatización docente",
                                      started_at=datetime.utcnow(),
                                      last_activity_at=datetime.utcnow())
            db.session.add(cs_blanca)
            db.session.flush()
            report["created"].append({"cs_id": cs_blanca.id, "session_num": cs_blanca.session_num, "topic": cs_blanca.topic})
        else:
            report["skipped"].append({"reason": "cs_blanca_existed", "cs_id": cs_blanca.id})
        cs1.topic = "La generación de cursos con IA en base a estándares de competencia"
        cs1.last_activity_at = datetime.utcnow()
        report["updated"].append({"cs_id": cs1.id, "topic": cs1.topic})
        if sf56.course_session_id != cs1.id:
            sf56.course_session_id = cs1.id
            report["updated"].append({"sf_id": 56, "moved_to_cs": cs1.id})
        if sf61.course_session_id != cs_oswaldo.id:
            sf61.course_session_id = cs_oswaldo.id
            report["updated"].append({"sf_id": 61, "moved_to_cs": cs_oswaldo.id})
        if sf62.course_session_id != cs_blanca.id:
            sf62.course_session_id = cs_blanca.id
            report["updated"].append({"sf_id": 62, "moved_to_cs": cs_blanca.id})
        if sf63.course_session_id != cs_blanca.id:
            sf63.course_session_id = cs_blanca.id
            report["updated"].append({"sf_id": 63, "moved_to_cs": cs_blanca.id})
        ch5 = ChatHistory.query.get(5)
        if ch5 is not None and ch5.user_id == target_user_id and ch5.course_session_id != cs_blanca.id:
            ch5.course_session_id = cs_blanca.id
            report["updated"].append({"ch_id": 5, "moved_to_cs": cs_blanca.id})
        empty_css = CourseSession.query.filter_by(user_id=target_user_id).all()
        for _cs in empty_css:
            if _cs.id in (cs1.id, cs_oswaldo.id, cs_blanca.id):
                continue
            _docs = StoredFile.query.filter_by(user_id=target_user_id, course_session_id=_cs.id).count()
            _msgs = ChatHistory.query.filter_by(user_id=target_user_id, course_session_id=_cs.id).count()
            if _docs == 0 and _msgs == 0:
                report["deleted"].append({"cs_id": _cs.id, "session_num": _cs.session_num})
                db.session.delete(_cs)
            else:
                report["skipped"].append({"reason": "cs_not_empty", "cs_id": _cs.id, "docs": _docs, "msgs": _msgs})
        CourseSession.query.filter_by(user_id=target_user_id, is_active=True).update({'is_active': False})
        cs1.is_active = True
        target.active_course_session_id = cs1.id
        report["updated"].append({"active_cs_id": cs1.id})
        db.session.commit()
        logger.info(f"ADMIN_REALIGN_USER_COURSES_OK admin={current_user.id} target={target_user_id} report={report}")
        return jsonify({"success": True, "report": report})
    except Exception as e:
        try: db.session.rollback()
        except Exception: pass
        logger.warning(f"ADMIN_REALIGN_USER_COURSES_FAIL admin={current_user.id} err={type(e).__name__}: {e}")
        return jsonify({"success": False, "error": f"{type(e).__name__}: {e}", "report": report}), 500


@app.route("/api/docs_meta", methods=["GET"])
@login_required
def api_docs_meta():
    try:
        rows = StoredFile.query.filter_by(user_id=current_user.id, file_category='document').all()
        cs_cache = {}
        meta = {}
        for sf in rows:
            cs_id = sf.course_session_id
            if cs_id and cs_id not in cs_cache:
                _cs = CourseSession.query.get(cs_id)
                if _cs is not None:
                    try:
                        _t = _sanitize_course_name_from_message(_cs.topic or '') or (_cs.topic or '')
                    except Exception:
                        _t = _cs.topic or ''
                    cs_cache[cs_id] = {'cs_num': _cs.session_num, 'cs_id': _cs.id, 'topic': _t}
            info = cs_cache.get(cs_id) if cs_id else None
            meta[sf.filename] = info or {'cs_num': None, 'cs_id': None, 'topic': ''}
        return jsonify({"success": True, "meta": meta})
    except Exception as e:
        logger.warning(f"api_docs_meta_fail user={current_user.id} err={type(e).__name__}: {e}")
        return jsonify({"success": False, "meta": {}}), 200


@app.route("/api/course_sessions", methods=["GET"])
@login_required
def api_course_sessions():
    try:
        try:
            _auto_split_user_course_sessions(current_user.id)
        except Exception as _e_split:
            logger.warning(f"course_sessions_split_skip user={current_user.id} err={type(_e_split).__name__}")
        sessions = CourseSession.query.filter_by(user_id=current_user.id).all()
        sessions.sort(key=lambda c: c.session_num or 0, reverse=True)
        out = []
        any_persisted = False
        for cs in sessions:
            doc_count = StoredFile.query.filter_by(user_id=current_user.id, course_session_id=cs.id, file_category='document').count()
            msg_rows = ChatHistory.query.filter_by(user_id=current_user.id, course_session_id=cs.id).all()
            msg_count = 0
            elements_seen = set()
            for r in msg_rows:
                elements_seen.add(r.element_num)
                try:
                    import json as _j
                    arr = _j.loads(r.messages_json or '[]')
                    if isinstance(arr, list):
                        msg_count += len(arr)
                except Exception:
                    pass
            try:
                started_str = _fmt_cdmx(cs.started_at) if cs.started_at else ''
            except Exception:
                started_str = ''
            try:
                last_str = _fmt_cdmx(cs.last_activity_at) if cs.last_activity_at else ''
            except Exception:
                last_str = ''
            try:
                _topic_clean = _sanitize_course_name_from_message(cs.topic or '') or (cs.topic or '')
            except Exception:
                _topic_clean = cs.topic or ''
            try:
                _looks_dirty = (not _topic_clean) or _topic_clean.strip().lower().startswith(('genera', 'crea', 'diseña', 'disena', 'haz'))
                if doc_count > 0 and _looks_dirty:
                    _real = _extract_course_name_from_cs_docs(current_user.id, cs.id)
                    if _real:
                        _topic_clean = _real
                        try:
                            cs.topic = _real
                            any_persisted = True
                        except Exception:
                            pass
            except Exception:
                pass
            out.append({
                'id': cs.id,
                'session_num': cs.session_num,
                'topic': _topic_clean,
                'is_active': bool(cs.is_active),
                'started_at': started_str,
                'last_activity_at': last_str,
                'document_count': doc_count,
                'message_count': msg_count,
                'elements_with_chat': sorted(list(elements_seen)),
            })
        if any_persisted:
            try:
                db.session.commit()
                logger.info(f"course_sessions_topic_autoheal user={current_user.id}")
            except Exception:
                try: db.session.rollback()
                except Exception: pass
        return jsonify({"success": True, "sessions": out,
                        "active_id": current_user.active_course_session_id,
                        "total": len(out)})
    except Exception as e:
        try: db.session.rollback()
        except Exception: pass
        logger.warning(f"api_course_sessions_fail user={current_user.id} err={type(e).__name__}: {e}")
        return jsonify({"success": False, "sessions": [], "active_id": None, "total": 0}), 200


@app.route("/api/load_course/<int:cs_id>", methods=["POST"])
@login_required
def api_load_course(cs_id):
    cs = CourseSession.query.get(cs_id)
    if cs is None or cs.user_id != current_user.id:
        return jsonify({"success": False, "error": "not_found"}), 404
    ok = _restore_course_session_state(current_user, cs_id)
    try:
        track_event('Curso', 'Recargar Curso', user_id=current_user.id,
                    extra_data={'course_session_id': cs_id, 'session_num': cs.session_num})
    except Exception:
        pass
    return jsonify({"success": bool(ok), "course_session_id": cs_id,
                    "session_num": cs.session_num, "topic": cs.topic or ''})

@app.route("/dashboard-ce")
@login_required
def dashboard_ce():
    if not current_user.is_premium:
        track_event('Navegacion', 'Acceso Denegado CE', user_id=current_user.id, extra_data={'tier': current_user.tier})
        flash('Esta sección es exclusiva para usuarios con plan PREMIUM.', 'error')
        return redirect(url_for('home'))
    track_event('Navegacion', 'Vista Centro Evaluador', user_id=current_user.id, extra_data={'tier': current_user.tier})
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    candidatos = Candidato.query.filter_by(ce_user_id=current_user.id).order_by(Candidato.id.desc()).all()
    return render_template("dashboard_ce.html",
                           title="Centro Evaluador IA",
                           active_page="dashboard_ce",
                           ce_profile=ce_profile,
                           candidatos=candidatos,
                           normative_accepted=current_user.normative_agreement_accepted)

@app.route("/api/ce/update_profile", methods=["POST"])
@login_required
def update_ce_profile():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado. Plan PREMIUM requerido."}), 403
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"error": "Datos inválidos"}), 400
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    if not ce_profile:
        ce_profile = CEProfile(user_id=current_user.id)
        db.session.add(ce_profile)
    ce_profile.ce_name = (data.get("ce_name") or "").strip()[:200]
    ce_profile.ce_key = (data.get("ce_key") or "").strip()[:100]
    ce_profile.evaluator_name = (data.get("evaluator_name") or "").strip()[:200]
    db.session.commit()
    return jsonify({"success": True, "message": "Perfil del Centro Evaluador actualizado correctamente."})

@app.route("/api/ce/upload_logo", methods=["POST"])
@login_required
def upload_ce_logo():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    if 'logo' not in request.files:
        return jsonify({"error": "No se envió ningún archivo"}), 400
    file = request.files['logo']
    if not file.filename:
        return jsonify({"error": "Archivo vacío"}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ('png', 'jpg', 'jpeg'):
        return jsonify({"error": "Solo se permiten archivos .png, .jpg o .jpeg"}), 400
    safe_filename = f"logo_ce_{current_user.id}.{ext}"
    save_path = os.path.join("static", "uploads", "logos", safe_filename)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    file.save(save_path)
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    if not ce_profile:
        ce_profile = CEProfile(user_id=current_user.id)
        db.session.add(ce_profile)
    ce_profile.logo_path = save_path
    db.session.commit()
    return jsonify({"success": True, "logo_url": "/" + save_path, "message": "Logotipo actualizado correctamente."})

@app.route("/api/upload_course_logo", methods=["POST"])
def upload_course_logo():
    """Sube un logotipo opcional (punto 10 del formulario) para insertarlo en la
    portada de la presentación y de los manuales generados. Disponible para
    usuarios autenticados y anónimos; la ruta se guarda en la sesión."""
    import uuid as _uuid
    if 'logo' not in request.files:
        return jsonify({"error": "No se envió ningún archivo"}), 400
    file = request.files['logo']
    if not file.filename:
        return jsonify({"error": "Archivo vacío"}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ('png', 'jpg', 'jpeg'):
        return jsonify({"error": "Solo se permiten archivos .png, .jpg o .jpeg"}), 400
    try:
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
    except Exception:
        size = 0
    if size > 5 * 1024 * 1024:
        return jsonify({"error": "El logotipo no debe exceder 5 MB."}), 400
    if current_user.is_authenticated:
        key = f"u{current_user.id}"
    else:
        anon_key = session.get('course_logo_uid')
        if not anon_key:
            anon_key = _uuid.uuid4().hex[:12]
            session['course_logo_uid'] = anon_key
        key = f"a{anon_key}"
    safe_filename = f"course_logo_{key}.{ext}"
    save_path = os.path.join("static", "uploads", "logos", safe_filename)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    try:
        file.save(save_path)
    except Exception:
        return jsonify({"error": "No se pudo guardar el logotipo. Intenta de nuevo."}), 500
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(save_path) as _img:
            _img.verify()
    except Exception:
        try:
            os.remove(save_path)
        except Exception:
            pass
        return jsonify({"error": "El archivo no es una imagen válida (usa PNG o JPG)."}), 400
    session['course_logo_path'] = save_path
    session.modified = True
    return jsonify({"success": True, "logo_url": "/" + save_path, "message": "Logotipo cargado correctamente."})

@app.route("/api/ce/candidatos", methods=["POST"])
@login_required
def crear_candidato():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"error": "Datos inválidos"}), 400
    nombre = (data.get("nombre_completo") or "").strip()
    apellidos = (data.get("apellidos") or "").strip()
    if not nombre:
        return jsonify({"error": "El nombre es obligatorio"}), 400
    curp = (data.get("curp") or "").strip().upper()[:18]
    cand = Candidato(
        ce_user_id=current_user.id,
        nombre_completo=nombre[:200],
        apellidos=apellidos[:200],
        curp=curp
    )
    db.session.add(cand)
    db.session.commit()
    return jsonify({
        "success": True,
        "candidato": {
            "id": cand.id,
            "nombre_completo": cand.nombre_completo,
            "apellidos": cand.apellidos,
            "curp": cand.curp,
            "estatus_autodiagnostico": cand.estatus_autodiagnostico,
            "estatus_plan": cand.estatus_plan,
            "estatus_dictamen": cand.estatus_dictamen
        }
    })

@app.route("/api/ce/candidatos/<int:cand_id>/toggle", methods=["POST"])
@login_required
def toggle_candidato_estatus(cand_id):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    cand = Candidato.query.filter_by(id=cand_id, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    data = request.get_json(silent=True)
    if not data or 'field' not in data:
        return jsonify({"error": "Datos inválidos"}), 400
    field = data['field']
    if field not in ('estatus_autodiagnostico', 'estatus_plan', 'estatus_dictamen'):
        return jsonify({"error": "Campo inválido"}), 400
    current_val = getattr(cand, field)
    setattr(cand, field, not current_val)
    db.session.commit()
    return jsonify({"success": True, "value": getattr(cand, field)})

@app.route("/api/ce/candidatos/<int:cand_id>", methods=["DELETE"])
@login_required
def eliminar_candidato(cand_id):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    cand = Candidato.query.filter_by(id=cand_id, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    db.session.delete(cand)
    db.session.commit()
    return jsonify({"success": True})

PORTAFOLIO_MAX_PDF_BYTES = 50 * 1024 * 1024
PORTAFOLIO_MAX_VIDEO_BYTES = 600 * 1024 * 1024
PORTAFOLIO_ALLOWED_PDF_MIMES = {'application/pdf'}
PORTAFOLIO_ALLOWED_VIDEO_MIMES = {'video/mp4', 'video/quicktime', 'video/webm', 'video/x-matroska', 'application/octet-stream'}
PORTAFOLIO_VALID_ESTANDARES = {'EC0301', 'EC0217.01'}
YOUTUBE_URL_RE = re.compile(r'(?:youtube\.com/(?:watch\?v=|shorts/|embed/|live/)|youtu\.be/)([A-Za-z0-9_-]{11})')

def _log_portafolio_attempt(ce_user_id, candidato_id, estandar, tipo, filename, size_bytes, success, error_code='', error_msg=''):
    try:
        att = PortafolioUploadAttempt(
            ce_user_id=ce_user_id,
            candidato_id=candidato_id,
            estandar=(estandar or '')[:20],
            tipo=(tipo or '')[:20],
            filename=(filename or '')[:300],
            size_bytes=int(size_bytes or 0),
            success=bool(success),
            error_code=(error_code or '')[:80],
            error_msg=error_msg,
            user_agent=(request.headers.get('User-Agent') or '')[:300]
        )
        db.session.add(att)
        db.session.commit()
        if not success:
            try:
                track_event('portafolio', 'upload_failed', user_id=ce_user_id, extra_data={
                    'estandar': estandar, 'tipo': tipo, 'error_code': error_code,
                    'filename': filename, 'size_bytes': size_bytes
                })
            except Exception:
                pass
    except Exception as e:
        print(f"[portafolio] no se pudo registrar intento: {e}")

def _leer_transcripcion_archivo(fs):
    """Extrae texto plano de un archivo de transcripción subido (.txt o .docx).

    Devuelve el texto (str) o '' si no hay archivo o no se pudo leer. No lanza
    excepciones: cualquier error se traduce a cadena vacía para que la capa de
    ruta muestre un mensaje amable en español.
    """
    if not fs or not getattr(fs, 'filename', ''):
        return ''
    nombre = (fs.filename or '').lower()
    try:
        if nombre.endswith('.docx'):
            import io
            import docx
            doc = docx.Document(io.BytesIO(fs.read()))
            return '\n'.join(p.text for p in doc.paragraphs).strip()
        data = fs.read()
        if not data:
            return ''
        for enc in ('utf-8', 'utf-8-sig', 'latin-1'):
            try:
                return data.decode(enc).strip()
            except Exception:
                continue
        return ''
    except Exception:
        return ''


@app.route("/api/ce/portafolio/upload", methods=["POST"])
@login_required
def portafolio_upload():
    if not current_user.is_premium:
        return jsonify({"error": "Solo disponible para usuarios Premium del Centro Evaluador"}), 403
    try:
        cand_id = int(request.form.get('candidato_id') or 0)
    except (ValueError, TypeError):
        return jsonify({"error": "ID de candidato inválido"}), 400
    estandar = (request.form.get('estandar') or '').strip()
    if estandar not in PORTAFOLIO_VALID_ESTANDARES:
        return jsonify({"error": "Estándar inválido. Use EC0301 o EC0217.01"}), 400
    cand = Candidato.query.filter_by(id=cand_id, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado o no pertenece a tu Centro"}), 404

    pdf_file = request.files.get('pdf')
    transcripcion_text = (request.form.get('transcripcion') or '').strip()
    transcripcion_file = request.files.get('transcripcion_file')
    if not transcripcion_text and transcripcion_file:
        transcripcion_text = _leer_transcripcion_archivo(transcripcion_file)

    if not pdf_file or not pdf_file.filename:
        _log_portafolio_attempt(current_user.id, cand_id, estandar, 'pdf', '', 0, False, 'pdf_missing', 'No se adjuntó PDF del portafolio')
        return jsonify({"error": "Debes adjuntar el PDF del portafolio integrado"}), 400

    if estandar == 'EC0217.01' and not transcripcion_text:
        return jsonify({"error": "Para EC0217.01 debes pegar la transcripción de la sesión o subir un archivo .txt o .docx con la transcripción"}), 400
    if estandar == 'EC0301' and (transcripcion_text or (transcripcion_file and transcripcion_file.filename)):
        return jsonify({"error": "EC0301 no requiere transcripción. Solo el PDF del portafolio."}), 400

    pdf_bytes = pdf_file.read()
    pdf_size = len(pdf_bytes)
    pdf_mime = (pdf_file.mimetype or '').lower()
    if pdf_mime not in PORTAFOLIO_ALLOWED_PDF_MIMES and not (pdf_file.filename or '').lower().endswith('.pdf'):
        _log_portafolio_attempt(current_user.id, cand_id, estandar, 'pdf', pdf_file.filename, pdf_size, False, 'pdf_mime_invalid', f'mime={pdf_mime}')
        return jsonify({"error": "El archivo debe ser un PDF válido"}), 400
    if pdf_size == 0:
        _log_portafolio_attempt(current_user.id, cand_id, estandar, 'pdf', pdf_file.filename, 0, False, 'pdf_empty')
        return jsonify({"error": "El PDF está vacío"}), 400
    if pdf_size > PORTAFOLIO_MAX_PDF_BYTES:
        _log_portafolio_attempt(current_user.id, cand_id, estandar, 'pdf', pdf_file.filename, pdf_size, False, 'pdf_too_large', f'limit={PORTAFOLIO_MAX_PDF_BYTES}')
        return jsonify({"error": f"PDF excede el límite de {PORTAFOLIO_MAX_PDF_BYTES // (1024*1024)} MB. Si necesitas un límite mayor, avísanos."}), 413

    try:
        ev = PortafolioEvaluacion(
            candidato_id=cand.id,
            ce_user_id=current_user.id,
            estandar=estandar,
            status='pendiente'
        )
        if estandar == 'EC0217.01':
            ev.video_transcripcion = transcripcion_text
            ev.video_duracion_seg = 0
            ev.video_idioma = 'es'
            ev.video_preview_seg = 0
            ev.video_procesado_at = datetime.utcnow()
        db.session.add(ev)
        db.session.flush()

        db.session.add(PortafolioArchivo(
            portafolio_id=ev.id, tipo='pdf', filename=(pdf_file.filename or 'portafolio.pdf')[:300],
            contenido=pdf_bytes, mime=(pdf_mime or 'application/pdf')[:80], size_bytes=pdf_size
        ))
        db.session.commit()
        _log_portafolio_attempt(current_user.id, cand_id, estandar, 'pdf', pdf_file.filename, pdf_size, True)
        if estandar == 'EC0217.01':
            _log_portafolio_attempt(current_user.id, cand_id, estandar, 'transcripcion', '', len(transcripcion_text), True)
        try:
            track_event('portafolio', 'upload_success', user_id=current_user.id, extra_data={
                'estandar': estandar, 'candidato_id': cand.id, 'evaluacion_id': ev.id,
                'has_transcripcion': bool(estandar == 'EC0217.01')
            })
        except Exception:
            pass
        return jsonify({
            "success": True,
            "evaluacion": {
                "id": ev.id, "estandar": ev.estandar, "status": ev.status,
                "created_at": ev.created_at.isoformat() if ev.created_at else None,
                "tiene_pdf": True, "tiene_transcripcion": bool(estandar == 'EC0217.01')
            },
            "mensaje": "Portafolio cargado correctamente. Ya puedes auditar la conformación y emitir el dictamen."
        })
    except Exception as e:
        db.session.rollback()
        _log_portafolio_attempt(current_user.id, cand_id, estandar, 'pdf', pdf_file.filename if pdf_file else '', pdf_size, False, 'db_error', str(e)[:500])
        app.logger.exception("Error guardando portafolio (user=%s, cand=%s, estandar=%s)", current_user.id, cand_id, estandar)
        return jsonify({"error": "No pudimos guardar el portafolio en este momento. Intenta de nuevo en unos minutos."}), 500

@app.route("/api/ce/portafolio/list/<int:cand_id>")
@login_required
def portafolio_list(cand_id):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    cand = Candidato.query.filter_by(id=cand_id, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    evals = PortafolioEvaluacion.query.filter_by(candidato_id=cand.id).order_by(PortafolioEvaluacion.created_at.desc()).all()
    out = []
    for ev in evals:
        archivos = [{"id": a.id, "tipo": a.tipo, "filename": a.filename, "size_bytes": a.size_bytes,
                     "source_url": a.source_url, "categoria": a.categoria,
                     "orden_oficial": a.orden_oficial, "descripcion": a.descripcion} for a in ev.archivos]
        resultado = None
        if ev.resultado_json:
            try: resultado = json.loads(ev.resultado_json)
            except Exception: resultado = None
        out.append({
            "id": ev.id, "estandar": ev.estandar, "status": ev.status,
            "created_at": ev.created_at.isoformat() if ev.created_at else None,
            "cost_mxn": ev.cost_mxn, "cost_usd": ev.cost_usd,
            "tokens_total": (ev.tokens_prompt or 0) + (ev.tokens_completion or 0),
            "tokens_prompt": ev.tokens_prompt, "tokens_completion": ev.tokens_completion,
            "model_used": ev.model_used,
            "tiene_integrado": bool(ev.portafolio_integrado_pdf),
            "integrado_size": ev.portafolio_integrado_size,
            "hash_sha256": ev.hash_sha256,
            "autorizado": ev.autorizado_por_evaluador,
            "autorizado_at": ev.autorizado_at.isoformat() if ev.autorizado_at else None,
            "resultado": resultado,
            "dictamen_final": ev.dictamen_final or '',
            "tiene_dictamen": bool(ev.dictamen_pdf),
            "dictamen_pdf_size": ev.dictamen_pdf_size,
            "dictamen_generado_at": ev.dictamen_generado_at.isoformat() if ev.dictamen_generado_at else None,
            "dictamen_cost_mxn": ev.dictamen_cost_mxn,
            "dictamen_cost_usd": ev.dictamen_cost_usd,
            "dictamen_tokens": (ev.dictamen_tokens_prompt or 0) + (ev.dictamen_tokens_completion or 0),
            "dictamen_resumen": (lambda: (json.loads(ev.dictamen_json).get('porcentaje_cumplimiento') if ev.dictamen_json else None))(),
            "video_procesado": bool(ev.video_transcripcion),
            "video_procesado_at": ev.video_procesado_at.isoformat() if ev.video_procesado_at else None,
            "video_duracion_seg": ev.video_duracion_seg or 0,
            "video_idioma": ev.video_idioma or '',
            "video_cost_mxn": ev.video_cost_mxn or 0.0,
            "video_cost_usd": ev.video_cost_usd or 0.0,
            "video_modelo": ev.video_modelo or '',
            "video_preview": ((ev.video_transcripcion or '')[:240]) if ev.video_transcripcion else '',
            "archivos": archivos
        })
    return jsonify({"evaluaciones": out})

@app.route("/api/ce/portafolio/<int:pid>", methods=["DELETE"])
@login_required
def portafolio_delete(pid):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    ev = PortafolioEvaluacion.query.filter_by(id=pid, ce_user_id=current_user.id).first()
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    db.session.delete(ev)
    db.session.commit()
    return jsonify({"success": True})

def _portafolio_owned(pid):
    return PortafolioEvaluacion.query.filter_by(id=pid, ce_user_id=current_user.id).first()

@app.route("/api/ce/portafolio/<int:pid>/auditar", methods=["POST"])
@login_required
def portafolio_auditar(pid):
    if not current_user.is_premium:
        return jsonify({"error": "Solo Premium"}), 403
    ev = _portafolio_owned(pid)
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    if ev.estandar not in ('EC0301', 'EC0217.01'):
        return jsonify({"error": "Estándar no soportado para auditoría IA"}), 400
    pdf_arch = PortafolioArchivo.query.filter_by(portafolio_id=ev.id, tipo='pdf').filter(PortafolioArchivo.categoria.in_(['', 'portafolio_inicial'])).order_by(PortafolioArchivo.created_at.asc()).first()
    if not pdf_arch or not pdf_arch.contenido:
        return jsonify({"error": "No hay PDF inicial para auditar"}), 400
    try:
        ev.status = 'auditando'
        db.session.commit()
        if ev.estandar == 'EC0301':
            from portafolio_ai_ec0301 import auditar_portafolio_ec0301
            resultado = auditar_portafolio_ec0301(
                pdf_arch.contenido, ev.candidato,
                evaluador_nombre=(current_user.nombre or current_user.email or '')
            )
        else:
            from portafolio_ai_ec0217 import auditar_portafolio_ec0217
            video_arch = PortafolioArchivo.query.filter_by(portafolio_id=ev.id).filter(PortafolioArchivo.tipo.in_(['video', 'youtube_url'])).first()
            video_meta = None
            if video_arch:
                video_meta = {'duration_seg': video_arch.duracion_seg or ev.video_duracion_seg or 0}
            resultado = auditar_portafolio_ec0217(
                pdf_arch.contenido, ev.candidato,
                evaluador_nombre=(current_user.nombre or current_user.email or ''),
                tiene_video=bool(video_arch), video_metadata=video_meta
            )
        meta = resultado.pop('_meta', {})
        if not pdf_arch.categoria:
            pdf_arch.categoria = 'portafolio_inicial'
            pdf_arch.orden_oficial = 0
            pdf_arch.descripcion = 'PDF integrado entregado por el evaluador'
        db.session.commit()
        PortafolioEvaluacion.query.filter_by(id=ev.id).update({
            'resultado_json': json.dumps(resultado, ensure_ascii=False),
            'tokens_prompt': PortafolioEvaluacion.tokens_prompt + int(meta.get('tokens_prompt', 0)),
            'tokens_completion': PortafolioEvaluacion.tokens_completion + int(meta.get('tokens_completion', 0)),
            'cost_usd': PortafolioEvaluacion.cost_usd + float(meta.get('cost_usd', 0.0)),
            'cost_mxn': PortafolioEvaluacion.cost_mxn + float(meta.get('cost_mxn', 0.0)),
            'model_used': meta.get('model_used', '') or ev.model_used,
            'status': 'auditado',
        }, synchronize_session=False)
        db.session.commit()
        try:
            track_event('portafolio', 'auditoria_completa', user_id=current_user.id,
                        extra_data={'pid': ev.id, 'cost_mxn': meta.get('cost_mxn', 0),
                                    'completitud': resultado.get('porcentaje_completitud')})
        except Exception:
            pass
        return jsonify({"success": True, "resultado": resultado, "meta": meta})
    except Exception as e:
        db.session.rollback()
        try:
            ev2 = _portafolio_owned(pid)
            if ev2:
                ev2.status = 'error'
                ev2.error_msg = str(e)[:500]
                db.session.commit()
        except Exception:
            pass
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Error en auditoría IA: {str(e)[:200]}"}), 500

@app.route("/api/ce/portafolio/<int:pid>/adjuntar_faltante", methods=["POST"])
@login_required
def portafolio_adjuntar_faltante(pid):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    ev = _portafolio_owned(pid)
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    categoria = (request.form.get('categoria') or '').strip()[:60]
    descripcion = (request.form.get('descripcion') or '').strip()[:500]
    try:
        orden = int(request.form.get('orden_oficial') or 999)
    except (ValueError, TypeError):
        orden = 999
    if not categoria:
        return jsonify({"error": "Falta categoría del documento"}), 400
    f = request.files.get('archivo')
    if not f or not f.filename:
        return jsonify({"error": "Adjunta el archivo PDF"}), 400
    data = f.read()
    size = len(data)
    if size == 0:
        return jsonify({"error": "Archivo vacío"}), 400
    if size > PORTAFOLIO_MAX_PDF_BYTES:
        return jsonify({"error": f"PDF excede {PORTAFOLIO_MAX_PDF_BYTES // (1024*1024)} MB"}), 413
    mime = (f.mimetype or '').lower()
    if mime not in PORTAFOLIO_ALLOWED_PDF_MIMES and not (f.filename or '').lower().endswith('.pdf'):
        return jsonify({"error": "El archivo debe ser PDF"}), 400
    existing = PortafolioArchivo.query.filter_by(portafolio_id=ev.id, categoria=categoria).first()
    if existing and existing.categoria not in ('portafolio_inicial',):
        existing.filename = (f.filename or 'doc.pdf')[:300]
        existing.contenido = data
        existing.size_bytes = size
        existing.mime = (mime or 'application/pdf')[:80]
        existing.orden_oficial = orden
        existing.descripcion = descripcion
    else:
        db.session.add(PortafolioArchivo(
            portafolio_id=ev.id, tipo='pdf', filename=(f.filename or 'doc.pdf')[:300],
            contenido=data, mime=(mime or 'application/pdf')[:80], size_bytes=size,
            categoria=categoria, orden_oficial=orden, descripcion=descripcion
        ))
    ev.portafolio_integrado_pdf = None
    ev.portafolio_integrado_size = 0
    ev.hash_sha256 = ''
    ev.autorizado_por_evaluador = False
    ev.autorizado_at = None
    ev.firma_evaluador_json = None
    db.session.commit()
    return jsonify({"success": True, "categoria": categoria, "size_bytes": size})

@app.route("/api/ce/portafolio/archivo/<int:aid>", methods=["DELETE"])
@login_required
def portafolio_archivo_delete(aid):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    arch = PortafolioArchivo.query.get(aid)
    if not arch:
        return jsonify({"error": "Archivo no encontrado"}), 404
    ev = PortafolioEvaluacion.query.filter_by(id=arch.portafolio_id, ce_user_id=current_user.id).first()
    if not ev:
        return jsonify({"error": "No autorizado"}), 403
    if arch.categoria == 'portafolio_inicial':
        return jsonify({"error": "No puedes eliminar el PDF base. Elimina la evaluación completa si es necesario."}), 400
    db.session.delete(arch)
    ev.portafolio_integrado_pdf = None
    ev.portafolio_integrado_size = 0
    ev.hash_sha256 = ''
    ev.autorizado_por_evaluador = False
    ev.autorizado_at = None
    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/ce/portafolio/<int:pid>/generar_integrado", methods=["POST"])
@login_required
def portafolio_generar_integrado(pid):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    ev = _portafolio_owned(pid)
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    try:
        from portafolio_integrador import integrar_portafolio
        archivos = PortafolioArchivo.query.filter_by(portafolio_id=ev.id, tipo='pdf').filter(PortafolioArchivo.categoria != '').order_by(PortafolioArchivo.orden_oficial.asc(), PortafolioArchivo.created_at.asc()).all()
        if not archivos:
            return jsonify({"error": "No hay documentos para integrar. Adjunta primero el PDF base y los faltantes."}), 400
        secciones = []
        for i, a in enumerate(archivos, start=1):
            orden = a.orden_oficial if a.orden_oficial < 999 else i
            secciones.append((orden, a.categoria, a.descripcion or a.filename, a.contenido))
        cand = ev.candidato
        cand_nombre = f"{cand.nombre_completo} {cand.apellidos}".strip()
        pdf_bytes, sha = integrar_portafolio(
            secciones,
            candidato_nombre=cand_nombre,
            candidato_curp=cand.curp,
            evaluador_nombre=(current_user.nombre or current_user.email or ''),
            estandar=ev.estandar,
            centro_evaluador=getattr(current_user, 'razon_social_ce', '') or '',
            evaluador_cedula='',
            incluir_autorizacion=False,
        )
        ev.portafolio_integrado_pdf = pdf_bytes
        ev.portafolio_integrado_size = len(pdf_bytes)
        ev.hash_sha256 = sha
        ev.autorizado_por_evaluador = False
        ev.autorizado_at = None
        ev.firma_evaluador_json = None
        ev.status = 'integrado'
        db.session.commit()
        return jsonify({"success": True, "size_bytes": len(pdf_bytes), "hash_sha256": sha,
                        "secciones": len(secciones),
                        "mensaje": "Vista preliminar generada. Revisa, marca la casilla de autorización y firma para sellar."})
    except Exception as e:
        db.session.rollback()
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Error al integrar: {str(e)[:200]}"}), 500

@app.route("/api/ce/portafolio/<int:pid>/autorizar", methods=["POST"])
@login_required
def portafolio_autorizar(pid):
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    ev = _portafolio_owned(pid)
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    data = request.get_json(silent=True) or {}
    evaluador_nombre = (data.get('evaluador_nombre') or current_user.nombre or '').strip()
    evaluador_cedula = (data.get('evaluador_cedula') or '').strip()[:60]
    acepta = bool(data.get('acepta_responsabilidad'))
    if not evaluador_nombre:
        return jsonify({"error": "Falta el nombre del evaluador"}), 400
    if not acepta:
        return jsonify({"error": "Debes aceptar la responsabilidad de la autorización"}), 400
    if not ev.portafolio_integrado_pdf:
        return jsonify({"error": "Primero genera el portafolio integrado"}), 400
    try:
        from portafolio_integrador import integrar_portafolio
        archivos = PortafolioArchivo.query.filter_by(portafolio_id=ev.id, tipo='pdf').filter(PortafolioArchivo.categoria != '').order_by(PortafolioArchivo.orden_oficial.asc(), PortafolioArchivo.created_at.asc()).all()
        secciones = []
        for i, a in enumerate(archivos, start=1):
            orden = a.orden_oficial if a.orden_oficial < 999 else i
            secciones.append((orden, a.categoria, a.descripcion or a.filename, a.contenido))
        cand = ev.candidato
        cand_nombre = f"{cand.nombre_completo} {cand.apellidos}".strip()
        pdf_bytes, sha = integrar_portafolio(
            secciones, candidato_nombre=cand_nombre, candidato_curp=cand.curp,
            evaluador_nombre=evaluador_nombre, estandar=ev.estandar,
            centro_evaluador=getattr(current_user, 'razon_social_ce', '') or '',
            evaluador_cedula=evaluador_cedula,
            incluir_autorizacion=True,
        )
        ev.portafolio_integrado_pdf = pdf_bytes
        ev.portafolio_integrado_size = len(pdf_bytes)
        ev.hash_sha256 = sha
        ev.autorizado_por_evaluador = True
        ev.autorizado_at = datetime.utcnow()
        ev.firma_evaluador_json = json.dumps({
            'evaluador_nombre': evaluador_nombre, 'evaluador_cedula': evaluador_cedula,
            'user_id': current_user.id, 'ip': request.remote_addr,
            'user_agent': (request.headers.get('User-Agent') or '')[:300]
        }, ensure_ascii=False)
        ev.status = 'autorizado'
        db.session.commit()
        try:
            track_event('portafolio', 'autorizado', user_id=current_user.id,
                        extra_data={'pid': ev.id, 'hash': sha[:16]})
        except Exception:
            pass
        return jsonify({"success": True, "hash_sha256": sha, "autorizado_at": ev.autorizado_at.isoformat()})
    except Exception as e:
        db.session.rollback()
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Error al autorizar: {str(e)[:200]}"}), 500

@app.route("/api/ce/portafolio/<int:pid>/descargar_integrado")
@login_required
def portafolio_descargar_integrado(pid):
    if not current_user.is_premium:
        return "No autorizado", 403
    ev = _portafolio_owned(pid)
    if not ev or not ev.portafolio_integrado_pdf:
        return "Portafolio integrado no disponible", 404
    cand = ev.candidato
    fname_base = f"Portafolio_{ev.estandar}_{(cand.nombre_completo + '_' + cand.apellidos).replace(' ', '_')}"
    if ev.autorizado_por_evaluador:
        fname_base += "_AUTORIZADO"
    resp = send_file(io.BytesIO(ev.portafolio_integrado_pdf),
                     mimetype='application/pdf', as_attachment=True,
                     download_name=f"{fname_base}.pdf")
    resp.headers['Cache-Control'] = 'private, no-store, max-age=0'
    return resp

@app.route("/api/ce/portafolio/<int:pid>/dictaminar", methods=["POST"])
@login_required
def portafolio_dictaminar(pid):
    if not current_user.is_premium:
        return jsonify({"error": "Solo Premium"}), 403
    ev = _portafolio_owned(pid)
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    if ev.estandar not in ('EC0301', 'EC0217.01'):
        return jsonify({"error": "Estándar no soportado para dictamen"}), 400
    if not ev.autorizado_por_evaluador or not ev.portafolio_integrado_pdf:
        return jsonify({"error": "Primero debes integrar y autorizar el portafolio antes de emitir el dictamen de competencia"}), 400
    if ev.estandar == 'EC0217.01' and not (ev.video_transcripcion or '').strip():
        return jsonify({"error": "Primero agrega la transcripción de la sesión de impartición antes del dictamen EC0217.01"}), 400
    try:
        from portafolio_integrador import generar_dictamen_pdf
        ev.status = 'dictaminando'
        db.session.commit()
        if ev.estandar == 'EC0301':
            from dictamen_ai_ec0301 import dictaminar_competencia_ec0301
            dictamen = dictaminar_competencia_ec0301(
                bytes(ev.portafolio_integrado_pdf), ev.candidato,
                evaluador_nombre=(current_user.nombre or current_user.email or '')
            )
        else:
            from dictamen_ai_ec0217 import dictaminar_competencia_ec0217
            dictamen = dictaminar_competencia_ec0217(
                bytes(ev.portafolio_integrado_pdf),
                ev.video_transcripcion or '',
                ev.candidato,
                evaluador_nombre=(current_user.nombre or current_user.email or ''),
                video_duracion_seg=ev.video_duracion_seg or 0
            )
        meta = dictamen.pop('_meta', {})
        cand = ev.candidato
        cand_nombre = f"{cand.nombre_completo} {cand.apellidos}".strip()
        pdf_bytes = generar_dictamen_pdf(
            dictamen, candidato_nombre=cand_nombre, candidato_curp=cand.curp,
            evaluador_nombre=(current_user.nombre or current_user.email or ''),
            estandar=ev.estandar,
            centro_evaluador=getattr(current_user, 'razon_social_ce', '') or '',
            portafolio_hash=ev.hash_sha256 or ''
        )
        PortafolioEvaluacion.query.filter_by(id=ev.id).update({
            'dictamen_json': json.dumps(dictamen, ensure_ascii=False),
            'dictamen_final': (dictamen.get('dictamen_final') or '')[:40],
            'dictamen_pdf': pdf_bytes,
            'dictamen_pdf_size': len(pdf_bytes),
            'dictamen_generado_at': datetime.utcnow(),
            'dictamen_tokens_prompt': PortafolioEvaluacion.dictamen_tokens_prompt + int(meta.get('tokens_prompt', 0)),
            'dictamen_tokens_completion': PortafolioEvaluacion.dictamen_tokens_completion + int(meta.get('tokens_completion', 0)),
            'dictamen_cost_usd': PortafolioEvaluacion.dictamen_cost_usd + float(meta.get('cost_usd', 0.0)),
            'dictamen_cost_mxn': PortafolioEvaluacion.dictamen_cost_mxn + float(meta.get('cost_mxn', 0.0)),
            'status': 'dictaminado',
        }, synchronize_session=False)
        db.session.commit()
        try:
            track_event('portafolio', 'dictamen_emitido', user_id=current_user.id,
                        extra_data={'pid': ev.id, 'dictamen': dictamen.get('dictamen_final'),
                                    'pct': dictamen.get('porcentaje_cumplimiento'),
                                    'cost_mxn': meta.get('cost_mxn', 0)})
        except Exception:
            pass
        return jsonify({"success": True, "dictamen": dictamen, "meta": meta,
                        "pdf_size": len(pdf_bytes)})
    except Exception as e:
        db.session.rollback()
        try:
            ev2 = _portafolio_owned(pid)
            if ev2:
                ev2.status = 'error'
                ev2.error_msg = str(e)[:500]
                db.session.commit()
        except Exception:
            pass
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Error en dictamen IA: {str(e)[:200]}"}), 500

@app.route("/api/ce/portafolio/<int:pid>/procesar_video", methods=["POST"])
@login_required
def portafolio_procesar_video(pid):
    ev = _portafolio_owned(pid)
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    is_e5_preview = (ev.e5_mode == 'preview' and not ev.e5_unlocked)
    if not (current_user.is_premium or ev.e5_unlocked or is_e5_preview):
        return jsonify({"error": "Requiere PREMIUM o evaluación E5 (EC0217.01) activada en /evaluar-mi-clase"}), 403
    if ev.estandar != 'EC0217.01':
        return jsonify({"error": "El procesamiento de video sólo aplica a EC0217.01"}), 400
    video_arch = PortafolioArchivo.query.filter_by(portafolio_id=ev.id, tipo='video').first()
    yt_arch = PortafolioArchivo.query.filter_by(portafolio_id=ev.id, tipo='youtube_url').first()
    if not video_arch and not yt_arch:
        return jsonify({"error": "No hay video MP4 ni URL de YouTube vinculados a esta evaluación"}), 400
    PREVIEW_SEG = 600
    is_upgrade = (ev.e5_unlocked and (ev.video_preview_seg or 0) > 0 and (ev.video_transcripcion or '').strip())
    if is_e5_preview and (ev.video_transcripcion or '').strip():
        return jsonify({"error": "Pre-dictamen ya generado. Adquiere E5 para análisis completo."}), 400
    if ev.e5_unlocked and ev.status in ('video_procesado', 'dictaminado') and not is_upgrade:
        return jsonify({"error": "Video ya procesado completamente.", "already_processed": True}), 400
    if is_upgrade and ev.video_procesado_at and (ev.video_duracion_seg or 0) >= (ev.video_preview_seg or 0) and ev.status == 'video_procesado':
        return jsonify({"error": "Continuación ya procesada.", "already_processed": True}), 400
    if is_upgrade:
        _total_known = (yt_arch.duracion_seg if yt_arch else 0) or 0
        if _total_known and (ev.video_preview_seg or 0) >= _total_known - 5:
            PortafolioEvaluacion.query.filter_by(id=ev.id).update({
                'status': 'video_procesado',
                'video_procesado_at': datetime.utcnow(),
            }, synchronize_session=False)
            db.session.commit()
            return jsonify({"success": True, "skipped": True,
                            "reason": "Video más corto que el preview; nada por transcribir.",
                            "segundos": 0, "cost_mxn": 0, "num_chunks": 0,
                            "preview": (ev.video_transcripcion or '')[:300]})
    try:
        from video_processor import (descargar_youtube, extraer_audio, transcribir_audio,
                                      obtener_transcripcion_captions, VideoError)
        ev.status = 'procesando_video'
        db.session.commit()

        _audio_max_seconds = None
        result = None
        if video_arch and video_arch.contenido:
            video_bytes = bytes(video_arch.contenido)
            source_ext = (video_arch.filename.rsplit('.', 1)[-1] if '.' in video_arch.filename else 'mp4').lower()[:5]
            video_meta = {'duration_seg': getattr(video_arch, 'duracion_seg', 0) or 0, 'source': 'mp4_upload'}
            # Pre-dictamen gratuito por mp4: solo transcribimos los primeros 10 min.
            if is_e5_preview:
                _audio_max_seconds = PREVIEW_SEG
        else:
            # 1) Primero intentamos los subtítulos/transcripción oficiales de YouTube.
            #    Van por un endpoint distinto al de descarga de medios, así que suelen
            #    funcionar aunque YouTube bloquee la descarga desde nuestros servidores,
            #    y NO consumen Whisper (costo 0). Sólo el tramo necesario (preview/upgrade).
            _cap_start = 0 if is_e5_preview else ((ev.video_preview_seg or 0) if is_upgrade else None)
            _cap_end = PREVIEW_SEG if is_e5_preview else None
            try:
                result = obtener_transcripcion_captions(
                    yt_arch.source_url, start_sec=_cap_start, end_sec=_cap_end, idioma_pref='es')
                video_meta = {'duration_seg': result.get('total_duration_seg', 0) or 0,
                              'source': 'youtube_captions', 'title': ''}
                if not is_upgrade:
                    yt_arch.duracion_seg = video_meta['duration_seg']
                    db.session.commit()
            except VideoError as _cap_err:
                app.logger.info("Captions YouTube no disponibles (%s); usaremos descarga de audio.",
                                str(_cap_err)[:160])
                result = None
            # 2) Fallback: descargar el audio y transcribir con Whisper (como antes).
            if result is None:
                if is_e5_preview:
                    video_bytes, ydl_meta = descargar_youtube(yt_arch.source_url, start_sec=0, end_sec=PREVIEW_SEG)
                elif is_upgrade:
                    video_bytes, ydl_meta = descargar_youtube(yt_arch.source_url, start_sec=ev.video_preview_seg)
                else:
                    video_bytes, ydl_meta = descargar_youtube(yt_arch.source_url)
                source_ext = ydl_meta.get('ext') or 'mp4'
                video_meta = {'duration_seg': ydl_meta.get('duration_seg', 0),
                              'source': 'youtube', 'title': ydl_meta.get('title', '')}
                if not is_upgrade:
                    yt_arch.duracion_seg = video_meta['duration_seg']
                    yt_arch.size_bytes = ydl_meta.get('size_bytes', 0)
                    db.session.commit()

        _dur_check = max(int(video_meta.get('duration_seg') or 0), 0)
        if not is_e5_preview and not is_upgrade and _dur_check > 8280:
            ev.status = 'error'
            ev.error_msg = f"Video {_dur_check//60} min excede el máximo permitido (138 min = 120 min ±15%)"
            db.session.commit()
            return jsonify({"error": f"El video dura {_dur_check//60} min. El EC0217.01 establece 120 min ±15% (máx. 138 min)."}), 400
        # Si no obtuvimos la transcripción por captions, descargamos el audio y usamos Whisper.
        if result is None:
            audio_bytes, audio_dur = extraer_audio(video_bytes, source_ext=source_ext,
                                                   max_seconds=_audio_max_seconds)
            del video_bytes
            result = transcribir_audio(audio_bytes, idioma_hint='es',
                                       duration_seg=audio_dur or 0)
            del audio_bytes

        nuevo_texto = result.get('texto', '') or ''
        # El modo incremental (continuar desde el preview) solo aplica a YouTube,
        # donde descargamos por tramos. Para mp4 tenemos el archivo completo, así que
        # al desbloquear (upgrade) se transcribe el video completo de una vez.
        if is_upgrade and not video_arch:
            transcripcion_final = ((ev.video_transcripcion or '').rstrip() +
                                   f"\n\n[--- Continuación desde min {ev.video_preview_seg//60} ---]\n\n" +
                                   nuevo_texto)
            duracion_total = (ev.video_preview_seg or 0) + (result.get('segundos', 0) or 0)
            segundos_transcritos_total = (ev.video_segundos_transcritos or 0) + (result.get('segundos', 0) or 0)
            new_status = 'video_procesado'
        elif is_e5_preview:
            transcripcion_final = nuevo_texto
            duracion_total = result.get('segundos', 0)
            segundos_transcritos_total = result.get('segundos', 0)
            new_status = 'preview_procesado'
        else:
            transcripcion_final = nuevo_texto
            duracion_total = max(result.get('segundos', 0), video_meta.get('duration_seg', 0))
            segundos_transcritos_total = result.get('segundos', 0)
            new_status = 'video_procesado'

        update_dict = {
            'video_transcripcion': transcripcion_final,
            'video_duracion_seg': duracion_total,
            'video_idioma': result.get('idioma', '')[:10],
            'video_segundos_transcritos': segundos_transcritos_total,
            'video_cost_usd': PortafolioEvaluacion.video_cost_usd + float(result.get('cost_usd', 0.0)),
            'video_cost_mxn': PortafolioEvaluacion.video_cost_mxn + float(result.get('cost_mxn', 0.0)),
            'video_procesado_at': datetime.utcnow(),
            'video_modelo': result.get('model', '')[:40],
            'status': new_status,
        }
        if is_e5_preview:
            update_dict['video_preview_seg'] = result.get('segundos', 0) or PREVIEW_SEG
        PortafolioEvaluacion.query.filter_by(id=ev.id).update(update_dict, synchronize_session=False)
        db.session.commit()
        try:
            track_event('portafolio', 'video_transcrito', user_id=current_user.id,
                        extra_data={'pid': ev.id, 'segundos': result.get('segundos', 0),
                                    'cost_mxn': result.get('cost_mxn', 0)})
        except Exception:
            pass
        return jsonify({"success": True,
                        "segundos": result.get('segundos', 0),
                        "idioma": result.get('idioma', ''),
                        "cost_mxn": result.get('cost_mxn', 0),
                        "num_chunks": result.get('num_chunks', 1),
                        "preview": (result.get('texto', '') or '')[:300]})
    except VideoError as e:
        db.session.rollback()
        try:
            ev2 = _portafolio_owned(pid)
            if ev2:
                ev2.status = 'error'; ev2.error_msg = str(e)[:500]; db.session.commit()
        except Exception:
            pass
        return jsonify({"error": str(e)[:400], "puede_subir_archivo": True}), 400
    except Exception as e:
        db.session.rollback()
        try:
            ev2 = _portafolio_owned(pid)
            if ev2:
                ev2.status = 'error'; ev2.error_msg = str(e)[:500]; db.session.commit()
        except Exception:
            pass
        import traceback; traceback.print_exc()
        return jsonify({"error": "No pudimos procesar el video en este momento. Intenta de nuevo en unos minutos o sube el archivo de video directamente."}), 500

@app.route("/api/ce/portafolio/<int:pid>/descargar_dictamen")
@login_required
def portafolio_descargar_dictamen(pid):
    ev = _portafolio_owned(pid)
    if not ev:
        return "Dictamen no disponible", 404
    if not (current_user.is_premium or ev.e5_unlocked):
        return "No autorizado", 403
    if not ev.dictamen_pdf:
        return "Dictamen no disponible", 404
    cand = ev.candidato
    fname = f"Dictamen_{ev.estandar}_{(cand.nombre_completo + '_' + cand.apellidos).replace(' ', '_')}.pdf"
    resp = send_file(io.BytesIO(ev.dictamen_pdf),
                     mimetype='application/pdf', as_attachment=True,
                     download_name=fname)
    resp.headers['Cache-Control'] = 'private, no-store, max-age=0'
    return resp

@app.route("/evaluar-mi-clase")
@app.route("/diagnostico-mi-clase")
def evaluar_mi_clase_landing():
    is_anon = not current_user.is_authenticated
    has_full_access = (not is_anon) and (current_user.is_premium or current_user.has_alacarte(5))
    ec0217_saldo = (current_user.ec0217_grants or 0) if not is_anon else 0
    mp4_authorized = (not is_anon) and (current_user.is_premium or current_user.has_alacarte(5) or ec0217_saldo > 0)
    preview_used = (not is_anon) and bool(current_user.e5_preview_used_at)
    canonical_path = '/diagnostico-mi-clase'
    return render_template("evaluar_mi_clase.html",
                           title="Diagnostica tu clase contra el EC0217.01 — SEP/CONOCER",
                           active_page="evaluar_mi_clase",
                           is_anon=is_anon,
                           has_full_access=has_full_access,
                           mp4_authorized=mp4_authorized,
                           ec0217_saldo=ec0217_saldo,
                           preview_used=preview_used,
                           canonical_url=(os.environ.get('PUBLIC_BASE_URL','https://pertinentia.com') + canonical_path),
                           stripe_link_e5=STRIPE_LINKS.get('ALACARTE_E5', ''))

@app.route("/recursos/ec0217-ficha-estandar.pdf")
def recurso_ec0217_ficha():
    """Descarga pública de la ficha oficial del estándar EC0217.01 (CONOCER)."""
    try:
        track_event('Recursos', 'Descarga Ficha EC0217.01',
                    user_id=(current_user.id if current_user.is_authenticated else None))
    except Exception:
        pass
    path = os.path.join('plantillas', 'EC0217.01_Ficha_Estandar.pdf')
    if not os.path.exists(path):
        return ('Recurso no disponible temporalmente.', 404)
    return send_file(path, as_attachment=False, download_name='EC0217.01_Ficha_Estandar.pdf',
                     mimetype='application/pdf', max_age=86400)

@app.route("/evaluar-mi-clase/start", methods=["POST"])
def evaluar_mi_clase_start():
    # DESACTIVADO: el flujo por YouTube fue reemplazado por la entrada de
    # transcripción (evaluar_mi_clase_start_texto). Se conserva el código
    # original debajo por referencia, pero la ruta redirige al flujo nuevo.
    return redirect(url_for('evaluar_mi_clase_landing'))
    youtube_url = (request.form.get('youtube_url') or '').strip()[:500]
    if not youtube_url or 'youtu' not in youtube_url.lower():
        flash('Pega una URL válida de YouTube de tu clase.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))

    if not current_user.is_authenticated:
        session['pending_e5_url'] = youtube_url
        try:
            track_event('autoevaluacion_e5', 'gancho_anon_submit',
                        extra_data={'url': youtube_url[:100]})
        except Exception:
            pass
        flash('Crea tu cuenta gratuita para recibir tu pre-dictamen IA (primeros 10 min de tu clase).', 'info')
        return redirect(url_for('registro', **{'from': 'evaluar_mi_clase'}))

    mode = None
    unlocked_via = ''
    if current_user.is_premium:
        mode = 'full'; unlocked_via = 'premium'
    elif current_user.has_alacarte(5):
        if current_user.use_alacarte(5):
            mode = 'full'; unlocked_via = 'alacarte_e5'
    elif not current_user.e5_preview_used_at:
        mode = 'preview'
    else:
        flash('Ya usaste tu pre-dictamen gratuito. Adquiere E5 ($349) para tu evaluación completa.', 'info')
        return redirect(url_for('evaluar_mi_clase_landing'))

    cand = Candidato.query.filter_by(ce_user_id=current_user.id, curp='AUTOEVAL_E5').first()
    if not cand:
        cand = Candidato(
            ce_user_id=current_user.id,
            nombre_completo=(current_user.email or 'Auto-evaluación')[:200],
            apellidos='(Auto-evaluación E5)', curp='AUTOEVAL_E5',
        )
        db.session.add(cand); db.session.commit()
    ev = PortafolioEvaluacion(
        candidato_id=cand.id, ce_user_id=current_user.id,
        estandar='EC0217.01', status='pendiente_video',
        e5_mode=mode,
        e5_unlocked=(mode == 'full'),
        e5_unlocked_at=(datetime.utcnow() if mode == 'full' else None),
        e5_unlocked_via=unlocked_via,
    )
    db.session.add(ev); db.session.commit()
    db.session.add(PortafolioArchivo(
        portafolio_id=ev.id, tipo='youtube_url',
        filename='youtube_link', source_url=youtube_url, size_bytes=0,
    ))
    if mode == 'preview':
        current_user.e5_preview_used_at = datetime.utcnow()
    db.session.commit()
    try:
        track_event('autoevaluacion_e5', 'iniciada', user_id=current_user.id,
                    extra_data={'pid': ev.id, 'mode': mode, 'via': unlocked_via})
    except Exception:
        pass
    return redirect(url_for('evaluar_mi_clase_vista', pid=ev.id))

@app.route("/evaluar-mi-clase/start-mp4", methods=["POST"])
@login_required
def evaluar_mi_clase_start_mp4():
    """Inicia una evaluación EC0217.01 COMPLETA subiendo un video mp4 (no YouTube).
    Disponible para usuarios autorizados: Premium, crédito a la carta E5, o saldo
    autorizado por el administrador (descuenta 1 y lo registra en la bitácora)."""
    # DESACTIVADO: la carga de video mp4 fue reemplazada por la entrada de
    # transcripción (evaluar_mi_clase_start_texto). Se conserva el código original
    # debajo por referencia, pero la ruta redirige al flujo nuevo.
    return redirect(url_for('evaluar_mi_clase_landing'))
    video_file = request.files.get('video')
    if not video_file or not video_file.filename:
        flash('Adjunta el archivo de video (mp4) de tu clase.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))
    if not video_file.filename.lower().endswith(('.mp4', '.mov', '.webm', '.mkv', '.m4v')):
        flash('Formato no válido. Sube un video mp4, mov, webm, mkv o m4v.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))
    video_bytes = video_file.read()
    video_size = len(video_bytes)
    if video_size == 0:
        flash('El archivo de video está vacío.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))
    if video_size > PORTAFOLIO_MAX_VIDEO_BYTES:
        flash(f'El video supera el límite de {PORTAFOLIO_MAX_VIDEO_BYTES // (1024*1024)} MB. '
              'Comprime el video o reduce su resolución e inténtalo de nuevo.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))

    # Determinar vía/modo de acceso SIN consumir todavía (el consumo se hace atómico abajo).
    # Usuarios con acceso full (premium / a la carta / autorización admin) reciben la
    # evaluación completa; un usuario gratuito sin pre-dictamen previo recibe el
    # pre-dictamen (primeros 10 min) también por mp4, no solo por YouTube.
    if current_user.is_premium:
        unlocked_via = 'premium'; mp4_mode = 'full'
    elif current_user.has_alacarte(5):
        unlocked_via = 'alacarte_e5'; mp4_mode = 'full'
    elif current_user.ec0217_authorized:
        unlocked_via = 'admin_grant'; mp4_mode = 'full'
    elif not current_user.e5_preview_used_at:
        unlocked_via = ''; mp4_mode = 'preview'
    else:
        flash('Ya usaste tu pre-dictamen gratuito. Adquiere la evaluación completa '
              '($349) o solicita autorización al administrador.', 'info')
        return redirect(url_for('evaluar_mi_clase_landing'))

    cand = Candidato.query.filter_by(ce_user_id=current_user.id, curp='AUTOEVAL_E5').first()
    if not cand:
        cand = Candidato(
            ce_user_id=current_user.id,
            nombre_completo=(current_user.email or 'Auto-evaluación')[:200],
            apellidos='(Auto-evaluación E5)', curp='AUTOEVAL_E5',
        )
        db.session.add(cand); db.session.commit()

    # Transacción única: el descuento de crédito, la creación de la evaluación,
    # el archivo de video y la bitácora se confirman juntos o se revierten juntos.
    # Así nunca se pierde un crédito sin evaluación ni queda bitácora inconsistente.
    try:
        if unlocked_via == 'alacarte_e5':
            res = db.session.execute(
                db.text('UPDATE "user" SET alacarte_e5 = alacarte_e5 - 1 '
                        'WHERE id = :uid AND alacarte_e5 > 0'), {"uid": current_user.id})
            if res.rowcount == 0:
                db.session.rollback()
                flash('No tienes una evaluación EC0217.01 disponible. Adquiérela ($349) '
                      'o solicita autorización al administrador.', 'info')
                return redirect(url_for('evaluar_mi_clase_landing'))
        elif unlocked_via == 'admin_grant':
            res = db.session.execute(
                db.text('UPDATE "user" SET ec0217_grants = ec0217_grants - 1 '
                        'WHERE id = :uid AND ec0217_grants > 0'), {"uid": current_user.id})
            if res.rowcount == 0:
                db.session.rollback()
                flash('No tienes una evaluación EC0217.01 disponible. Adquiérela ($349) '
                      'o solicita autorización al administrador.', 'info')
                return redirect(url_for('evaluar_mi_clase_landing'))

        ev = PortafolioEvaluacion(
            candidato_id=cand.id, ce_user_id=current_user.id,
            estandar='EC0217.01', status='pendiente_video',
            e5_mode=mp4_mode, e5_unlocked=(mp4_mode == 'full'),
            e5_unlocked_at=(datetime.utcnow() if mp4_mode == 'full' else None),
            e5_unlocked_via=unlocked_via,
        )
        db.session.add(ev)
        db.session.flush()  # obtiene ev.id sin confirmar la transacción

        db.session.add(PortafolioArchivo(
            portafolio_id=ev.id, tipo='video',
            filename=video_file.filename[:300],
            contenido=video_bytes,
            mime=(video_file.mimetype or 'video/mp4')[:80],
            size_bytes=video_size,
        ))

        if mp4_mode == 'preview':
            current_user.e5_preview_used_at = datetime.utcnow()

        if unlocked_via == 'admin_grant':
            saldo_after = db.session.execute(
                db.text('SELECT COALESCE(ec0217_grants,0) FROM "user" WHERE id = :uid'),
                {"uid": current_user.id}).scalar() or 0
            db.session.add(Ec0217Ledger(
                user_id=current_user.id, tipo='consume', cantidad=-1,
                saldo_after=saldo_after, portafolio_id=ev.id,
                nota=f'Evaluación EC0217.01 (mp4) iniciada — eval #{ev.id}'))

        db.session.commit()
        db.session.refresh(current_user)
        new_pid = ev.id
    except Exception:
        db.session.rollback()
        flash('No pudimos iniciar tu evaluación en este momento. Tu crédito no fue '
              'descontado. Intenta de nuevo en unos minutos.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))

    try:
        track_event('autoevaluacion_e5', 'iniciada_mp4', user_id=current_user.id,
                    extra_data={'pid': new_pid, 'mode': mp4_mode, 'via': unlocked_via,
                                'video_mb': round(video_size / (1024*1024), 1)})
    except Exception:
        pass
    return redirect(url_for('evaluar_mi_clase_vista', pid=new_pid))

@app.route("/evaluar-mi-clase/start-texto", methods=["POST"])
@login_required
def evaluar_mi_clase_start_texto():
    """Inicia una evaluación EC0217.01 a partir de la TRANSCRIPCIÓN de la clase
    (texto pegado o archivo .txt/.docx). Reemplaza por completo el flujo de video
    YouTube/MP4: no requiere descarga ni transcripción automática, el texto se
    guarda directamente y la evaluación avanza al dictamen.

    Mismo modelo de acceso que el flujo anterior: Premium / crédito a la carta E5 /
    autorización del administrador reciben evaluación completa; un usuario gratuito
    sin pre-dictamen previo recibe el pre-dictamen (Lista previa + Encuadre)."""
    transcripcion = (request.form.get('transcripcion') or '').strip()
    archivo = request.files.get('archivo')
    if not transcripcion and archivo:
        transcripcion = _leer_transcripcion_archivo(archivo)
    if not transcripcion:
        flash('Pega la transcripción de tu clase o sube un archivo .txt o .docx con la transcripción.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))
    if len(transcripcion) < 80:
        flash('La transcripción es muy corta para evaluarse. Pega el texto completo de tu clase.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))

    # Determinar vía/modo de acceso SIN consumir todavía (el consumo se hace atómico abajo).
    if current_user.is_premium:
        unlocked_via = 'premium'; modo = 'full'
    elif current_user.has_alacarte(5):
        unlocked_via = 'alacarte_e5'; modo = 'full'
    elif current_user.ec0217_authorized:
        unlocked_via = 'admin_grant'; modo = 'full'
    elif not current_user.e5_preview_used_at:
        unlocked_via = ''; modo = 'preview'
    else:
        flash('Ya usaste tu pre-dictamen gratuito. Adquiere la evaluación completa '
              '($349) o solicita autorización al administrador.', 'info')
        return redirect(url_for('evaluar_mi_clase_landing'))

    cand = Candidato.query.filter_by(ce_user_id=current_user.id, curp='AUTOEVAL_E5').first()
    if not cand:
        cand = Candidato(
            ce_user_id=current_user.id,
            nombre_completo=(current_user.email or 'Auto-evaluación')[:200],
            apellidos='(Auto-evaluación E5)', curp='AUTOEVAL_E5',
        )
        db.session.add(cand); db.session.commit()

    # Transacción única: el descuento de crédito, la creación de la evaluación y la
    # bitácora se confirman juntos o se revierten juntos.
    try:
        if unlocked_via == 'alacarte_e5':
            res = db.session.execute(
                db.text('UPDATE "user" SET alacarte_e5 = alacarte_e5 - 1 '
                        'WHERE id = :uid AND alacarte_e5 > 0'), {"uid": current_user.id})
            if res.rowcount == 0:
                db.session.rollback()
                flash('No tienes una evaluación EC0217.01 disponible. Adquiérela ($349) '
                      'o solicita autorización al administrador.', 'info')
                return redirect(url_for('evaluar_mi_clase_landing'))
        elif unlocked_via == 'admin_grant':
            res = db.session.execute(
                db.text('UPDATE "user" SET ec0217_grants = ec0217_grants - 1 '
                        'WHERE id = :uid AND ec0217_grants > 0'), {"uid": current_user.id})
            if res.rowcount == 0:
                db.session.rollback()
                flash('No tienes una evaluación EC0217.01 disponible. Adquiérela ($349) '
                      'o solicita autorización al administrador.', 'info')
                return redirect(url_for('evaluar_mi_clase_landing'))

        ev = PortafolioEvaluacion(
            candidato_id=cand.id, ce_user_id=current_user.id,
            estandar='EC0217.01',
            status=('video_procesado' if modo == 'full' else 'preview_procesado'),
            e5_mode=modo, e5_unlocked=(modo == 'full'),
            e5_unlocked_at=(datetime.utcnow() if modo == 'full' else None),
            e5_unlocked_via=unlocked_via,
        )
        ev.video_transcripcion = transcripcion
        ev.video_duracion_seg = 0
        ev.video_idioma = 'es'
        ev.video_preview_seg = 0
        ev.video_procesado_at = datetime.utcnow()
        db.session.add(ev)
        db.session.flush()

        if modo == 'preview':
            current_user.e5_preview_used_at = datetime.utcnow()

        if unlocked_via == 'admin_grant':
            saldo_after = db.session.execute(
                db.text('SELECT COALESCE(ec0217_grants,0) FROM "user" WHERE id = :uid'),
                {"uid": current_user.id}).scalar() or 0
            db.session.add(Ec0217Ledger(
                user_id=current_user.id, tipo='consume', cantidad=-1,
                saldo_after=saldo_after, portafolio_id=ev.id,
                nota=f'Evaluación EC0217.01 (transcripción) iniciada — eval #{ev.id}'))

        db.session.commit()
        db.session.refresh(current_user)
        new_pid = ev.id
    except Exception:
        db.session.rollback()
        flash('No pudimos iniciar tu evaluación en este momento. Tu crédito no fue '
              'descontado. Intenta de nuevo en unos minutos.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))

    try:
        track_event('autoevaluacion_e5', 'iniciada_texto', user_id=current_user.id,
                    extra_data={'pid': new_pid, 'mode': modo, 'via': unlocked_via,
                                'chars': len(transcripcion)})
    except Exception:
        pass
    return redirect(url_for('evaluar_mi_clase_vista', pid=new_pid))

@app.route("/evaluar-mi-clase/<int:pid>")
@login_required
def evaluar_mi_clase_vista(pid):
    ev = PortafolioEvaluacion.query.filter_by(
        id=pid, ce_user_id=current_user.id, estandar='EC0217.01'
    ).first()
    if not ev:
        flash('Evaluación no encontrada.', 'error')
        return redirect(url_for('evaluar_mi_clase_landing'))
    yt_arch = PortafolioArchivo.query.filter_by(
        portafolio_id=ev.id, tipo='youtube_url'
    ).first()
    is_unlocked = current_user.is_premium or ev.e5_unlocked
    is_preview = (ev.e5_mode == 'preview' and not ev.e5_unlocked)
    has_e5 = is_unlocked
    dictamen = None
    if ev.dictamen_json:
        try:
            dictamen = json.loads(ev.dictamen_json)
        except Exception:
            dictamen = None
    return render_template("evaluar_mi_clase_eval.html",
                           title="Tu Evaluación EC0217.01",
                           active_page="evaluar_mi_clase",
                           ev=ev, yt_url=(yt_arch.source_url if yt_arch else ''),
                           is_unlocked=is_unlocked, is_preview=is_preview, has_e5=has_e5,
                           dictamen=dictamen,
                           stripe_link_e5=STRIPE_LINKS.get('ALACARTE_E5', ''))

_IEC0217_INSTRUMENT_CACHE = None

def _iec0217_instrument_ref(max_chars=42000):
    """Texto del Instrumento de Evaluación EC0217.01 (códigos REALES de reactivos,
    p. ej. 65.1/11-D1E2, 138.1/6-D8E2), cacheado en memoria para anclar el dictamen IA."""
    global _IEC0217_INSTRUMENT_CACHE
    if _IEC0217_INSTRUMENT_CACHE is not None:
        return _IEC0217_INSTRUMENT_CACHE
    txt = ''
    try:
        from pypdf import PdfReader
        path = os.path.join('normatividad', 'Instrumento_Evaluacion_EC0217.01.pdf')
        reader = PdfReader(path)
        parts = []
        total = 0
        for page in reader.pages:
            t = page.extract_text() or ''
            if t:
                parts.append(t)
                total += len(t)
                if total > max_chars:
                    break
        txt = '\n'.join(parts)[:max_chars]
    except Exception:
        app.logger.exception('No se pudo leer el Instrumento EC0217.01')
        txt = ''
    _IEC0217_INSTRUMENT_CACHE = txt
    return txt


@app.route("/api/eval-mi-clase/<int:pid>/dictamen-rapido", methods=["POST"])
@login_required
def evaluar_mi_clase_dictamen_rapido(pid):
    ev = PortafolioEvaluacion.query.filter_by(
        id=pid, ce_user_id=current_user.id, estandar='EC0217.01'
    ).first()
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    if not (ev.video_transcripcion or '').strip():
        return jsonify({"error": "Primero registra la transcripción de tu clase."}), 400
    is_full = current_user.is_premium or ev.e5_unlocked
    is_preview = (ev.e5_mode == 'preview' and not ev.e5_unlocked)
    if not (is_full or is_preview):
        return jsonify({"error": "Esta evaluación no tiene acceso E5 activo."}), 403
    try:
        from google import genai
        from google.genai import types
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        transcripcion = (ev.video_transcripcion or '')[:60000]
        instrumento_ref = _iec0217_instrument_ref()
        reglas_g3 = (
            "REGLAS DE EVALUACIÓN (OBLIGATORIAS):\n"
            "1. Usa ÚNICAMENTE los códigos REALES de reactivo del Instrumento de Evaluación "
            "EC0217.01 incluido abajo (formato como '65.1/11-D1E2', '138.1/6-D8E2'). NUNCA "
            "inventes códigos ni uses abreviaturas propias (prohibido LP.x, EN.x, DE.x, CI.x).\n"
            "2. Evalúa EXCLUSIVAMENTE lo que aparezca escrito en la TRANSCRIPCIÓN recibida. No "
            "supongas, no infieras ni inventes evidencia que no esté en el texto.\n"
            "3. Antes de marcar un reactivo como 'no_cumplido', confirma que la transcripción SÍ "
            "permite evaluarlo y que la evidencia explícita NO aparece. Si existe cualquier cita "
            "textual que lo respalde (síntesis pedida a participantes, mención de la próxima "
            "sesión, despedida, direccionamiento a plataforma, etc.), márcalo 'cumplido' con esa "
            "cita.\n"
            "4. Marca como 'no_evaluable' todo reactivo que requiera observación PRESENCIAL física "
            "o productos documentales que NO pueden constatarse en una transcripción de audio "
            "(p. ej. condiciones físicas del aula, materiales impresos, lenguaje corporal, lista "
            "de asistencia firmada, formatos en papel). NUNCA los marques como 'no_cumplido'.\n"
            "5. El porcentaje de cobertura se calcula SOLO sobre reactivos evaluables por "
            "transcripción (estado 'cumplido' o 'no_cumplido'), nunca sobre los 'no_evaluable'.\n"
        )
        json_spec = (
            "Devuelve ÚNICAMENTE JSON válido con esta forma:\n"
            "{\"dictamen_global\":\"Competente\"|\"Todavía No Competente\","
            "\"resumen_global\":\"2-3 líneas honestas sobre lo evidenciado en la transcripción\","
            "\"reactivos\":[{\"id\":\"<código real, ej 65.1/11-D1E2>\","
            "\"momento\":\"LISTA_PREVIA|ENCUADRE|DESARROLLO|CIERRE\","
            "\"descripcion\":\"criterio del reactivo en breve\","
            "\"estado\":\"cumplido\"|\"no_cumplido\"|\"no_evaluable\","
            "\"evidencia\":\"cita textual LITERAL de la transcripción que lo respalda; vacío si "
            "no_cumplido o no_evaluable\","
            "\"recomendacion\":\"qué mejorar (si no_cumplido) o por qué no es evaluable\"}]}\n"
        )
        if is_preview:
            prompt = (
                "Eres Evaluador Acreditado CONOCER del estándar EC0217.01 "
                "'Impartición de cursos de formación del capital humano de manera presencial grupal'. "
                "Recibirás la TRANSCRIPCIÓN de una clase. En este PRE-DICTAMEN gratuito evalúa "
                "SOLO los reactivos de la 'Lista de verificación previa a la impartición' y del "
                "'Encuadre' (apertura de la sesión) del Instrumento de Evaluación EC0217.01.\n\n"
                + reglas_g3 + "\n" + json_spec +
                "\nRegla de dictamen: 'Competente' solo si (reactivos cumplidos / reactivos "
                "evaluables) >= 0.80. Sé exigente y honesto — esto es un pre-dictamen IA, no el "
                "dictamen oficial.\n\n"
                "=== INSTRUMENTO DE EVALUACIÓN EC0217.01 (códigos reales de reactivos) ===\n"
                f"{instrumento_ref}\n\n"
                f"=== TRANSCRIPCIÓN DE LA CLASE ===\n{transcripcion}"
            )
        else:
            prompt = (
                "Eres Evaluador Acreditado CONOCER del estándar EC0217.01 "
                "'Impartición de cursos de formación del capital humano de manera presencial grupal'. "
                "Recibirás la TRANSCRIPCIÓN COMPLETA de una clase. Evalúa, reactivo por reactivo, "
                "la Guía de Observación del Instrumento de Evaluación EC0217.01 a lo largo de las "
                "etapas: Lista previa, Encuadre/apertura, Desarrollo y Cierre.\n\n"
                + reglas_g3 + "\n" + json_spec +
                "\nRegla de dictamen: 'Competente' solo si (reactivos cumplidos / reactivos "
                "evaluables) >= 0.80. Sé exigente y honesto.\n\n"
                "=== INSTRUMENTO DE EVALUACIÓN EC0217.01 (códigos reales de reactivos) ===\n"
                f"{instrumento_ref}\n\n"
                f"=== TRANSCRIPCIÓN DE LA CLASE ===\n{transcripcion}"
            )
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=16000,
            )
        )
        raw = response.text if response.text else '{}'
        dictamen = json.loads(raw)
        # --- Normalización Grupo 3: cobertura SOLO sobre reactivos evaluables ---
        reactivos = dictamen.get('reactivos') or []
        def _estado_reactivo(r):
            e = (r.get('estado') or '').strip().lower().replace(' ', '_')
            if e in ('cumplido', 'no_cumplido', 'no_evaluable'):
                return e
            if r.get('cumple') is True:
                return 'cumplido'
            if r.get('cumple') is False:
                return 'no_cumplido'
            return 'no_evaluable'
        for r in reactivos:
            r['estado'] = _estado_reactivo(r)
            r['cumple'] = (r['estado'] == 'cumplido')
        n_eval = sum(1 for r in reactivos if r['estado'] in ('cumplido', 'no_cumplido'))
        n_cumpl = sum(1 for r in reactivos if r['estado'] == 'cumplido')
        dictamen['reactivos'] = reactivos
        dictamen['total_reactivos'] = len(reactivos)
        dictamen['reactivos_evaluables'] = n_eval
        dictamen['reactivos_cumplidos'] = n_cumpl
        dictamen['porcentaje_cobertura'] = int(round(n_cumpl / n_eval * 100)) if n_eval else 0
        dictamen['dictamen_global'] = ('Competente' if (n_eval and (n_cumpl / n_eval) >= 0.80)
                                       else 'Todavía No Competente')
        ev.dictamen_json = json.dumps(dictamen, ensure_ascii=False)
        ev.dictamen_final = (dictamen.get('dictamen_global') or '')[:40]
        ev.dictamen_generado_at = datetime.utcnow()
        ev.status = ('predictamen_generado' if is_preview else 'dictaminado')
        db.session.commit()
        try:
            track_event('autoevaluacion_e5',
                        ('predictamen_generado' if is_preview else 'dictamen_generado'),
                        user_id=current_user.id,
                        extra_data={'pid': ev.id, 'dictamen': dictamen.get('dictamen_global'),
                                    'pct': dictamen.get('porcentaje_cobertura'),
                                    'mode': ev.e5_mode, 'via': ev.e5_unlocked_via})
        except Exception:
            pass
        if is_full:
            payload = dictamen
        else:
            # FREE/preview: nunca exponer el desglose completo por la API.
            payload = {
                'dictamen_global': dictamen.get('dictamen_global'),
                'porcentaje_cobertura': dictamen.get('porcentaje_cobertura'),
                'resumen_global': dictamen.get('resumen_global'),
                'reactivos_evaluables': dictamen.get('reactivos_evaluables'),
                'reactivos_cumplidos': dictamen.get('reactivos_cumplidos'),
                'reactivos': (dictamen.get('reactivos') or [])[:3],
            }
        return jsonify({"success": True, "dictamen": payload,
                        "unlocked": is_full, "mode": ev.e5_mode})
    except Exception:
        db.session.rollback()
        app.logger.exception("Error en dictamen rápido EC0217.01")
        return jsonify({"error": "No pudimos generar tu dictamen en este momento. "
                                 "Vuelve a intentarlo en unos minutos."}), 500

@app.route("/api/eval-mi-clase/<int:pid>/upgrade-to-full", methods=["POST"])
@login_required
def evaluar_mi_clase_upgrade_to_full(pid):
    ev = PortafolioEvaluacion.query.filter_by(
        id=pid, ce_user_id=current_user.id, estandar='EC0217.01'
    ).first()
    if not ev:
        return jsonify({"error": "Evaluación no encontrada"}), 404
    if ev.e5_unlocked:
        return jsonify({"success": True, "already_unlocked": True,
                        "needs_video_continuation": False})
    lock_result = db.session.execute(
        db.text("UPDATE portafolio_evaluacion SET e5_unlocked = TRUE, e5_unlocked_at = NOW(), "
                "e5_mode = 'full', e5_unlocked_via = 'pending' "
                "WHERE id = :pid AND e5_unlocked = FALSE"),
        {"pid": ev.id}
    )
    db.session.commit()
    if lock_result.rowcount == 0:
        db.session.refresh(ev)
        return jsonify({"success": True, "already_unlocked": True,
                        "needs_video_continuation": False})
    unlocked_via = ''
    if current_user.is_premium:
        unlocked_via = 'premium'
    elif current_user.has_alacarte(5):
        if current_user.use_alacarte(5):
            unlocked_via = 'alacarte_e5'
    if not unlocked_via:
        db.session.execute(
            db.text("UPDATE portafolio_evaluacion SET e5_unlocked = FALSE, e5_unlocked_at = NULL, "
                    "e5_mode = 'preview', e5_unlocked_via = '' WHERE id = :pid"),
            {"pid": ev.id}
        )
        db.session.commit()
        return jsonify({"error": "Sin crédito E5 disponible. Adquiere el servicio en Stripe."}), 402
    db.session.execute(
        db.text("UPDATE portafolio_evaluacion SET e5_unlocked_via = :via, "
                "dictamen_json = NULL, dictamen_final = '' WHERE id = :pid"),
        {"via": unlocked_via, "pid": ev.id}
    )
    db.session.commit()
    db.session.refresh(ev)
    try:
        track_event('autoevaluacion_e5', 'upgrade_to_full', user_id=current_user.id,
                    extra_data={'pid': ev.id, 'via': unlocked_via,
                                'preview_seg_previo': ev.video_preview_seg or 0})
    except Exception:
        pass
    return jsonify({"success": True, "via": unlocked_via,
                    "needs_video_continuation": (ev.video_preview_seg or 0) > 0})

@app.route("/api/eval-mi-clase/<int:pid>/descargar-pdf", methods=["GET"])
@login_required
def evaluar_mi_clase_descargar_pdf(pid):
    ev = PortafolioEvaluacion.query.filter_by(
        id=pid, ce_user_id=current_user.id, estandar='EC0217.01'
    ).first()
    if not ev:
        abort(404)
    is_full = current_user.is_premium or ev.e5_unlocked
    if not is_full:
        flash("Para descargar el informe en PDF necesitas la evaluación completa.", "warning")
        return redirect(url_for('evaluar_mi_clase_vista', pid=ev.id))
    if not ev.dictamen_json:
        flash("Aún no hay un dictamen generado para descargar.", "warning")
        return redirect(url_for('evaluar_mi_clase_vista', pid=ev.id))
    try:
        dictamen = json.loads(ev.dictamen_json)
    except Exception:
        flash("No se pudo leer el dictamen. Vuelve a generarlo.", "danger")
        return redirect(url_for('evaluar_mi_clase_vista', pid=ev.id))
    cand = ev.candidato
    cand_nombre = ''
    if cand:
        cand_nombre = f"{cand.nombre_completo or ''} {cand.apellidos or ''}".strip()
    curso_nombre = getattr(ev, 'curso_nombre', '') or getattr(ev, 'nombre_curso', '') or ''
    try:
        from portafolio_integrador import generar_pdf_dictamen_rapido
        pdf_bytes = generar_pdf_dictamen_rapido(
            dictamen, candidato_nombre=cand_nombre,
            curso_nombre=curso_nombre, estandar='EC0217.01')
    except Exception:
        app.logger.exception("Error generando PDF de dictamen rápido")
        flash("No pudimos generar el PDF en este momento. Inténtalo de nuevo más tarde.", "danger")
        return redirect(url_for('evaluar_mi_clase_vista', pid=ev.id))
    try:
        track_event('autoevaluacion_e5', 'pdf_descargado', user_id=current_user.id,
                    extra_data={'pid': ev.id})
    except Exception:
        pass
    return send_file(io.BytesIO(pdf_bytes), mimetype='application/pdf',
                     as_attachment=True,
                     download_name=f"dictamen_ec0217_{ev.id}.pdf")

@app.route("/api/ads/active")
def ads_active():
    from datetime import datetime
    today = datetime.utcnow().date()
    ads = AdCampaign.query.filter(
        AdCampaign.status == 'Activo',
        AdCampaign.starts_at <= today,
        AdCampaign.ends_at >= today
    ).all()
    return jsonify([{
        "id": ad.id,
        "title": ad.title,
        "description": ad.description,
        "bg_gradient": ad.bg_gradient,
        "target_url": ad.target_url,
        "image_url": ad.image_url
    } for ad in ads])

@app.route("/api/ads/availability")
def ads_availability():
    from datetime import datetime, timedelta
    today = datetime.utcnow().date()
    max_concurrent = int(Config.get('MAX_CONCURRENT_ADS', '5'))
    saturated = []
    active_statuses = ['Activo', 'Pendiente de Revisión']
    candidates = AdCampaign.query.filter(
        AdCampaign.status.in_(active_statuses),
        AdCampaign.ends_at >= today
    ).all()
    for i in range(60):
        check_date = today + timedelta(days=i)
        count = sum(1 for ad in candidates if ad.starts_at and ad.ends_at and ad.starts_at <= check_date <= ad.ends_at)
        if count >= max_concurrent:
            saturated.append(check_date.isoformat())
    return jsonify({"saturated_dates": saturated, "max_concurrent": max_concurrent})

@app.route("/api/ads/impression/<int:ad_id>", methods=["POST"])
def ad_impression(ad_id):
    ad = db.session.get(AdCampaign, ad_id)
    if ad and ad.status == 'Activo':
        ad.impressions = (ad.impressions or 0) + 1
        db.session.commit()
    return jsonify({"ok": True})

@app.route("/api/ads/click/<int:ad_id>")
def ad_click(ad_id):
    ad = db.session.get(AdCampaign, ad_id)
    if ad:
        ad.clicks = (ad.clicks or 0) + 1
        db.session.commit()
        return redirect(ad.target_url)
    return redirect("/")

@app.route("/api/ads/upload-image", methods=["POST"])
@login_required
def ads_upload_image():
    import uuid as _uuid
    upload_dir = os.path.join(app.static_folder, 'uploads', 'ads')
    os.makedirs(upload_dir, exist_ok=True)

    if 'image' not in request.files:
        return jsonify({"error": "No se envió ninguna imagen."}), 400
    file = request.files['image']
    if not file.filename:
        return jsonify({"error": "Archivo vacío."}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ('png', 'jpg', 'jpeg'):
        return jsonify({"error": "Solo se aceptan imágenes PNG, JPG o JPEG."}), 400

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > 2 * 1024 * 1024:
        return jsonify({"error": "La imagen no debe superar 2MB."}), 400

    filename = f"{_uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(upload_dir, filename)
    file.save(filepath)
    image_url = f"/static/uploads/ads/{filename}"
    return jsonify({"success": True, "image_url": image_url})

@app.route("/api/ads/create-session", methods=["POST"])
@login_required
def ads_create_session():
    import base64
    from datetime import datetime, timedelta
    data = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    target_url = (data.get("target_url") or "").strip()
    starts_at_str = (data.get("starts_at") or "").strip()
    total_cost = float(data.get("total_cost") or 0)

    if not title or not target_url:
        return jsonify({"error": "Título y URL destino son obligatorios."}), 400
    if total_cost < 500:
        return jsonify({"error": "El presupuesto mínimo es de $500 MXN."}), 400
    if not starts_at_str:
        return jsonify({"error": "La fecha de inicio es obligatoria."}), 400

    try:
        starts_at = datetime.strptime(starts_at_str, "%Y-%m-%d").date()
    except ValueError:
        return jsonify({"error": "Formato de fecha inválido."}), 400

    if starts_at < datetime.now().date():
        return jsonify({"error": "La fecha de inicio no puede ser en el pasado."}), 400

    try:
        cost_per_day = float(Config.get('COST_PER_DAY_ADS', '150'))
        if cost_per_day <= 0:
            cost_per_day = 150.0
    except (ValueError, TypeError):
        cost_per_day = 150.0
    num_days = max(1, int(total_cost / cost_per_day))
    ends_at = starts_at + timedelta(days=num_days)

    max_concurrent = int(Config.get('MAX_CONCURRENT_ADS', '5'))
    active_statuses = ['Activo', 'Pendiente de Revisión']
    existing = AdCampaign.query.filter(
        AdCampaign.status.in_(active_statuses),
        AdCampaign.ends_at >= starts_at,
        AdCampaign.starts_at <= ends_at
    ).all()
    for i in range(num_days + 1):
        check_date = starts_at + timedelta(days=i)
        count = sum(1 for ad in existing if ad.starts_at and ad.ends_at and ad.starts_at <= check_date <= ad.ends_at)
        if count >= max_concurrent:
            return jsonify({"error": f"Inventario agotado en la fecha {check_date.isoformat()}. Máximo {max_concurrent} campañas simultáneas. Ajusta tus fechas o presupuesto."}), 400

    signature_b64 = data.get("signature_image")
    if total_cost > 10000 and not signature_b64:
        return jsonify({"error": "Para montos superiores a $10,000 MXN se requiere firma digital del contrato."}), 400

    signature_bytes = None
    if signature_b64:
        try:
            if ',' in signature_b64:
                signature_b64 = signature_b64.split(',', 1)[1]
            signature_bytes = base64.b64decode(signature_b64)
        except Exception:
            return jsonify({"error": "Imagen de firma inválida."}), 400

    if not stripe.api_key:
        logger.error("STRIPE_SECRET_KEY not configured — cannot create checkout session.")
        return jsonify({"error": "Pasarela de pago no configurada. Contacta al administrador."}), 500

    try:
        success_url = url_for('publicitar', paid=1, _external=True)
        cancel_url = url_for('publicitar', cancelled=1, _external=True)
        amount_cents = int(total_cost * 100)

        ad = AdCampaign(
            title=title,
            target_url=target_url,
            description=(data.get("description") or "").strip(),
            total_cost=total_cost,
            starts_at=starts_at,
            ends_at=ends_at,
            advertiser_email=current_user.email,
            advertiser_id=current_user.id,
            status='Pendiente Pago',
            signature_image=signature_bytes,
            signing_ip=request.remote_addr,
            image_url=(data.get("image_url") or "").strip() or None
        )
        db.session.add(ad)
        db.session.flush()

        checkout_session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{
                "price_data": {
                    "currency": "mxn",
                    "product_data": {
                        "name": f"Campaña Publicitaria: {title}",
                        "description": f"Publicidad digital en Pertinentia® — {num_days} días (del {starts_at} al {ends_at})",
                    },
                    "unit_amount": amount_cents,
                },
                "quantity": 1,
            }],
            mode="payment",
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                "ad_id": str(ad.id),
                "type": "advertisement",
            },
            customer_email=current_user.email,
        )
        ad.stripe_session_id = checkout_session.id
        db.session.commit()
        return jsonify({"success": True, "checkout_url": checkout_session.url})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Stripe session creation failed: {e}")
        return jsonify({"error": "Error al procesar el pago. Verifica tu conexión e intenta de nuevo."}), 500

@app.route("/api/ads/delete/<int:ad_id>", methods=["POST"])
@login_required
def ads_delete(ad_id):
    ad = db.session.get(AdCampaign, ad_id)
    if not ad or ad.advertiser_id != current_user.id:
        return jsonify({"error": "Campaña no encontrada."}), 404
    if ad.status != 'Pendiente Pago':
        return jsonify({"error": "Esta campaña ya ha sido procesada o está activa y no puede ser eliminada."}), 400
    db.session.delete(ad)
    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/ads/report-issue", methods=["POST"])
@login_required
def ads_report_issue():
    data = request.get_json(silent=True) or {}
    issue_type = (data.get("type") or "").strip()
    description = (data.get("description") or "").strip()
    if not description:
        return jsonify({"error": "Describe el problema."}), 400
    issue = AdIssue(
        user_id=current_user.id,
        user_email=current_user.email,
        issue_type=issue_type or 'otro',
        description=description
    )
    db.session.add(issue)
    db.session.commit()
    logger.info(f"Ad issue reported: user={current_user.id}, type={issue_type}")
    import threading
    _app = app._get_current_object()
    _user_email = current_user.email
    _user_name = current_user.full_name
    _issue_type = issue_type
    _description = description
    _user_id_capture = current_user.id
    def _notify_admin():
        with _app.app_context():
            log_id = None
            try:
                smtp_server = os.environ.get('SMTP_SERVER')
                smtp_port = int(os.environ.get('SMTP_PORT', 587))
                smtp_email = os.environ.get('SMTP_EMAIL')
                smtp_password = os.environ.get('SMTP_PASSWORD')
                log_id = _log_email_attempt(
                    email_type='ad_issue',
                    recipient='soporte@pertinentia.com',
                    sender=smtp_email,
                    subject=f'[Incidencia Publicidad] {_issue_type} — {_user_email}',
                    user_id=_user_id_capture,
                )
                if not all([smtp_server, smtp_email, smtp_password]):
                    _log_email_result(log_id, False, error="SMTP credentials not configured")
                    return
                msg = MIMEMultipart('mixed')
                msg['Subject'] = f'[Incidencia Publicidad] {_issue_type} — {_user_email}'
                msg['From'] = smtp_email
                msg['To'] = 'soporte@pertinentia.com'
                html_body = f"""
                <div style="font-family:Arial,sans-serif;max-width:550px;margin:0 auto;padding:20px;">
                    <h2 style="color:#d97706;">Nueva Incidencia de Publicidad</h2>
                    <p><strong>Anunciante:</strong> {_user_name} ({_user_email})</p>
                    <p><strong>Tipo:</strong> {_issue_type}</p>
                    <p><strong>Descripci&oacute;n:</strong></p>
                    <p style="background:#fefce8;padding:12px;border-radius:8px;border:1px solid #fde68a;">{_description}</p>
                    <hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0;">
                    <p style="color:#9ca3af;font-size:0.75em;">PertinentIA — Sistema de Incidencias Publicitarias</p>
                </div>
                """
                msg.attach(MIMEText(html_body, 'html'))
                with smtplib.SMTP(smtp_server, smtp_port) as server:
                    server.starttls()
                    server.login(smtp_email, smtp_password)
                    server.sendmail(smtp_email, 'soporte@pertinentia.com', msg.as_string())
                logger.info(f"Ad issue notification sent to soporte@pertinentia.com for user {_user_email}")
                _log_email_result(log_id, True, smtp_response="250 Accepted by Titan")
            except Exception as e:
                logger.warning(f"Failed to send ad issue notification: {e}")
                _log_email_result(log_id, False, error=e)
    threading.Thread(target=_notify_admin, daemon=True).start()
    my_issues = AdIssue.query.filter_by(user_id=current_user.id).order_by(AdIssue.created_at.desc()).limit(10).all()
    issues_list = [{"id": i.id, "type": i.issue_type, "description": i.description, "status": i.status, "date": i.created_at.strftime('%d/%m/%Y') if i.created_at else ''} for i in my_issues]
    return jsonify({"success": True, "issues": issues_list})

@app.route("/api/ads/my-issues")
@login_required
def ads_my_issues():
    issues = AdIssue.query.filter_by(user_id=current_user.id).order_by(AdIssue.created_at.desc()).limit(20).all()
    return jsonify({"issues": [{"id": i.id, "type": i.issue_type, "description": i.description, "status": i.status, "date": i.created_at.strftime('%d/%m/%Y') if i.created_at else ''} for i in issues]})

@app.route("/api/admin/ads", methods=["POST"])
@login_required
def admin_create_ad():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    target_url = (data.get("target_url") or "").strip()
    if not title or not target_url:
        return jsonify({"error": "Título y URL destino son obligatorios."}), 400
    ad = AdCampaign(
        title=title,
        description=(data.get("description") or "").strip(),
        bg_gradient=(data.get("bg_gradient") or "from-primary/15 to-primary/5").strip(),
        target_url=target_url,
        advertiser_email=(data.get("advertiser_email") or "").strip(),
        status='Pendiente'
    )
    db.session.add(ad)
    db.session.commit()
    return jsonify({"success": True, "id": ad.id})

@app.route("/api/admin/ads/<int:ad_id>/status", methods=["POST"])
@login_required
def admin_toggle_ad_status(ad_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    ad = db.session.get(AdCampaign, ad_id)
    if not ad:
        return jsonify({"error": "Anuncio no encontrado"}), 404
    data = request.get_json(silent=True) or {}
    new_status = (data.get("status") or "").strip()
    if new_status not in ("Activo", "Pausado", "Pendiente"):
        return jsonify({"error": "Estatus inválido. Valores: Activo, Pausado, Pendiente"}), 400
    ad.status = new_status
    db.session.commit()
    return jsonify({"success": True, "status": ad.status})

@app.route("/api/admin/ads/<int:ad_id>", methods=["DELETE"])
@login_required
def admin_delete_ad(ad_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    ad = db.session.get(AdCampaign, ad_id)
    if not ad:
        return jsonify({"error": "Anuncio no encontrado"}), 404
    db.session.delete(ad)
    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/admin/export_atribucion")
@login_required
def admin_export_atribucion():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    import csv
    from io import StringIO
    from flask import Response
    days = request.args.get('days', type=int) or 0
    where_t = ''
    params = {}
    if days > 0:
        where_t = "AND created_at >= NOW() - (:dd || ' days')::interval"
        params['dd'] = str(days)
    rows = db.session.execute(db.text(f'''
        SELECT id, email, full_name, tier, created_at,
               utm_source, utm_medium, utm_campaign, utm_term, utm_content, utm_id, utm_landing,
               chat_usage_count, free_downloads_used
        FROM "user"
        WHERE (utm_source IS NOT NULL OR utm_campaign IS NOT NULL OR utm_content IS NOT NULL)
          {where_t}
        ORDER BY created_at DESC
    '''), params).fetchall()
    buf = StringIO()
    w = csv.writer(buf)
    w.writerow(['id','email','nombre','plan','registro_cdmx',
                'utm_source','utm_medium','utm_campaign','utm_term','utm_content','utm_id','utm_landing',
                'chats_usados','descargas_free'])
    for r in rows:
        w.writerow([r.id, r.email, r.full_name or '', r.tier, _fmt_cdmx(r.created_at, '%Y-%m-%d %H:%M'),
                    r.utm_source or '', r.utm_medium or '', r.utm_campaign or '',
                    r.utm_term or '', r.utm_content or '', r.utm_id or '', r.utm_landing or '',
                    r.chat_usage_count or 0, r.free_downloads_used or 0])
    return Response(buf.getvalue(), mimetype='text/csv; charset=utf-8',
                    headers={'Content-Disposition': 'attachment; filename=atribucion_utm.csv'})


@app.route("/admin/atribucion")
@login_required
def admin_atribucion():
    if not current_user.is_admin:
        track_event('Seguridad', 'Acceso Denegado Admin', user_id=current_user.id,
                    extra_data={'email': current_user.email, 'ruta': '/admin/atribucion'})
        flash('Acceso restringido.', 'error')
        return redirect(url_for('home'))

    sort = (request.args.get('sort') or 'registros').strip()
    group_by = (request.args.get('group') or 'content').strip()
    days = request.args.get('days', type=int) or 0

    group_col_map = {
        'content': 'utm_content',
        'term': 'utm_term',
        'campaign': 'utm_campaign',
        'medium': 'utm_medium',
        'source': 'utm_source',
        'landing': 'utm_landing',
    }
    group_col = group_col_map.get(group_by, 'utm_content')

    sort_sql = {
        'registros': 'registros DESC NULLS LAST',
        'pro': 'pro_paid DESC NULLS LAST',
        'activacion': 'activacion_pct DESC NULLS LAST',
        'reciente': 'ultimo_registro DESC NULLS LAST',
    }.get(sort, 'registros DESC NULLS LAST')

    where_time = ''
    params = {}
    if days and days > 0:
        where_time = 'AND u.created_at >= NOW() - (:dd || \' days\')::interval'
        params['dd'] = str(days)

    agg_sql = db.text(f'''
        SELECT
            COALESCE(NULLIF(u.{group_col}, ''), '(sin atribución)') AS bucket,
            u.utm_source, u.utm_medium, u.utm_campaign, u.utm_term, u.utm_content, u.utm_landing,
            COUNT(*) AS registros,
            SUM(CASE WHEN u.tier IN ('PRO','PREMIUM','MULTI','PRO_MULTI') THEN 1 ELSE 0 END) AS pro_paid,
            SUM(CASE WHEN u.chat_usage_count > 0 OR u.free_downloads_used > 0
                       OR u.alacarte_e1 + u.alacarte_e2 + u.alacarte_e3 + u.alacarte_e4 > 0
                     THEN 1 ELSE 0 END) AS activados,
            ROUND(100.0 * SUM(CASE WHEN u.chat_usage_count > 0 OR u.free_downloads_used > 0
                       OR u.alacarte_e1 + u.alacarte_e2 + u.alacarte_e3 + u.alacarte_e4 > 0
                     THEN 1 ELSE 0 END) / NULLIF(COUNT(*),0), 1) AS activacion_pct,
            MAX(u.created_at) AS ultimo_registro,
            MIN(u.created_at) AS primer_registro
        FROM "user" u
        WHERE (u.utm_source IS NOT NULL OR u.utm_campaign IS NOT NULL OR u.utm_content IS NOT NULL)
          {where_time}
        GROUP BY u.{group_col}, u.utm_source, u.utm_medium, u.utm_campaign, u.utm_term, u.utm_content, u.utm_landing
        ORDER BY {sort_sql}
        LIMIT 200
    ''')
    try:
        agg_rows = db.session.execute(agg_sql, params).fetchall()
    except Exception as _e_agg:
        logger.warning(f"admin_atribucion agg fail: {_e_agg}")
        agg_rows = []

    detail_sql = db.text(f'''
        SELECT id, email, full_name, tier, created_at,
               utm_source, utm_medium, utm_campaign, utm_term, utm_content, utm_id, utm_landing
        FROM "user"
        WHERE (utm_source IS NOT NULL OR utm_campaign IS NOT NULL OR utm_content IS NOT NULL)
          {where_time.replace('u.', '')}
        ORDER BY created_at DESC
        LIMIT 500
    ''')
    try:
        detail_rows = db.session.execute(detail_sql, params).fetchall()
    except Exception as _e_det:
        logger.warning(f"admin_atribucion detail fail: {_e_det}")
        detail_rows = []

    try:
        total_with_utm = db.session.execute(db.text(
            'SELECT COUNT(*) FROM "user" WHERE utm_source IS NOT NULL OR utm_campaign IS NOT NULL OR utm_content IS NOT NULL'
        )).scalar() or 0
        total_users = db.session.execute(db.text('SELECT COUNT(*) FROM "user"')).scalar() or 0
    except Exception:
        total_with_utm = 0
        total_users = 0

    return render_template('admin_atribucion.html',
                           title='Atribución de Campañas',
                           active_page='admin',
                           agg_rows=agg_rows,
                           detail_rows=detail_rows,
                           sort=sort, group_by=group_by, days=days,
                           total_with_utm=total_with_utm, total_users=total_users,
                           fmt_cdmx=_fmt_cdmx)


@app.route("/admin-crm")
@login_required
def admin_crm():
    if not current_user.is_admin:
        track_event('Seguridad', 'Acceso Denegado Admin', user_id=current_user.id, extra_data={'email': current_user.email, 'tier': current_user.tier})
        flash('Usuario no autorizado.', 'error')
        return redirect(url_for('home'))
    page = request.args.get('page', 1, type=int)
    per_page = 20
    f_q = (request.args.get('q', '') or '').strip()
    f_tier = (request.args.get('tier', '') or '').strip().upper()
    f_wa = (request.args.get('wa', '') or '').strip().lower()
    f_mkt = (request.args.get('mkt', '') or '').strip().lower()
    f_act = (request.args.get('act', '') or '').strip().lower()
    _u_query = User.query
    if f_q:
        like = f"%{f_q}%"
        try:
            qid = int(f_q)
        except (TypeError, ValueError):
            qid = None
        from sqlalchemy import or_ as _or_
        cond = _or_(
            User.email.ilike(like),
            User.full_name.ilike(like),
            User.whatsapp.ilike(like),
        )
        if qid is not None:
            cond = _or_(cond, User.id == qid)
        _u_query = _u_query.filter(cond)
    if f_tier in ('FREE', 'PRO', 'PREMIUM', 'MULTI', 'PRO_MULTI'):
        _u_query = _u_query.filter(User.tier == f_tier)
    if f_wa == 'con':
        _u_query = _u_query.filter(User.whatsapp.isnot(None), User.whatsapp != '')
    elif f_wa == 'sin':
        from sqlalchemy import or_ as _or_w
        _u_query = _u_query.filter(_or_w(User.whatsapp.is_(None), User.whatsapp == ''))
    if f_mkt == 'si':
        _u_query = _u_query.filter(User.marketing_consent.is_(True))
    elif f_mkt == 'no':
        from sqlalchemy import or_ as _or_m
        _u_query = _u_query.filter(_or_m(User.marketing_consent.is_(False), User.marketing_consent.is_(None)))
    if f_act == 'con_ia':
        _u_query = _u_query.filter(User.chat_usage_count > 0)
    elif f_act == 'sin_ia':
        from sqlalchemy import or_ as _or_a
        _u_query = _u_query.filter(_or_a(User.chat_usage_count == 0, User.chat_usage_count.is_(None)))
    elif f_act == 'rebote':
        _u_query = _u_query.filter(User.email_bounced_at.isnot(None))
    elif f_act == 'contactado':
        _u_query = _u_query.filter(User.ultimo_contacto_at.isnot(None))
    elif f_act == 'no_contactado':
        _u_query = _u_query.filter(User.ultimo_contacto_at.is_(None))
    users_pagination = _u_query.order_by(User.id.desc()).paginate(page=page, per_page=per_page, error_out=False)
    users = users_pagination.items
    total_users = users_pagination.total
    user_filters = {'q': f_q, 'tier': f_tier, 'wa': f_wa, 'mkt': f_mkt, 'act': f_act}
    from urllib.parse import quote as _qs_quote
    user_filters_qs = '&'.join(f'{k}={_qs_quote(v)}' for k, v in user_filters.items() if v)
    referrer_map = {}
    referred_ids = [u.referred_by for u in users if u.referred_by]
    if referred_ids:
        referrers = User.query.filter(User.id.in_(set(referred_ids))).all()
        referrer_map = {r.id: r.full_name for r in referrers}
    pro_signal_counts = {}
    _uid_list = [u.id for u in users]
    if _uid_list:
        try:
            from sqlalchemy import func as _sfunc
            _sig_rows = db.session.query(
                UserEvent.user_id, UserEvent.event_action, _sfunc.count(UserEvent.id)
            ).filter(
                UserEvent.user_id.in_(_uid_list),
                UserEvent.event_action.in_(['PRO Curso Agotado', 'PRO Lock Fallo'])
            ).group_by(UserEvent.user_id, UserEvent.event_action).all()
            for _uid_r, _act_r, _cnt_r in _sig_rows:
                _d = pro_signal_counts.setdefault(_uid_r, {'agotado': 0, 'fallo': 0})
                if _act_r == 'PRO Curso Agotado':
                    _d['agotado'] = _cnt_r
                elif _act_r == 'PRO Lock Fallo':
                    _d['fallo'] = _cnt_r
        except Exception:
            try:
                db.session.rollback()
            except Exception:
                pass
            pro_signal_counts = {}
    all_leads = AffiliateLead.query.order_by(AffiliateLead.id.desc()).all()
    paid_leads = [l for l in all_leads if l.status == 'Pagado']
    comision_pct = Config.get('COMISION_PORCENTAJE', '0.30')
    stripe_wh_db = Config.get('STRIPE_WEBHOOK_SECRET', '')
    stripe_wh_env = STRIPE_WEBHOOK_SECRET
    stripe_wh_effective = stripe_wh_db or stripe_wh_env
    if stripe_wh_db and len(stripe_wh_db) > 6:
        stripe_wh_masked = '••••••' + stripe_wh_db[-6:]
    elif stripe_wh_db:
        stripe_wh_masked = '••••••'
    else:
        stripe_wh_masked = ''
    stripe_wh_source = 'bd' if stripe_wh_db else ('env' if stripe_wh_env else '')
    stripe_checkout_aff = Config.get('STRIPE_CHECKOUT_URL_AFFILIATE', '')
    cost_per_day_ads = Config.get('COST_PER_DAY_ADS', '150')
    max_concurrent_ads = Config.get('MAX_CONCURRENT_ADS', '5')
    max_concurrent_ai = Config.get('MAX_CONCURRENT_AI', '3')

    from datetime import datetime, timedelta
    now_utc = datetime.utcnow()
    sat_24h = BetaMetric.query.filter(
        BetaMetric.metric_type == 'ai_saturation_reject',
        BetaMetric.created_at >= now_utc - timedelta(hours=24)
    ).all()
    sat_7d = BetaMetric.query.filter(
        BetaMetric.metric_type == 'ai_saturation_reject',
        BetaMetric.created_at >= now_utc - timedelta(days=7)
    ).count()
    sat_count_24h = len(sat_24h)
    sat_by_hour = {}
    for s in sat_24h:
        h = s.created_at.strftime('%H:00') if s.created_at else '??'
        sat_by_hour[h] = sat_by_hour.get(h, 0) + 1
    sat_peak_hour = max(sat_by_hour, key=sat_by_hour.get) if sat_by_hour else None
    sat_peak_count = sat_by_hour.get(sat_peak_hour, 0) if sat_peak_hour else 0

    current_max = int(max_concurrent_ai)
    if sat_count_24h == 0:
        sat_suggestion = f"Sin rechazos en 24h. Tu l\u00edmite actual ({current_max}) es adecuado."
        sat_level = "ok"
    elif sat_count_24h <= 3:
        sat_suggestion = f"{sat_count_24h} rechazos en 24h — nivel bajo, monitorea. Si se repite ma\u00f1ana, sube a {current_max + 1}."
        sat_level = "low"
    elif sat_count_24h <= 10:
        sat_suggestion = f"{sat_count_24h} rechazos en 24h (pico: {sat_peak_count} a las {sat_peak_hour}). Sube el l\u00edmite a {min(current_max + 2, 10)}."
        sat_level = "medium"
    else:
        sat_suggestion = f"\u00a1{sat_count_24h} rechazos en 24h! Pico: {sat_peak_count} a las {sat_peak_hour}. Sube urgente a {min(current_max + 3, 10)}."
        sat_level = "high"

    ad_campaigns = AdCampaign.query.order_by(AdCampaign.id.desc()).all()
    ads_total_revenue = sum(ad.total_cost or 0 for ad in ad_campaigns if ad.stripe_session_id)
    ads_paid_count = sum(1 for ad in ad_campaigns if ad.stripe_session_id)
    ads_pending_count = sum(1 for ad in ad_campaigns if not ad.stripe_session_id and ad.status in ('Pendiente', 'Pendiente Pago'))
    ads_active_count = sum(1 for ad in ad_campaigns if ad.status == 'Activo')
    ads_by_advertiser = {}
    for ad in ad_campaigns:
        key = ad.advertiser_email or (f'user_{ad.advertiser_id}' if ad.advertiser_id else 'Admin')
        if key not in ads_by_advertiser:
            ads_by_advertiser[key] = {'count': 0, 'revenue': 0, 'impressions': 0, 'clicks': 0}
        ads_by_advertiser[key]['count'] += 1
        ads_by_advertiser[key]['revenue'] += (ad.total_cost or 0) if ad.stripe_session_id else 0
        ads_by_advertiser[key]['impressions'] += ad.impressions or 0
        ads_by_advertiser[key]['clicks'] += ad.clicks or 0
    ghost_result = db.session.query(
        db.func.count(db.func.distinct(UserEvent.ip_address))
    ).filter(
        UserEvent.event_category == 'IA',
        UserEvent.user_id.is_(None),
        UserEvent.ip_address.isnot(None),
        UserEvent.ip_address != ''
    ).scalar() or 0
    ghost_count = ghost_result
    orphan_count = UserEvent.query.filter(
        UserEvent.event_category == 'IA',
        UserEvent.user_id.is_(None)
    ).count()
    converted_count = UserEvent.query.filter_by(
        event_category='Lead', event_action='Documento Anónimo Vinculado'
    ).count()
    total_funnel = ghost_count + converted_count
    conversion_rate = round((converted_count / total_funnel * 100), 1) if total_funnel > 0 else 0
    security_events = UserEvent.query.filter(
        UserEvent.event_category == 'Seguridad'
    ).order_by(UserEvent.timestamp.desc()).limit(20).all()
    affiliate_commissions = AffiliateCommission.query.order_by(AffiliateCommission.created_at.desc()).limit(50).all()
    ad_issues = AdIssue.query.order_by(AdIssue.created_at.desc()).limit(50).all()
    broken_banners = []
    for ad in ad_campaigns:
        if ad.image_url and ad.status in ('Activo', 'Pendiente de Revisión'):
            img_path = ad.image_url.lstrip('/')
            if not os.path.isfile(img_path):
                broken_banners.append({"id": ad.id, "title": ad.title, "image_url": ad.image_url, "advertiser": ad.advertiser_email or f'user_{ad.advertiser_id}'})
    # --- Estadísticas de impacto para XPRIZE: docentes activos y participantes ---
    import json as _json_stats
    _real_courses = CourseSession.query.filter(CourseSession.master_doc.isnot(None)).all()
    stats_cursos_reales = len(_real_courses)
    stats_docentes_unicos = len(set(cs.user_id for cs in _real_courses if cs.user_id))
    stats_total_participantes = 0
    for _cs in _real_courses:
        try:
            if _cs.course_info_json:
                _ci_stats = _json_stats.loads(_cs.course_info_json)
                _num_p = _ci_stats.get('num_participantes')
                if _num_p:
                    stats_total_participantes += int(_num_p)
        except Exception:
            pass
    return render_template("admin.html",
                           stats_cursos_reales=stats_cursos_reales,
                           stats_docentes_unicos=stats_docentes_unicos,
                           stats_total_participantes=stats_total_participantes,
                           title="Panel de Administración",
                           ghost_count=ghost_count,
                           orphan_count=orphan_count,
                           converted_count=converted_count,
                           conversion_rate=conversion_rate,
                           security_events=security_events,
                           active_page="admin",
                           users=users,
                           all_leads=all_leads,
                           paid_leads=paid_leads,
                           comision_pct=comision_pct,
                           stripe_wh_masked=stripe_wh_masked,
                           stripe_wh_configured=bool(stripe_wh_effective),
                           stripe_wh_source=stripe_wh_source,
                           stripe_checkout_aff=stripe_checkout_aff,
                           cost_per_day_ads=cost_per_day_ads,
                           max_concurrent_ads=max_concurrent_ads,
                           max_concurrent_ai=max_concurrent_ai,
                           sat_count_24h=sat_count_24h,
                           sat_7d=sat_7d,
                           sat_suggestion=sat_suggestion,
                           sat_level=sat_level,
                           sat_peak_hour=sat_peak_hour,
                           sat_peak_count=sat_peak_count,
                           ad_campaigns=ad_campaigns,
                           ads_total_revenue=ads_total_revenue,
                           ads_paid_count=ads_paid_count,
                           ads_pending_count=ads_pending_count,
                           ads_active_count=ads_active_count,
                           ads_by_advertiser=ads_by_advertiser,
                           affiliate_commissions=affiliate_commissions,
                           referrer_map=referrer_map,
                           tier_labels=TIER_LABELS,
                           users_pagination=users_pagination,
                           total_users=total_users,
                           user_filters=user_filters,
                           user_filters_qs=user_filters_qs,
                           ad_issues=ad_issues,
                           pro_signal_counts=pro_signal_counts,
                           broken_banners=broken_banners)

@app.route("/api/admin/email_logs")
@login_required
def admin_email_logs():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    from datetime import timedelta
    status = (request.args.get('status') or '').strip()
    etype = (request.args.get('type') or '').strip()
    user_id_filter = request.args.get('user_id', type=int)
    days = request.args.get('days', default=30, type=int)
    limit = min(request.args.get('limit', default=200, type=int), 1000)
    q = EmailLog.query
    if status in ('pending', 'sent', 'failed'):
        q = q.filter(EmailLog.status == status)
    if etype:
        q = q.filter(EmailLog.email_type == etype)
    if user_id_filter:
        q = q.filter(EmailLog.user_id == user_id_filter)
    if days and days > 0:
        cutoff = datetime.utcnow() - timedelta(days=days)
        q = q.filter(EmailLog.attempted_at >= cutoff)
    rows = q.order_by(EmailLog.id.desc()).limit(limit).all()
    user_ids = list({r.user_id for r in rows if r.user_id})
    user_map = {}
    if user_ids:
        for u in User.query.filter(User.id.in_(user_ids)).all():
            user_map[u.id] = {'email': u.email, 'name': u.full_name, 'tier': u.tier}
    today_cutoff = datetime.utcnow() - timedelta(hours=24)
    failed_24h = EmailLog.query.filter(EmailLog.status == 'failed', EmailLog.attempted_at >= today_cutoff).count()
    pending_24h = EmailLog.query.filter(EmailLog.status == 'pending', EmailLog.attempted_at >= today_cutoff).count()
    sent_24h = EmailLog.query.filter(EmailLog.status == 'sent', EmailLog.attempted_at >= today_cutoff).count()
    items = []
    for r in rows:
        items.append({
            'id': r.id,
            'user_id': r.user_id,
            'user_info': user_map.get(r.user_id),
            'recipient': r.recipient_email,
            'sender': r.sender_email,
            'type': r.email_type,
            'subject': r.subject,
            'document_filename': r.document_filename,
            'status': r.status,
            'error_message': r.error_message,
            'smtp_response': r.smtp_response,
            'attempted_at': _fmt_cdmx(r.attempted_at, '%d/%m/%Y %H:%M'),
            'sent_at': _fmt_cdmx(r.sent_at, '%d/%m/%Y %H:%M'),
            'retry_count': r.retry_count or 0,
            'opened_at': _fmt_cdmx(getattr(r, 'opened_at', None), '%d/%m/%Y %H:%M'),
            'open_count': getattr(r, 'open_count', 0) or 0,
        })
    return jsonify({
        'items': items,
        'totals_24h': {'sent': sent_24h, 'failed': failed_24h, 'pending': pending_24h},
        'count': len(items),
    })

@app.route("/api/track/stripe_click", methods=["POST"])
@login_required
def api_track_stripe_click():
    try:
        data = request.get_json(silent=True) or {}
        plan = (data.get('plan') or '')[:80]
        href = (data.get('href') or '')[:300]
        track_event('Funnel', 'Stripe Click', user_id=current_user.id, extra_data={'plan': plan, 'href': href})
    except Exception as _e:
        logger.warning(f"stripe_click track error: {_e}")
    return ('', 204)

_pixel_recent_hits = {}
_pixel_lock = threading.Lock()

def _pixel_should_persist(log_id, ttl_seconds=60):
    """In-process throttle: si el mismo log_id ya se actualizó en los últimos
    `ttl_seconds`, no volvemos a tocar la DB. Mitiga DoS sobre el pixel."""
    now = _time.time()
    with _pixel_lock:
        last = _pixel_recent_hits.get(log_id, 0)
        if now - last < ttl_seconds:
            return False
        _pixel_recent_hits[log_id] = now
        if len(_pixel_recent_hits) > 5000:
            cutoff = now - ttl_seconds
            for k in [k for k, v in _pixel_recent_hits.items() if v < cutoff]:
                _pixel_recent_hits.pop(k, None)
    return True

@app.route("/e/o/<token>.gif")
@app.route("/e/o/<token>")
def email_open_pixel(token):
    resp = Response(reengagement.OPEN_PIXEL_GIF, mimetype='image/gif')
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    log_id = reengagement.parse_open_token(app, token)
    if log_id and _pixel_should_persist(log_id):
        def _persist(lid):
            try:
                with app.app_context(), db.engine.begin() as conn:
                    conn.execute(db.text(
                        "UPDATE email_log SET opened_at = COALESCE(opened_at, NOW()), "
                        "open_count = COALESCE(open_count, 0) + 1 WHERE id = :lid"
                    ), {"lid": lid})
            except Exception as _e:
                logger.warning(f"email_open_pixel: persist fail log {lid}: {_e}")
        threading.Thread(target=_persist, args=(log_id,), daemon=True).start()
    return resp

@app.route("/api/admin/reengagement/preview")
@login_required
def admin_reengagement_preview():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    data = reengagement.preview_all(db, User, UserEvent, UserPurchase, EmailLog, sample=8)
    return jsonify({"campaigns": data, "now_utc": datetime.utcnow().isoformat()})

@app.route("/api/admin/outreach/u16_u17/preview/<int:user_id>")
@login_required
def admin_outreach_u16_u17_preview(user_id):
    """Preview HTML del outreach personalizado a U16 o U17 (one-shot 2026-05-25)."""
    if not current_user.is_admin:
        return ("No autorizado", 403)
    import outreach_u16_u17 as ou
    data = ou.build_preview(user_id, app=app)
    if not data:
        return ("Usuario no es objetivo de este outreach (solo 16 o 17).", 404)
    # Render directo del HTML del correo con un header de metadatos arriba.
    header = (f'<div style="background:#1f2937;color:#fff;padding:10px 16px;'
              f'border-radius:8px;margin:0 auto 12px;max-width:620px;font-family:Arial;font-size:0.88em;">'
              f'<strong>Para:</strong> {data["first_name"]} &lt;{data["recipient_email"]}&gt; · '
              f'<strong>WhatsApp:</strong> {data["whatsapp_display"]}<br>'
              f'<strong>Asunto:</strong> {data["subject"]}</div>')
    return (f'<!doctype html><html><body style="margin:0;background:#f3f4f6;padding:14px 0;">'
            f'{header}{data["html"]}</body></html>')

@app.route("/api/admin/outreach/u16_u17/send", methods=["POST"])
@login_required
def admin_outreach_u16_u17_send():
    """Envía el outreach a U16, U17 o ambos. Requiere ?confirm=YES en body."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    payload = request.get_json(silent=True) or {}
    if (payload.get('confirm') or '').strip() != 'YES':
        return jsonify({"error": "Falta confirmación (envía confirm='YES')."}), 400
    which = (payload.get('which') or '').strip().lower()
    if which not in ('u16', 'u17', 'both'):
        return jsonify({"error": "Parámetro 'which' debe ser u16, u17 o both."}), 400
    import outreach_u16_u17 as ou
    from reengagement import _send_smtp
    targets = [16] if which == 'u16' else ([17] if which == 'u17' else [16, 17])
    results = []
    for uid in targets:
        try:
            r = ou.send(uid, app=app, send_smtp_fn=_send_smtp,
                        log_attempt_fn=_log_email_attempt,
                        log_result_fn=_log_email_result)
        except Exception as e:
            logger.exception(f"outreach_u16_u17.send failed for {uid}")
            r = {'sent': False, 'reason': f'exception: {e}'}
        r['user_id'] = uid
        results.append(r)
        # Telemetría manual para tener trazabilidad fuera del EmailLog.
        try:
            track_event('Outreach', 'send_u16_u17', user_id=current_user.id,
                        extra_data={'target_user_id': uid, 'sent': r.get('sent'),
                                    'reason': r.get('reason')})
        except Exception:
            pass
    return jsonify({'results': results, 'now_utc': datetime.utcnow().isoformat()})

@app.route("/api/admin/reengagement/preview_html/<campaign_key>")
@login_required
def admin_reengagement_preview_html(campaign_key):
    if not current_user.is_admin:
        return ("No autorizado", 403)
    if campaign_key not in reengagement.CAMPAIGNS:
        return ("Campaña desconocida", 404)
    data = reengagement.render_preview_html(app, campaign_key)
    if not data:
        return ("No se pudo renderizar", 500)
    variants_html = ''
    if data.get('variants') and len(data['variants']) > 1:
        items = ''.join(f'<li style="margin:3px 0;">{v}</li>' for v in data['variants'])
        variants_html = f'<div style="background:#fef3c7;border:1px solid #fde68a;padding:8px 12px;margin:0 0 12px;border-radius:6px;font-size:0.85em;"><strong>A/B subject rotation activa:</strong><ul style="margin:4px 0 0 18px;padding:0;">{items}</ul></div>'
    return (f'<!doctype html><html><head><meta charset="utf-8"><title>Preview · {campaign_key}</title></head>'
            f'<body style="margin:0;background:#f3f4f6;padding:20px;font-family:Arial,sans-serif;">'
            f'<div style="max-width:680px;margin:0 auto;">'
            f'<div style="background:#1f2937;color:#fff;padding:10px 16px;border-radius:8px 8px 0 0;font-size:0.9em;"><strong>Asunto:</strong> {data["subject"]}</div>'
            f'<div style="background:#fff;padding:18px;border-radius:0 0 8px 8px;">{variants_html}{data["html"]}</div>'
            f'</div></body></html>')

@app.route("/api/admin/reengagement/metrics")
@login_required
def admin_reengagement_metrics():
    """Mini-dashboard de re-engagement: open rate y totales por campaña en 7/30/90 días."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    from datetime import timedelta
    out = {}
    now = datetime.utcnow()
    for key in reengagement.CAMPAIGNS.keys():
        per_window = {}
        for label, days in (('7d', 7), ('30d', 30), ('90d', 90)):
            cutoff = now - timedelta(days=days)
            base = EmailLog.query.filter(EmailLog.email_type == key, EmailLog.attempted_at >= cutoff)
            sent = base.filter(EmailLog.status == 'sent').count()
            failed = base.filter(EmailLog.status == 'failed').count()
            opened = base.filter(EmailLog.status == 'sent', EmailLog.opened_at.isnot(None)).count()
            open_rate = round((opened / sent * 100), 1) if sent else 0.0
            per_window[label] = {'sent': sent, 'failed': failed, 'opened': opened, 'open_rate_pct': open_rate}
        # Recomendación automática simple.
        rec = ''
        s30 = per_window['30d']['sent']
        o30 = per_window['30d']['open_rate_pct']
        if s30 == 0:
            rec = 'Sin envíos en 30d — revisa segmentación o el gate enabled_from.'
        elif o30 < 8:
            rec = f'Open rate bajo ({o30}%) — prueba A/B subject o ajusta hora de envío.'
        elif o30 >= 25:
            rec = f'Excelente open rate ({o30}%) — considera escalar volumen.'
        else:
            rec = f'Open rate sano ({o30}%) — mantener.'
        out[key] = {'label': reengagement.CAMPAIGNS[key]['label'], 'windows': per_window, 'recommendation': rec}
    return jsonify({'campaigns': out, 'now_utc': now.isoformat()})

@app.route("/api/admin/reengagement/recent_sends")
@login_required
def admin_reengagement_recent_sends():
    """Últimos 100 envíos de campañas reengage_*/value_* con nombre, email, campaña y estado."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    keys = list(reengagement.CAMPAIGNS.keys())
    rows = EmailLog.query.filter(EmailLog.email_type.in_(keys)).order_by(EmailLog.id.desc()).limit(100).all()
    user_ids = list({r.user_id for r in rows if r.user_id})
    user_map = {}
    if user_ids:
        for u in User.query.filter(User.id.in_(user_ids)).all():
            user_map[u.id] = {'name': u.full_name or '', 'email': u.email, 'tier': u.tier}
    items = []
    for r in rows:
        info = user_map.get(r.user_id) or {}
        cfg = reengagement.CAMPAIGNS.get(r.email_type, {})
        items.append({
            'id': r.id,
            'user_id': r.user_id,
            'name': info.get('name', ''),
            'email': r.recipient_email or info.get('email', ''),
            'tier': info.get('tier', ''),
            'campaign_key': r.email_type,
            'campaign_label': cfg.get('label', r.email_type),
            'subject': r.subject,
            'status': r.status,
            'attempted_at': _fmt_cdmx(r.attempted_at, '%d/%m/%Y %H:%M'),
            'opened_at': _fmt_cdmx(getattr(r, 'opened_at', None), '%d/%m/%Y %H:%M'),
        })
    return jsonify({'items': items, 'count': len(items)})

_reengage_run_state = {'running': False, 'started_at': None, 'last_results': None, 'mode': None}
_reengage_run_lock = threading.Lock()

def _reengage_background_run(campaign_key, max_per):
    with app.app_context():
        try:
            if campaign_key:
                users = reengagement.find_eligible_users(db, User, UserEvent, UserPurchase, EmailLog, campaign_key, limit=max_per)
                sent = 0; skipped = 0; errors = []
                for u in users:
                    try:
                        r = reengagement.send_campaign_email(app, db, EmailLog, u, campaign_key, _log_email_attempt, _log_email_result)
                        if r.get('sent'): sent += 1
                        else: skipped += 1
                    except Exception as _e:
                        errors.append(f"user_id={u.id}: {_e}")
                results = [{'campaign_key': campaign_key, 'sent': sent, 'skipped': skipped, 'errors': errors[:5]}]
            else:
                results = reengagement.run_all_campaigns(app, db, User, UserEvent, UserPurchase, EmailLog, _log_email_attempt, _log_email_result, max_per_campaign=max_per)
            _reengage_run_state['last_results'] = results
            logger.info(f"Re-engagement manual run completado: {results}")
        except Exception as _e:
            logger.error(f"Re-engagement manual run error: {_e}")
            _reengage_run_state['last_results'] = [{'error': str(_e)}]
        finally:
            _reengage_run_state['running'] = False

@app.route("/api/admin/reengagement/run", methods=["POST"])
@login_required
def admin_reengagement_run():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    payload = request.get_json(silent=True) or {}
    campaign_key = (payload.get('campaign_key') or '').strip() or None
    max_per = int(payload.get('max_per_campaign') or 200)
    if campaign_key and campaign_key not in reengagement.CAMPAIGNS:
        return jsonify({"error": f"Campaña desconocida: {campaign_key}"}), 400
    with _reengage_run_lock:
        if _reengage_run_state['running']:
            return jsonify({"error": "Ya hay una corrida en curso", "started_at": _reengage_run_state['started_at']}), 409
        _reengage_run_state['running'] = True
        _reengage_run_state['started_at'] = datetime.utcnow().isoformat()
        _reengage_run_state['mode'] = campaign_key or 'ALL'
        _reengage_run_state['last_results'] = None
    threading.Thread(target=_reengage_background_run, args=(campaign_key, max_per), daemon=True).start()
    return jsonify({"queued": True, "mode": campaign_key or 'ALL', "started_at": _reengage_run_state['started_at'],
                    "message": "Corrida iniciada en background. Recarga 'Estado de corrida' en unos segundos."})

@app.route("/api/admin/reengagement/status")
@login_required
def admin_reengagement_status():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    return jsonify({
        'running': _reengage_run_state['running'],
        'started_at': _reengage_run_state['started_at'],
        'mode': _reengage_run_state['mode'],
        'last_results': _reengage_run_state['last_results'],
    })

@app.route("/api/admin/email_logs/<int:log_id>/retry", methods=["POST"])
@login_required
def admin_email_log_retry(log_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    log = db.session.get(EmailLog, log_id)
    if not log:
        return jsonify({"error": "Log no encontrado"}), 404
    if log.email_type != 'document':
        return jsonify({"error": f"El reintento automático solo está disponible para correos de tipo 'document'. Este log es '{log.email_type}'."}), 400
    if not log.document_filename:
        return jsonify({"error": "Este log no tiene nombre de archivo asociado"}), 400
    filepath = os.path.join("generated_docs", log.document_filename)
    try:
        log.retry_count = (log.retry_count or 0) + 1
        db.session.commit()
    except Exception:
        db.session.rollback()
    ok = _send_document_email(log.recipient_email, log.document_filename, filepath, user_id=log.user_id)
    return jsonify({'success': bool(ok), 'retry_count': log.retry_count})

@app.route("/api/admin/resend_documents/<int:user_id>", methods=["POST"])
@login_required
def admin_resend_documents(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    target = db.session.get(User, user_id)
    if not target:
        return jsonify({"error": "Usuario no encontrado"}), 404
    sf = StoredFile.query.filter_by(user_id=user_id).order_by(StoredFile.id.desc()).first()
    if not sf:
        return jsonify({"error": "Este usuario no tiene documentos generados"}), 404
    filepath = os.path.join("generated_docs", sf.filename)
    ok = _send_document_email(target.email, sf.filename, filepath, user_id=user_id)
    return jsonify({'success': bool(ok), 'recipient': target.email, 'filename': sf.filename})

def _admin_append_generated_file(user_id, cs_id, element_num, filename):
    """Agrega `filename` al campo generated_files del ChatHistory del usuario
    (elemento `element_num`, sesion de curso `cs_id`) para que el boton de descarga
    aparezca en su cuenta. Crea la fila si no existe. Commit con rollback en error."""
    import json as _jh
    row = None
    if cs_id:
        row = ChatHistory.query.filter_by(user_id=user_id, element_num=element_num, course_session_id=cs_id).first()
    if row is None:
        # Crear la fila para la sesion objetivo (no reutilizar otra sesion: el boton de
        # descarga del elemento se lee filtrando por la sesion activa del usuario).
        row = ChatHistory(user_id=user_id, element_num=element_num, messages_json='[]', course_session_id=cs_id)
        db.session.add(row)
    files = []
    if row.generated_files:
        try:
            files = _jh.loads(row.generated_files)
            if not isinstance(files, list):
                files = []
        except Exception:
            files = []
    if filename not in files:
        files.append(filename)
    row.generated_files = _jh.dumps(files[-20:], ensure_ascii=False)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        raise

@app.route("/api/admin/users/<int:user_id>/restore_carta", methods=["POST"])
@login_required
def admin_restore_carta(user_id):
    """Regenera la Carta Descriptiva de un usuario a partir de los datos guardados de
    su curso y la deja disponible para descarga EN SU CUENTA (persiste StoredFile +
    actualiza ChatHistory.generated_files). Para casos en que la generacion original
    fallo silenciosamente y el usuario quedo sin su entregable."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    target = db.session.get(User, user_id)
    if not target:
        return jsonify({"error": "Usuario no encontrado"}), 404

    cs = None
    if target.active_course_session_id:
        cs = CourseSession.query.get(target.active_course_session_id)
    if cs is None or not (cs.course_info_json or cs.master_doc or cs.topic):
        cs = (CourseSession.query.filter_by(user_id=user_id)
              .filter(db.or_(CourseSession.course_info_json.isnot(None),
                             CourseSession.master_doc.isnot(None)))
              .order_by(CourseSession.last_activity_at.desc().nullslast())
              .first())
    if cs is None:
        return jsonify({"error": "El usuario no tiene datos de curso para regenerar la Carta Descriptiva."}), 404

    raw_topic = (cs.topic or "").strip()
    first_line = raw_topic.split("\n")[0].strip() if raw_topic else ""
    course_name = _strip_course_label_prefix(first_line) or first_line or "[Tema por definir]"

    ci_text = ""
    if cs.course_info_json:
        try:
            import json as _jrc
            ci = _jrc.loads(cs.course_info_json)
            if isinstance(ci, dict):
                ci_text = "\n".join(str(v).strip() for v in ci.values() if v and str(v).strip())
        except Exception:
            ci_text = ""
    if not ci_text:
        ci_text = raw_topic or (cs.master_doc or "")[:2000]

    msg = ("Genera la carta descriptiva completa de mi curso con todos los campos "
           "requeridos por el EC0301: informacion general, objetivo general, objetivos "
           "particulares (cognitivo, psicomotor, afectivo), requerimientos, estrategias "
           "de evaluacion, y las 3 etapas (apertura, desarrollo, cierre) con tiempos, "
           "tecnicas y materiales.")
    try:
        resp = chat_with_ai(1, msg, [], {}, course_info_text=(ci_text or None), user_tier=target.tier)
        clean = _strip_conversational_preamble(resp)
        sections = _parse_response_to_sections(clean)
        if not sections:
            return jsonify({"error": "No se pudo estructurar la Carta Descriptiva (respuesta vacia)."}), 500
        filepath = generate_custom_docx('Carta_Descriptiva', sections, course_name=course_name, user_tier=target.tier)
        directory = os.path.dirname(filepath)
        prefixed = f"u{user_id}_{os.path.basename(filepath)}"
        new_path = os.path.join(directory, prefixed)
        try:
            os.rename(filepath, new_path)
            filepath = new_path
        except Exception:
            pass
        target_cs_id = target.active_course_session_id or cs.id
        _persist_file_to_db(filepath, user_id, course_session_id=target_cs_id)
        fname = os.path.basename(filepath)
        _admin_append_generated_file(user_id, target_cs_id, 1, fname)
        logger.info(f"ADMIN_RESTORE_CARTA admin={current_user.id} user={user_id} cs={cs.id} file={fname}")
        return jsonify({"success": True, "filename": fname, "recipient": target.email,
                        "course_name": course_name})
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.warning(f"ADMIN_RESTORE_CARTA fail user={user_id} err={type(e).__name__}: {e}")
        return jsonify({"error": "No se pudo regenerar la Carta Descriptiva. Intenta de nuevo."}), 500

@app.route("/api/admin/ad-issues/<int:issue_id>", methods=["POST"])
@login_required
def admin_update_issue(issue_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True) or {}
    issue = db.session.get(AdIssue, issue_id)
    if not issue:
        return jsonify({"error": "Incidencia no encontrada"}), 404
    new_status = data.get("status", "").strip()
    if new_status in ('Abierto', 'En Proceso', 'Resuelto'):
        issue.status = new_status
        db.session.commit()
    return jsonify({"success": True})

@app.route("/admin/update_config", methods=["POST"])
@login_required
def admin_update_config():
    if not current_user.is_admin:
        flash('Usuario no autorizado.', 'error')
        return redirect(url_for('home'))
    key = (request.form.get('key') or '').strip()
    value = (request.form.get('value') or '').strip()
    if not key:
        flash('Clave es obligatoria.', 'error')
        return redirect(url_for('admin_crm'))
    if not value and key not in ('STRIPE_CHECKOUT_URL_AFFILIATE',):
        flash('El valor es obligatorio.', 'error')
        return redirect(url_for('admin_crm'))
    if key == 'COMISION_PORCENTAJE':
        try:
            val_float = float(value)
            if val_float < 0 or val_float > 1:
                flash('El porcentaje debe estar entre 0 y 1.', 'error')
                return redirect(url_for('admin_crm'))
        except ValueError:
            flash('El valor debe ser un número decimal válido.', 'error')
            return redirect(url_for('admin_crm'))
    if key == 'COST_PER_DAY_ADS':
        try:
            val_float = float(value)
            if val_float <= 0:
                flash('La tarifa diaria debe ser un número positivo.', 'error')
                return redirect(url_for('admin_crm'))
        except ValueError:
            flash('La tarifa diaria debe ser un número válido.', 'error')
            return redirect(url_for('admin_crm'))
    if key == 'MAX_CONCURRENT_ADS':
        try:
            val_int = int(value)
            if val_int < 1:
                flash('El límite de anuncios simultáneos debe ser al menos 1.', 'error')
                return redirect(url_for('admin_crm'))
        except ValueError:
            flash('El límite debe ser un número entero válido.', 'error')
            return redirect(url_for('admin_crm'))
    if key == 'STRIPE_WEBHOOK_SECRET':
        if not value.startswith('whsec_'):
            flash('El Secreto de Firma debe comenzar con "whsec_".', 'error')
            return redirect(url_for('admin_crm'))
    if key == 'STRIPE_CHECKOUT_URL_AFFILIATE' and value and not value.startswith('https://'):
        flash('La URL de Stripe debe comenzar con "https://".', 'error')
        return redirect(url_for('admin_crm'))
    Config.set(key, value)
    display_val = ('••••••' + value[-6:]) if key == 'STRIPE_WEBHOOK_SECRET' and len(value) > 6 else value
    flash(f'Configuración "{key}" actualizada correctamente → {display_val}', 'success')
    return redirect(url_for('admin_crm'))

@app.route("/api/admin/toggle_tier", methods=["POST"])
@login_required
def toggle_tier():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True)
    if not data or not data.get("user_id"):
        return jsonify({"error": "Datos inválidos"}), 400
    try:
        uid = int(data["user_id"])
    except (ValueError, TypeError):
        return jsonify({"error": "ID inválido"}), 400
    user = db.session.get(User, uid)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    valid_tiers = ('FREE', 'PRO', 'PRO_5', 'PREMIUM')
    new_tier = (data.get("new_tier") or "").upper()
    if new_tier not in valid_tiers:
        return jsonify({"error": "Plan inválido. Valores permitidos: FREE, PRO, PRO_5, PREMIUM"}), 400
    notify = bool(data.get("notify", True))
    reason = (data.get("reason") or "").strip()[:500]
    prev_tier = user.tier
    prev_credits = user.pro_courses_remaining or 0
    if new_tier == 'PRO_5':
        target_tier = 'PRO'
        target_credits = 5
        applied_label = 'PRO_MULTICURSO'
    elif new_tier == 'PRO':
        target_tier = 'PRO'
        target_credits = max(user.pro_courses_remaining or 0, 1)
        applied_label = 'PRO_PROJECT'
    elif new_tier == 'PREMIUM':
        target_tier = 'PREMIUM'
        target_credits = user.pro_courses_remaining or 0
        applied_label = 'PREMIUM'
    else:
        target_tier = 'FREE'
        target_credits = user.pro_courses_remaining or 0
        applied_label = 'FREE'

    if prev_tier == target_tier and prev_credits == target_credits:
        return jsonify({
            "success": True, "new_tier": user.tier, "user_id": user.id,
            "credits": user.pro_courses_remaining or 0, "applied_as": applied_label,
            "no_change": True, "notified": False,
            "message": "Sin cambios: el usuario ya tiene ese plan y crédito."
        })

    user.tier = target_tier
    user.pro_courses_remaining = target_credits
    if applied_label != 'PREMIUM':
        user.pro_active_course = None
    db.session.commit()

    new_credits = user.pro_courses_remaining or 0
    credits_delta = new_credits - prev_credits

    if applied_label in ('PRO_PROJECT', 'PRO_MULTICURSO', 'PREMIUM') and credits_delta > 0:
        try:
            grant = UserPurchase(
                user_id=user.id,
                payment_type=applied_label,
                amount_mxn=0.0,
                credits_granted=credits_delta if applied_label != 'PREMIUM' else 0,
                stripe_session_id=None,
                source='admin_grant'
            )
            db.session.add(grant)
            db.session.commit()
            logger.info(f"Admin grant recorded: admin={current_user.id} target={user.id} type={applied_label} credits_delta={credits_delta}")
        except Exception as e:
            logger.error(f"Failed to record admin grant: {e}")
            db.session.rollback()
    elif applied_label == 'PREMIUM' and prev_tier != 'PREMIUM':
        try:
            grant = UserPurchase(
                user_id=user.id, payment_type='PREMIUM', amount_mxn=0.0,
                credits_granted=0, stripe_session_id=None, source='admin_grant'
            )
            db.session.add(grant)
            db.session.commit()
        except Exception as e:
            logger.error(f"Failed to record PREMIUM admin grant: {e}")
            db.session.rollback()

    notified = False
    if notify:
        try:
            _dispatch_welcome_email_async(user, prev_tier, user.tier, prev_credits, new_credits, applied_label, reason=reason or '')
            notified = True
        except Exception as e:
            logger.error(f"Failed to spawn tier-change email thread: {e}")

    track_event('Admin', 'Cambio de Tier', user_id=current_user.id, extra_data={
        'target_user_id': user.id,
        'target_email': user.email,
        'prev_tier': prev_tier,
        'prev_credits': prev_credits,
        'new_tier': user.tier,
        'new_credits': new_credits,
        'applied_as': applied_label,
        'notify_requested': notify,
        'notify_dispatched': notified,
        'reason': reason or None
    })
    return jsonify({
        "success": True, "new_tier": user.tier, "user_id": user.id,
        "credits": new_credits, "applied_as": applied_label,
        "notified": notified, "no_change": False
    })

@app.route("/api/admin/export_users")
@login_required
def export_users_csv():
    if not current_user.is_admin:
        abort(403)
    include_bounced = request.args.get('include_bounced') == '1'
    q = User.query.order_by(User.id)
    if not include_bounced:
        q = q.filter(User.email_bounced_at.is_(None))
    users = q.all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['id', 'email', 'full_name', 'first_name', 'apellido_paterno', 'apellido_materno', 'tier', 'whatsapp', 'marketing_consent'])
    for u in users:
        writer.writerow([u.id, u.email, u.full_name, u.first_name or '', u.apellido_paterno or '', u.apellido_materno or '', u.tier, u.whatsapp or '', 'Sí' if u.marketing_consent else 'No'])
    fname = 'pertinentia_usuarios_con_rebotados.csv' if include_bounced else 'pertinentia_usuarios.csv'
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={fname}'}
    )

@app.route("/api/admin/export_bounced")
@login_required
def export_bounced_csv():
    if not current_user.is_admin:
        abort(403)
    users = User.query.filter(User.email_bounced_at.isnot(None)).order_by(User.email_bounced_at.desc()).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['id', 'email', 'full_name', 'tier', 'bounce_type', 'bounced_at', 'note'])
    for u in users:
        writer.writerow([
            u.id, u.email, u.full_name, u.tier,
            u.email_bounce_type or '',
            u.email_bounced_at.strftime('%Y-%m-%d %H:%M') if u.email_bounced_at else '',
            u.email_bounce_note or ''
        ])
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment; filename=pertinentia_rebotados.csv'}
    )

@app.route("/api/admin/export_lead_segmentation")
@login_required
def export_lead_segmentation_csv():
    """CSV con segmentación automática de leads para campañas WhatsApp.
    Calculado SIEMPRE en vivo a partir de user_event, stored_file y user_purchases.
    """
    if not current_user.is_admin:
        abort(403)

    sql = text("""
        WITH ev AS (
            SELECT
                user_id,
                COUNT(*) FILTER (WHERE event_action = 'Generación IA Iniciada') AS n_gen,
                COUNT(*) FILTER (WHERE event_action = 'Descarga Exitosa') AS n_dl_ok,
                COUNT(*) FILTER (WHERE event_action = 'Descarga Bloqueada') AS n_dl_block,
                COUNT(*) FILTER (WHERE event_action = 'Vista Precios') AS n_view_precios,
                COUNT(*) FILTER (WHERE event_action = 'Vista Elemento') AS n_view_elem,
                COUNT(*) FILTER (WHERE event_action = 'Stripe Click') AS n_stripe_click,
                MAX(timestamp) AS last_event_ts,
                COUNT(*) AS n_total_events
            FROM user_event
            WHERE user_id IS NOT NULL
            GROUP BY user_id
        ),
        files AS (
            SELECT user_id, COUNT(*) AS n_files, MAX(created_at) AS last_file_ts
            FROM stored_file
            GROUP BY user_id
        ),
        purch AS (
            SELECT user_id, COUNT(*) AS n_purchases, MAX(created_at) AS last_purchase_ts
            FROM user_purchases
            GROUP BY user_id
        )
        SELECT
            u.id, u.email, u.full_name, u.whatsapp, u.tier,
            u.chat_usage_count,
            u.created_at,
            u.email_bounced_at IS NOT NULL AS bounced,
            u.marketing_consent,
            COALESCE(ev.n_gen, 0) AS n_gen,
            COALESCE(ev.n_dl_ok, 0) AS n_dl_ok,
            COALESCE(ev.n_dl_block, 0) AS n_dl_block,
            COALESCE(ev.n_view_precios, 0) AS n_view_precios,
            COALESCE(ev.n_view_elem, 0) AS n_view_elem,
            COALESCE(ev.n_stripe_click, 0) AS n_stripe_click,
            COALESCE(ev.n_total_events, 0) AS n_total_events,
            ev.last_event_ts,
            COALESCE(files.n_files, 0) AS n_files,
            files.last_file_ts,
            COALESCE(purch.n_purchases, 0) AS n_purchases,
            purch.last_purchase_ts
        FROM "user" u
        LEFT JOIN ev ON ev.user_id = u.id
        LEFT JOIN files ON files.user_id = u.id
        LEFT JOIN purch ON purch.user_id = u.id
        ORDER BY u.id
    """)
    rows = db.session.execute(sql).fetchall()

    PAID_TIERS = {'PRO', 'PREMIUM', 'MULTI', 'PRO_MULTI'}

    def _classify(r):
        _email_lc = (r.email or '').lower()
        _admin_lc = (ADMIN_EMAIL or '').lower()
        is_internal = bool(
            _email_lc and (
                _email_lc.endswith('@pertinentia.com')
                or (_admin_lc and _email_lc == _admin_lc)
            )
        )
        if is_internal:
            return ('EXCLUIR', 'Cuenta interna/admin')
        if r.bounced:
            return ('EXCLUIR', 'Email rebotado')
        if not r.whatsapp or not r.whatsapp.strip():
            return ('SIN_WHATSAPP', 'No tiene WhatsApp registrado')

        tier = (r.tier or 'FREE').upper()
        n_files = r.n_files
        n_gen = r.n_gen
        n_view_precios = r.n_view_precios
        n_view_elem = r.n_view_elem

        if tier in PAID_TIERS:
            if n_files >= 3:
                return ('A_PAGO_ACTIVO', f'Cliente {tier} usando ({n_files} docs, {r.n_dl_ok} descargas)')
            return ('B_PAGO_SUBUSO', f'Cliente {tier} subutilizado ({n_files} docs)')

        if n_gen >= 1 and n_files == 0:
            return ('E_BUG_GENERO_SIN_ARCHIVO', f'{n_gen} intentos de generación, 0 archivos — posible falla técnica')
        if n_files >= 1 and (n_view_precios >= 1 or r.n_stripe_click >= 1):
            return ('C_HOT_LEAD', f'{n_files} doc(s) + {n_view_precios} vista(s) precios + {r.n_stripe_click} clic(s) Stripe')
        if n_files >= 1:
            return ('D_PROBO_NO_VOLVIO', f'{n_files} doc(s) descargado(s), no volvió a precios')
        if n_view_elem >= 1:
            return ('F_ENTRO_SIN_GENERAR', f'Entró al módulo {n_view_elem} vez(es), no escribió')
        return ('F_FRIO_TOTAL', 'Solo se registró, no entró al módulo')

    SEGMENT_LABEL = {
        'A_PAGO_ACTIVO': 'A · Cliente activo',
        'B_PAGO_SUBUSO': 'B · Pagó pero subusó',
        'C_HOT_LEAD': 'C · Hot lead',
        'D_PROBO_NO_VOLVIO': 'D · Probó y no volvió',
        'E_BUG_GENERO_SIN_ARCHIVO': 'E · Falla técnica (recuperar)',
        'F_ENTRO_SIN_GENERAR': 'F · Entró sin generar',
        'F_FRIO_TOTAL': 'F · Frío total',
        'EXCLUIR': 'EXCLUIR',
        'SIN_WHATSAPP': 'Sin WhatsApp',
    }
    PRIORITY = {
        'B_PAGO_SUBUSO': 1, 'E_BUG_GENERO_SIN_ARCHIVO': 2,
        'C_HOT_LEAD': 3, 'A_PAGO_ACTIVO': 4,
        'D_PROBO_NO_VOLVIO': 5, 'F_ENTRO_SIN_GENERAR': 6, 'F_FRIO_TOTAL': 7,
        'SIN_WHATSAPP': 8, 'EXCLUIR': 9,
    }

    now = datetime.utcnow()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        'id', 'segmento', 'segmento_label', 'motivo',
        'nombre', 'email', 'whatsapp',
        'tier', 'created_at', 'dias_desde_registro',
        'last_activity', 'dias_desde_ult_actividad',
        'n_archivos', 'n_generaciones', 'n_descargas_ok', 'n_descargas_bloqueadas',
        'n_vistas_precios', 'n_vistas_elemento', 'n_stripe_click',
        'n_compras', 'marketing_consent',
        'snapshot_generado_en'
    ])

    classified = []
    for r in rows:
        seg, motivo = _classify(r)
        classified.append((seg, motivo, r))

    classified.sort(key=lambda x: (PRIORITY.get(x[0], 99), x[2].id))

    snapshot_iso = now.strftime('%Y-%m-%d %H:%M:%S UTC')
    for seg, motivo, r in classified:
        last_ev = r.last_event_ts or r.last_file_ts or r.last_purchase_ts
        days_reg = (now - r.created_at).days if r.created_at else ''
        days_act = (now - last_ev).days if last_ev else ''
        writer.writerow([
            r.id, seg, SEGMENT_LABEL.get(seg, seg), motivo,
            r.full_name or '', r.email or '', r.whatsapp or '',
            r.tier or '',
            r.created_at.strftime('%Y-%m-%d %H:%M') if r.created_at else '',
            days_reg,
            last_ev.strftime('%Y-%m-%d %H:%M') if last_ev else '',
            days_act,
            r.n_files, r.n_gen, r.n_dl_ok, r.n_dl_block,
            r.n_view_precios, r.n_view_elem, r.n_stripe_click,
            r.n_purchases, 'Sí' if r.marketing_consent else 'No',
            snapshot_iso
        ])

    fname = f'pertinentia_segmentacion_leads_{now.strftime("%Y%m%d_%H%M")}.csv'
    return Response(
        output.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename={fname}'}
    )

@app.route("/api/admin/users/<int:user_id>/bounce", methods=["POST"])
@login_required
def toggle_user_bounce(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    payload = request.get_json(silent=True) or {}
    action = (payload.get("action") or "").lower()
    bounce_type = (payload.get("bounce_type") or "hard").lower()
    note = (payload.get("note") or "")[:255]
    if bounce_type not in ("hard", "soft"):
        bounce_type = "hard"
    if action == "clear":
        user.email_bounced_at = None
        user.email_bounce_type = None
        user.email_bounce_note = None
    else:
        user.email_bounced_at = datetime.utcnow()
        user.email_bounce_type = bounce_type
        user.email_bounce_note = note or None
    db.session.commit()
    try:
        track_event("marketing", "admin_mark_bounce", user_id=user.id,
                    extra_data={"action": action or "set", "type": bounce_type, "by_admin_id": current_user.id})
    except Exception:
        pass
    return jsonify({
        "success": True,
        "bounced": user.email_bounced_at is not None,
        "bounce_type": user.email_bounce_type,
        "bounced_at": user.email_bounced_at.isoformat() if user.email_bounced_at else None,
        "note": user.email_bounce_note
    })

@app.route("/api/admin/users/<int:user_id>/contacto", methods=["POST"])
@login_required
def update_user_contacto(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    payload = request.get_json(silent=True) or {}
    action = (payload.get("action") or "set").lower()
    if action == "clear":
        user.ultimo_contacto_at = None
        user.notas_contacto = None
    else:
        notas_raw = payload.get("notas")
        notas = (notas_raw or "").strip()[:4000] if isinstance(notas_raw, str) else None
        user.ultimo_contacto_at = datetime.utcnow()
        if notas is not None:
            user.notas_contacto = notas or None
    db.session.commit()
    try:
        track_event("crm", "admin_mark_contacto", user_id=user.id,
                    extra_data={"action": action, "by_admin_id": current_user.id})
    except Exception:
        pass
    return jsonify({
        "success": True,
        "ultimo_contacto_at": user.ultimo_contacto_at.isoformat() if user.ultimo_contacto_at else None,
        "notas_contacto": user.notas_contacto or ""
    })

CRM_ALLOWED_CHANNELS = {'whatsapp', 'email', 'llamada', 'sms', 'reunion', 'otro'}
CRM_ALLOWED_DIRECTIONS = {'enviado', 'recibido', 'nota'}

@app.route("/api/admin/users/<int:user_id>/crm_log", methods=["POST"])
@login_required
def admin_crm_log_interaction(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    payload = request.get_json(silent=True) or {}
    channel = (payload.get("channel") or "").strip().lower()
    direction = (payload.get("direction") or "").strip().lower()
    text = (payload.get("text") or "").strip()
    extra = payload.get("extra") if isinstance(payload.get("extra"), dict) else {}
    if channel not in CRM_ALLOWED_CHANNELS:
        return jsonify({"error": f"Canal no permitido. Usa: {', '.join(sorted(CRM_ALLOWED_CHANNELS))}"}), 400
    if direction not in CRM_ALLOWED_DIRECTIONS:
        return jsonify({"error": f"Direccion no permitida. Usa: {', '.join(sorted(CRM_ALLOWED_DIRECTIONS))}"}), 400
    if not text:
        return jsonify({"error": "El texto de la interaccion es obligatorio"}), 400
    text = text[:5000]
    action_label = f"{channel.capitalize()} {direction.capitalize()}"
    metadata = {
        "canal": channel,
        "direccion": direction,
        "texto": text,
        "registrado_por_admin_id": current_user.id,
        "registrado_por_email": current_user.email,
    }
    if extra:
        metadata["extra"] = extra
    try:
        track_event('CRM', action_label, user_id=user.id, extra_data=metadata)
        user.ultimo_contacto_at = datetime.utcnow()
        existing_notas = (user.notas_contacto or '').strip()
        snippet = text[:140] + ('…' if len(text) > 140 else '')
        new_line = f"[{datetime.utcnow().strftime('%Y-%m-%d %H:%M')}] {action_label}: {snippet}"
        user.notas_contacto = (new_line + ("\n" + existing_notas if existing_notas else ""))[:4000]
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.warning(f"CRM_LOG_FAIL user_id={user_id} error={type(e).__name__}: {e}")
        return jsonify({"error": "No se pudo registrar"}), 500
    return jsonify({"success": True, "action": action_label})

@app.route("/api/admin/users/<int:user_id>/crm_history", methods=["GET"])
@login_required
def admin_crm_history(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    rows = UserEvent.query.filter_by(user_id=user_id, event_category='CRM').order_by(UserEvent.timestamp.desc()).limit(100).all()
    out = []
    for r in rows:
        try:
            meta = json.loads(r.metadata_json) if r.metadata_json else {}
        except Exception:
            meta = {}
        out.append({
            "id": r.id,
            "action": r.event_action,
            "timestamp": r.timestamp.isoformat() if r.timestamp else None,
            "canal": meta.get("canal"),
            "direccion": meta.get("direccion"),
            "texto": meta.get("texto", ""),
            "extra": meta.get("extra", {}),
        })
    return jsonify({"user_id": user_id, "email": user.email, "interactions": out, "count": len(out)})

def _seed_crm_magda_one_shot():
    try:
        with app.app_context():
            existing = UserEvent.query.filter_by(user_id=45, event_category='CRM').first()
            if existing:
                return
            magda = db.session.get(User, 45)
            if not magda:
                return
            sent_meta = {
                "canal": "whatsapp",
                "direccion": "enviado",
                "texto": "Hola Magda, te saluda Arturo de PertinentIA. Notamos que creaste tu cuenta y trabajaste en una carta descriptiva. Queremos estar seguros: ¿pudiste descargar tu documento? ¿algo no quedó claro? Si te late, te puedo enviar por aquí mismo el documento que generaste para que le des una revisada con calma. Sin compromiso. ¿Te lo mando?",
                "registrado_por_admin_id": 1,
                "registrado_por_email": "seed",
                "extra": {"contexto": "Seguimiento Día del Maestro 2026", "fecha_real": "2026-05-12 14:39 CDMX"}
            }
            recv_meta = {
                "canal": "whatsapp",
                "direccion": "recibido",
                "texto": "Si. No me permitió descargarlo y quería checar si me conviene, porque soy Evaluadora independiente",
                "registrado_por_admin_id": 1,
                "registrado_por_email": "seed",
                "extra": {
                    "perfil_revelado": "Evaluadora independiente EC0301/EC0217.01",
                    "intent_compra": "alto",
                    "bug_reportado": "no_pudo_descargar",
                    "fecha_real": "2026-05-12 14:50 CDMX",
                    "siguiente_paso": "Enviar archivo .docx + ofrecer cupon MAESTRO15 con vencimiento domingo"
                }
            }
            ev1 = UserEvent(
                user_id=45, event_category='CRM', event_action='Whatsapp Enviado',
                url='whatsapp://crm-seed', metadata_json=json.dumps(sent_meta, ensure_ascii=False),
                ip_address='0.0.0.0', user_agent='CRM-Seed',
                timestamp=datetime(2026, 5, 12, 20, 39, 0)
            )
            ev2 = UserEvent(
                user_id=45, event_category='CRM', event_action='Whatsapp Recibido',
                url='whatsapp://crm-seed', metadata_json=json.dumps(recv_meta, ensure_ascii=False),
                ip_address='0.0.0.0', user_agent='CRM-Seed',
                timestamp=datetime(2026, 5, 12, 20, 50, 0)
            )
            db.session.add_all([ev1, ev2])
            magda.ultimo_contacto_at = datetime(2026, 5, 12, 20, 50, 0)
            existing_notas = (magda.notas_contacto or '').strip()
            seed_note = "[2026-05-12 14:50 CDMX] WhatsApp recibido: 'No me permitió descargarlo, soy Evaluadora independiente, quiero revisar si me conviene' — INTENT ALTO, perfil PREMIUM ideal, bug descarga reportado."
            magda.notas_contacto = (seed_note + ("\n" + existing_notas if existing_notas else ""))[:4000]
            db.session.commit()
            logger.info(f"CRM_SEED_MAGDA_OK user_id=45 events_inserted=2")
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.warning(f"CRM_SEED_MAGDA_SKIP error={type(e).__name__}: {e}")

try:
    _seed_crm_magda_one_shot()
except Exception:
    pass

FUNNEL_ALLOWED_ACTIONS = {
    'click_descargar', 'click_personalizar', 'descarga_fallida'
}

@app.route("/api/funnel/event", methods=["POST"])
def funnel_event():
    data = request.get_json(silent=True) or {}
    action = (data.get("action") or "").strip()
    if action not in FUNNEL_ALLOWED_ACTIONS:
        return jsonify({"error": "action no permitida"}), 400
    extras = data.get("extra_data") or {}
    if not isinstance(extras, dict):
        extras = {}
    safe_extras = {}
    for k, v in list(extras.items())[:10]:
        sk = str(k)[:40]
        if isinstance(v, (str, int, float, bool)) or v is None:
            sv = v if not isinstance(v, str) else v[:200]
            safe_extras[sk] = sv
    safe_extras['tier'] = current_user.tier if current_user.is_authenticated else 'ANON'
    uid = current_user.id if current_user.is_authenticated else None
    try:
        track_event('Funnel', action, user_id=uid, extra_data=safe_extras)
    except Exception:
        return jsonify({"error": "tracking_failed"}), 500
    return jsonify({"success": True})

@app.route("/api/admin/users/<int:user_id>", methods=["DELETE"])
@login_required
def delete_user(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    if user.email == ADMIN_EMAIL:
        return jsonify({"error": "No se puede eliminar al administrador"}), 400
    try:
        UserEvent.query.filter_by(user_id=user_id).delete()
        StoredFile.query.filter_by(user_id=user_id).delete()
        EvaluationProcess.query.filter_by(user_id=user_id).delete()
        Candidato.query.filter_by(ce_user_id=user_id).delete()
        CEProfile.query.filter_by(user_id=user_id).delete()
        try: ChatHistory.query.filter_by(user_id=user_id).delete()
        except Exception: pass
        try: ChatSpec.query.filter_by(user_id=user_id).delete()
        except Exception: pass
        try: EmailLog.query.filter_by(user_id=user_id).update({EmailLog.user_id: None})
        except Exception: pass
        try: AffiliateVideoView.query.filter_by(sponsor_id=user_id).delete()
        except Exception: pass
        try: UserPurchase.query.filter_by(user_id=user_id).delete()
        except Exception: pass
        try: AffiliateLead.query.filter_by(sponsor_id=user_id).delete()
        except Exception: pass
        AffiliateCommission.query.filter_by(sponsor_id=user_id).delete()
        AffiliateCommission.query.filter_by(user_id=user_id).update({AffiliateCommission.user_id: None})
        AdIssue.query.filter_by(user_id=user_id).delete()
        BetaFeedback.query.filter_by(user_id=user_id).update({BetaFeedback.user_id: None})
        BetaMetric.query.filter_by(user_id=user_id).update({BetaMetric.user_id: None})
        AdCampaign.query.filter(AdCampaign.advertiser_id == user_id).update({AdCampaign.advertiser_id: None})
        user_prefix = f"u{user_id}_"
        docs_dir = "generated_docs"
        if os.path.isdir(docs_dir):
            for fname in os.listdir(docs_dir):
                if fname.startswith(user_prefix):
                    try:
                        os.remove(os.path.join(docs_dir, fname))
                    except OSError:
                        pass
        db.session.delete(user)
        db.session.commit()
        logger.info(f"Admin deleted user {user_id} ({user.email})")
        return jsonify({"success": True, "deleted_id": user_id})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting user {user_id}: {e}")
        return jsonify({"error": f"Error al eliminar: {str(e)}"}), 500

@app.route("/api/admin/users/reset-sequence", methods=["POST"])
@login_required
def admin_reset_user_sequence():
    """Reajusta user_id_seq al MAX(id) actual para que el siguiente registro
    use MAX+1 (rellena 'huecos' al inicio cuando los IDs altos fueron borrados).
    NOTA: NO reutiliza IDs que estén entre el min y el max ocupado — solo
    ajusta el contador hacia abajo si la sequence está adelantada."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    try:
        from sqlalchemy import text  # noqa: F401  (already imported globally)
        row = db.session.execute(text('SELECT COALESCE(MAX(id), 0) AS max_id FROM "user"')).fetchone()
        max_id = int(row[0]) if row else 0
        seq_row = db.session.execute(text("SELECT last_value, is_called FROM user_id_seq")).fetchone()
        prev_last = int(seq_row[0]) if seq_row else None
        next_val = max_id + 1
        db.session.execute(text("SELECT setval('user_id_seq', :v, true)"), {"v": max(max_id, 1)})
        db.session.commit()
        logger.info(f"ADMIN_RESET_USER_SEQ admin={current_user.id} max_id={max_id} prev_last={prev_last} next_will_be={next_val}")
        return jsonify({
            "success": True,
            "max_id_actual": max_id,
            "sequence_anterior": prev_last,
            "siguiente_id_sera": next_val
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"reset_user_sequence error: {e}")
        return jsonify({"error": f"Error: {str(e)}"}), 500

@app.route("/api/admin/users/<int:user_id>/marketing_consent", methods=["POST"])
@login_required
def toggle_marketing_consent(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    from datetime import datetime
    user.marketing_consent = not user.marketing_consent
    user.marketing_consent_source = 'ADMIN'
    if not user.marketing_consent:
        if user.mkt_unsubscribed_at is None:
            user.mkt_unsubscribed_at = datetime.utcnow()
    else:
        user.mkt_unsubscribed_at = None
    db.session.commit()
    try:
        track_event(
            "marketing",
            "admin_toggle_consent",
            user_id=user.id,
            extra_data={"new": bool(user.marketing_consent), "by_admin_id": current_user.id}
        )
    except Exception:
        pass
    return jsonify({
        "success": True,
        "marketing_consent": user.marketing_consent,
        "marketing_consent_source": user.marketing_consent_source,
        "user_id": user.id,
        "unsubscribed_at": user.mkt_unsubscribed_at.isoformat() if user.mkt_unsubscribed_at else None
    })

@app.route("/api/evaluator/accept_normative", methods=["POST"])
@login_required
def accept_normative():
    if not current_user.is_premium and not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    current_user.normative_agreement_accepted = True
    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/evaluator/download_iec")
@login_required
def download_iec():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    iec_path = os.path.join("normatividad", "Instrumento Evaluación EC0301.pdf")
    if not os.path.isfile(iec_path):
        return jsonify({"error": "Archivo IEC no encontrado"}), 404
    return send_file(iec_path, as_attachment=True, download_name="Instrumento_Evaluacion_EC0301.pdf",
                     mimetype="application/pdf")

@app.route("/api/evaluator/generar_plan", methods=["POST"])
@login_required
def generar_plan_evaluacion_route():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True)
    if not data or not data.get("candidato_id"):
        return jsonify({"error": "Datos inválidos"}), 400
    try:
        cid = int(data["candidato_id"])
    except (ValueError, TypeError):
        return jsonify({"error": "ID de candidato inválido"}), 400
    cand = Candidato.query.filter_by(id=cid, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    if not ce_profile or not ce_profile.ce_name or not ce_profile.ce_key:
        return jsonify({"error": "Por favor, completa la configuración de tu Centro para incluir tus datos en el documento."}), 400
    nombre_full = f"{cand.nombre_completo} {cand.apellidos}".strip()
    datos_candidato = {
        "nombre": nombre_full,
        "correo": "N/A",
        "whatsapp": "N/A"
    }
    datos_evaluador = {
        "nombre": ce_profile.evaluator_name if ce_profile.evaluator_name else current_user.full_name,
        "centro": ce_profile.ce_name,
        "clave_conocer": ce_profile.ce_key
    }
    fechas = {
        "fecha_gabinete": data.get("fecha_gabinete", "Por definir"),
        "fecha_campo": data.get("fecha_campo", "Por definir"),
        "fecha_emision": data.get("fecha_emision", "Por definir")
    }
    try:
        logo = ce_profile.logo_path if ce_profile else None
        plan_data, filled_doc = ai_generate_plan_evaluacion(datos_candidato, datos_evaluador, fechas, logo_path=logo)
        buffer = io.BytesIO()
        filled_doc.save(buffer)
        buffer.seek(0)
        safe_name = re.sub(r'[^\w\s-]', '', nombre_full).strip().replace(' ', '_')
        filename = f"Plan_Evaluacion_{safe_name}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        app.logger.error(f"Error generando plan de evaluación: {str(e)}")
        return jsonify({"error": "Error interno al generar el plan de evaluación"}), 500

@app.route("/api/evaluator/generar_portadas", methods=["POST"])
@login_required
def generar_portadas_route():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True)
    if not data or not data.get("candidato_id"):
        return jsonify({"error": "Datos inválidos"}), 400
    try:
        cid = int(data["candidato_id"])
    except (ValueError, TypeError):
        return jsonify({"error": "ID de candidato inválido"}), 400
    cand = Candidato.query.filter_by(id=cid, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    nombre_full = f"{cand.nombre_completo} {cand.apellidos}".strip()
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    if not ce_profile or not ce_profile.ce_name or not ce_profile.ce_key:
        return jsonify({"error": "Por favor, completa la configuración de tu Centro para incluir tus datos en el documento."}), 400
    datos_candidato = {"nombre": nombre_full, "correo": "N/A"}
    datos_evaluador = {
        "nombre": ce_profile.evaluator_name if ce_profile.evaluator_name else current_user.full_name,
        "clave_conocer": ce_profile.ce_key
    }
    try:
        logo = ce_profile.logo_path if ce_profile else None
        filled_doc = fill_portada_template(datos_candidato, datos_evaluador, logo_path=logo)
        buffer = io.BytesIO()
        filled_doc.save(buffer)
        buffer.seek(0)
        safe_name = re.sub(r'[^\w\s-]', '', nombre_full).strip().replace(' ', '_')
        return send_file(buffer, as_attachment=True, download_name=f"Portadas_{safe_name}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        app.logger.error(f"Error generando portadas: {str(e)}")
        return jsonify({"error": "Error interno al generar las portadas"}), 500

@app.route("/api/evaluator/generar_servicio", methods=["POST"])
@login_required
def generar_servicio_route():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True)
    if not data or not data.get("candidato_id"):
        return jsonify({"error": "Datos inválidos"}), 400
    try:
        cid = int(data["candidato_id"])
    except (ValueError, TypeError):
        return jsonify({"error": "ID de candidato inválido"}), 400
    cand = Candidato.query.filter_by(id=cid, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    nombre_full = f"{cand.nombre_completo} {cand.apellidos}".strip()
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    if not ce_profile or not ce_profile.ce_name or not ce_profile.ce_key:
        return jsonify({"error": "Por favor, completa la configuración de tu Centro para incluir tus datos en el documento."}), 400
    datos_candidato = {"nombre": nombre_full}
    fechas = {
        "fecha_emision": data.get("fecha_emision", ""),
        "fecha_campo": data.get("fecha_campo", "")
    }
    try:
        logo = ce_profile.logo_path if ce_profile else None
        filled_doc = fill_servicio_template(datos_candidato, fechas, logo_path=logo)
        buffer = io.BytesIO()
        filled_doc.save(buffer)
        buffer.seek(0)
        safe_name = re.sub(r'[^\w\s-]', '', nombre_full).strip().replace(' ', '_')
        return send_file(buffer, as_attachment=True, download_name=f"Servicio_Usuarios_{safe_name}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        app.logger.error(f"Error generando servicio a usuarios: {str(e)}")
        return jsonify({"error": "Error interno al generar formato de servicio"}), 500

@app.route("/api/evaluator/generar_encuesta", methods=["POST"])
@login_required
def generar_encuesta_route():
    if not current_user.is_premium:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True)
    if not data or not data.get("candidato_id"):
        return jsonify({"error": "Datos inválidos"}), 400
    try:
        cid = int(data["candidato_id"])
    except (ValueError, TypeError):
        return jsonify({"error": "ID de candidato inválido"}), 400
    cand = Candidato.query.filter_by(id=cid, ce_user_id=current_user.id).first()
    if not cand:
        return jsonify({"error": "Candidato no encontrado"}), 404
    nombre_full = f"{cand.nombre_completo} {cand.apellidos}".strip()
    ce_profile = CEProfile.query.filter_by(user_id=current_user.id).first()
    if not ce_profile or not ce_profile.ce_name or not ce_profile.ce_key:
        return jsonify({"error": "Por favor, completa la configuración de tu Centro para incluir tus datos en el documento."}), 400
    datos_candidato = {"nombre": nombre_full}
    fechas = {
        "fecha_emision": data.get("fecha_emision", ""),
        "fecha_campo": data.get("fecha_campo", "")
    }
    try:
        logo = ce_profile.logo_path if ce_profile else None
        filled_doc = fill_encuesta_template(datos_candidato, fechas, logo_path=logo)
        buffer = io.BytesIO()
        filled_doc.save(buffer)
        buffer.seek(0)
        safe_name = re.sub(r'[^\w\s-]', '', nombre_full).strip().replace(' ', '_')
        return send_file(buffer, as_attachment=True, download_name=f"Encuesta_Satisfaccion_{safe_name}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        app.logger.error(f"Error generando encuesta: {str(e)}")
        return jsonify({"error": "Error interno al generar la encuesta"}), 500

@app.route("/api/admin/users/<int:user_id>/reset_course", methods=["POST"])
@login_required
def admin_reset_user_course(user_id):
    """Admin-only: dispara reset_active_course_state() para el usuario indicado.

    Mismo helper que usa el webhook de Stripe y el botón "Iniciar Nuevo Curso"
    del usuario. Útil para soporte: cuando un cliente reporta que su curso
    sigue contaminado por un demo o borrador previo y no quiere/puede hacerlo
    él mismo. Para un usuario PRO con candado puesto DEVUELVE 1 crédito (el que se
    consumió por el curso que se descarta), para rehabilitar la generación. El
    snapshot-guard contempla este cambio intencional.
    """
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    try:
        summary = reset_active_course_state(user.id, reason=f'admin_manual_by_{current_user.id}', restore_pro_credit=True)
    except Exception as _e:
        logger.exception(f'[admin_reset_user_course] failed admin={current_user.id} target={user_id}')
        return jsonify({"error": f"Excepción: {type(_e).__name__}: {_e}"}), 500
    if not summary or not summary.get('ok'):
        return jsonify({"error": "Reset no completado",
                        "details": (summary or {}).get('details', [])}), 500
    try:
        track_event('Admin', 'reset_user_course', user_id=current_user.id,
                    extra_data={'target_user_id': user.id,
                                'new_cs_id': summary.get('new_cs_id'),
                                'purged_demo': summary.get('purged_demo'),
                                'credit_snapshot_match': summary.get('credit_snapshot_match'),
                                'pro_credit_restored': summary.get('pro_credit_restored', False)})
    except Exception:
        pass
    try:
        db.session.refresh(user)
    except Exception:
        pass
    return jsonify({"success": True,
                    "new_course_session_id": summary.get('new_cs_id'),
                    "new_course_session_num": summary.get('new_cs_num'),
                    "purged_demo": summary.get('purged_demo'),
                    "credit_snapshot_match": summary.get('credit_snapshot_match'),
                    "pro_credit_restored": summary.get('pro_credit_restored', False),
                    "pro_courses_remaining": (user.pro_courses_remaining or 0)})


@app.route("/api/admin/users/<int:user_id>/reset_password", methods=["POST"])
@login_required
def reset_user_password(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    user.password_hash = generate_password_hash('Temporal123!')
    db.session.commit()
    return jsonify({"success": True, "message": "Contraseña reseteada a Temporal123!"})

@app.route("/api/admin/users/create", methods=["POST"])
@login_required
def admin_create_user():
    """Admin crea una cuenta nueva. A diferencia del registro público, permite
    correo (y WhatsApp) repetidos: varias cuentas DISTINTAS pueden compartir el
    mismo correo/WhatsApp. Regla anti-confusión: cuando el correo ya existe, la
    combinación nombre+apellido paterno+apellido materno y la contraseña deben ser
    DISTINTAS de las cuentas existentes con ese correo (la contraseña desambigua
    el inicio de sesión). NO inicia sesión como la nueva cuenta (preserva la sesión
    del admin)."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    data = request.get_json(silent=True) or {}
    first_name = (data.get('first_name') or '').strip()[:80]
    apellido_paterno = (data.get('apellido_paterno') or '').strip()[:80]
    apellido_materno = (data.get('apellido_materno') or '').strip()[:80]
    email = (data.get('email') or '').strip().lower()
    whatsapp = (data.get('whatsapp') or '').strip()[:20]
    password = data.get('password') or ''
    tier = (data.get('tier') or 'FREE').strip().upper()
    if tier not in ('FREE', 'PRO', 'PREMIUM'):
        tier = 'FREE'

    if not first_name or not apellido_paterno or not email or not whatsapp or not password:
        return jsonify({"error": "Nombre, apellido paterno, correo, WhatsApp y contraseña son obligatorios."}), 400
    if len(password) < 6:
        return jsonify({"error": "La contraseña debe tener al menos 6 caracteres."}), 400
    if '@' not in email or '.' not in email.split('@')[-1]:
        return jsonify({"error": "Correo electrónico inválido."}), 400
    if email == ADMIN_EMAIL:
        return jsonify({"error": "Este correo está reservado para el administrador."}), 400

    full_name = " ".join(p for p in [first_name, apellido_paterno, apellido_materno] if p).strip()[:150]

    # Si el correo ya existe, exigir que esta cuenta sea claramente distinta.
    existing = User.query.filter_by(email=email).all()
    if existing:
        def _norm(u):
            return (
                (u.first_name or '').strip().lower(),
                (u.apellido_paterno or '').strip().lower(),
                (u.apellido_materno or '').strip().lower(),
            )
        nueva_ident = (first_name.lower(), apellido_paterno.lower(), apellido_materno.lower())
        for u in existing:
            if _norm(u) == nueva_ident:
                return jsonify({"error": f"Ya existe una cuenta con ese correo y el mismo nombre/apellidos (cuenta #{u.id}). Usa nombre o apellidos distintos para diferenciarla."}), 409
            if u.check_password(password):
                return jsonify({"error": f"La contraseña coincide con otra cuenta del mismo correo (cuenta #{u.id}). Usa una contraseña distinta: la contraseña es la que distingue el inicio de sesión."}), 409

    try:
        user = User(
            email=email,
            full_name=full_name,
            first_name=first_name,
            apellido_paterno=apellido_paterno,
            apellido_materno=(apellido_materno or None),
            whatsapp=whatsapp,
            tier=tier,
            terms_accepted=True,
            marketing_consent=False,
            marketing_consent_source='ADMIN_CREATE',
        )
        user.set_password(password)
        db.session.add(user)
        db.session.flush()
        db.session.add(EvaluationProcess(user_id=user.id))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.error(f"admin_create_user error: {type(e).__name__}: {e}")
        return jsonify({"error": "No se pudo crear la cuenta. Revisa los datos e inténtalo de nuevo."}), 500

    _comparte = len(existing) > 0
    logger.info(f"admin_create_user: cuenta #{user.id} creada por admin {current_user.email} (email={email}, comparte_correo={_comparte}, tier={tier})")
    return jsonify({
        "success": True,
        "user_id": user.id,
        "comparte_correo": _comparte,
        "message": (f"Cuenta #{user.id} creada ({full_name}). "
                    + (f"Comparte el correo {email} con {len(existing)} cuenta(s) existente(s); inicia sesión con su contraseña propia." if _comparte
                       else f"Cuenta nueva con correo {email}."))
    })

@app.route("/api/admin/users/<int:user_id>/grant_ec0217", methods=["POST"])
@login_required
def admin_grant_ec0217(user_id):
    """Admin autoriza N evaluaciones EC0217.01 (mp4) a un usuario (suma al saldo),
    o lo revoca (cantidad=0). Registra todo en la bitácora Ec0217Ledger."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    data = request.get_json(silent=True) or {}
    try:
        cantidad = int(data.get('cantidad', -1))
    except (TypeError, ValueError):
        return jsonify({"error": "Cantidad inválida"}), 400
    if cantidad < 0 or cantidad > 999:
        return jsonify({"error": "La cantidad debe estar entre 0 (revocar) y 999"}), 400
    nota = (data.get('nota') or '').strip()[:300]
    try:
        if cantidad == 0:
            prev = db.session.execute(
                db.text('SELECT COALESCE(ec0217_grants,0) FROM "user" WHERE id = :uid'),
                {"uid": user.id}).scalar() or 0
            db.session.execute(db.text('UPDATE "user" SET ec0217_grants = 0 WHERE id = :uid'),
                               {"uid": user.id})
            nuevo_saldo = 0
            entry = Ec0217Ledger(user_id=user.id, tipo='revoke', cantidad=-prev,
                                 saldo_after=0, admin_email=current_user.email,
                                 nota=nota or 'Revocación de saldo EC0217.01')
            msg = f"Saldo EC0217.01 de {user.email} revocado (0 evaluaciones)."
        else:
            db.session.execute(
                db.text('UPDATE "user" SET ec0217_grants = COALESCE(ec0217_grants,0) + :n WHERE id = :uid'),
                {"n": cantidad, "uid": user.id})
            nuevo_saldo = db.session.execute(
                db.text('SELECT COALESCE(ec0217_grants,0) FROM "user" WHERE id = :uid'),
                {"uid": user.id}).scalar() or 0
            entry = Ec0217Ledger(user_id=user.id, tipo='grant', cantidad=cantidad,
                                 saldo_after=nuevo_saldo, admin_email=current_user.email,
                                 nota=nota or f'Autorización de {cantidad} evaluación(es) EC0217.01 (mp4)')
            msg = (f"Autorizadas {cantidad} evaluación(es) EC0217.01 a {user.email}. "
                   f"Saldo total: {nuevo_saldo}.")
        db.session.add(entry)
        db.session.commit()
        db.session.refresh(user)
    except Exception:
        db.session.rollback()
        return jsonify({"error": "No se pudo actualizar el saldo. Intenta de nuevo."}), 500
    try:
        track_event('admin_ec0217', ('revoke' if cantidad == 0 else 'grant'),
                    user_id=current_user.id,
                    extra_data={'target_user': user.id, 'cantidad': cantidad,
                                'saldo_after': user.ec0217_grants})
    except Exception:
        pass
    return jsonify({"success": True, "saldo": user.ec0217_grants, "message": msg})

@app.route("/api/admin/users/<int:user_id>/ec0217_ledger", methods=["GET"])
@login_required
def admin_ec0217_ledger(user_id):
    """Bitácora EC0217.01 de un usuario: autorizaciones y servicios consumidos
    con día y hora, para resolver cualquier reclamo/aclaración."""
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    rows = (Ec0217Ledger.query.filter_by(user_id=user.id)
            .order_by(Ec0217Ledger.created_at.desc()).limit(300).all())
    tipo_lbl = {'grant': 'Autorización', 'consume': 'Servicio usado', 'revoke': 'Revocación'}
    movimientos = [{
        "fecha": (r.created_at.strftime('%d/%m/%Y %H:%M') if r.created_at else ''),
        "tipo": tipo_lbl.get(r.tipo, r.tipo),
        "cantidad": r.cantidad,
        "saldo_after": r.saldo_after,
        "admin": r.admin_email or '',
        "portafolio_id": r.portafolio_id,
        "nota": r.nota or '',
    } for r in rows]
    return jsonify({"success": True, "saldo": user.ec0217_grants or 0,
                    "email": user.email, "movimientos": movimientos})

@app.route("/api/admin/users/<int:user_id>/send_nudge_first_doc", methods=["POST"])
@login_required
def send_nudge_first_doc(user_id):
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
    if user.email_bounced_at:
        return jsonify({"error": "Email marcado como rebotado, no se envía"}), 400
    if not user.marketing_consent:
        return jsonify({"error": "Usuario optó por no recibir comunicaciones de marketing"}), 400
    gen_count = db.session.query(BetaMetric).filter(
        BetaMetric.user_id == user.id,
        BetaMetric.metric_type == 'generation_complete'
    ).count()
    if gen_count > 0:
        return jsonify({
            "error": f"El usuario ya generó {gen_count} documento(s). Este nudge es para usuarios sin generaciones."
        }), 400
    prev_nudge = db.session.query(EmailLog).filter(
        EmailLog.user_id == user.id,
        EmailLog.email_type == 'first_doc_nudge',
        EmailLog.status == 'sent'
    ).order_by(EmailLog.id.desc()).first()
    if prev_nudge:
        when = _fmt_cdmx(prev_nudge.sent_at or prev_nudge.attempted_at, '%d/%m/%Y %H:%M')
        return jsonify({
            "error": f"Ya se envió un nudge a este usuario el {when} CDMX. No se reenvía para evitar duplicados."
        }), 400
    ok = _send_first_doc_nudge_email(user)
    if not ok:
        return jsonify({"error": "Falló el envío del correo (revisa SMTP/logs)"}), 500
    try:
        track_event("email", "first_doc_nudge_sent", user_id=user.id,
                    extra_data={"by_admin": current_user.id})
    except Exception:
        pass
    return jsonify({"success": True, "message": f"Nudge enviado a {user.email}"})

@app.route("/perfil")
@login_required
def perfil():
    return render_template("perfil.html", title="Mi Perfil", active_page="perfil")

@app.route("/api/user/marketing_preferences", methods=["POST"])
@login_required
def update_marketing_preferences():
    from datetime import datetime
    data = request.get_json(silent=True) or {}
    raw = data.get("marketing_consent")
    if isinstance(raw, bool):
        new_value = raw
    elif isinstance(raw, str):
        new_value = raw.lower() in ("true", "1", "yes", "si", "sí", "on")
    else:
        return jsonify({"error": "Parámetro marketing_consent inválido"}), 400
    previous = current_user.marketing_consent
    current_user.marketing_consent = new_value
    current_user.marketing_consent_source = 'USER'
    if not new_value:
        current_user.mkt_unsubscribed_at = datetime.utcnow()
    else:
        current_user.mkt_unsubscribed_at = None
    db.session.commit()
    try:
        track_event(
            "marketing",
            "user_self_optout" if not new_value else "user_self_optin",
            user_id=current_user.id,
            extra_data={"previous": bool(previous), "new": bool(new_value), "source": "perfil"}
        )
    except Exception:
        pass
    return jsonify({
        "success": True,
        "marketing_consent": current_user.marketing_consent,
        "unsubscribed_at": current_user.mkt_unsubscribed_at.isoformat() if current_user.mkt_unsubscribed_at else None
    })

@app.route("/api/user/change_password", methods=["POST"])
@login_required
def change_password():
    nueva = request.form.get("nueva_contraseña", "")
    confirmar = request.form.get("confirmar_contraseña", "")
    if not nueva or not confirmar:
        flash('Todos los campos son obligatorios.', 'error')
    elif nueva != confirmar:
        flash('Las contraseñas no coinciden.', 'error')
    elif len(nueva) < 6:
        flash('La contraseña debe tener al menos 6 caracteres.', 'error')
    else:
        current_user.password_hash = generate_password_hash(nueva)
        db.session.commit()
        flash('Contraseña actualizada correctamente.', 'success')
    return redirect(url_for('perfil'))

def _dedupe_row_cells(row):
    """Devuelve los textos de las celdas de un row, deduplicando ÚNICAMENTE celdas
    combinadas (merged) reales en Word.

    python-docx expone una celda merged como múltiples elementos `cell` en `row.cells`,
    pero todos ellos comparten el MISMO elemento XML subyacente (`cell._tc`). Usamos
    identidad de objeto XML para detectar merges reales y NO degradar tablas legítimas
    con celdas distintas que casualmente tengan el mismo texto (ej. 'Sí | Sí | No').
    """
    seen_tc_ids = set()
    out = []
    for cell in row.cells:
        try:
            tc_id = id(cell._tc)
        except Exception:
            tc_id = None
        if tc_id is not None and tc_id in seen_tc_ids:
            continue
        if tc_id is not None:
            seen_tc_ids.add(tc_id)
        txt = cell.text.strip()
        if txt:
            out.append(txt)
    return out

def _extract_full_docx_text(fpath):
    import docx
    doc = docx.Document(fpath)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(_dedupe_row_cells(row))
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


_MODULO_PRODUCTS = [
    'Carta_Descriptiva', 'Contrato_de_Aprendizaje', 'Objetivo_General',
    'Evaluacion_Diagnostica', 'Evaluacion_Formativa', 'Evaluacion_Sumativa',
    'Lista_de_Cotejo', 'Guia_de_Observacion', 'Hojas_de_Respuestas',
    'Encuesta_Satisfaccion', 'Manual_del_Instructor', 'Manual_del_Participante',
]

@app.route("/api/collect_user_docs", methods=["POST"])
@login_required
def collect_user_docs():
    """Recopila SOLO los productos generados en el curso (course_session) activo.

    Fuente confiable acotada al curso: ChatHistory.generated_files de los elementos 1-3
    del course_session activo. El contenido se lee de StoredFile (autoritativo) con
    respaldo a disco. Esto evita la contaminación entre cursos (los nombres de archivo
    se reutilizan entre cursos, por lo que filtrar por nombre o por StoredFile global
    mezclaba documentos de otros cursos).
    """
    import json as _json_cu
    prefix = f"u{current_user.id}_"
    docs_dir = "generated_docs"
    _no_docs = "No se encontraron documentos generados para tu curso activo. Primero genera tus productos en los Módulos 1, 2 y 3."

    active_cs_id = current_user.active_course_session_id
    if not active_cs_id:
        return jsonify({"error": _no_docs}), 400

    wanted = []
    seen = set()
    for el in (1, 2, 3):
        row = ChatHistory.query.filter_by(
            user_id=current_user.id, element_num=el, course_session_id=active_cs_id
        ).first()
        if not row or not row.generated_files:
            continue
        try:
            files = _json_cu.loads(row.generated_files)
        except Exception:
            files = []
        if not isinstance(files, list):
            continue
        for fn in files:
            if not isinstance(fn, str) or fn in seen or not fn.endswith('.docx'):
                continue
            base = fn.replace(prefix, '').replace('.docx', '')
            if base.startswith('Diagnostico') or base.startswith('Autodiagnostico'):
                continue
            seen.add(fn)
            wanted.append(fn)

    collected = {}
    for fn in wanted:
        base = fn.replace(prefix, '').replace('.docx', '')
        label = base.replace('_', ' ')
        text = ''
        # Acotar al curso activo: nunca servir un StoredFile etiquetado a OTRO curso.
        # Se aceptan filas con course_session_id NULL (legado) o igual al curso activo.
        sf = (StoredFile.query
              .filter_by(user_id=current_user.id, filename=fn)
              .filter((StoredFile.course_session_id == active_cs_id) |
                      (StoredFile.course_session_id.is_(None)))
              .first())
        if sf and sf.content:
            try:
                from docx import Document as DocxDocument
                doc_obj = DocxDocument(io.BytesIO(sf.content))
                parts = [p.text for p in doc_obj.paragraphs if p.text.strip()]
                for table in doc_obj.tables:
                    for row in table.rows:
                        row_text = ' | '.join(_dedupe_row_cells(row))
                        if row_text:
                            parts.append(row_text)
                text = '\n'.join(parts)
            except Exception:
                text = ''
        if not text.strip():
            fpath = os.path.join(docs_dir, fn)
            if os.path.isfile(fpath):
                try:
                    text = _extract_full_docx_text(fpath)
                except Exception:
                    text = ''
        if text.strip():
            collected[label] = text

    if not collected:
        return jsonify({"error": _no_docs}), 400
    combined = ""
    for label, text in collected.items():
        combined += f"\n\n===== DOCUMENTO: {label} =====\n{text}"
    return jsonify({"text": combined, "doc_count": len(collected), "doc_names": list(collected.keys())})


_DIAG_PRODUCT_TABLES = {
    1: [2, 3, 4, 5, 6, 7, 8, 9, 10],
    2: [13, 14, 15, 16, 17],
    3: [19, 20, 21, 22, 23, 24, 25, 26, 27, 28],
}
_diag_catalog_cache = None


def _diag_uniq_cells(row):
    seen = set()
    out = []
    for c in row.cells:
        cid = id(c._tc)
        if cid in seen:
            continue
        seen.add(cid)
        out.append(c)
    return out


def _diag_markable_rows(table):
    """Filas de PRODUCTOS marcables: reactivo en col0, columnas SI/NO vacías."""
    out = []
    for ridx, r in enumerate(table.rows):
        u = _diag_uniq_cells(r)
        if len(u) < 3:
            continue
        c0, c1, c2 = u[0], u[1], u[2]
        if c1.text.strip().upper() == 'SI' or c2.text.strip().upper() == 'NO':
            continue
        if not c0.text.strip():
            continue
        if c1.text.strip() or c2.text.strip():
            continue
        out.append((ridx, c0.text.strip()))
    return out


def _get_diag_product_catalog():
    """Catálogo (cacheado) de los 138 reactivos de PRODUCTOS con código E{el}P##."""
    global _diag_catalog_cache
    if _diag_catalog_cache is not None:
        return _diag_catalog_cache
    from docx import Document as DocxDocument
    template_path = os.path.join("plantillas", "Diagnóstico_EC0301.docx")
    catalog = []
    try:
        doc = DocxDocument(template_path)
        for el in (1, 2, 3):
            counter = 0
            for ti in _DIAG_PRODUCT_TABLES[el]:
                table = doc.tables[ti]
                for (ridx, text) in _diag_markable_rows(table):
                    counter += 1
                    catalog.append({
                        "code": f"E{el}P{counter:02d}",
                        "element": el,
                        "ti": ti,
                        "ridx": ridx,
                        "text": text,
                    })
    except Exception:
        logger.warning("diag product catalog build failed", exc_info=True)
        return []
    _diag_catalog_cache = catalog
    return catalog


def _build_audit_catalog_prompt():
    """Bloque para anteponer al prompt de auditoría: pide marcar reactivo por reactivo."""
    catalog = _get_diag_product_catalog()
    if not catalog:
        return ""
    lines = [
        "INSTRUCCIONES DE MARCADO REACTIVO POR REACTIVO (OBLIGATORIO):",
        "Evalua CADA UNO de los siguientes reactivos de PRODUCTOS de forma individual contra la evidencia textual de los documentos proporcionados.",
        "Para cada codigo responde exactamente \"SI\" (existe evidencia textual explicita) o \"NO\" (no hay evidencia, o el documento correspondiente no fue proporcionado).",
        "NO evalues conocimientos ni actitud: esas secciones las completa el candidato, dejalas fuera del marcado.",
        "Devuelve OBLIGATORIAMENTE dentro del bloque JSON un objeto \"reactivos\" que incluya TODOS los siguientes codigos con su valor SI o NO:",
        "",
    ]
    cur = None
    for item in catalog:
        if item["element"] != cur:
            cur = item["element"]
            lines.append(f"ELEMENTO {cur} - PRODUCTOS:")
        lines.append(f"{item['code']}: {item['text']}")
    return "\n".join(lines)


def _rebuild_diag_header(doc, logo_path):
    """Etapa 1 (Autodiagnóstico): encabezado con 3 recuadros editables.
    Quita los logos por defecto embebidos en la plantilla y arma una tabla 1x3:
    IZQ 'Inserta aquí el logo de la Red CONOCER', CENTRO el logo que subió el
    usuario en el elemento 1 (vacío si no lo subió), DER 'Inserta aquí el logo
    de tu OC/ECE'. Todo queda editable en el .docx descargable.
    Solo aplica al Autodiagnóstico; degradación silenciosa si algo falla."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        valid_logo = bool(logo_path) and os.path.exists(logo_path)
    except Exception:
        valid_logo = False

    def _set_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            e = OxmlElement('w:' + edge)
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '4')
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), 'BFBFBF')
            borders.append(e)
        tcPr.append(borders)

    def _placeholder(cell, text):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _set_cell_borders(cell)

    def _build(header):
        for child in list(header._element):
            header._element.remove(child)
        tbl = header.add_table(rows=1, cols=3, width=Inches(6.5))
        tbl.autofit = False
        try:
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            pass
        cells = tbl.rows[0].cells
        for c, w in zip(cells, (Inches(2.1), Inches(2.3), Inches(2.1))):
            c.width = w
            try:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass
        _placeholder(cells[0], 'Inserta aquí el logo de la Red CONOCER')
        pc = cells[1].paragraphs[0]
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if valid_logo:
            try:
                pc.add_run().add_picture(logo_path, height=Inches(0.6))
            except Exception:
                pass
        _placeholder(cells[2], 'Inserta aquí el logo de tu OC/ECE')
        header.add_paragraph()

    for section in doc.sections:
        for attr in ('header', 'first_page_header', 'even_page_header'):
            try:
                h = getattr(section, attr)
                if h.is_linked_to_previous:
                    continue
                _build(h)
            except Exception:
                pass


def _norm_label(s):
    """Normaliza una etiqueta para emparejar sin acentos/mayúsculas."""
    import unicodedata
    s = (s or "").strip().lower()
    s = unicodedata.normalize('NFKD', s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return " ".join(s.split())


# Etapa 2: llaves canónicas de Datos personales (sirven Autodiagnóstico y Ficha).
DATOS_PERSONALES_KEYS = (
    'nombre_completo', 'curp', 'domicilio', 'ultimo_grado',
    'tel_casa', 'tel_celular', 'correo', 'fecha_aplicacion',
    'lugar_nacimiento', 'nacionalidad', 'genero', 'fecha_nacimiento',
)

# Mapa etiqueta-de-tabla (normalizada) -> llave canónica, para la Tabla 1
# de "Datos personales" del Autodiagnóstico EC0301.
_DIAG_DATOS_LABELMAP = {
    'nombre completo': 'nombre_completo',
    'curp': 'curp',
    'domicilio': 'domicilio',
    'ultimo grado de estudios': 'ultimo_grado',
    'telefono de casa': 'tel_casa',
    'telefono de celular': 'tel_celular',
    'correo electronico': 'correo',
    'fecha de aplicacion': 'fecha_aplicacion',
}


def _fill_datos_personales(doc, datos):
    """Etapa 2: vierte los Datos personales en las celdas blancas (columna
    derecha) de la Tabla 'Datos personales' del Autodiagnóstico. Solo llena
    lo que el usuario proporcionó; si un campo viene vacío se deja en blanco.
    Identifica la tabla por su encabezado para no depender del índice."""
    if not datos:
        return
    try:
        for table in doc.tables:
            if not table.rows:
                continue
            first = _norm_label(table.rows[0].cells[0].text)
            if not first.startswith('datos personales'):
                continue
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                key = _DIAG_DATOS_LABELMAP.get(_norm_label(cells[0].text))
                if not key:
                    continue
                val = (datos.get(key) or '').strip()
                if not val:
                    continue
                cells[1].text = val
            break
    except Exception as _dp_e:
        logger.warning(f'[_fill_datos_personales] no se pudo llenar Tabla 1: {_dp_e}')


def _get_datos_personales():
    """Devuelve los Datos personales: de la BD para usuarios autenticados
    (persistencia server-side, sin PII en cookie), o de la sesión para
    usuarios anónimos. Devuelve dict o None."""
    try:
        if current_user.is_authenticated:
            row = DatosPersonales.query.filter_by(user_id=current_user.id).first()
            if row:
                return {k: (getattr(row, k, '') or '') for k in DATOS_PERSONALES_KEYS}
            return None
    except Exception as _gd_e:
        logger.warning(f'[_get_datos_personales] {_gd_e}')
    return session.get('datos_personales')


def _ficha_row_unique_cells(row):
    """Devuelve [(start_col_idx, cell)] deduplicando celdas combinadas por _tc."""
    out = []
    last = None
    for ci, c in enumerate(row.cells):
        tcid = id(c._tc)
        if tcid != last:
            out.append((ci, c))
            last = tcid
    return out


def _fill_ficha_registro(doc, datos):
    """Etapa 4: vierte los Datos personales en las celdas blancas de la Ficha de
    Registro del Candidato (plantillas/Ficha_Registro_Candidato.docx).
    Reglas (según especificación del dueño):
      - Nombre Completo / Lugar de Nacimiento / Nacionalidad / CURP / Género /
        Fecha de Nacimiento: celda blanca a la DERECHA de la etiqueta.
      - Domicilio: recuadro blanco DEBAJO de 'Domicilio Particular'.
      - Correo / Teléfono / Teléfono Celular: recuadro blanco JUSTO ARRIBA de
        cada etiqueta (alineado por columna).
      - Estándar de Competencia = EC0301; Fecha (sup. der.) = fecha de aplicación.
    Campos vacíos ('') se dejan en blanco (respeta 'Lo editaré manualmente').
    No lanza excepción: ante cualquier desajuste, omite el campo afectado."""
    datos = datos or {}
    try:
        tables = doc.tables
        # --- Tabla superior: Estándar de Competencia / Fecha ---
        try:
            cells0 = _ficha_row_unique_cells(tables[0].rows[0])
            for idx, (_ci, c) in enumerate(cells0):
                lab = _norm_label(c.text)
                if idx + 1 >= len(cells0):
                    continue
                nxt = cells0[idx + 1][1]
                if lab.startswith('estandar de competencia') and not nxt.text.strip():
                    nxt.text = 'EC0301'
                elif lab.startswith('fecha') and not nxt.text.strip():
                    v = datos.get('fecha_aplicacion', '')
                    if v:
                        nxt.text = v
        except Exception:
            pass
        # --- Tabla principal de datos personales ---
        tmain = None
        for t in tables:
            txt = ' '.join(c.text for r in t.rows for c in r.cells)
            if 'Nombre Completo' in txt and 'CURP' in txt:
                tmain = t
                break
        if tmain is None:
            return
        rows = tmain.rows
        left_map = {
            'nombre completo': 'nombre_completo',
            'lugar de nacimiento': 'lugar_nacimiento',
            'nacionalidad': 'nacionalidad',
            'curp': 'curp',
            'genero': 'genero',
            'fecha de nacimiento': 'fecha_nacimiento',
        }
        domicilio_row_idx = None
        contacts_label_row_idx = None
        for ri, row in enumerate(rows):
            uc = _ficha_row_unique_cells(row)
            labels_norm = [_norm_label(c.text) for (_ci, c) in uc]
            # Etiqueta-izquierda -> valor a la derecha (misma fila).
            for idx, (_ci, c) in enumerate(uc):
                lab = _norm_label(c.text)
                key = None
                for lk, kk in left_map.items():
                    if lab.startswith(lk):
                        key = kk
                        break
                if key and idx + 1 < len(uc):
                    target = uc[idx + 1][1]
                    v = datos.get(key, '')
                    if v and not target.text.strip():
                        target.text = v
            if any(l.startswith('domicilio particular') for l in labels_norm):
                domicilio_row_idx = ri
            if any(l.startswith('e-mail') or l.startswith('telefono') for l in labels_norm):
                contacts_label_row_idx = ri
        # Domicilio -> primera celda blanca de la fila SIGUIENTE a la etiqueta.
        if domicilio_row_idx is not None and domicilio_row_idx + 1 < len(rows):
            v = datos.get('domicilio', '')
            if v:
                for ci, c in _ficha_row_unique_cells(rows[domicilio_row_idx + 1]):
                    if ci >= 1 and not c.text.strip():
                        c.text = v
                        break
        # Contactos -> fila ANTERIOR a las etiquetas, alineado por columna de inicio.
        if contacts_label_row_idx is not None and contacts_label_row_idx - 1 >= 0:
            colkey = {}
            for ci, c in _ficha_row_unique_cells(rows[contacts_label_row_idx]):
                lab = _norm_label(c.text)
                if lab.startswith('e-mail'):
                    colkey[ci] = 'correo'
                elif lab.startswith('telefono celular'):
                    colkey[ci] = 'tel_celular'
                elif lab.startswith('telefono'):
                    colkey[ci] = 'tel_casa'
            for ci, c in _ficha_row_unique_cells(rows[contacts_label_row_idx - 1]):
                if ci in colkey and not c.text.strip():
                    v = datos.get(colkey[ci], '')
                    if v:
                        c.text = v
    except Exception as _f_e:
        logger.warning(f'[_fill_ficha_registro] no se pudo llenar la Ficha: {_f_e}')


def _fill_diagnostico_template(ai_data, course_topic=None):
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from collections import defaultdict
    template_path = os.path.join("plantillas", "Diagnóstico_EC0301.docx")
    if not os.path.exists(template_path):
        return None
    doc = DocxDocument(template_path)

    def _set_cell(cell, txt, bold=False, center=False):
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].text = txt
            for r in p.runs[1:]:
                r.text = ''
            p.runs[0].bold = bold
        else:
            r = p.add_run(txt)
            r.bold = bold

    reactivos = ai_data.get("reactivos") or {}
    if not isinstance(reactivos, dict):
        reactivos = {}
    have_map = len(reactivos) > 0
    catalog = _get_diag_product_catalog()

    # Conteos agregados por elemento, usados SOLO como respaldo si falta el mapa per-reactivo.
    fallback_si = {
        1: int(ai_data.get("e1_productos_si", 0) or 0),
        2: int(ai_data.get("e2_productos_si", 0) or 0),
        3: int(ai_data.get("e3_productos_si", 0) or 0),
    }
    fallback_used = {1: 0, 2: 0, 3: 0}

    by_table = defaultdict(list)
    for item in catalog:
        by_table[item["ti"]].append(item)

    counts = {1: {"si": 0, "no": 0}, 2: {"si": 0, "no": 0}, 3: {"si": 0, "no": 0}}
    _truthy = {"SI", "SÍ", "S", "YES", "Y", "1", "TRUE", "CUMPLE"}

    for ti, items in by_table.items():
        table = doc.tables[ti]
        for item in items:
            el = item["element"]
            try:
                row = table.rows[item["ridx"]]
            except Exception:
                continue
            u = _diag_uniq_cells(row)
            if len(u) < 3:
                continue
            si_cell, no_cell = u[1], u[2]
            if have_map:
                raw = reactivos.get(item["code"])
                verdict = "SI" if (isinstance(raw, str) and raw.strip().upper() in _truthy) else "NO"
            else:
                if fallback_used[el] < fallback_si[el]:
                    verdict = "SI"
                    fallback_used[el] += 1
                else:
                    verdict = "NO"
            if verdict == "SI":
                _set_cell(si_cell, "X", bold=True, center=True)
                counts[el]["si"] += 1
            else:
                _set_cell(no_cell, "X", bold=True, center=True)
                counts[el]["no"] += 1

    # VALORACIÓN (tabla 31): llenar SOLO filas de PRODUCTOS. CONOCIMIENTOS/ACTITUD y el
    # TOTAL global se dejan en blanco (los completa el candidato); sólo se fija el denominador.
    val_table = doc.tables[31]

    def _fill_val_full(ri, s, n, total):
        u = _diag_uniq_cells(val_table.rows[ri])
        if len(u) > 4:
            _set_cell(u[2], str(s), center=True)
            _set_cell(u[3], str(n), center=True)
            _set_cell(u[4], str(total), center=True)

    def _fill_val_total_only(ri, total):
        u = _diag_uniq_cells(val_table.rows[ri])
        if len(u) > 4:
            _set_cell(u[4], str(total), center=True)

    _fill_val_full(2, counts[1]["si"], counts[1]["no"], 47)    # E1 productos
    _fill_val_full(7, counts[2]["si"], counts[2]["no"], 31)    # E2 productos
    _fill_val_full(11, counts[3]["si"], counts[3]["no"], 60)   # E3 productos
    _fill_val_total_only(3, 4)     # E1 conocimientos (SI/NO en blanco)
    _fill_val_total_only(4, 1)     # E1 actitud (SI/NO en blanco)
    _fill_val_total_only(8, 2)     # E2 conocimientos (SI/NO en blanco)
    _fill_val_total_only(12, 145)  # TOTAL global (SI/NO en blanco hasta completar)

    productos_si = counts[1]["si"] + counts[2]["si"] + counts[3]["si"]

    # CONCLUSIÓN (tabla 29): el dictamen final depende de CONOCIMIENTOS/ACTITUD => pendiente.
    concl_table = doc.tables[29]
    concl_table.rows[2].cells[1].text = (
        f"Productos verificados: {productos_si} de 138. "
        f"Pendiente: complete las secciones de CONOCIMIENTOS y ACTITUD para obtener el dictamen final."
    )
    concl_table.rows[3].cells[1].text = ""
    if len(concl_table.rows[2].cells) > 2:
        concl_table.rows[2].cells[2].text = ""

    # Etapa 1: encabezado con logos editables (solo Autodiagnóstico).
    try:
        _rebuild_diag_header(doc, session.get('course_logo_path'))
    except Exception as _hdr_e:
        logger.warning(f'[_fill_diagnostico_template] header rebuild fallo: {_hdr_e}')

    # Etapa 2: vierte los Datos personales en la Tabla 1 (si existen).
    try:
        _fill_datos_personales(doc, _get_datos_personales())
    except Exception as _dp_e:
        logger.warning(f'[_fill_diagnostico_template] datos personales fallo: {_dp_e}')

    out_dir = "generated_docs"
    os.makedirs(out_dir, exist_ok=True)
    if current_user.is_authenticated:
        out_name = f"u{current_user.id}_Diagnostico_EC0301.docx"
    else:
        topic_slug = _sanitize_topic_for_filename(course_topic)
        if topic_slug:
            out_name = f"Diagnostico_EC0301_{topic_slug}.docx"
        else:
            out_name = "anon_Diagnostico_EC0301.docx"
    out_path = os.path.join(out_dir, out_name)
    doc.save(out_path)
    return out_name


@app.route("/api/save_datos_personales", methods=["POST"])
def save_datos_personales():
    """Etapa 2: guarda los Datos personales del candidato/usuario en sesión.
    Sirven para el Autodiagnóstico (Tabla 1) y, más adelante, la Ficha de
    Registro y el Portafolio de Evidencias (E5). Campos vacíos => en blanco."""
    try:
        payload = request.get_json(silent=True) or {}
    except Exception:
        payload = {}
    _limits = {
        'nombre_completo': 200, 'curp': 18, 'domicilio': 300, 'ultimo_grado': 150,
        'tel_casa': 30, 'tel_celular': 30, 'correo': 150, 'fecha_aplicacion': 40,
        'lugar_nacimiento': 150, 'nacionalidad': 80, 'genero': 40, 'fecha_nacimiento': 40,
    }
    datos = {}
    for key in DATOS_PERSONALES_KEYS:
        val = payload.get(key, '')
        if not isinstance(val, str):
            val = str(val) if val is not None else ''
        val = val.strip()[:_limits.get(key, 200)]
        if key == 'curp':
            val = val.upper()
        datos[key] = val
    if current_user.is_authenticated:
        # Persistencia server-side: evita guardar PII (CURP, domicilio, etc.)
        # en la cookie de sesión del cliente.
        try:
            row = DatosPersonales.query.filter_by(user_id=current_user.id).first()
            if not row:
                row = DatosPersonales(user_id=current_user.id)
                db.session.add(row)
            for key in DATOS_PERSONALES_KEYS:
                setattr(row, key, datos.get(key, ''))
            db.session.commit()
            session.pop('datos_personales', None)
        except Exception as _save_e:
            db.session.rollback()
            logger.warning(f'[save_datos_personales] DB upsert fallo: {_save_e}')
            session['datos_personales'] = datos
            session.modified = True
    else:
        session['datos_personales'] = datos
        session.modified = True
    filled = sum(1 for v in datos.values() if v)
    return jsonify({"success": True, "filled": filled})


# Etapa 3: extracción de CURP desde el PDF oficial (gob.mx).
# Formato CURP: 4 letras + 6 dígitos (AAMMDD) + H/M + 2 letras estado +
# 3 consonantes + 1 alfanumérico (homoclave) + 1 dígito verificador = 18.
_CURP_RE = re.compile(r'\b([A-Z]{4}\d{6}[HM][A-Z]{5}[0-9A-Z]\d)\b')


def _extract_curp_from_text(text):
    if not text:
        return ''
    up = text.upper()
    m = _CURP_RE.search(up)
    if m:
        return m.group(1)
    # Respaldo: buscar 18 caracteres tras la etiqueta "CURP".
    m2 = re.search(r'CURP[^A-Z0-9]{0,10}([A-Z0-9]{18})', up)
    if m2 and _CURP_RE.match(m2.group(1)):
        return m2.group(1)
    return ''


@app.route("/api/curp_pdf", methods=["POST"])
def curp_pdf():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No se envió ningún archivo."}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({"success": False, "error": "Archivo vacío."}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext != 'pdf':
        return jsonify({"success": False, "error": "Sube el PDF oficial de tu CURP (formato .pdf)."}), 400
    try:
        data = file.read()
    except Exception:
        return jsonify({"success": False, "error": "No se pudo leer el archivo."}), 400
    if not data:
        return jsonify({"success": False, "error": "El archivo está vacío."}), 400
    if len(data) > 5 * 1024 * 1024:
        return jsonify({"success": False, "error": "El PDF es demasiado grande (máx. 5 MB)."}), 400
    text = ""
    try:
        import pypdf, io as _io
        reader = pypdf.PdfReader(_io.BytesIO(data))
        pages = []
        for page in reader.pages:
            pt = page.extract_text()
            if pt:
                pages.append(pt)
        text = '\n'.join(pages)
    except Exception as _pdf_e:
        logger.warning(f'[curp_pdf] PDF parse fallo: {_pdf_e}')
        return jsonify({"success": False, "error": "No se pudo leer el PDF. Captura tu CURP manualmente."}), 200
    curp = _extract_curp_from_text(text)
    stored = False
    # Persistir el PDF para el Portafolio (E5) solo para usuarios autenticados.
    if current_user.is_authenticated:
        try:
            fname = f'curp_{current_user.id}.pdf'
            cs_id_for_file = current_user.active_course_session_id
            existing = StoredFile.query.filter_by(user_id=current_user.id, filename=fname).first()
            if existing:
                existing.content = data
                existing.content_type = 'application/pdf'
                existing.file_category = 'curp_pdf'
                if cs_id_for_file and not existing.course_session_id:
                    existing.course_session_id = cs_id_for_file
            else:
                db.session.add(StoredFile(
                    user_id=current_user.id, filename=fname, content=data,
                    content_type='application/pdf', file_category='curp_pdf',
                    course_session_id=cs_id_for_file))
            db.session.commit()
            stored = True
            if curp:
                row = DatosPersonales.query.filter_by(user_id=current_user.id).first()
                if not row:
                    row = DatosPersonales(user_id=current_user.id)
                    db.session.add(row)
                row.curp = curp
                db.session.commit()
        except Exception as _st_e:
            db.session.rollback()
            logger.warning(f'[curp_pdf] persistencia fallo: {_st_e}')
    if curp:
        return jsonify({"success": True, "curp": curp, "stored": stored})
    return jsonify({"success": True, "curp": "", "stored": stored,
                    "message": "Guardamos tu PDF, pero no pudimos leer la CURP automáticamente. Captúrala manualmente."})


@app.route("/api/descargar_ficha_registro", methods=["GET"])
def descargar_ficha_registro():
    """Etapa 4 (Opción B): entrega la Ficha de Registro del Candidato en Word,
    ya llena con los Datos personales guardados, para que el evaluador la revise,
    ajuste y la cargue al Portafolio de Evidencias. Campos vacíos quedan en blanco
    (respeta 'Lo editaré manualmente')."""
    datos = _get_datos_personales() or {}
    template_path = os.path.join("plantillas", "Ficha_Registro_Candidato.docx")
    if not os.path.exists(template_path):
        return jsonify({"error": "La plantilla de la Ficha no está disponible."}), 404
    try:
        import docx as _docx
        doc = _docx.Document(template_path)
        _fill_ficha_registro(doc, datos)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer, as_attachment=True,
            download_name="Ficha_de_Registro_del_Candidato.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as _fr_e:
        logger.warning(f'[descargar_ficha_registro] {_fr_e}')
        return jsonify({"error": "No se pudo generar la Ficha de Registro."}), 500


@app.route("/api/descargar_curp_pdf", methods=["GET"])
@login_required
def descargar_curp_pdf():
    """Etapa 4 (Opción B): descarga el PDF oficial de CURP que el usuario subió,
    para anexarlo en el Portafolio de Evidencias justo después de la Ficha."""
    fname = f'curp_{current_user.id}.pdf'
    sf = StoredFile.query.filter_by(user_id=current_user.id, filename=fname).first()
    if not sf or not sf.content:
        return jsonify({"error": "No tienes un PDF de CURP guardado."}), 404
    return send_file(
        io.BytesIO(sf.content), as_attachment=True,
        download_name="CURP.pdf", mimetype="application/pdf")


@app.route("/api/extract_text", methods=["POST"])
@login_required
def extract_text():
    if 'file' not in request.files:
        return jsonify({"error": "No se envió ningún archivo"}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({"error": "Archivo vacío"}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    text = ""
    try:
        if ext == 'txt':
            text = file.read().decode('utf-8', errors='replace')
        elif ext == 'docx':
            import docx
            doc = docx.Document(file)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(_dedupe_row_cells(row))
                    if row_text:
                        parts.append(row_text)
            text = '\n'.join(parts)
        elif ext == 'pdf':
            import pypdf
            reader = pypdf.PdfReader(file)
            pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = '\n'.join(pages)
        else:
            return jsonify({"error": "Formato no soportado. Usa .docx, .pdf o .txt"}), 400
    except Exception as e:
        return jsonify({"error": f"Error al procesar el archivo: {str(e)}"}), 500
    if not text.strip():
        return jsonify({"error": "No se pudo extraer texto del archivo."}), 400
    return jsonify({"text": text})

def _ensure_anon_sid():
    import uuid as _uuid
    sid = session.get('anon_sid')
    if not sid:
        sid = _uuid.uuid4().hex
        session['anon_sid'] = sid
    return sid

def _save_full_chat_history(element_num, messages, course_topic_val=None, generated_file=None):
    try:
        import json as _json_ch
        if not isinstance(messages, list):
            return
        trimmed = messages[-40:] if len(messages) > 40 else messages
        clean = []
        for m in trimmed:
            if not isinstance(m, dict):
                continue
            r = m.get('role')
            c = m.get('content', '')
            if r in ('user', 'assistant', 'system') and isinstance(c, str):
                clean.append({'role': r, 'content': c[:8000]})
        payload = _json_ch.dumps(clean, ensure_ascii=False)
        active_cs_id = None
        if current_user.is_authenticated:
            try:
                active_cs_id = current_user.active_course_session_id or _get_or_create_active_course_session(current_user).id
            except Exception:
                active_cs_id = None
            if active_cs_id is None:
                logger.warning(f"_save_full_chat_history skipped user={current_user.id} elem={element_num}: no active_course_session_id")
                return
            row = ChatHistory.query.filter_by(user_id=current_user.id, element_num=element_num, course_session_id=active_cs_id).first()
        else:
            sid = _ensure_anon_sid()
            row = ChatHistory.query.filter_by(session_id=sid, element_num=element_num, user_id=None).first()
        _safe_topic = None
        if course_topic_val:
            _safe_topic = _sanitize_course_name_from_message(course_topic_val) or None
            if not _safe_topic:
                logger.info(f"_save_full_chat_history: course_topic_val rejected by sanitizer user={current_user.id if current_user.is_authenticated else 'anon'} preview={course_topic_val[:60]!r}")
        if row is None:
            row = ChatHistory(
                user_id=current_user.id if current_user.is_authenticated else None,
                session_id=None if current_user.is_authenticated else _ensure_anon_sid(),
                element_num=element_num,
                messages_json=payload,
                course_topic=(_safe_topic or '')[:300] or None,
                course_session_id=active_cs_id,
            )
            db.session.add(row)
        else:
            row.messages_json = payload
            if _safe_topic and not row.course_topic:
                row.course_topic = _safe_topic[:300]
            if current_user.is_authenticated and active_cs_id and not row.course_session_id:
                row.course_session_id = active_cs_id
        if current_user.is_authenticated and active_cs_id:
            try:
                cs = CourseSession.query.get(active_cs_id)
                if cs is not None:
                    if _safe_topic and not cs.topic:
                        cs.topic = _safe_topic[:300]
                    cs.last_activity_at = datetime.utcnow()
            except Exception:
                pass
        if generated_file:
            existing_files = []
            if row.generated_files:
                try:
                    existing_files = _json_ch.loads(row.generated_files)
                    if not isinstance(existing_files, list):
                        existing_files = []
                except Exception:
                    existing_files = []
            if generated_file not in existing_files:
                existing_files.append(generated_file)
                row.generated_files = _json_ch.dumps(existing_files[-20:], ensure_ascii=False)
        db.session.commit()
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.warning(f"_save_full_chat_history failed: {type(e).__name__}: {e}")

def _load_full_chat_history(element_num):
    try:
        import json as _json_ch
        if current_user.is_authenticated:
            active_cs_id = current_user.active_course_session_id
            if not active_cs_id:
                return {'messages': [], 'course_topic': None, 'generated_files': []}
            row = ChatHistory.query.filter_by(user_id=current_user.id, element_num=element_num, course_session_id=active_cs_id).first()
        else:
            sid = session.get('anon_sid')
            if not sid:
                return {'messages': [], 'course_topic': None, 'generated_files': []}
            row = ChatHistory.query.filter_by(session_id=sid, element_num=element_num, user_id=None).first()
        if not row:
            return {'messages': [], 'course_topic': None, 'generated_files': []}
        try:
            msgs = _json_ch.loads(row.messages_json or '[]')
            if not isinstance(msgs, list):
                msgs = []
        except Exception:
            msgs = []
        try:
            files = _json_ch.loads(row.generated_files) if row.generated_files else []
            if not isinstance(files, list):
                files = []
        except Exception:
            files = []
        return {'messages': msgs, 'course_topic': row.course_topic, 'generated_files': files}
    except Exception as e:
        logger.warning(f"_load_full_chat_history failed: {type(e).__name__}: {e}")
        return {'messages': [], 'course_topic': None, 'generated_files': []}

def _transfer_chat_history_to_user(user_id, anon_sid):
    if not anon_sid or not user_id:
        return 0
    try:
        rows = ChatHistory.query.filter_by(session_id=anon_sid, user_id=None).all()
        if not rows:
            return 0
        target_cs_id = None
        try:
            target_user = User.query.get(user_id)
            if target_user is not None:
                target_cs = _get_or_create_active_course_session(target_user)
                if target_cs is not None:
                    target_cs_id = target_cs.id
        except Exception as _ecs:
            logger.warning(f"_transfer_chat_history_to_user: could not resolve active course session for user {user_id}: {_ecs}")
        if target_cs_id is None:
            logger.warning(f"_transfer_chat_history_to_user aborted: no active_course_session_id for user {user_id}; {len(rows)} anon rows left intact for retry")
            try: db.session.rollback()
            except Exception: pass
            return 0
        moved = 0
        for r in rows:
            existing = ChatHistory.query.filter_by(user_id=user_id, element_num=r.element_num, course_session_id=target_cs_id).first()
            if existing:
                try:
                    import json as _json_ch
                    cur = _json_ch.loads(existing.messages_json or '[]')
                    incoming = _json_ch.loads(r.messages_json or '[]')
                    merged = (cur + incoming)[-40:]
                    existing.messages_json = _json_ch.dumps(merged, ensure_ascii=False)
                    if not existing.course_topic and r.course_topic:
                        existing.course_topic = r.course_topic
                    if r.generated_files:
                        existing.generated_files = r.generated_files
                except Exception:
                    pass
                db.session.delete(r)
            else:
                r.user_id = user_id
                r.session_id = None
                r.course_session_id = target_cs_id
            moved += 1
        db.session.commit()
        return moved
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.warning(f"_transfer_chat_history_to_user failed: {type(e).__name__}: {e}")
        return 0

def _scan_non_inclusive_language(text):
    """Escaneo simple de lenguaje no inclusivo (articulo masculino generico +
    sustantivo de rol). Solo informa, nunca modifica el archivo del usuario.
    Devuelve una lista de numeros de linea (1-indexed) donde aparece."""
    import re as _re_incl
    pattern = _re_incl.compile(r'\b(el|los)\s+(participante|alumno|estudiante|profesor|docente|instructor)s?\b(?!\s+o\s+(la\s+)?(alumna|alumnas|profesora|profesoras|instructora|instructoras|docenta|docentas))', _re_incl.IGNORECASE)
    hits = []
    for _i, _line in enumerate(text.split('\n'), 1):
        if pattern.search(_line):
            hits.append(_i)
            if len(hits) >= 8:
                break
    return hits

def _save_chat_spec(element_num, content, filenames_str):
    from sqlalchemy.exc import IntegrityError
    if current_user.is_authenticated:
        filt = {'user_id': current_user.id, 'element_num': element_num}
        new_kwargs = {'user_id': current_user.id, 'element_num': element_num,
                      'content': content, 'filenames': filenames_str}
    else:
        sid = _ensure_anon_sid()
        filt = {'session_id': sid, 'user_id': None, 'element_num': element_num}
        new_kwargs = {'session_id': sid, 'element_num': element_num,
                      'content': content, 'filenames': filenames_str}
    row = ChatSpec.query.filter_by(**filt).first()
    if row:
        row.content = content
        row.filenames = filenames_str
        db.session.commit()
        return
    try:
        db.session.add(ChatSpec(**new_kwargs))
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        row = ChatSpec.query.filter_by(**filt).first()
        if row:
            row.content = content
            row.filenames = filenames_str
            db.session.commit()

def _get_chat_spec(element_num):
    if current_user.is_authenticated:
        row = ChatSpec.query.filter_by(user_id=current_user.id, element_num=element_num).first()
        if row and row.content:
            return row.content
    sid = session.get('anon_sid')
    if sid:
        row = ChatSpec.query.filter_by(session_id=sid, user_id=None, element_num=element_num).first()
        if row and row.content:
            return row.content
    return session.get(f'user_specs_e{element_num}', '')

def _clear_chat_spec(element_num):
    try:
        if current_user.is_authenticated:
            ChatSpec.query.filter_by(user_id=current_user.id, element_num=element_num).delete()
        sid = session.get('anon_sid')
        if sid:
            ChatSpec.query.filter_by(session_id=sid, user_id=None, element_num=element_num).delete()
        db.session.commit()
    except Exception:
        db.session.rollback()
    session.pop(f'user_specs_e{element_num}', None)

_ANON_UPLOAD_RL = {}
_ANON_UPLOAD_RL_LOCK = threading.Lock()
_ANON_LAST_PURGE = [0.0]

def _anon_rate_limit_check(ip, max_per_hour=10):
    import time as _t
    now = _t.time()
    cutoff = now - 3600
    with _ANON_UPLOAD_RL_LOCK:
        bucket = [t for t in _ANON_UPLOAD_RL.get(ip, []) if t > cutoff]
        if len(bucket) >= max_per_hour:
            return False
        bucket.append(now)
        _ANON_UPLOAD_RL[ip] = bucket
        if len(_ANON_UPLOAD_RL) > 5000:
            for k in list(_ANON_UPLOAD_RL.keys()):
                if not _ANON_UPLOAD_RL[k] or _ANON_UPLOAD_RL[k][-1] < cutoff:
                    _ANON_UPLOAD_RL.pop(k, None)
        return True

def _purge_stale_anon_specs():
    import time as _t
    now = _t.time()
    if now - _ANON_LAST_PURGE[0] < 3600:
        return
    _ANON_LAST_PURGE[0] = now
    try:
        deleted = db.session.execute(db.text(
            "DELETE FROM chat_spec WHERE user_id IS NULL "
            "AND created_at < NOW() - INTERVAL '24 hours'")).rowcount
        db.session.commit()
        if deleted:
            logger.info(f"chat_spec purge: deleted {deleted} stale anon rows")
    except Exception as _e:
        db.session.rollback()
        logger.warning(f"chat_spec purge failed: {_e}")

def _cocreation_session_key(element_num):
    return f'cocreation_state_e{int(element_num)}'

def _get_cocreation_state(element_num):
    """States: 'pending' (uploaded, awaiting preview), 'shown', 'confirmed', 'refining', 'skipped', or None."""
    return session.get(_cocreation_session_key(element_num))

def _set_cocreation_state(element_num, state):
    session[_cocreation_session_key(element_num)] = state
    session.modified = True

def _clear_cocreation_state(element_num):
    session.pop(_cocreation_session_key(element_num), None)
    session.modified = True

def _cocreation_skip_session_key():
    return 'cocreation_skip_session'

def _is_cocreation_skipped_session():
    return bool(session.get(_cocreation_skip_session_key()))

def _set_cocreation_skipped_session():
    session[_cocreation_skip_session_key()] = True
    session.modified = True

def _sanitize_filename_for_output(name):
    if not name:
        return ''
    s = str(name)
    for _bad in ('<', '>', '"', "'", '&', '\n', '\r', '\t', '\x00'):
        s = s.replace(_bad, '')
    return s[:200]

@app.route("/api/upload_specs", methods=["POST"])
def upload_specs():
    element_num = request.form.get('element_num', '1')
    try:
        element_num = int(element_num)
    except ValueError:
        element_num = 1
    if element_num not in (1, 2, 3, 4):
        return jsonify({"error": "Elemento inválido"}), 400
    files = request.files.getlist('files')
    if not files or len(files) == 0:
        return jsonify({"error": "No se enviaron archivos"}), 400
    is_anon = not current_user.is_authenticated
    user_tier_for_limits = 'ANON'
    if is_anon:
        ip = (request.headers.get('X-Forwarded-For', request.remote_addr) or '').split(',')[0].strip()
        if not _anon_rate_limit_check(ip):
            return jsonify({"error": "Has subido demasiados archivos. Espera un momento o crea tu cuenta gratis."}), 429
        _purge_stale_anon_specs()
        MAX_FILES = 1
        MAX_SIZE_MB = 2
        MAX_CHARS_PER_FILE = 6000
        MAX_PAGES_PDF = 10
        user_tier_for_limits = 'ANON'
    elif current_user.tier == 'FREE':
        MAX_FILES = 1
        MAX_SIZE_MB = 3
        MAX_CHARS_PER_FILE = 8000
        MAX_PAGES_PDF = 15
        user_tier_for_limits = 'FREE'
    else:
        MAX_FILES = 3
        MAX_SIZE_MB = 5
        MAX_CHARS_PER_FILE = 15000
        MAX_PAGES_PDF = 30
        user_tier_for_limits = current_user.tier
    if len(files) > MAX_FILES:
        if user_tier_for_limits == 'ANON':
            msg = "Sin registrarte solo puedes subir 1 archivo. Crea tu cuenta FREE para acceso completo al Módulo 1; con PRO podrás subir hasta 3 archivos."
        elif user_tier_for_limits == 'FREE':
            msg = "En tu plan FREE puedes subir 1 archivo (es una probada). PRO permite subir hasta 3 archivos."
        else:
            msg = f"Máximo {MAX_FILES} archivos permitidos"
        return jsonify({"error": msg, "tier": user_tier_for_limits, "upsell": user_tier_for_limits == 'FREE'}), 400
    all_text = []
    filenames = []
    truncation_events = []
    for f in files:
        if not f.filename:
            continue
        f.seek(0, 2)
        size = f.tell()
        f.seek(0)
        size_mb = size / (1024 * 1024)
        safe_filename = _sanitize_filename_for_output(f.filename)
        if size > MAX_SIZE_MB * 1024 * 1024:
            if user_tier_for_limits == 'FREE':
                _emit_subnorm_metric_server(
                    'cocreation_file_size_rejected',
                    value_int=int(size_mb * 1024),
                    value_text=(f"FREE limit {MAX_SIZE_MB}MB exceeded by {safe_filename} ({size_mb:.2f}MB)")[:500],
                    element_num=element_num,
                    user_id=current_user.id,
                    session_id=None,
                )
                return jsonify({
                    "error": f"Tu archivo '{safe_filename}' pesa {size_mb:.1f} MB. En tu plan FREE el límite es {MAX_SIZE_MB} MB (es una probada). PRO acepta archivos hasta 5 MB.",
                    "tier": "FREE",
                    "upsell": True,
                    "limit_kind": "size_mb",
                    "limit_value": MAX_SIZE_MB,
                    "your_value_mb": round(size_mb, 2),
                }), 413
            else:
                return jsonify({"error": f"El archivo '{safe_filename}' excede {MAX_SIZE_MB}MB"}), 400
        ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
        text = ""
        original_chars = 0
        original_pages = 0
        used_pages = 0
        try:
            if ext == 'txt':
                raw = f.read().decode('utf-8', errors='replace')
                original_chars = len(raw)
                text = raw[:MAX_CHARS_PER_FILE]
            elif ext == 'docx':
                import docx
                doc = docx.Document(f)
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(_dedupe_row_cells(row))
                        if row_text:
                            parts.append(row_text)
                full_text = '\n'.join(parts)
                original_chars = len(full_text)
                text = full_text[:MAX_CHARS_PER_FILE]
            elif ext == 'pdf':
                import pypdf
                reader = pypdf.PdfReader(f)
                original_pages = len(reader.pages)
                used_pages = min(original_pages, MAX_PAGES_PDF)
                pages = []
                for page in reader.pages[:MAX_PAGES_PDF]:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                full_text = '\n'.join(pages)
                if not full_text.strip():
                    try:
                        from google.cloud import vision
                        f.seek(0)
                        _vision_client = vision.ImageAnnotatorClient()
                        _input_config = vision.InputConfig(content=f.read(), mime_type='application/pdf')
                        _feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)
                        _vision_request = vision.AnnotateFileRequest(input_config=_input_config, features=[_feature])
                        _vision_response = _vision_client.batch_annotate_files(requests=[_vision_request])
                        _ocr_parts = []
                        for _img_resp in _vision_response.responses[0].responses:
                            if _img_resp.full_text_annotation and _img_resp.full_text_annotation.text:
                                _ocr_parts.append(_img_resp.full_text_annotation.text)
                        full_text = '\n'.join(_ocr_parts)
                        logger.info(f"vision_ocr_fallback_used file={safe_filename!r} chars_recovered={len(full_text)}")
                    except Exception as _e_ocr:
                        logger.warning(f"vision_ocr_fallback_fail file={safe_filename!r} err={type(_e_ocr).__name__}: {_e_ocr}")
                original_chars = len(full_text)
                text = full_text[:MAX_CHARS_PER_FILE]
            elif ext == 'xlsx':
                try:
                    from openpyxl import load_workbook
                    wb = load_workbook(filename=f, read_only=True, data_only=True)
                    parts = []
                    MAX_SHEETS = 10
                    MAX_ROWS_PER_SHEET = 200
                    for sh_idx, sheet in enumerate(wb.worksheets[:MAX_SHEETS]):
                        parts.append(f"--- Hoja: {sheet.title} ---")
                        row_count = 0
                        for row in sheet.iter_rows(values_only=True):
                            if row_count >= MAX_ROWS_PER_SHEET:
                                parts.append(f"... (hoja truncada a {MAX_ROWS_PER_SHEET} filas)")
                                break
                            cells = [str(c) if c is not None else '' for c in row]
                            if any(c.strip() for c in cells):
                                parts.append(' | '.join(cells))
                                row_count += 1
                    full_text = '\n'.join(parts)
                    original_chars = len(full_text)
                    text = full_text[:MAX_CHARS_PER_FILE]
                except Exception as _e_xlsx:
                    logger.warning(f"upload_specs xlsx parse fail file={safe_filename!r} err={type(_e_xlsx).__name__}")
                    return jsonify({"error": f"No se pudo leer el archivo Excel '{safe_filename}'. Gu\u00e1rdalo como .xlsx o exp\u00f3rtalo a PDF/DOCX."}), 400
            else:
                return jsonify({"error": f"Formato '{ext}' no soportado en '{safe_filename}'. Usa .docx, .pdf, .txt o .xlsx"}), 400
        except Exception as e:
            return jsonify({"error": f"Error procesando '{safe_filename}': {str(e)}"}), 500
        if text.strip():
            all_text.append(f"===== DOCUMENTO: {safe_filename} =====\n{text.strip()}")
            filenames.append(safe_filename)
            chars_truncated = original_chars > len(text)
            pages_truncated = (ext == 'pdf' and original_pages > used_pages)
            if chars_truncated or pages_truncated:
                truncation_events.append({
                    'filename': safe_filename,
                    'original_chars': original_chars,
                    'used_chars': len(text),
                    'original_pages': original_pages if ext == 'pdf' else None,
                    'used_pages': used_pages if ext == 'pdf' else None,
                    'tier_limit_chars': MAX_CHARS_PER_FILE,
                    'tier_limit_pages': MAX_PAGES_PDF if ext == 'pdf' else None,
                })
    if not all_text:
        return jsonify({"error": "No se pudo extraer texto de los archivos"}), 400
    combined = '\n\n'.join(all_text)
    _non_inclusive_lines = []
    if session.get('contexto_institucional') == 'uam':
        try:
            _non_inclusive_lines = _scan_non_inclusive_language(combined)
        except Exception:
            _non_inclusive_lines = []
    try:
        _save_chat_spec(element_num, combined, ', '.join(filenames)[:780])
    except Exception as e:
        db.session.rollback()
        logger.error(f"upload_specs save failed: {e}")
        return jsonify({"error": "Error guardando especificaciones"}), 500
    _set_cocreation_state(element_num, 'pending')
    resp = {"success": True, "filenames": filenames, "chars": len(combined), "tier": user_tier_for_limits, "cocreation_eligible": not _is_cocreation_skipped_session()}
    if _non_inclusive_lines:
        _lineas_str = ', '.join(str(n) for n in _non_inclusive_lines[:5])
        resp["inclusive_language_warning"] = f"Detectamos posible lenguaje no inclusivo en tu documento subido (líneas {_lineas_str}) — considera actualizarlo."
    if truncation_events and user_tier_for_limits == 'FREE':
        import json as _json_trunc
        for ev in truncation_events:
            _emit_subnorm_metric_server(
                'cocreation_file_truncated',
                value_int=int(ev.get('original_chars') or 0),
                value_text=_json_trunc.dumps(ev, ensure_ascii=False)[:500],
                element_num=element_num,
                user_id=current_user.id,
                session_id=None,
            )
        resp['truncation_info'] = {
            'tier': 'FREE',
            'events': truncation_events,
            'tier_limit_chars': MAX_CHARS_PER_FILE,
            'pro_limit_chars': 15000,
        }
    return jsonify(resp)

@app.route("/api/clear_specs", methods=["POST"])
def clear_specs():
    data = request.json or {}
    element_num = data.get('element_num', 1)
    try:
        element_num = int(element_num)
    except (ValueError, TypeError):
        element_num = 1
    if element_num not in (1, 2, 3, 4):
        return jsonify({"error": "Elemento inválido"}), 400
    _clear_chat_spec(element_num)
    _clear_cocreation_state(element_num)
    return jsonify({"success": True})

@app.route("/api/cocreation/preview", methods=["POST"])
def api_cocreation_preview():
    data = request.get_json(silent=True) or {}
    element_num = data.get('element_num', 1)
    try:
        element_num = int(element_num)
    except (ValueError, TypeError):
        element_num = 1
    if element_num not in (1, 2, 3, 4):
        return jsonify({"error": "Elemento inválido"}), 400
    if _is_cocreation_skipped_session():
        return jsonify({"skipped_session": True, "state": "skipped"}), 200
    user_specs = _get_chat_spec(element_num)
    if not user_specs or not user_specs.strip():
        return jsonify({"error": "No hay archivo cargado para revisar"}), 400
    current_state = _get_cocreation_state(element_num)
    if current_state in ('confirmed', 'skipped'):
        return jsonify({"already_resolved": True, "state": current_state}), 200
    user_tier = current_user.tier if current_user.is_authenticated else 'ANON'
    uid = current_user.id if current_user.is_authenticated else None
    sid = None if uid else _ensure_anon_sid()
    try:
        from ai_helper import generate_cocreation_preview
        preview_text = generate_cocreation_preview(element_num, user_specs, user_tier)
    except Exception as e:
        logger.error(f"cocreation preview generation failed: {e}")
        try:
            _emit_subnorm_metric_server(
                'cocreation_preview_failed',
                value_int=len(user_specs or ''),
                value_text=f"tier={user_tier}|reason=exception|err={str(e)[:200]}",
                element_num=element_num,
                user_id=uid,
                session_id=sid,
            )
        except Exception:
            pass
        return jsonify({"error": "No se pudo generar la vista previa de cocreación. Puedes continuar normalmente."}), 500
    if not preview_text or not preview_text.strip():
        try:
            _emit_subnorm_metric_server(
                'cocreation_preview_failed',
                value_int=len(user_specs or ''),
                value_text=f"tier={user_tier}|reason=empty_response",
                element_num=element_num,
                user_id=uid,
                session_id=sid,
            )
        except Exception:
            pass
        return jsonify({"error": "La vista previa de cocreación vino vacía. Continúa normalmente."}), 500
    _set_cocreation_state(element_num, 'shown')
    _emit_subnorm_metric_server(
        'cocreation_preview_shown',
        value_int=len(user_specs),
        value_text=f"tier={user_tier}|chars={len(user_specs)}",
        element_num=element_num,
        user_id=uid,
        session_id=sid,
    )
    return jsonify({
        "success": True,
        "preview_text": preview_text,
        "tier": user_tier,
        "state": "shown",
    })

@app.route("/api/cocreation/decision", methods=["POST"])
def api_cocreation_decision():
    data = request.get_json(silent=True) or {}
    element_num = data.get('element_num', 1)
    try:
        element_num = int(element_num)
    except (ValueError, TypeError):
        element_num = 1
    if element_num not in (1, 2, 3, 4):
        return jsonify({"error": "Elemento inválido"}), 400
    decision = (data.get('decision') or '').strip().lower()
    if decision not in ('confirm', 'refine', 'skip'):
        return jsonify({"error": "decision inválida"}), 400
    state_map = {'confirm': 'confirmed', 'refine': 'refining', 'skip': 'skipped'}
    new_state = state_map[decision]
    _set_cocreation_state(element_num, new_state)
    if decision == 'skip':
        _set_cocreation_skipped_session()
    uid = current_user.id if current_user.is_authenticated else None
    sid = None if uid else _ensure_anon_sid()
    metric_type_map = {
        'confirm': 'cocreation_decision_confirm',
        'refine': 'cocreation_decision_refine',
        'skip': 'cocreation_decision_skip',
    }
    _emit_subnorm_metric_server(
        metric_type_map[decision],
        value_int=int(element_num),
        value_text=f"new_state={new_state}",
        element_num=element_num,
        user_id=uid,
        session_id=sid,
    )
    return jsonify({"success": True, "state": new_state})

@app.route("/api/ai_status")
def api_ai_status():
    with _ai_queue_lock:
        waiting = _ai_queue_waiting[0]
        active = _ai_active_count[0]
    max_ai = int(Config.get('MAX_CONCURRENT_AI', '3'))
    return jsonify({"queue_waiting": waiting, "max_concurrent": max_ai, "active": active})

@app.route("/api/beta/feedback", methods=["POST"])
def api_beta_feedback():
    data = request.get_json(silent=True) or {}
    rating = data.get("rating")
    if rating is not None:
        try:
            rating = int(rating)
            if rating < 1 or rating > 5:
                rating = None
        except (ValueError, TypeError):
            rating = None
    comment = (data.get("comment") or "").strip()[:2000]
    page = (data.get("page") or "")[:200]
    element_num = data.get("element_num")
    if element_num is not None:
        try:
            element_num = int(element_num)
            if element_num < 1 or element_num > 4:
                element_num = None
        except (ValueError, TypeError):
            element_num = None
    generated_file = (data.get("generated_file") or "")[:300]
    feedback_type = data.get("feedback_type", "widget")
    if feedback_type not in ('widget', 'post_gen'):
        feedback_type = 'widget'
    if not rating and not comment:
        return jsonify({"error": "Sin datos"}), 400
    uid = current_user.id if current_user.is_authenticated else None
    sid = _ensure_anon_sid()[:100] if not uid else None
    fb = BetaFeedback(
        user_id=uid, session_id=sid, rating=rating, comment=comment,
        page=page, element_num=element_num, generated_file=generated_file,
        feedback_type=feedback_type
    )
    db.session.add(fb)
    db.session.commit()
    return jsonify({"ok": True, "id": fb.id})

@app.route("/api/beta/metric", methods=["POST"])
def api_beta_metric():
    data = request.get_json(silent=True) or {}
    VALID_METRIC_TYPES = ('page_enter', 'generation_complete', 'download_click', 'page_leave_no_action', 'time_to_first_gen', 'confirm_modal_shown', 'confirm_modal_confirmed', 'confirm_modal_cancelled', 'subnorm_negotiation_shown', 'subnorm_user_duration_confirmed', 'subnorm_default_120_accepted', 'cocreation_file_truncated', 'cocreation_file_size_rejected', 'cocreation_upsell_clicked', 'cocreation_preview_shown', 'cocreation_decision_confirm', 'cocreation_decision_refine', 'cocreation_decision_skip', 'cocreation_preview_failed')
    metric_type = (data.get("type") or "")[:50]
    if metric_type not in VALID_METRIC_TYPES:
        return jsonify({"error": "type inválido"}), 400
    uid = current_user.id if current_user.is_authenticated else None
    sid = _ensure_anon_sid()[:100] if not uid else None
    el_num = data.get("element_num")
    if el_num is not None:
        try:
            el_num = int(el_num)
            if el_num < 1 or el_num > 4:
                el_num = None
        except (ValueError, TypeError):
            el_num = None
    v_int = data.get("value_int")
    if v_int is not None:
        try:
            v_int = int(v_int)
            if v_int < 0 or v_int > 86400:
                v_int = None
        except (ValueError, TypeError):
            v_int = None
    m = BetaMetric(
        user_id=uid, session_id=sid, metric_type=metric_type,
        element_num=el_num,
        value_int=v_int,
        value_text=(data.get("value_text") or "")[:500]
    )
    db.session.add(m)
    db.session.commit()
    return jsonify({"ok": True})

@app.route("/api/admin/beta_feedback")
@login_required
def api_admin_beta_feedback():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    feedbacks = BetaFeedback.query.order_by(BetaFeedback.created_at.desc()).limit(200).all()
    result = []
    for fb in feedbacks:
        user_name = ""
        if fb.user_id:
            u = db.session.get(User, fb.user_id)
            user_name = u.full_name if u else f"User #{fb.user_id}"
        result.append({
            "id": fb.id, "user": user_name or "Anónimo", "user_id": fb.user_id,
            "rating": fb.rating, "comment": fb.comment, "page": fb.page,
            "element_num": fb.element_num, "generated_file": fb.generated_file,
            "type": fb.feedback_type,
            "date": fb.created_at.strftime("%Y-%m-%d %H:%M") if fb.created_at else ""
        })
    return jsonify(result)

@app.route("/api/admin/beta_metrics")
@login_required
def api_admin_beta_metrics():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    total_sessions = db.session.query(db.func.count(db.func.distinct(
        db.case((BetaMetric.user_id.isnot(None), db.cast(BetaMetric.user_id, db.String)),
                 else_=BetaMetric.session_id)
    ))).scalar() or 0
    sessions_logueadas = db.session.query(db.func.count(db.func.distinct(BetaMetric.user_id))).filter(
        BetaMetric.user_id.isnot(None)
    ).scalar() or 0
    sessions_anonimas = max(int(total_sessions) - int(sessions_logueadas), 0)
    sessions_breakdown = {
        'total': int(total_sessions),
        'logueados': int(sessions_logueadas),
        'anonimos': int(sessions_anonimas),
    }
    breakdown_rows = db.session.query(
        BetaMetric.metric_type,
        db.func.count().label('total'),
        db.func.sum(db.case((BetaMetric.user_id.isnot(None), 1), else_=0)).label('logueados'),
        db.func.sum(db.case((BetaMetric.user_id.is_(None), 1), else_=0)).label('anonimos'),
    ).group_by(BetaMetric.metric_type).all()
    breakdown = {}
    for mt, total, log, ano in breakdown_rows:
        breakdown[mt] = {
            'total': int(total or 0),
            'logueados': int(log or 0),
            'anonimos': int(ano or 0),
        }
    page_enters = breakdown.get('page_enter', {}).get('total', 0)
    generations = breakdown.get('generation_complete', {}).get('total', 0)
    downloads = breakdown.get('download_click', {}).get('total', 0)
    abandonments = breakdown.get('page_leave_no_action', {}).get('total', 0)
    avg_time_rows = db.session.query(db.func.avg(BetaMetric.value_int)).filter(
        BetaMetric.metric_type == 'time_to_first_gen'
    ).scalar()
    avg_time = round(avg_time_rows) if avg_time_rows else None
    by_element = db.session.query(
        BetaMetric.element_num, BetaMetric.metric_type, db.func.count()
    ).filter(BetaMetric.element_num.isnot(None)).group_by(
        BetaMetric.element_num, BetaMetric.metric_type
    ).all()
    element_data = {}
    for el, mt, cnt in by_element:
        if el not in element_data:
            element_data[el] = {}
        element_data[el][mt] = int(cnt)
    element_table = []
    for el in sorted(element_data.keys()):
        ed = element_data[el]
        entradas = ed.get('page_enter', 0)
        rebotes = ed.get('page_leave_no_action', 0)
        gens = ed.get('generation_complete', 0)
        dls = ed.get('download_click', 0)
        rebote_pct = round((rebotes / entradas) * 100, 1) if entradas else 0.0
        conv_pct = round((dls / entradas) * 100, 1) if entradas else 0.0
        element_table.append({
            'element_num': el,
            'entradas': entradas,
            'rebotes': rebotes,
            'rebote_pct': rebote_pct,
            'generaciones': gens,
            'descargas': dls,
            'conversion_pct': conv_pct,
        })
    return jsonify({
        "total_sessions": total_sessions,
        "page_enters": page_enters,
        "generations": generations,
        "downloads": downloads,
        "abandonments": abandonments,
        "avg_time_to_first_gen_seconds": avg_time,
        "by_element": element_data,
        "breakdown": breakdown,
        "sessions_breakdown": sessions_breakdown,
        "element_table": element_table,
    })


@app.route("/api/admin/insights")
@login_required
def api_admin_insights():
    if not current_user.is_admin:
        return jsonify({"error": "No autorizado"}), 403
    try:
        cards = analytics_rules.evaluate_all(db, BetaMetric, BetaFeedback, User)
    except Exception as e:
        logger.exception("insights eval failed")
        return jsonify({"error": f"Error evaluando reglas: {type(e).__name__}"}), 500
    severity_counts = {'critical': 0, 'warning': 0, 'info': 0}
    for c in cards:
        sev = c.get('severity', 'info')
        if sev in severity_counts:
            severity_counts[sev] += 1
    return jsonify({
        "cards": cards,
        "summary": severity_counts,
        "generated_at": datetime.utcnow().isoformat(),
    })

@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.json
    element_num = data.get("element_num", 1)
    message = data.get("message", "")
    history = data.get("history", [])
    course_topic = data.get("course_topic", None)
    template_style = data.get("template_style", "corporativo")
    user_specs = data.get("user_specs", "")
    user_considerations = (data.get("user_considerations", "") or "").strip()[:2000]
    contexto_institucional = (data.get("contexto_institucional", "") or "").strip()
    if contexto_institucional:
        session['contexto_institucional'] = contexto_institucional
    is_quick_action = bool(data.get("is_quick_action", False))

    if not message:
        return jsonify({"error": "Mensaje vacío"})

    uid = current_user.id if current_user.is_authenticated else None
    track_event('IA', 'Generación IA Iniciada', user_id=uid, extra_data={'modulo': element_num, 'tier': current_user.tier if current_user.is_authenticated else 'ANON'})
    try:
        _msg_len = len(message or '')
        _has_specs = bool((user_specs or '').strip())
        track_event('Funnel', 'Submit Personalizacion', user_id=uid, extra_data={
            'modulo': element_num,
            'tier': current_user.tier if current_user.is_authenticated else 'ANON',
            'mensaje_chars': _msg_len,
            'tiene_specs': _has_specs
        })
    except Exception:
        pass

    is_authenticated = current_user.is_authenticated
    anon_is_last_turn = False

    if element_num in (2, 3):
        if is_authenticated:
            _has_element_access = current_user.tier in ('PRO', 'PREMIUM') or current_user.has_alacarte(element_num)
            if not _has_element_access:
                return jsonify({
                    "error": "El plan gratuito solo incluye acceso al Módulo 1 (Carta Descriptiva). Para acceder a los demás módulos, elige un plan PRO, PREMIUM o adquiere el elemento a la carta.",
                    "redirect": "/precios"
                })
        if is_authenticated and not session.get('master_doc', '') and getattr(current_user, 'active_course_session_id', None):
            try:
                _acs_recover = db.session.get(CourseSession, current_user.active_course_session_id)
                if _acs_recover and _acs_recover.master_doc:
                    session['master_doc'] = _acs_recover.master_doc
                    if _acs_recover.topic and not session.get('master_doc_topic'):
                        session['master_doc_topic'] = _acs_recover.topic
            except Exception:
                pass
        if not session.get('master_doc', ''):
            _producto = "los Instrumentos de Evaluación" if element_num == 2 else "los Manuales del Curso"
            return jsonify({
                "response": f"📋 **Para que {_producto} sean coherentes con tu curso**, te recomendamos generar primero tu **Carta Descriptiva** en el Elemento 1.\n\nEs el documento maestro que define tema, objetivos, tiempos y técnicas. Una vez listo, {_producto} se generan alineados a él, sin contradicciones de contenido ni de duración.\n\n👉 **Cambia al [Elemento 1 — Carta Descriptiva](/elemento/1)** y escribe el tema de tu curso para comenzar. Tu primera Carta Descriptiva es gratuita.",
                "generated_file": None,
                "detected_topic": None,
                "requires_carta": True
            })

    free_is_last_turn = False
    alacarte_active = False
    if is_authenticated and current_user.tier == 'PRO':
        db.session.refresh(current_user)
        remaining = current_user.pro_courses_remaining or 0
        active = (current_user.pro_active_course or '').strip()
        if remaining <= 0 and not active:
            track_event('Conversion', 'PRO Curso Agotado', user_id=current_user.id,
                        extra_data={'tipo': 'posible_escalamiento_tier'})
            return jsonify({
                "error": "Has usado el curso incluido en tu plan PRO. Para diseñar otro curso, adquiere un nuevo PRO por curso o cambia a PREMIUM para generación ilimitada.",
                "redirect": "/precios"
            })
    should_inc_free = False
    should_inc_anon = False
    should_use_alacarte = False
    if is_authenticated and current_user.tier == 'FREE':
        if current_user.has_alacarte(element_num):
            should_use_alacarte = True
        elif element_num in (2, 3, 4):
            return jsonify({
                "error": "El plan gratuito solo incluye acceso al Módulo 1 (Carta Descriptiva). Para acceder a los demás módulos, elige un plan PRO, PREMIUM o adquiere el elemento a la carta.",
                "redirect": "/precios"
            })
        elif current_user.chat_usage_count >= 3:
            flash('Prueba gratuita agotada. Elige un plan para continuar.', 'error')
            return jsonify({"error": "Prueba gratuita agotada. Elige un plan para continuar.", "redirect": "/precios"})
        else:
            should_inc_free = True
    elif not is_authenticated:
        anon_count = session.get('anon_chat_count', 0)
        if anon_count >= 3:
            return jsonify({
                "response": "Has alcanzado el l\u00edmite de tu prueba gratuita (3 consultas). Reg\u00edstrate gratis para seguir generando contenido.",
                "generated_file": None,
                "detected_topic": None,
                "user_tier": "FREE",
                "limit_warning": True,
                "is_anonymous": True
            })
        should_inc_anon = True

    if element_num != 4:
        extracted = _extract_course_info_from_message(message)
        if extracted:
            _merge_course_info(extracted)
        if not _is_course_info_complete() and session.get('course_info_paso0_shown'):
            missing_now = _get_missing_course_fields()
            if missing_now:
                conv_ctx = ""
                try:
                    recent = history[-6:] if history else []
                    conv_lines = []
                    for _hm in recent:
                        _role = _hm.get('role', 'user')
                        _cnt = (_hm.get('content') or '')[:500]
                        conv_lines.append(f"{_role}: {_cnt}")
                    conv_ctx = "\n".join(conv_lines)
                except Exception:
                    pass
                ai_extracted = extract_course_info_ai(
                    message=message,
                    missing_fields=missing_now,
                    captured_fields=_get_course_info(),
                    conversation_context=conv_ctx,
                )
                if ai_extracted:
                    for _k, _v in list(ai_extracted.items()):
                        if _v == '__BACK_REF__':
                            if _k == 'nombre_curso':
                                _fallback = course_topic or _extract_course_name(history)
                                if _fallback:
                                    ai_extracted[_k] = _fallback
                                else:
                                    del ai_extracted[_k]
                            else:
                                del ai_extracted[_k]
                        elif _v == '__SKIP__':
                            ai_extracted[_k] = COURSE_INFO_SKIP
                    _merge_course_info(ai_extracted)
        if not _get_course_info().get('nombre_curso'):
            _topic_fb = course_topic or _extract_course_name(history)
            if _topic_fb:
                _merge_course_info({'nombre_curso': _topic_fb})

        gen_intent = _is_generation_intent(message, is_quick_action=is_quick_action)
        if not _is_course_info_complete() and gen_intent:
            pending = session.get('pending_action')
            if is_quick_action and not pending:
                session['pending_action'] = {
                    'message': message,
                    'element_num': element_num,
                    'template_style': template_style,
                }
            first_time = not session.get('course_info_paso0_shown', False)
            missing = _get_missing_course_fields()
            response_text = _paso0_message(missing=missing, first_time=first_time)
            session['course_info_paso0_shown'] = True
            return jsonify({
                "response": response_text,
                "generated_file": None,
                "detected_topic": course_topic,
                "course_info_pending": True,
            })

        pending = session.get('pending_action')
        if pending and isinstance(pending, dict) and _is_course_info_complete():
            session.pop('pending_action', None)
            message = pending.get('message') or message
            element_num = pending.get('element_num') or element_num
            template_style = pending.get('template_style') or template_style
            is_quick_action = True

        ci = _get_course_info()
        if ci.get('nombre_curso') and ci['nombre_curso'] != COURSE_INFO_SKIP and not course_topic:
            course_topic = ci['nombre_curso']

    try:
        if not course_topic:
            course_topic = _extract_course_name(history)

        is_command = any(kw in message.lower() for kw in ["genera", "redacta", "checklist", "evalua", "muestra", "objetivo", "etapa"])

        if not course_topic and not is_command and len(message) < 150:
            _low = message.lower().strip()
            _is_question = ('?' in _low) or any(_low.startswith(p) for p in ('qu\u00e9', 'que ', 'qu\u00e9 ', 'c\u00f3mo', 'como ', 'c\u00f3mo ', 'por qu\u00e9', 'porqu\u00e9', 'porque', 'd\u00f3nde', 'donde', 'cu\u00e1ndo', 'cuando', 'cu\u00e1l', 'cual', 'qui\u00e9n', 'quien', 'para qu\u00e9'))
            _is_greeting = any(_low.startswith(g) for g in ('hola', 'buenas', 'buenos', 'gracias', 'ok ', 'okay', 'vale', 's\u00ed', 'si ', 'no ', 'no.', 'ayuda', 'expl\u00edcame', 'explicame', 'explica '))
            if not _is_question and not _is_greeting:
                course_topic = message.strip()

        prev_master_topic = session.get('master_doc_topic', '')
        if course_topic and prev_master_topic and _normalize_topic(course_topic) != _normalize_topic(prev_master_topic) and is_authenticated and current_user.tier in ('PRO', 'PREMIUM'):
            # Tema genuinamente distinto en cuenta PRO/PREMIUM sin haber dado clic
            # en "Iniciar Nuevo Curso": se asume que el usuario quiere un curso
            # nuevo. Se preserva el curso anterior intacto (mismo mecanismo que el
            # boton manual) y se abre uno nuevo limpio para el tema recibido.
            try:
                reset_active_course_state(current_user.id, reason='auto_topic_change')
            except Exception:
                pass
            for _k in ('course_info', 'master_doc', 'master_doc_topic', 'pending_action',
                       'pending_document', 'pending_e5_url', 'course_info_paso0_shown',
                       'subnorm_pending', 'cocreation_skip_session', 'course_logo_path'):
                session.pop(_k, None)
            for _k in [k for k in list(session.keys()) if isinstance(k, str) and k.startswith('pending_')]:
                session.pop(_k, None)
            for _e in (1, 2, 3, 4):
                session.pop(f'cocreation_state_e{_e}', None)
            session.modified = True
            try:
                logger.info(f"AUTO_NEW_COURSE_ON_TOPIC_CHANGE user_id={current_user.id} prev='{prev_master_topic}' new='{course_topic}'")
            except Exception:
                pass
            prev_master_topic = ''
        if course_topic and prev_master_topic and _normalize_topic(course_topic) != _normalize_topic(prev_master_topic):
            # Si el usuario autenticado tiene un curso ACTIVO con Carta Descriptiva ya
            # persistida, NO purgamos el master_doc por una discrepancia del course_topic
            # recibido: ese valor puede venir de un localStorage desactualizado del curso
            # anterior, y borrar el master_doc bloquearia Contrato/Lista/IEC/Manual del
            # curso vigente. El cambio real de curso se hace con "Iniciar Nuevo Curso".
            _active_master = None
            _acs_guard = None
            if is_authenticated and getattr(current_user, 'active_course_session_id', None):
                try:
                    _acs_guard = db.session.get(CourseSession, current_user.active_course_session_id)
                    if _acs_guard and _acs_guard.master_doc:
                        _active_master = _acs_guard.master_doc
                except Exception:
                    _active_master = None
            if _active_master:
                if not session.get('master_doc'):
                    session['master_doc'] = _active_master
                if not session.get('master_doc_topic') and getattr(_acs_guard, 'topic', None):
                    session['master_doc_topic'] = _acs_guard.topic
            else:
                session.pop('master_doc', None)
                session.pop('master_doc_topic', None)
                try:
                    _uid_log = current_user.id if is_authenticated else None
                    logger.info(f"NEW_TOPIC_RESET user_id={_uid_log} prev='{prev_master_topic}' new='{course_topic}'")
                    track_event('Funnel', 'Tema Nuevo Detectado', user_id=_uid_log, extra_data={
                        'tema_anterior': prev_master_topic[:80],
                        'tema_nuevo': course_topic[:80]
                    })
                except Exception:
                    pass

        if is_authenticated and current_user.tier == 'PRO':
            active = (current_user.pro_active_course or '').strip()
            _ci_lock = _get_course_info() or {}
            _lock_name = (_ci_lock.get('nombre_curso') or '').strip()
            if (not _lock_name) or _is_placeholder_value(_lock_name) or _lock_name == COURSE_INFO_SKIP:
                _lock_name = _sanitize_course_name_from_message(course_topic or '') or (course_topic or '')
            _lock_name = (_lock_name or '').strip()[:300]
            if active and _lock_name and _normalize_topic(_lock_name) != _normalize_topic(active):
                return jsonify({
                    "error": f"Tu plan PRO está activo para el curso '{active}'. Para generar otro curso, adquiere un nuevo PRO por curso o actualiza a PREMIUM.",
                    "redirect": "/precios",
                    "pro_locked_course": active
                })
            if active and not course_topic:
                course_topic = active
            if not active and _lock_name:
                try:
                    rows = db.session.execute(
                        db.text('UPDATE "user" SET pro_active_course = :topic, pro_courses_remaining = pro_courses_remaining - 1 WHERE id = :uid AND pro_courses_remaining > 0 AND (pro_active_course IS NULL OR pro_active_course = \'\')'),
                        {"topic": _lock_name, "uid": current_user.id}
                    )
                    db.session.commit()
                except Exception as _lock_err:
                    try:
                        db.session.rollback()
                    except Exception:
                        pass
                    logger.error(f"PRO_LOCK_FAIL user={current_user.id} name_len={len(_lock_name)} err={type(_lock_err).__name__}: {_lock_err}")
                    track_event('Error', 'PRO Lock Fallo', user_id=current_user.id,
                                extra_data={'tipo': 'fallo_tecnico', 'name_len': len(_lock_name)})
                    return jsonify({
                        "error": "No pudimos iniciar tu curso en este momento. Tu crédito sigue intacto. Inténtalo de nuevo en unos segundos; si el problema continúa, escríbenos y lo resolvemos."
                    })
                if rows.rowcount == 0:
                    track_event('Conversion', 'PRO Curso Agotado', user_id=current_user.id,
                                extra_data={'tipo': 'posible_escalamiento_tier'})
                    return jsonify({
                        "error": "Has usado el curso incluido en tu plan PRO. Para diseñar otro curso, adquiere un nuevo PRO por curso o cambia a PREMIUM para generación ilimitada.",
                        "redirect": "/precios"
                    })
                db.session.refresh(current_user)
                logger.info(f"PRO course locked for user {current_user.id}: '{_lock_name}', remaining={current_user.pro_courses_remaining}")

        prompt_to_send = message
        if course_topic:
            if message.strip() == course_topic:
                if element_num == 1:
                    _ci_horas_raw = (session.get('course_info') or {}).get('num_horas', '').strip()
                    _ci_horas = _ci_horas_raw if not _is_placeholder_value(_ci_horas_raw) else ''
                    _dur_clause = f"de {_ci_horas}" if _ci_horas else "respetando la REGLA SOBERANA DE DURACION de tus instrucciones de sistema (120 min default si el usuario no prescribio otra duracion)"
                    prompt_to_send = f"El tema del curso es: '{course_topic}'. Actúa de inmediato. Genera la Carta Descriptiva completa {_dur_clause} con los encabezados Markdown exactos que se te pidieron en tus instrucciones de sistema. No me hagas preguntas, comienza a generar las tablas y los objetivos."
                elif element_num == 2:
                    prompt_to_send = f"El tema del curso es: '{course_topic}'. Genera los instrumentos de evaluación requeridos."
                elif element_num == 3:
                    prompt_to_send = f"El tema del curso es: '{course_topic}'. Genera la estructura del Manual del Instructor y Participante."
            elif is_command:
                prompt_to_send = f"El tema del curso es: '{course_topic}'. {message}"

        if element_num != 4 and prompt_to_send != message and not _is_course_info_complete():
            if not session.get('pending_action'):
                session['pending_action'] = {
                    'message': message,
                    'element_num': element_num,
                    'template_style': template_style,
                }
            first_time = not session.get('course_info_paso0_shown', False)
            missing = _get_missing_course_fields()
            response_text = _paso0_message(missing=missing, first_time=first_time)
            session['course_info_paso0_shown'] = True
            return jsonify({
                "response": response_text,
                "generated_file": None,
                "detected_topic": course_topic,
                "course_info_pending": True,
            })

        docs = get_reference_docs()

        is_contrato_msg = "contrato" in message.lower() and element_num == 1
        is_lista_req_msg = ("lista de verificaci" in message.lower() or "lista de requerim" in message.lower()) and element_num == 1
        master_doc_val = session.get('master_doc', '')
        if not master_doc_val and is_authenticated and getattr(current_user, 'active_course_session_id', None):
            # Red de seguridad: la Carta Descriptiva del curso ACTIVO es la fuente
            # autoritativa. Si la sesion perdio el master_doc (p.ej. localStorage del
            # curso anterior), lo rehidratamos antes de decidir bloquear Contrato/Lista.
            try:
                _acs_md = db.session.get(CourseSession, current_user.active_course_session_id)
                if _acs_md and _acs_md.master_doc:
                    master_doc_val = _acs_md.master_doc
                    session['master_doc'] = _acs_md.master_doc
                    if _acs_md.topic and not session.get('master_doc_topic'):
                        session['master_doc_topic'] = _acs_md.topic
            except Exception:
                pass
        if (is_contrato_msg or is_lista_req_msg) and not master_doc_val:
            _producto = "Contrato de Aprendizaje" if is_contrato_msg else "Lista de Verificación de Requerimientos"
            return jsonify({
                "response": f"📋 **Para garantizar la congruencia de todos tus productos**, te recomiendo generar primero tu **Carta Descriptiva**.\n\nEscribe el tema de tu curso (por ejemplo: *Repostería Básica*, *Soldadura Industrial*, *Primeros Auxilios*) y generaré tu Carta Descriptiva completa. Una vez lista, podrás generar el {_producto} con toda la información alineada.\n\n👉 **Escribe tu tema abajo para comenzar.**",
                "generated_file": None,
                "detected_topic": course_topic,
                "requires_carta": True
            })
        master_doc = master_doc_val if element_num in (2, 3) or is_contrato_msg or is_lista_req_msg else None

        max_concurrent_ai = int(Config.get('MAX_CONCURRENT_AI', '3'))

        with _ai_condition:
            _ai_queue_waiting[0] += 1
            waiting_count = _ai_queue_waiting[0]

        acquired = False
        with _ai_condition:
            _ai_queue_waiting[0] = max(0, _ai_queue_waiting[0] - 1)
            if _ai_active_count[0] < max_concurrent_ai:
                _ai_active_count[0] += 1
                acquired = True
            else:
                _ai_condition.wait(timeout=15)
                if _ai_active_count[0] < max_concurrent_ai:
                    _ai_active_count[0] += 1
                    acquired = True

        if not acquired:
            logger.warning(f"AI concurrency limit — too many concurrent generations (active={_ai_active_count[0]}, max={max_concurrent_ai}, waiting={waiting_count})")
            try:
                sat_metric = BetaMetric(
                    user_id=uid,
                    metric_type='ai_saturation_reject',
                    element_num=element_num,
                    value_int=max_concurrent_ai,
                    value_text=f'active={_ai_active_count[0]},waiting={waiting_count}'
                )
                db.session.add(sat_metric)
                db.session.commit()
            except Exception:
                db.session.rollback()
            return jsonify({
                "error": "El sistema está procesando varias solicitudes simultáneas. Por favor intenta de nuevo en unos segundos.",
                "retry": True
            })

        try:
            specs_for_ai = user_specs or _get_chat_spec(element_num)
            current_tier = current_user.tier if is_authenticated else None
            course_info_text = _course_info_for_prompt() if element_num != 4 else None
            if element_num != 4 and course_topic and not _is_course_info_complete():
                if not session.get('pending_action'):
                    session['pending_action'] = {
                        'message': message,
                        'element_num': element_num,
                        'template_style': template_style,
                    }
                first_time = not session.get('course_info_paso0_shown', False)
                missing = _get_missing_course_fields()
                paso0_text = _paso0_message(missing=missing, first_time=first_time)
                session['course_info_paso0_shown'] = True
                return jsonify({
                    "response": paso0_text,
                    "generated_file": None,
                    "detected_topic": course_topic,
                    "course_info_pending": True,
                })
            if should_use_alacarte:
                used = current_user.use_alacarte(element_num)
                if not used:
                    return jsonify({
                        "error": "Tu crédito A la Carta para este módulo no está disponible. Adquiere otro o actualiza tu plan.",
                        "redirect": "/precios"
                    })
                alacarte_active = True
            if should_inc_free:
                current_user.chat_usage_count += 1
                db.session.commit()
                if current_user.chat_usage_count >= 3:
                    free_is_last_turn = True
            elif should_inc_anon:
                anon_count = session.get('anon_chat_count', 0)
                session['anon_chat_count'] = anon_count + 1
                anon_is_last_turn = (anon_count + 1 >= 3)
            subnorm_pending = session.get('subnorm_pending')
            if subnorm_pending and element_num == 1:
                decision = _classify_subnorm_user_decision(message)
                _uid_for_metric = current_user.id if is_authenticated else None
                _sid_for_metric = _ensure_anon_sid()[:100] if not is_authenticated else None
                if decision == 'confirmed':
                    confirmed_min = int(subnorm_pending.get('minutes_requested') or 0)
                    if confirmed_min > 0:
                        ci = session.get('course_info') or {}
                        ci['num_horas'] = f"{confirmed_min} minutos"
                        session['course_info'] = ci
                        ts_decision = datetime.utcnow().isoformat()
                        import json as _json_subnorm
                        _vt = _json_subnorm.dumps({
                            'minutes_requested': confirmed_min,
                            'minutes_chosen': confirmed_min,
                            'ts_proposed': subnorm_pending.get('ts_proposed'),
                            'ts_decided': ts_decision,
                            'topic': subnorm_pending.get('topic') or course_topic or '',
                        }, ensure_ascii=False)
                        _emit_subnorm_metric_server(
                            'subnorm_user_duration_confirmed',
                            value_int=confirmed_min, value_text=_vt,
                            element_num=element_num,
                            user_id=_uid_for_metric, session_id=_sid_for_metric,
                        )
                        prompt_to_send = (
                            f"[CONFIRMACION DE DURACION SUB-NORMATIVA: el usuario confirmo expresamente "
                            f"{confirmed_min} minutos bajo su responsabilidad sobre el cumplimiento normativo, "
                            f"registrado en el sistema el {ts_decision}. Procede AHORA a generar el curso COMPLETO "
                            f"a {confirmed_min} minutos exactos respetando todas las reglas estructurales (objetivos, "
                            f"tecnicas, formato, encabezados Markdown obligatorios). Distribuye Apertura/Desarrollo/"
                            f"Cierre proporcionalmente sumando exactamente {confirmed_min} minutos.]\n\n"
                            + prompt_to_send
                        )
                        session.pop('subnorm_pending', None)
                elif decision == 'declined':
                    ts_decision = datetime.utcnow().isoformat()
                    import json as _json_subnorm
                    _vt = _json_subnorm.dumps({
                        'minutes_requested': int(subnorm_pending.get('minutes_requested') or 0),
                        'minutes_chosen': 120,
                        'ts_proposed': subnorm_pending.get('ts_proposed'),
                        'ts_decided': ts_decision,
                        'topic': subnorm_pending.get('topic') or course_topic or '',
                    }, ensure_ascii=False)
                    _emit_subnorm_metric_server(
                        'subnorm_default_120_accepted',
                        value_int=120, value_text=_vt,
                        element_num=element_num,
                        user_id=_uid_for_metric, session_id=_sid_for_metric,
                    )
                    ci = session.get('course_info') or {}
                    ci['num_horas'] = "120 minutos (2 horas)"
                    session['course_info'] = ci
                    prompt_to_send = (
                        f"[CONFIRMACION DE DURACION: el usuario acepto la duracion minima normativa de 120 minutos, "
                        f"registrado en el sistema el {ts_decision}. Procede AHORA a generar el curso COMPLETO a "
                        f"120 minutos exactos respetando todas las reglas estructurales.]\n\n"
                        + prompt_to_send
                    )
                    session.pop('subnorm_pending', None)
            if user_considerations:
                prompt_to_send = (
                    f"[CONSIDERACIONES DEL USUARIO — DE CUMPLIMIENTO OBLIGATORIO PARA LA IA]\n{user_considerations}\n[FIN CONSIDERACIONES]\n\n"
                    + (prompt_to_send or "")
                )
            if element_num == 4:
                _audit_catalog = _build_audit_catalog_prompt()
                if _audit_catalog:
                    prompt_to_send = _audit_catalog + "\n\n=== DOCUMENTOS A AUDITAR ===\n" + (prompt_to_send or "")
            _pres_modalidad = "ejecutiva" if "ejecutiva" in _normalize(message.lower()) else "facilitacion"
            response_text = chat_with_ai(element_num, prompt_to_send, history, docs, master_doc=master_doc, user_specs=specs_for_ai, user_tier=current_tier, course_info_text=course_info_text, presentacion_modalidad=_pres_modalidad, contexto_institucional=session.get('contexto_institucional'))
        finally:
            with _ai_condition:
                _ai_active_count[0] = max(0, _ai_active_count[0] - 1)
                _ai_condition.notify()

        subnorm_negotiation = _detect_subnorm_negotiation(response_text) if element_num == 1 else None
        if subnorm_negotiation and not session.get('subnorm_pending'):
            _minutes_req = subnorm_negotiation['minutes_requested']
            _ts_proposed = datetime.utcnow().isoformat()
            session['subnorm_pending'] = {
                'minutes_requested': _minutes_req,
                'topic': course_topic or '',
                'ts_proposed': _ts_proposed,
            }
            import json as _json_subnorm
            _vt = _json_subnorm.dumps({
                'minutes_requested': _minutes_req,
                'topic': course_topic or '',
                'ts_proposed': _ts_proposed,
            }, ensure_ascii=False)
            _emit_subnorm_metric_server(
                'subnorm_negotiation_shown',
                value_int=_minutes_req, value_text=_vt,
                element_num=element_num,
                user_id=current_user.id if is_authenticated else None,
                session_id=_ensure_anon_sid()[:100] if not is_authenticated else None,
            )
        response_text = _strip_subnorm_marker(response_text)
        if subnorm_negotiation:
            return jsonify({
                "response": response_text,
                "generated_file": None,
                "detected_topic": course_topic,
                "subnorm_pending": True,
            })

        if is_authenticated:
            estimated_tokens = (len(prompt_to_send) + len(response_text)) // 4
            if element_num == 3 and len(response_text) > 15000:
                estimated_tokens = estimated_tokens * 4
            current_user.add_token_usage(current_user.tier, estimated_tokens)
            db.session.commit()

        user_tier = current_user.tier if is_authenticated else "FREE"
        effective_tier = "PRO" if alacarte_active else user_tier
        result = {"response": response_text, "generated_file": None, "detected_topic": None, "user_tier": user_tier}

        if course_topic:
            result["detected_topic"] = course_topic

        NO_DOCX_TITLES = ["Objetivo_General", "Objetivos_Particulares", "Tecnicas_Instruccionales", "Etapas_del_Curso"]
        is_audit = element_num == 4
        is_info_only = is_command and any(kw in message.lower() for kw in ["checklist", "muestra"])
        is_contrato = "contrato" in message.lower() and element_num == 1
        is_lista_req = ("lista de verificaci" in message.lower() or "lista de requerim" in message.lower()) and element_num == 1

        _msg_norm_lower = _normalize(message.lower())
        is_presentacion = element_num == 3 and "presentacion" in _msg_norm_lower and "manual" not in _msg_norm_lower

        if is_contrato and len(response_text) > 50:
            try:
                course_name = _canonical_course_name(course_topic, message)
                user_name = current_user.full_name if is_authenticated else ""
                _ci_canon = _course_info_to_dict_for_docx() or {}
                def _ci_clean(_v):
                    return '' if _is_placeholder_value(_v) else _v
                master_doc_text = session.get('master_doc', '')
                duracion_curso = "120 minutos (2 horas)"
                if master_doc_text:
                    dur_match = re.search(r'DURACIÓN:\s*(.+)', master_doc_text)
                    if dur_match:
                        duracion_curso = dur_match.group(1).strip()

                objetivo_curso = ""
                criterios_evaluacion = ""
                derechos_participante = []
                lines = response_text.split("\n")
                for idx, line in enumerate(lines):
                    line_lower = line.lower().strip()
                    if "objetivo" in line_lower and ("general" in line_lower or "curso" in line_lower):
                        after_colon = line.split(":", 1)[1].strip() if ":" in line else ""
                        if after_colon and len(after_colon) > 10:
                            objetivo_curso = after_colon
                        elif idx + 1 < len(lines) and lines[idx + 1].strip() and not lines[idx + 1].strip().startswith("#"):
                            objetivo_curso = lines[idx + 1].strip()
                    if "criterio" in line_lower and "evaluaci" in line_lower:
                        after_colon = line.split(":", 1)[1].strip() if ":" in line else ""
                        if after_colon and len(after_colon) > 10:
                            criterios_evaluacion = after_colon
                        elif idx + 1 < len(lines) and lines[idx + 1].strip() and not lines[idx + 1].strip().startswith("#"):
                            criterios_evaluacion = lines[idx + 1].strip()

                derechos_participante = _extract_list_items(response_text, "derecho")
                if not derechos_participante:
                    derechos_participante = [
                        "Recibir retroalimentación oportuna sobre su desempeño en cada momento de evaluación.",
                        "Solicitar aclaración de dudas al instructor en cualquier momento de la sesión.",
                        "Conocer los criterios y momentos de evaluación desde el inicio del curso.",
                        "Solicitar repetir la evaluación en caso de no alcanzar el resultado esperado.",
                        "Recibir trato respetuoso y equitativo durante toda la sesión.",
                    ]

                content_data = {
                    "course_info": {
                        "nombre_curso": f"Curso-Taller de {course_name}",
                        "nombre_disenador": _ci_clean(_ci_canon.get('nombre_disenador')) or user_name or "[Nombre del Diseñador]",
                        "nombre_instructor": _ci_clean(_ci_canon.get('nombre_instructor')) or user_name or "[Nombre del Instructor]",
                        "duracion": duracion_curso,
                        "lugar": _ci_clean(_ci_canon.get('lugar')) or "[Lugar de impartición]",
                        "horario": "[Horario]",
                        "fecha": _ci_clean(_ci_canon.get('periodo_imparticion')) or "[Fecha de impartición]",
                    },
                    "contenido_completo": response_text,
                    "compromisos_instructor": _extract_list_items(response_text, "instructor"),
                    "compromisos_participante": _extract_list_items(response_text, "participante"),
                    "objetivo_curso": objetivo_curso,
                    "criterios_evaluacion": criterios_evaluacion,
                    "derechos_participante": derechos_participante,
                }
                filepath, error = generate_from_template("contrato_aprendizaje", content_data, user_tier=effective_tier, logo_path=session.get('course_logo_path'))
                if filepath and not error:
                    if is_authenticated:
                        filepath = _prefix_user_file(filepath)
                        _persist_file_to_db(filepath, current_user.id)
                    else:
                        filepath = _prefix_anon_file(filepath, course_topic=course_topic)
                        session['pending_document'] = os.path.basename(filepath)
                    result["generated_file"] = os.path.basename(filepath)
            except Exception as e:
                logger.error(f"Contrato generation failed: {e}")
        elif is_lista_req and len(response_text) > 50:
            try:
                items_by_section = _parse_lista_req_sections(response_text)
                if any(items_by_section.values()):
                    ci_raw = _course_info_to_dict_for_docx()
                    def _ci_clean(_v):
                        return '' if _is_placeholder_value(_v) else _v
                    course_info_for_lista = {
                        "nombre_curso": _canonical_course_name(course_topic, message),
                        "nombre_disenador": _ci_clean(ci_raw.get("nombre_disenador", "")),
                        "nombre_instructor": _ci_clean(ci_raw.get("nombre_instructor", "")),
                        "duracion": _ci_clean(ci_raw.get("num_horas", "")),
                        "fecha": _ci_clean(ci_raw.get("periodo_imparticion", "")),
                        "lugar": _ci_clean(ci_raw.get("lugar", "")),
                        "horario": "",
                    }
                    content_data = {
                        "course_info": course_info_for_lista,
                        "items_by_section": items_by_section,
                    }
                    filepath, error = generate_from_template("lista_requerimientos", content_data, user_tier=effective_tier, logo_path=session.get('course_logo_path'))
                    if filepath and not error:
                        if is_authenticated:
                            filepath = _prefix_user_file(filepath)
                            _persist_file_to_db(filepath, current_user.id)
                        else:
                            filepath = _prefix_anon_file(filepath, course_topic=course_topic)
                            session['pending_document'] = os.path.basename(filepath)
                        result["generated_file"] = os.path.basename(filepath)
            except Exception as e:
                logger.error(f"Lista requerimientos generation failed: {e}")
        elif is_presentacion and len(response_text) > 200:
            try:
                course_name = _canonical_course_name(course_topic, message)
                _pres_modalidad = "ejecutiva" if "ejecutiva" in _msg_norm_lower else "facilitacion"
                slides, _curso_datos = _parse_slides_from_response(response_text)
                if slides:
                    filepath = generate_pptx_from_slides(slides, course_name=course_name, modalidad=_pres_modalidad, logo_path=session.get('course_logo_path'), curso_datos=_curso_datos)
                    try:
                        _pex_uid = current_user.id if is_authenticated else None
                        _pex_rows = pop_last_pexels_usage()
                        if _pex_rows:
                            with db.session.begin_nested():
                                for _pex in _pex_rows:
                                    db.session.add(PexelsUsage(
                                        user_id=_pex_uid,
                                        query=(_pex.get('query') or '')[:300],
                                        photo_id=str(_pex.get('photo_id') or '')[:40],
                                        photographer=(_pex.get('photographer') or '')[:200],
                                        photographer_url=(_pex.get('photographer_url') or '')[:500],
                                        photo_url=(_pex.get('url') or '')[:500],
                                    ))
                            db.session.commit()
                    except Exception as _pex_err:
                        try:
                            db.session.rollback()
                        except Exception:
                            pass
                        logger.warning(f"PexelsUsage persist failed: {_pex_err}")
                    if is_authenticated:
                        filepath = _prefix_user_file(filepath)
                        _persist_file_to_db(filepath, current_user.id)
                    else:
                        filepath = _prefix_anon_file(filepath, course_topic=course_topic)
                        session['pending_document'] = os.path.basename(filepath)
                    result["generated_file"] = os.path.basename(filepath)
            except Exception:
                pass
        elif is_audit and len(response_text) > 100:
            try:
                import json as json_lib
                json_match = re.search(r'```json\s*(\{.*\})\s*```', response_text, re.DOTALL)
                if not json_match:
                    json_match = re.search(r'(\{.*"reactivos"\s*:\s*\{.*\}.*\})', response_text, re.DOTALL)
                if not json_match:
                    json_match = re.search(r'(\{[^{}]*"total_si"\s*:\s*\d+[^{}]*\})', response_text, re.DOTALL)
                if json_match:
                    ai_data = json_lib.loads(json_match.group(1))
                    for k in ["total_si", "total_no"]:
                        ai_data[k] = int(ai_data.get(k, 0))
                    for k in ["e1_productos_si", "e1_productos_no", "e1_conocimientos_si", "e1_conocimientos_no",
                              "e1_actitud_si", "e1_actitud_no", "e2_productos_si", "e2_productos_no",
                              "e2_conocimientos_si", "e2_conocimientos_no", "e3_productos_si", "e3_productos_no"]:
                        ai_data[k] = int(ai_data.get(k, 0))
                    ai_data["porcentaje"] = float(ai_data.get("porcentaje", 0))
                    computed_si = (ai_data["e1_productos_si"] + ai_data["e1_conocimientos_si"] + ai_data["e1_actitud_si"] +
                                   ai_data["e2_productos_si"] + ai_data["e2_conocimientos_si"] +
                                   ai_data["e3_productos_si"])
                    if ai_data["total_si"] != computed_si:
                        ai_data["total_si"] = computed_si
                        ai_data["total_no"] = 145 - computed_si
                        ai_data["porcentaje"] = round((computed_si / 145) * 100, 1)
                    out_name = _fill_diagnostico_template(ai_data, course_topic=_canonical_course_name(course_topic, message, fallback="EC0301"))
                    if out_name:
                        result["generated_file"] = out_name
                        if is_authenticated:
                            _persist_file_to_db(os.path.join("generated_docs", out_name), current_user.id)
                        else:
                            session['pending_document'] = out_name
                if not result.get("generated_file"):
                    course_name = _canonical_course_name(course_topic, message, fallback="EC0301")
                    sections = _parse_response_to_sections(response_text)
                    if sections:
                        filepath = generate_custom_docx("Diagnostico_EC0301", sections, course_name=course_name, user_tier=effective_tier, template_style=template_style, logo_path=session.get('course_logo_path'))
                        if is_authenticated:
                            filepath = _prefix_user_file(filepath)
                            _persist_file_to_db(filepath, current_user.id)
                        else:
                            filepath = _prefix_anon_file(filepath, course_topic=course_topic)
                            session['pending_document'] = os.path.basename(filepath)
                        result["generated_file"] = os.path.basename(filepath)
                        _fname = os.path.basename(filepath)
                        logger.info(f"DOC_GEN_OK element=4 doc_title=Diagnostico_EC0301 file_ext={_fname.rsplit('.',1)[-1] if '.' in _fname else 'unknown'} file_len={len(_fname)} sections={len(sections)} user_id={current_user.id if is_authenticated else 'anon'}")
                    else:
                        logger.warning(f"DOC_GEN_NO_SECTIONS element=4 doc_title=Diagnostico_EC0301 response_len={len(response_text)} topic_len={len(course_topic or '')} user_id={current_user.id if is_authenticated else 'anon'}")
            except Exception as e:
                logger.warning(f"DOC_GEN_FAIL element=4 doc_title=Diagnostico_EC0301 response_len={len(response_text)} topic_len={len(course_topic or '')} user_id={current_user.id if is_authenticated else 'anon'} error_type={type(e).__name__}")
                logger.debug("DOC_GEN_FAIL traceback element=4", exc_info=True)
        elif element_num in (1, 2, 3) and not is_audit and not is_info_only and len(response_text) > 200:
            try:
                course_name = _canonical_course_name(course_topic, message)
                if course_name == "[Tema por definir]":
                    logger.warning(f"COURSE_NAME_FALLBACK element={element_num} message_preview={message[:80]!r} user_id={current_user.id if is_authenticated else 'anon'}")
                doc_title = _determine_doc_title(element_num, message, len(history) == 0)
                if doc_title in NO_DOCX_TITLES:
                    logger.info(f"DOC_GEN_SKIP element={element_num} doc_title={doc_title} reason=intentional_partial_request user_id={current_user.id if is_authenticated else 'anon'}")
                else:
                    response_text_clean = _strip_conversational_preamble(response_text)
                    sections = _parse_response_to_sections(response_text_clean)
                    if sections:
                        filepath = generate_custom_docx(doc_title, sections, course_name=course_name, user_tier=effective_tier, template_style=template_style, logo_path=session.get('course_logo_path'))
                        if is_authenticated:
                            filepath = _prefix_user_file(filepath)
                            _persist_file_to_db(filepath, current_user.id)
                        else:
                            filepath = _prefix_anon_file(filepath, course_topic=course_topic)
                            session['pending_document'] = os.path.basename(filepath)
                        result["generated_file"] = os.path.basename(filepath)
                        result["refine_available"] = True
                        try:
                            _rc = session.get('refine_counts', {}) or {}
                            _rc[str(element_num)] = 0
                            session['refine_counts'] = _rc
                        except Exception:
                            pass
                        _fname = os.path.basename(filepath)
                        logger.info(f"DOC_GEN_OK element={element_num} doc_title={doc_title} file_ext={_fname.rsplit('.',1)[-1] if '.' in _fname else 'unknown'} file_len={len(_fname)} sections={len(sections)} user_id={current_user.id if is_authenticated else 'anon'}")
                        if is_authenticated and element_num == 1 and doc_title in ('Carta_Descriptiva', 'Elemento1_Carta_Descriptiva'):
                            try:
                                _disk_path = filepath if os.path.isfile(filepath) else os.path.join('generated_docs', os.path.basename(filepath))
                                _real_name = _extract_course_name_from_docx(_disk_path)
                                if _real_name:
                                    try:
                                        _real_clean = _sanitize_course_name_from_message(_real_name) or _real_name
                                    except Exception:
                                        _real_clean = _real_name
                                    try:
                                        _cs_active = CourseSession.query.get(current_user.active_course_session_id) if current_user.active_course_session_id else None
                                        if _cs_active is not None:
                                            _cs_active.topic = _real_clean
                                            _cs_active.last_activity_at = datetime.utcnow()
                                            db.session.commit()
                                    except Exception:
                                        try: db.session.rollback()
                                        except Exception: pass
                                    course_topic = _real_clean
                                    logger.info(f"COURSE_NAME_EXTRACTED user={current_user.id} cs={current_user.active_course_session_id} name={_real_clean[:80]!r}")
                            except Exception as _e_extract:
                                logger.warning(f"course_name_extract_fail user={current_user.id} err={type(_e_extract).__name__}")
                    else:
                        logger.warning(f"DOC_GEN_NO_SECTIONS element={element_num} doc_title={doc_title} response_len={len(response_text)} topic_len={len(course_topic or '')} user_id={current_user.id if is_authenticated else 'anon'}")
                if element_num == 1 and len(response_text) > 300 and not is_contrato and not is_lista_req:
                    _build_master_doc(response_text, course_topic)
                elif element_num == 1:
                    logger.warning(f"MASTER_DOC_SKIPPED element=1 response_len={len(response_text)} is_contrato={is_contrato} is_lista_req={is_lista_req} user_id={current_user.id if is_authenticated else 'anon'}")
            except Exception as e:
                logger.warning(f"DOC_GEN_FAIL element={element_num} response_len={len(response_text)} topic_len={len(course_topic or '')} user_id={current_user.id if is_authenticated else 'anon'} error_type={type(e).__name__}")
                logger.debug(f"DOC_GEN_FAIL traceback element={element_num}", exc_info=True)

        if is_authenticated and result.get("generated_file") and user_tier == 'FREE' and alacarte_active:
            remaining = getattr(current_user, f'alacarte_e{element_num}', 0)
            logger.info(f"Alacarte credit consumed (pre-reserved): user={current_user.id}, element={element_num}, remaining={remaining}")
            result["alacarte_remaining"] = remaining
            if remaining <= 0:
                result["alacarte_exhausted"] = True

        if is_authenticated and result.get("generated_file") and user_tier == 'FREE' and not alacarte_active:
            try:
                gen_file = result["generated_file"]
                gen_path = os.path.join("generated_docs", gen_file)
                import threading
                email_addr = current_user.email
                _user_id_capture = current_user.id
                _app = app._get_current_object()
                def _bg_send():
                    with _app.app_context():
                        try:
                            _send_document_email(email_addr, gen_file, gen_path, user_id=_user_id_capture)
                        except Exception as e:
                            logger.warning(f"Background email failed for {email_addr}: {e}")
                threading.Thread(target=_bg_send, daemon=True).start()
                result["email_queued"] = True
            except Exception as email_err:
                logger.warning(f"Failed to start doc email thread for {current_user.email}: {email_err}")

        if not is_authenticated and anon_is_last_turn:
            result["limit_warning"] = True
            result["is_anonymous"] = True

        if is_authenticated and free_is_last_turn:
            result["limit_warning"] = True
            result["response"] += "\n\n---\n\n⚠️ **Has utilizado tu última consulta gratuita.** Para seguir generando contenido y descargar documentos, visita la sección de Planes y Precios."

        result["doc_ready"] = bool(result.get("generated_file"))

        try:
            updated_history = list(history) if isinstance(history, list) else []
            updated_history.append({'role': 'user', 'content': message})
            updated_history.append({'role': 'assistant', 'content': response_text})
            _save_full_chat_history(
                element_num,
                updated_history,
                course_topic_val=course_topic,
                generated_file=result.get('generated_file'),
            )
        except Exception as _e_hist:
            logger.warning(f"chat_history_persist_fail user={current_user.id if is_authenticated else 'anon'} err={type(_e_hist).__name__}")

        return jsonify(result)
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.error(f"api_chat unhandled error user={current_user.id if is_authenticated else 'anon'}: {type(e).__name__}: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un problema técnico al procesar tu solicitud. Inténtalo de nuevo en unos momentos; si el problema persiste, escríbenos y lo revisamos."})

REFINE_MAX_CYCLES = 3
_REFINE_ACCION_DESC = {
    "Refinar": "REFINAR el contenido: dale mayor PRECISIÓN y CLARIDAD redactando con exactitud técnica, sin cambiar el alcance ni la duración del curso",
    "Fortalecer": "FORTALECER el contenido: eleva el NIVEL TÉCNICO y la PROFUNDIDAD de los conceptos, manteniendo la coherencia con la duración del curso",
    "Enriquecer": "ENRIQUECER el contenido: agrega CASOS PRÁCTICOS, EJEMPLOS, DINÁMICAS y EJERCICIOS aplicados al tema, sin alterar la estructura",
    "Optimizar": "OPTIMIZAR el contenido: hazlo más SIMPLE, BREVE y FÁCIL DE IMPARTIR, conservando todo lo esencial",
}
_REFINE_TARGET_LABELS = {
    "objetivos": "los Objetivos (general y particulares)",
    "temario": "el Temario / Contenido Temático",
    "actividades": "las Actividades, técnicas instruccionales y dinámicas",
    "evaluacion": "la Evaluación e instrumentos",
}


@app.route("/api/refine", methods=["POST"])
def api_refine():
    """Post-generación: mejora dirigida de un entregable conservando el resto.
    Endpoint AISLADO: no consume turnos gratuitos (chat_usage_count) ni cursos PRO.
    Tope propio de 3 ciclos por entregable (session['refine_counts'])."""
    data = request.json or {}
    try:
        element_num = int(data.get("element_num", 1) or 1)
    except (TypeError, ValueError):
        element_num = 1
    accion = (data.get("accion", "") or "").strip()
    objetivo_raw = (data.get("objetivo", "") or "").strip()
    detalle = (data.get("detalle", "") or "").strip()[:1000]
    history = data.get("history", []) or []
    course_topic = (data.get("course_topic") or "").strip() or None
    template_style = data.get("template_style", "corporativo") or "corporativo"

    if element_num not in (1, 2, 3):
        return jsonify({"error": "La mejora no está disponible para este módulo."})
    if accion not in _REFINE_ACCION_DESC:
        return jsonify({"error": "Selecciona qué tipo de mejora deseas aplicar."})

    is_authenticated = current_user.is_authenticated

    # Historial/documento persistido en el servidor (fuente de verdad post-generación).
    server_hist = {}
    try:
        server_hist = _load_full_chat_history(element_num) or {}
    except Exception:
        server_hist = {}
    has_generated = bool(server_hist.get('generated_files'))

    # --- Elegibilidad: Carta (1) = todos; IECs (2) y Manuales (3) = solo pago ---
    # Si el entregable ya se generó (server-side), el usuario conserva sus 3 ciclos de
    # mejora aunque su crédito "a la carta" ya se haya consumido al generarlo.
    if element_num in (2, 3):
        is_paid = is_authenticated and (
            current_user.tier in ("PRO", "PREMIUM") or current_user.has_alacarte(element_num)
        )
        if not (is_paid or has_generated):
            return jsonify({
                "error": "La mejora de este entregable está disponible en los planes PRO y PREMIUM. Mejora la Carta Descriptiva sin costo o actualiza tu plan para mejorar también este documento.",
                "refine_locked": True
            })

    # --- Tope de 3 ciclos por entregable ---
    counts = session.get('refine_counts', {}) or {}
    used = int(counts.get(str(element_num), 0))
    if used >= REFINE_MAX_CYCLES:
        return jsonify({
            "refine_exhausted": True,
            "refine_used": used,
            "refine_max": REFINE_MAX_CYCLES,
            "response": "Tu curso ya cuenta con la estructura suficiente para comenzar a impartirse. Te recomendamos avanzar a los entregables finales."
        })

    # --- Documento base = último mensaje 'assistant'. Preferimos el historial
    #     persistido en el servidor (post-generación verificable); si no existe,
    #     usamos el del cliente (consistente con /api/chat). ---
    def _pick_base(msgs):
        bd, bi = "", -1
        if isinstance(msgs, list):
            for i in range(len(msgs) - 1, -1, -1):
                h = msgs[i] if isinstance(msgs[i], dict) else {}
                if h.get('role') == 'assistant' and (h.get('content') or '').strip():
                    bd, bi = h['content'], i
                    break
        th = ""
        if bi > 0:
            for j in range(bi - 1, -1, -1):
                h = msgs[j] if isinstance(msgs[j], dict) else {}
                if h.get('role') == 'user' and (h.get('content') or '').strip():
                    th = h['content']
                    break
        return bd, th

    base_doc, title_hint = _pick_base(server_hist.get('messages') or [])
    if not base_doc or len(base_doc) < 200:
        base_doc, title_hint = _pick_base(history)
    if not base_doc or len(base_doc) < 200:
        return jsonify({"error": "No encontramos el documento a mejorar. Genera primero tu documento y vuelve a intentarlo."})

    if not course_topic:
        course_topic = server_hist.get('course_topic') or _extract_course_name(history)

    # --- Construcción del prompt de mejora (devuelve documento COMPLETO) ---
    accion_desc = _REFINE_ACCION_DESC[accion]
    obj_key = _normalize(objetivo_raw.lower())
    is_todo = (not objetivo_raw) or ("todo" in obj_key)
    if is_todo:
        alcance = ("Aplica la mejora a TODO el documento, conservando su estructura y sus encabezados.")
    else:
        target_label = None
        for k, lbl in _REFINE_TARGET_LABELS.items():
            if k in obj_key:
                target_label = lbl
                break
        if not target_label:
            target_label = f"la sección «{objetivo_raw}»"
        alcance = (
            f"Modifica ÚNICAMENTE {target_label}. El RESTO del documento debe quedar EXACTAMENTE IGUAL, "
            "palabra por palabra, sin cambios en las demás secciones."
        )

    refine_prompt = (
        "A continuación tienes un documento YA GENERADO. Tu única tarea es mejorarlo, NO empezar de cero.\n"
        f"ACCIÓN: {accion_desc}.\n"
        f"ALCANCE: {alcance}\n"
        + (f"INDICACIÓN ESPECÍFICA DEL USUARIO (cúmplela): {detalle}\n" if detalle else "")
        + "REGLAS OBLIGATORIAS: 1) Devuelve el DOCUMENTO COMPLETO actualizado, no un fragmento. "
          "2) Conserva EXACTAMENTE los mismos encabezados Markdown y la misma estructura/formato. "
          "3) No agregues comentarios, saludos ni explicaciones fuera del documento. "
          "4) Mantén la coherencia total del curso (tema, objetivos, tiempos y duración).\n\n"
        f"=== DOCUMENTO ACTUAL A MEJORAR ===\n{base_doc}"
    )

    docs = get_reference_docs()
    master_doc = session.get('master_doc', '') if element_num in (2, 3) else None
    current_tier = current_user.tier if is_authenticated else None
    course_info_text = _course_info_for_prompt()
    specs_for_ai = _get_chat_spec(element_num)

    # --- Control de concurrencia IA (mismo semáforo que /api/chat) ---
    max_concurrent_ai = int(Config.get('MAX_CONCURRENT_AI', '3'))
    with _ai_condition:
        _ai_queue_waiting[0] += 1
    acquired = False
    with _ai_condition:
        _ai_queue_waiting[0] = max(0, _ai_queue_waiting[0] - 1)
        if _ai_active_count[0] < max_concurrent_ai:
            _ai_active_count[0] += 1
            acquired = True
        else:
            _ai_condition.wait(timeout=15)
            if _ai_active_count[0] < max_concurrent_ai:
                _ai_active_count[0] += 1
                acquired = True
    if not acquired:
        return jsonify({
            "error": "El sistema está procesando varias solicitudes simultáneas. Por favor intenta de nuevo en unos segundos.",
            "retry": True
        })

    try:
        try:
            response_text = chat_with_ai(
                element_num, refine_prompt, history, docs,
                master_doc=master_doc, user_specs=specs_for_ai,
                user_tier=current_tier, course_info_text=course_info_text
            )
        finally:
            with _ai_condition:
                _ai_active_count[0] = max(0, _ai_active_count[0] - 1)
                _ai_condition.notify()

        response_text = _strip_subnorm_marker(response_text)

        if is_authenticated:
            try:
                estimated_tokens = (len(refine_prompt) + len(response_text)) // 4
                current_user.add_token_usage(current_user.tier, estimated_tokens)
                db.session.commit()
            except Exception:
                db.session.rollback()

        result = {"response": response_text, "generated_file": None, "detected_topic": course_topic}

        # --- Regeneración del documento completo (mismo motor que /api/chat) ---
        if len(response_text) > 200:
            try:
                course_name = _canonical_course_name(course_topic, title_hint)
                doc_title = _determine_doc_title(element_num, title_hint, is_first_topic=(element_num == 1))
                response_text_clean = _strip_conversational_preamble(response_text)
                sections = _parse_response_to_sections(response_text_clean)
                if sections:
                    effective_tier = "PRO" if (is_authenticated and current_user.has_alacarte(element_num)) else (current_tier or "FREE")
                    filepath = generate_custom_docx(doc_title, sections, course_name=course_name, user_tier=effective_tier, template_style=template_style, logo_path=session.get('course_logo_path'))
                    if is_authenticated:
                        filepath = _prefix_user_file(filepath)
                        _persist_file_to_db(filepath, current_user.id)
                    else:
                        filepath = _prefix_anon_file(filepath, course_topic=course_topic)
                        session['pending_document'] = os.path.basename(filepath)
                    result["generated_file"] = os.path.basename(filepath)
                    logger.info(f"REFINE_DOC_OK element={element_num} accion={accion} doc_title={doc_title} sections={len(sections)} user_id={current_user.id if is_authenticated else 'anon'}")
                else:
                    logger.warning(f"REFINE_NO_SECTIONS element={element_num} response_len={len(response_text)} user_id={current_user.id if is_authenticated else 'anon'}")
                if element_num == 1 and len(response_text) > 300:
                    _build_master_doc(response_text, course_topic)
                elif element_num == 1:
                    logger.warning(f"MASTER_DOC_SKIPPED element=1 response_len={len(response_text)} is_contrato={is_contrato} is_lista_req={is_lista_req} user_id={current_user.id if is_authenticated else 'anon'}")
            except Exception as _e_doc:
                logger.warning(f"REFINE_DOC_FAIL element={element_num} error_type={type(_e_doc).__name__}")
                logger.debug("REFINE_DOC_FAIL traceback", exc_info=True)

        # --- Incremento del contador de ciclos ---
        used += 1
        counts[str(element_num)] = used
        session['refine_counts'] = counts

        result["refine_used"] = used
        result["refine_max"] = REFINE_MAX_CYCLES
        result["refine_remaining"] = max(0, REFINE_MAX_CYCLES - used)
        result["refine_available"] = used < REFINE_MAX_CYCLES
        result["doc_ready"] = bool(result.get("generated_file"))

        try:
            updated_history = list(history) if isinstance(history, list) else []
            _obj_label = objetivo_raw or "Todo el curso"
            updated_history.append({'role': 'user', 'content': f"[Mejora · {accion} · {_obj_label}] {detalle}".strip()})
            updated_history.append({'role': 'assistant', 'content': response_text})
            _save_full_chat_history(element_num, updated_history, course_topic_val=course_topic, generated_file=result.get('generated_file'))
        except Exception as _e_hist:
            logger.warning(f"refine_history_persist_fail user={current_user.id if is_authenticated else 'anon'} err={type(_e_hist).__name__}")

        try:
            track_event('Mejora', f'Refine {accion}', user_id=current_user.id if is_authenticated else None,
                        extra_data={'element': element_num, 'objetivo': (objetivo_raw or 'todo')[:40], 'ciclo': used})
        except Exception:
            pass

        return jsonify(result)
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.error(f"api_refine unhandled error user={current_user.id if is_authenticated else 'anon'}: {type(e).__name__}: {e}", exc_info=True)
        return jsonify({"error": "Ocurrió un problema al aplicar la mejora. Inténtalo de nuevo en unos momentos; si el problema persiste, escríbenos y lo revisamos."})


@app.route("/api/generate_template", methods=["POST"])
@login_required
def api_generate_template():
    data = request.json
    template_key = data.get("template_key", "")
    content_data = data.get("content_data", {})

    try:
        _ut = current_user.tier if current_user.is_authenticated else "FREE"
        filepath, error = generate_from_template(template_key, content_data, user_tier=_ut, logo_path=session.get('course_logo_path'))
        if error:
            return jsonify({"error": error})
        filepath = _prefix_user_file(filepath)
        _persist_file_to_db(filepath, current_user.id)
        return jsonify({"success": True, "file": os.path.basename(filepath)})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/download/<filename>")
@login_required
def download_file(filename):
    user_prefix = f"u{current_user.id}_"
    if not filename.startswith(user_prefix) and not current_user.is_admin:
        abort(403)
    if not current_user.is_pro:
        if current_user.tier == 'FREE':
            alacarte_element = _detect_element_from_filename(filename)
            is_alacarte_download = alacarte_element and current_user.has_alacarte(alacarte_element)
            if not is_alacarte_download:
                has_own_file = _get_file_from_db(filename, user_id=current_user.id) is not None or \
                               os.path.isfile(os.path.join("generated_docs", filename))
                if current_user.chat_usage_count == 0 or not has_own_file:
                    track_event('Documento', 'Descarga Bloqueada', user_id=current_user.id, extra_data={'archivo': filename, 'razon': 'sin_generacion_ia'})
                    return jsonify({"error": "Primero debes generar un documento con la IA antes de descargar."}), 400
                rows = db.session.execute(
                    db.text('UPDATE "user" SET free_downloads_used = free_downloads_used + 1 WHERE id = :uid AND free_downloads_used < 1'),
                    {"uid": current_user.id}
                )
                db.session.commit()
                if rows.rowcount == 0:
                    track_event('Documento', 'Descarga Bloqueada', user_id=current_user.id, extra_data={'archivo': filename, 'razon': 'cuota_agotada'})
                    return jsonify({"error": "Ya usaste tu descarga gratuita. Actualiza a PRO para descargas ilimitadas."}), 403
                db.session.refresh(current_user)
        else:
            abort(403)
    track_event('Documento', 'Descarga Exitosa', user_id=current_user.id, extra_data={'archivo': filename, 'tier': current_user.tier})
    try:
        _mod_match = _re.search(r'Elemento(\d+)_', filename)
        _mod_num = int(_mod_match.group(1)) if _mod_match else None
        track_event('Funnel', 'Descarga Completada', user_id=current_user.id, extra_data={
            'archivo': filename, 'tier': current_user.tier, 'modulo': _mod_num
        })
    except Exception:
        pass
    disk_path = os.path.join("generated_docs", filename)
    if os.path.isfile(disk_path):
        # Sanitizar el nombre para el header Content-Disposition: algunos archivos
        # se generaron con saltos de línea / caracteres de control en el nombre
        # (arrastrados del texto del formulario), lo que rompía send_from_directory
        # con ValueError "Header values must not contain newline characters" (500).
        safe_name = re.sub(r'[\r\n\t\x00-\x1f\x7f]+', ' ', filename or 'documento').strip()
        safe_name = re.sub(r'\s+', ' ', safe_name) or 'documento'
        return send_file(disk_path, as_attachment=True, download_name=safe_name)
    sf = _get_file_from_db(filename, user_id=current_user.id)
    if not sf and current_user.is_admin:
        sf = _get_file_from_db(filename)
    if sf:
        # Sanitizar download_name: algunos nombres almacenados traían saltos de
        # línea / caracteres de control que rompen el header Content-Disposition
        # (ValueError: "Header values must not contain newline characters").
        safe_name = re.sub(r'[\r\n\t\x00-\x1f\x7f]+', ' ', filename or 'documento').strip()
        safe_name = re.sub(r'\s+', ' ', safe_name) or 'documento'
        return send_file(io.BytesIO(sf.content), download_name=safe_name,
                         as_attachment=True, mimetype=sf.content_type)
    abort(404)

@app.route("/api/customize_header/<filename>", methods=["POST"])
@login_required
def customize_header(filename):
    if ".." in filename or "/" in filename or "\\" in filename:
        abort(400)
    user_prefix = f"u{current_user.id}_"
    if not filename.startswith(user_prefix) and not current_user.is_admin:
        abort(403)
    data = request.get_json() or {}
    instructor = data.get('instructor', '').strip()
    lugar = data.get('lugar', '').strip()
    fecha = data.get('fecha', '').strip()
    horario = data.get('horario', '').strip()
    if not any([instructor, lugar, fecha, horario]):
        return jsonify({"ok": True})
    try:
        disk_path = os.path.join("generated_docs", filename)
        doc_bytes = None
        if os.path.isfile(disk_path):
            with open(disk_path, 'rb') as f:
                doc_bytes = f.read()
        else:
            sf = _get_file_from_db(filename, user_id=current_user.id)
            if sf:
                doc_bytes = sf.content
        if not doc_bytes:
            return jsonify({"error": "Archivo no encontrado"}), 404
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(doc_bytes))
        replacements = {}
        if instructor:
            replacements['[Nombre del Instructor]'] = instructor
            replacements['[Nombre del Diseñador]'] = instructor
        if lugar:
            replacements['[Lugar de impartición]'] = lugar
            replacements['[Lugar/Sede]'] = lugar
        if fecha:
            replacements['[Fecha de impartición]'] = fecha
            replacements['[Fecha]'] = fecha
        if horario:
            replacements['[Horario]'] = horario
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_val, new_val in replacements.items():
                        if old_val in cell.text:
                            for paragraph in cell.paragraphs:
                                full_text = ''.join(r.text for r in paragraph.runs)
                                if old_val in full_text:
                                    replaced = full_text.replace(old_val, new_val)
                                    if paragraph.runs:
                                        paragraph.runs[0].text = replaced
                                        for r in paragraph.runs[1:]:
                                            r.text = ""
        os.makedirs("generated_docs", exist_ok=True)
        doc.save(disk_path)
        _persist_file_to_db(disk_path, current_user.id)
        return jsonify({"ok": True})
    except Exception as e:
        logger.error(f"Header customization failed: {e}")
        return jsonify({"error": "No se pudo personalizar el encabezado. El documento se descargará sin cambios."}), 500

@app.route("/delete/<filename>", methods=["POST"])
@login_required
def delete_file(filename):
    if ".." in filename or "/" in filename or "\\" in filename:
        abort(400)
    user_prefix = f"u{current_user.id}_"
    if not filename.startswith(user_prefix) and not current_user.is_admin:
        abort(403)
    filepath = os.path.join("generated_docs", filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    _delete_file_from_db(filename, user_id=current_user.id)
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"success": True})
    return redirect("/documentos")

@app.route("/api/delete_bulk", methods=["POST"])
@login_required
def delete_bulk():
    data = request.get_json(silent=True)
    if not data or not isinstance(data.get("filenames"), list):
        return jsonify({"error": "Se requiere una lista de archivos"}), 400
    filenames = data["filenames"]
    if len(filenames) > 200:
        return jsonify({"error": "Demasiados archivos"}), 400
    user_prefix = f"u{current_user.id}_"
    deleted = []
    for filename in filenames:
        if not isinstance(filename, str) or ".." in filename or "/" in filename or "\\" in filename:
            continue
        if not filename.startswith(user_prefix) and not current_user.is_admin:
            continue
        filepath = os.path.join("generated_docs", filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        _delete_file_from_db(filename, user_id=current_user.id)
        deleted.append(filename)
    return jsonify({"success": True, "deleted": deleted, "count": len(deleted)})

def _extract_list_items(text, section_keyword):
    items = []
    in_section = False
    for line in text.split("\n"):
        line_lower = line.lower().strip()
        if section_keyword.lower() in line_lower and any(kw in line_lower for kw in ["compromiso", ":", "obligacion"]):
            in_section = True
            after_colon = line.split(":", 1)
            if len(after_colon) > 1 and after_colon[1].strip():
                item = after_colon[1].strip().lstrip("-").strip()
                if item and len(item) > 3:
                    items.append(item)
            continue
        if in_section:
            stripped = line.strip()
            if stripped.startswith("-") or stripped.startswith("*") or re.match(r'^\d+[\.\)]\s', stripped):
                item = re.sub(r'^[\-\*\d\.\)]+\s*', '', stripped)
                if item and len(item) > 3:
                    items.append(item)
            elif stripped and not stripped.startswith("#") and len(stripped) > 10 and not any(kw in line_lower for kw in ["compromiso", "firma", "nombre", "participante" if section_keyword.lower() == "instructor" else "instructor"]):
                items.append(stripped)
            elif items and stripped == "":
                in_section = False
            elif any(kw in line_lower for kw in ["compromiso", "firma", "nombre"]) and items:
                in_section = False
    return items

COURSE_INFO_FIELDS = [
    'nombre_curso',
    'disenador',
    'instructor',
    'lugar',
    'periodo_imparticion',
    'num_participantes',
    'num_horas',
]

COURSE_INFO_LABELS = {
    'nombre_curso': 'Nombre del curso',
    'disenador': 'Diseñador del curso',
    'instructor': 'Instructor que lo impartirá',
    'lugar': 'Lugar de impartición',
    'periodo_imparticion': 'Periodo de impartición (fecha y horario)',
    'num_participantes': 'Número de participantes',
    'num_horas': 'Número de horas',
}

COURSE_INFO_SKIP = '__POR_LLENAR__'

def _get_course_info():
    info = session.get('course_info')
    if not isinstance(info, dict):
        info = {}
    return info

def _set_course_info(info):
    session['course_info'] = info
    try:
        if current_user.is_authenticated:
            cs = _get_or_create_active_course_session(current_user)
            if cs is not None:
                import json as _json_ci
                cs.course_info_json = _json_ci.dumps(info, ensure_ascii=False)
                cs.last_activity_at = datetime.utcnow()
                db.session.commit()
    except Exception:
        try: db.session.rollback()
        except Exception: pass

def _lock_user_row(user_id):
    try:
        db.session.execute(db.text('SELECT id FROM "user" WHERE id = :uid FOR UPDATE'), {'uid': int(user_id)})
    except Exception as _elock:
        logger.warning(f"_lock_user_row failed for user {user_id}: {type(_elock).__name__}: {_elock}")
        raise

def _get_or_create_active_course_session(user):
    if user is None or not getattr(user, 'id', None):
        return None
    for _attempt in range(2):
        try:
            _lock_user_row(user.id)
            if user.active_course_session_id:
                cs = CourseSession.query.get(user.active_course_session_id)
                if cs is not None and cs.user_id == user.id:
                    db.session.commit()
                    return cs
            cs = CourseSession.query.filter_by(user_id=user.id, is_active=True).order_by(CourseSession.id.desc()).first()
            if cs is None:
                max_num = db.session.query(db.func.max(CourseSession.session_num)).filter_by(user_id=user.id).scalar() or 0
                cs = CourseSession(user_id=user.id, session_num=int(max_num) + 1, is_active=True)
                db.session.add(cs)
                db.session.flush()
            user.active_course_session_id = cs.id
            db.session.commit()
            return cs
        except Exception as _e:
            try: db.session.rollback()
            except Exception: pass
            try:
                db.session.refresh(user)
            except Exception:
                pass
            if _attempt == 0 and user.active_course_session_id:
                try:
                    cs2 = CourseSession.query.get(user.active_course_session_id)
                    if cs2 is not None:
                        return cs2
                except Exception:
                    pass
            if _attempt == 1:
                logger.warning(f"_get_or_create_active_course_session failed user={user.id}: {type(_e).__name__}: {_e}")
                return None

def reset_active_course_state(user_id, *, reason='manual', purge_master_only=True,
                              mark_frontend_signal=True, restore_pro_credit=False):
    """Hard-reset del estado del CURSO ACTIVO de un usuario.

    Diseñado para disparar tanto desde el webhook de Stripe (al confirmar pago
    FREE→PRO) como manualmente desde el botón "Iniciar Nuevo Curso".

    ► QUÉ HACE:
      1. Cierra (is_active=False) cualquier CourseSession activa del usuario.
      2. Purga JSON / texto del "curso actual": master_doc, course_info_json,
         topic — sólo en la sesión activa que se está cerrando (las sesiones
         históricas con chats/documentos conservan su contenido para auditoría).
      3. Si la sesión activa es is_demo=True, también la elimina por completo
         (no hay nada de valor real que conservar).
      4. Libera el "candado" pro_active_course (string lock) sin tocar saldos.
      5. Crea una CourseSession nueva en blanco y la marca como activa.
      6. Estampa user.needs_state_reset_at = ahora() para que el frontend
         purgue su localStorage/sessionStorage en la siguiente interacción.

    ► QUÉ **NO** HACE — BLINDAJE DE CRÉDITOS (ESTRICTO):
      - No modifica tier, pro_courses_remaining, alacarte_e1..e5,
        free_downloads_used, chat_usage_count, ni ningún saldo.
      - ÚNICA EXCEPCIÓN: si restore_pro_credit=True (solo lo activa la acción de
        admin "Reset curso activo") y el usuario es PRO con candado puesto, se
        DEVUELVE 1 crédito (pro_courses_remaining += 1) — el que se consumió para
        el curso que se descarta. El snapshot-guard contempla este cambio.
      - Al final compara snapshot pre/post y registra ERROR si difiere
        (defensa en profundidad ante futuros bugs).
      - Tolerante a NULL/JSON vacío: nunca falla si no hay nada que purgar.

    Args:
      user_id: ID del User al que se le hace reset.
      reason: etiqueta para auditoría ('stripe_upgrade'|'manual'|'demo_auto').
      purge_master_only: True (default) sólo purga master_doc/course_info_json
        de la sesión activa; False también limpia chat_history asociado a esa
        sesión (sólo si NO tiene documentos generados).
      mark_frontend_signal: True (default) estampa needs_state_reset_at.
      restore_pro_credit: False (default). Si True y el usuario es PRO con candado
        puesto, devuelve 1 crédito PRO al descartar el curso. Solo la acción de
        admin lo activa; el botón de usuario "Iniciar Nuevo Curso" NO.

    Returns:
      dict con: ok, reason, prev_active_cs_id, new_cs_id, new_cs_num,
                purged_demo, credit_snapshot_match, details.
    """
    result = {'ok': False, 'reason': reason, 'prev_active_cs_id': None,
              'new_cs_id': None, 'new_cs_num': None, 'purged_demo': False,
              'credit_snapshot_match': True, 'pro_credit_restored': False,
              'details': []}
    try:
        user = db.session.get(User, int(user_id))
    except Exception as _e:
        result['details'].append(f'user_lookup_failed: {_e}')
        return result
    if user is None:
        result['details'].append('user_not_found')
        return result

    try:
        # Fix architect #3: registrar fallo de lock (no enmascarar concurrencia).
        _row_lock_ok = True
        try:
            _lock_user_row(user.id)
        except Exception as _lock_e:
            _row_lock_ok = False
            result['details'].append(f'lock_failed: {type(_lock_e).__name__}: {_lock_e}')
            logger.warning(f'[reset_active_course_state] lock_failed user={user_id}: {_lock_e}')
        # Si vamos a DEVOLVER crédito (acción admin) y NO pudimos serializar la fila,
        # abortamos SIN mutar nada para evitar un doble-reembolso por concurrencia.
        # El admin puede reintentar de inmediato (estado intacto).
        if restore_pro_credit and not _row_lock_ok:
            result['details'].append('abort_no_lock_for_credit_refund')
            logger.warning(f'[reset_active_course_state] abort refund (no row lock) user={user_id}')
            return result
        try:
            db.session.refresh(user)
        except Exception:
            pass

        # 1) Snapshot DE CRÉDITOS antes de cualquier cambio (blindaje).
        credit_fields = ('tier', 'pro_courses_remaining',
                         'alacarte_e1', 'alacarte_e2', 'alacarte_e3',
                         'alacarte_e4', 'alacarte_e5',
                         'free_downloads_used', 'chat_usage_count')
        snap_before = {f: getattr(user, f, None) for f in credit_fields}

        # 2) Localizar sesión activa actual (si existe).
        prev_active = None
        try:
            if user.active_course_session_id:
                prev_active = db.session.get(CourseSession, user.active_course_session_id)
            if prev_active is None:
                prev_active = CourseSession.query.filter_by(
                    user_id=user.id, is_active=True
                ).order_by(CourseSession.session_num.desc()).first()
        except Exception as _e:
            result['details'].append(f'active_lookup_failed: {_e}')

        # 3) Si la activa es un demo, la borramos por completo. Si no, la
        #    cerramos y purgamos su JSON.
        if prev_active is not None:
            result['prev_active_cs_id'] = prev_active.id
            try:
                if getattr(prev_active, 'is_demo', False):
                    # Demo: limpieza total. Borrar chats vinculados (sin docs).
                    try:
                        ChatHistory.query.filter_by(
                            user_id=user.id, course_session_id=prev_active.id
                        ).delete(synchronize_session=False)
                    except Exception as _e:
                        result['details'].append(f'demo_chat_delete_skipped: {_e}')
                    db.session.delete(prev_active)
                    result['purged_demo'] = True
                else:
                    # Sesión real: purga JSON + topic (spec exige purga total
                    # del "curso actual"). NO borra histórico chats/docs salvo
                    # que purge_master_only=False y la sesión esté limpia.
                    prev_active.master_doc = None
                    prev_active.course_info_json = None
                    prev_active.topic = None  # ← purga topic SIEMPRE (fix architect #1)
                    # Tolerancia: si ya era NULL/vacío, no falla.
                    if not purge_master_only:
                        try:
                            # Sólo si no hay documentos generados en esa sesión
                            _docs = StoredFile.query.filter_by(
                                user_id=user.id, course_session_id=prev_active.id,
                                file_category='document'
                            ).count()
                            if _docs == 0:
                                ChatHistory.query.filter_by(
                                    user_id=user.id, course_session_id=prev_active.id
                                ).delete(synchronize_session=False)
                        except Exception as _e:
                            result['details'].append(f'chat_purge_skipped: {_e}')
                    prev_active.is_active = False
                    prev_active.last_activity_at = datetime.utcnow()
            except Exception as _e:
                result['details'].append(f'prev_active_mutation_failed: {_e}')

        # 4) Cerrar CUALQUIER otra sesión activa colgada del mismo user
        #    (defensa contra estados inconsistentes).
        try:
            CourseSession.query.filter_by(user_id=user.id, is_active=True).update(
                {'is_active': False}, synchronize_session=False)
        except Exception as _e:
            result['details'].append(f'deactivate_others_skipped: {_e}')

        # 5) Liberar candado de curso activo (string lock). NO es un crédito.
        had_pro_lock = False
        try:
            had_pro_lock = bool((user.pro_active_course or '').strip())
            if had_pro_lock:
                user.pro_active_course = None
        except Exception:
            pass

        # 5b) Reembolso de crédito PRO — SOLO acción de admin (restore_pro_credit).
        #     Un candado puesto implica que ya se consumió 1 crédito para ese curso;
        #     al descartarlo desde el panel admin devolvemos ese crédito para que el
        #     usuario (PRO-1 o PRO-5) pueda volver a generar. NO aplica al botón de
        #     usuario "Iniciar Nuevo Curso" (evita cursos ilimitados gratis).
        credit_restored = False
        if restore_pro_credit and had_pro_lock and (getattr(user, 'tier', None) == 'PRO'):
            try:
                user.pro_courses_remaining = (user.pro_courses_remaining or 0) + 1
                credit_restored = True
                result['pro_credit_restored'] = True
                result['details'].append(
                    f'pro_credit_restored_to={user.pro_courses_remaining}')
            except Exception as _e:
                result['details'].append(f'pro_credit_restore_failed: {_e}')

        # 6) Estampa señal frontend para que purgue localStorage/sessionStorage.
        if mark_frontend_signal:
            try:
                user.needs_state_reset_at = datetime.utcnow()
            except Exception as _e:
                result['details'].append(f'frontend_signal_skipped: {_e}')

        db.session.flush()

        # 7) Crear sesión nueva en blanco.
        try:
            max_num = db.session.query(db.func.max(CourseSession.session_num)).filter_by(
                user_id=user.id).scalar() or 0
            new_cs = CourseSession(user_id=user.id, session_num=int(max_num) + 1,
                                   is_active=True, topic=None,
                                   master_doc=None, course_info_json=None,
                                   is_demo=False)
            db.session.add(new_cs)
            db.session.flush()
            user.active_course_session_id = new_cs.id
            result['new_cs_id'] = new_cs.id
            result['new_cs_num'] = new_cs.session_num
        except Exception as _e:
            result['details'].append(f'create_new_cs_failed: {_e}')

        db.session.commit()

        # 8) Verificación de blindaje de créditos (post-commit, defensa profunda).
        try:
            db.session.refresh(user)
            snap_after = {f: getattr(user, f, None) for f in credit_fields}
            # El reembolso de crédito (acción admin) es un cambio INTENCIONAL: lo
            # incorporamos al esperado para no disparar el guard de mutación.
            expected = dict(snap_before)
            if credit_restored:
                expected['pro_courses_remaining'] = (snap_before.get('pro_courses_remaining') or 0) + 1
            if expected != snap_after:
                result['credit_snapshot_match'] = False
                diff = {k: (expected[k], snap_after[k]) for k in expected
                        if expected[k] != snap_after[k]}
                logger.error(f'[reset_active_course_state] CREDIT MUTATION DETECTED user={user.id} reason={reason} diff={diff}')
                result['details'].append(f'credit_diff={diff}')
        except Exception as _e:
            result['details'].append(f'snapshot_compare_skipped: {_e}')

        result['ok'] = True
        logger.info(f'[reset_active_course_state] user={user.id} reason={reason} '
                    f'prev_cs={result["prev_active_cs_id"]} new_cs={result["new_cs_id"]} '
                    f'purged_demo={result["purged_demo"]} credits_ok={result["credit_snapshot_match"]}')
    except Exception as _e:
        try: db.session.rollback()
        except Exception: pass
        result['details'].append(f'top_level_exception: {type(_e).__name__}: {_e}')
        logger.exception(f'reset_active_course_state failed user={user_id}')
    return result


def _create_new_course_session(user, topic=None):
    if user is None or not getattr(user, 'id', None):
        return None
    for _attempt in range(2):
        try:
            _lock_user_row(user.id)
            CourseSession.query.filter_by(user_id=user.id, is_active=True).update({'is_active': False})
            db.session.flush()
            max_num = db.session.query(db.func.max(CourseSession.session_num)).filter_by(user_id=user.id).scalar() or 0
            cs = CourseSession(user_id=user.id, session_num=int(max_num) + 1, is_active=True,
                               topic=(topic or None))
            db.session.add(cs)
            db.session.flush()
            user.active_course_session_id = cs.id
            db.session.commit()
            return cs
        except Exception as _e:
            try: db.session.rollback()
            except Exception: pass
            if _attempt == 1:
                logger.warning(f"_create_new_course_session failed user={user.id}: {type(_e).__name__}: {_e}")
                return None

def _persist_course_session_state(user):
    if user is None or not getattr(user, 'id', None):
        return False
    try:
        cs = _get_or_create_active_course_session(user)
        if cs is None:
            return False
        import json as _json_ps
        md = session.get('master_doc')
        if md:
            cs.master_doc = md[:8000]
        ci = session.get('course_info')
        if isinstance(ci, dict) and ci:
            cs.course_info_json = _json_ps.dumps(ci, ensure_ascii=False)
        topic = session.get('master_doc_topic')
        if topic and not cs.topic:
            cs.topic = topic[:300]
        cs.last_activity_at = datetime.utcnow()
        db.session.commit()
        return True
    except Exception:
        try: db.session.rollback()
        except Exception: pass
        return False

def _restore_course_session_state(user, cs_id):
    if user is None or not getattr(user, 'id', None):
        return False
    try:
        cs = CourseSession.query.get(int(cs_id))
        if cs is None or cs.user_id != user.id:
            return False
        _persist_course_session_state(user)
        CourseSession.query.filter_by(user_id=user.id, is_active=True).update({'is_active': False})
        cs.is_active = True
        user.active_course_session_id = cs.id
        cs.last_activity_at = datetime.utcnow()
        db.session.commit()
        for k in ('course_info', 'master_doc', 'master_doc_topic', 'pending_action',
                  'pending_document', 'subnorm_pending', 'cocreation_skip_session'):
            session.pop(k, None)
        for _e in (1, 2, 3, 4):
            session.pop(f'cocreation_state_e{_e}', None)
        if cs.master_doc:
            session['master_doc'] = cs.master_doc
        if cs.topic:
            session['master_doc_topic'] = cs.topic
        if cs.course_info_json:
            try:
                import json as _json_rs
                ci = _json_rs.loads(cs.course_info_json)
                if isinstance(ci, dict):
                    session['course_info'] = ci
            except Exception:
                pass
        session.modified = True
        return True
    except Exception:
        try: db.session.rollback()
        except Exception: pass
        return False

def _is_course_info_complete():
    info = _get_course_info()
    return bool(info.get('nombre_curso'))

def _get_missing_course_fields():
    info = _get_course_info()
    return [f for f in COURSE_INFO_FIELDS if not info.get(f)]

_LABEL_LIKE_VALUE_RE = re.compile(
    r'^(?:nombre|t[ií]tulo|tema|curso|dise[ñn]ador|dise[ñn]ado|instructor|impartido|imparte|'
    r'lugar|sede|ubicaci[oó]n|periodo|fecha|horario|participantes?|n[uú]mero|n[uú]m\.|no\.|duraci[oó]n)'
    r'\b[^:=\n]{0,40}?[:=]',
    re.IGNORECASE,
)

def _extract_course_info_from_message(text):
    found = {}
    if not text or not isinstance(text, str):
        return found
    raw = text.strip().lstrip('.•-*\u2022 ').strip()
    low = raw.lower()

    skip_patterns = [
        (r'no tengo (?:el |la |los |las )?nombre del curso', 'nombre_curso'),
        (r'no tengo (?:el |la )?dise[ñn]ador', 'disenador'),
        (r'no tengo (?:el |la )?instructor', 'instructor'),
        (r'no tengo (?:el |la )?lugar', 'lugar'),
        (r'no s[eé] (?:el |la )?lugar', 'lugar'),
        (r'no tengo (?:el |la |los |las )?(?:periodo|fecha|horario)', 'periodo_imparticion'),
        (r'no tengo (?:el |la |los |las )?(?:n[uú]mero de )?participantes?', 'num_participantes'),
        (r'no tengo (?:el |la |los |las )?(?:n[uú]mero de )?horas?', 'num_horas'),
        (r'no s[eé] (?:el |la )?dise[ñn]ador', 'disenador'),
        (r'no s[eé] (?:el |la )?instructor', 'instructor'),
        (r'no s[eé] (?:el |la |los |las )?(?:periodo|fecha|horario)', 'periodo_imparticion'),
    ]
    for pat, field in skip_patterns:
        if re.search(pat, low):
            found[field] = COURSE_INFO_SKIP

    field_keywords_re = (
        r'(?:nombre\s+del\s+curso|t[ií]tulo\s+del\s+curso|tema\s+del\s+curso|'
        r'dise[ñn]ador(?:a)?|dise[ñn]ado\s+por|'
        r'instructor(?:a)?|impartido\s+por|imparte|'
        r'lugar(?:\s+(?:de\s+impartici[oó]n|donde\s+se\s+impartir[áa]))?|sede|ubicaci[oó]n|'
        r'periodo(?:\s+de\s+impartici[oó]n)?|fecha(?:\s+(?:de\s+impartici[oó]n|y\s+horario))?|horario|'
        r'(?:n[uú]m(?:ero|\.)?(?:\s+de)?|no\.?(?:\s+de)?)\s*participantes?|participantes?|'
        r'(?:n[uú]m(?:ero|\.)?(?:\s+de)?|no\.?(?:\s+de)?)\s*horas?|duraci[oó]n)'
    )
    pieces = re.split(r'[;\n|]+', raw)
    extra_pieces = []
    for p in pieces:
        sub = re.split(r',\s*(?=' + field_keywords_re + r'\s*[:=\-])', p, flags=re.IGNORECASE)
        extra_pieces.extend(sub)
    chunks = [p.strip(' .,;\u2022\t') for p in extra_pieces if p.strip(' .,;\u2022\t')]

    kv_patterns = [
        ('nombre_curso', [
            r'^\s*(?:\d+[\.\)]\s*)?(?:nombre\s+del\s+curso|t[ií]tulo\s+del\s+curso|tema\s+del\s+curso|nombre|t[ií]tulo|tema|curso)\s*[:=\-]\s*(.+)$',
            r'(?:con|sobre)\s+(?:el\s+)?tema\s*[:=\-]\s*["\'\u201c\u00ab](.+?)["\'\u201d\u00bb]',
            r'(?:curso|taller|capacitaci[oó]n|formaci[oó]n)\b.*?(?:con|sobre)\s+(?:el\s+)?tema\s*:\s*(.+?)(?:\s*$)',
        ]),
        ('disenador', [
            r'^\s*(?:\d+[\.\)]\s*)?(?:dise[ñn]ador(?:a)?|dise[ñn]ado\s+por)\b[^:=\n]{0,40}?[:=\-]\s*(.+)$',
            r'^\s*(.+?)\s+(?:es\s+(?:el|la|quien)|fue\s+(?:el|la|quien)|fungi[oó]\s+como)\s+(?:el\s+|la\s+)?dise[ñn]ador(?:a)?(?:\s+(?:e|y)\s+instructor(?:a)?)?\s*\.?$',
        ]),
        ('instructor', [
            r'^\s*(?:\d+[\.\)]\s*)?(?:instructor(?:a)?|impartido\s+por|imparte|imparti[oó])\b[^:=\n]{0,40}?[:=\-]\s*(.+)$',
            r'^\s*(.+?)\s+(?:es\s+(?:el|la|quien)|fue\s+(?:el|la|quien)|fungi[oó]\s+como)\s+(?:el\s+|la\s+)?(?:dise[ñn]ador(?:a)?\s+(?:e|y)\s+)?instructor(?:a)?\s*\.?$',
            r'^\s*(.+?)\s+(?:lo|la|los)\s+(?:imparte|impartir[áa])\s*\.?$',
        ]),
        ('lugar', [
            r'^\s*(?:\d+[\.\)]\s*)?lugar\b[^:=\n]{0,40}?[:=\-]\s*(.+)$',
            r'^\s*(?:sede|ubicaci[oó]n)\s*[:=\-]\s*(.+)$',
        ]),
        ('periodo_imparticion', [
            r'^\s*(?:\d+[\.\)]\s*)?(?:periodo\s+de\s+impartici[oó]n|periodo|fecha\s+y\s+horario|fecha\s+de\s+impartici[oó]n|fecha|horario)\s*[:=\-]\s*(.+)$',
        ]),
        ('num_participantes', [
            r'^\s*(?:\d+[\.\)]\s*)?(?:n[uú]m(?:ero|\.)?(?:\s+de)?|no\.?(?:\s+de)?)?\s*participantes?\s*[:=\-]\s*(.+)$',
            r'^\s*participantes?\s*[:=\-]?\s*(\d{1,3})\b',
            r'(\d{1,3})\s+participantes?\b',
        ]),
        ('num_horas', [
            r'^\s*(?:\d+[\.\)]\s*)?(?:n[uú]m(?:ero|\.)?(?:\s+de)?|no\.?(?:\s+de)?)?\s*horas?\s*[:=\-]\s*(.+)$',
            r'^\s*duraci[oó]n\s*[:=\-]\s*(.+)$',
            r'(\d{1,3}(?:\.\d+)?)\s*horas?\b',
        ]),
    ]

    def _try_match(field, patterns, scope):
        if field in found:
            return False
        for pat in patterns:
            m = re.search(pat, scope, flags=re.IGNORECASE)
            if m:
                val = m.group(1).strip().strip('.,;:').strip()
                val = re.sub(r'\s+', ' ', val)
                if not val or len(val) > 200:
                    continue
                if field not in ('num_participantes', 'num_horas') and _LABEL_LIKE_VALUE_RE.match(val):
                    continue
                if field == 'num_participantes':
                    nm = re.search(r'(\d{1,3})', val)
                    if nm:
                        n = int(nm.group(1))
                        if n >= 1:
                            found[field] = str(max(n, 4)) if n < 4 else str(n)
                            return True
                elif field == 'num_horas':
                    nm = re.search(r'(\d{1,3}(?:\.\d+)?)', val)
                    if nm:
                        try:
                            h = float(nm.group(1))
                            if h >= 0.5:
                                found[field] = str(max(h, 2.0)) if h < 2.0 else (str(int(h)) if h == int(h) else str(h))
                                return True
                        except Exception:
                            pass
                else:
                    found[field] = val
                    return True
        return False

    for chunk in chunks:
        for field, patterns in kv_patterns:
            _try_match(field, patterns, chunk)

    for field, patterns in kv_patterns:
        _try_match(field, patterns, raw)

    return found

def _merge_course_info(new_fields):
    if not new_fields:
        return _get_course_info()
    info = _get_course_info()
    changed = False
    for k, v in new_fields.items():
        if k in COURSE_INFO_FIELDS and v and not info.get(k):
            info[k] = v
            changed = True
    if changed:
        _set_course_info(info)
    return info

def _course_info_for_prompt():
    info = _get_course_info()
    if not info:
        return None
    out_lines = []
    for f in COURSE_INFO_FIELDS:
        v = info.get(f)
        if v:
            label = COURSE_INFO_LABELS[f]
            display = '[Por llenar manualmente]' if v == COURSE_INFO_SKIP else v
            out_lines.append(f"- {label}: {display}")
    if not out_lines:
        return None
    return '\n'.join(out_lines)

def _course_info_to_dict_for_docx():
    info = _get_course_info()
    if not info:
        return {}
    def _val(field, placeholder):
        v = info.get(field)
        if not v or v == COURSE_INFO_SKIP:
            return placeholder
        return v
    return {
        'nombre_curso': _val('nombre_curso', '[Nombre del Curso]'),
        'nombre_disenador': _val('disenador', '[Nombre del Diseñador]'),
        'nombre_instructor': _val('instructor', '[Nombre del Instructor]'),
        'lugar': _val('lugar', '[Lugar de Impartición]'),
        'periodo_imparticion': _val('periodo_imparticion', '[Periodo de Impartición]'),
        'num_participantes': _val('num_participantes', '[Número de Participantes]'),
        'num_horas': _val('num_horas', '[Número de Horas]'),
    }

def _canonical_course_name(course_topic=None, message=None, fallback="[Tema por definir]"):
    """Fuente ÚNICA del nombre del curso para TODOS los productos descargables.

    Prioriza el dato estructurado del curso activo (session['course_info']) sobre el
    `course_topic` de la petición, que puede arrastrar el nombre de un curso ANTERIOR
    (localStorage del navegador). Así el título del archivo y el nombre dentro de
    cada documento corresponden SIEMPRE al curso en producción.
    """
    # 1) Dato estructurado del curso activo (verdad del servidor).
    try:
        ci = _course_info_to_dict_for_docx() or {}
        name = _strip_course_label_prefix((ci.get('nombre_curso') or '').strip())
        name = re.sub(r'[\r\n\t\x00-\x1f\x7f]+', ' ', name).strip()
        if name and not _is_placeholder_value(name):
            return name
    except Exception:
        pass
    # 2) course_topic de la petición, saneado (limpia prefijos de comando/etiqueta).
    if course_topic:
        cleaned = (_sanitize_course_name_from_message(course_topic) or '').strip()
        if not cleaned:
            cleaned = _strip_course_label_prefix(str(course_topic).strip())
        cleaned = re.sub(r'[\r\n\t\x00-\x1f\x7f]+', ' ', cleaned).strip()
        if cleaned and not _is_placeholder_value(cleaned):
            return cleaned
    # 3) Mensaje crudo, saneado.
    if message:
        cleaned = (_sanitize_course_name_from_message(message) or '').strip()
        if cleaned and not _is_placeholder_value(cleaned):
            return cleaned
    return fallback

def _paso0_message(missing=None, first_time=True):
    if missing is None:
        missing = COURSE_INFO_FIELDS
    intro_first = (
        "Antes de armar tus productos, ayúdame con estos datos. "
        "Los usaré tal cual los escribas en TODOS tus productos "
        "(Carta Descriptiva, Instrumentos, Manuales y Lista de Verificación). "
        "Si omites alguno, deberás llenarlo manualmente después en cada formato:"
    )
    if not first_time:
        captured = _get_course_info()
        captured_lines = []
        for _cf in COURSE_INFO_FIELDS:
            _cv = captured.get(_cf)
            if _cv and _cv != COURSE_INFO_SKIP:
                captured_lines.append(f"✓ {COURSE_INFO_LABELS[_cf]}: {_cv}")
        if captured_lines:
            intro_followup = "Gracias. Ya tengo:\n" + "\n".join(captured_lines) + "\n\nAún me faltan:"
        else:
            intro_followup = "Aún me faltan estos datos para continuar:"
    else:
        intro_followup = "Aún me faltan estos datos para continuar:"
    intro = intro_first if first_time else intro_followup

    bullets = []
    for f in missing:
        label = COURSE_INFO_LABELS[f]
        if f == 'num_participantes':
            label += " (mínimo 4 sugerido para fines de certificación; puedes indicar más)"
        elif f == 'num_horas':
            label += " (mínimo 2 sugeridas para fines de certificación; puedes indicar más)"
        bullets.append(f"{len(bullets)+1}. {label}")

    closing = (
        "\n\nPuedes dármelos todos juntos en un solo mensaje. "
        "Si no tienes alguno, escribe \"no tengo [el dato]\" y lo dejaré como campo por llenar para que lo completes a mano después."
    )
    return intro + "\n\n" + "\n".join(bullets) + closing

def _is_generation_intent(message, is_quick_action=False):
    if is_quick_action:
        return True
    if not message:
        return False
    low = message.lower()
    keywords = [
        'genera', 'redacta', 'crea', 'arma', 'prepara', 'dame',
        'carta descriptiva', 'manual del instructor', 'manual del participante',
        'instrumentos', 'evaluaci', 'objetivo', 'contrato', 'lista de verifica',
    ]
    return any(kw in low for kw in keywords)

LISTA_REQ_SECTION_KEYS = [
    ('instalaciones', ['instalaciones', 'mobiliario']),
    ('equipo_apoyo', ['equipo_de_apoyo', 'equipo de apoyo']),
    ('materiales_didacticos', ['materiales_didacticos', 'materiales didacticos', 'materiales didácticos', 'material didactico', 'material didáctico']),
    ('humanos', ['requerimientos_humanos', 'requerimientos humanos', 'humanos']),
    ('otros', ['otros_requerimientos', 'otros requerimientos', 'otros']),
]

def _parse_lista_req_sections(text):
    result = {k: [] for k, _ in LISTA_REQ_SECTION_KEYS}
    if not text:
        return result
    current_key = None
    for raw_line in text.split('\n'):
        line = raw_line.strip()
        if not line:
            continue
        is_header = line.startswith('#') or line.lower().startswith('## ')
        plain_low = re.sub(r'[^a-z0-9_\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\s]', '', line.lower()).strip()
        if is_header or any(line.lower().startswith(prefix) for prefix in ('## ', '# ', '**')):
            matched_key = None
            for key, kws in LISTA_REQ_SECTION_KEYS:
                if any(kw in plain_low for kw in kws):
                    matched_key = key
                    break
            if matched_key:
                current_key = matched_key
                continue
        if current_key:
            m = re.match(r'^\s*(?:\d+|\u2022|\-|\*)\s*[\.\-\)]?\s*(.+)$', line)
            if m:
                desc = m.group(1).strip().lstrip('.-) ').strip()
                desc = re.sub(r'\*+', '', desc).strip()
                if desc and not desc.startswith('#'):
                    result[current_key].append(desc)
    return result

def _extract_duracion_from_response(response_text):
    """Extrae la duración TOTAL del curso desde el response generado.

    Solo busca en los primeros 3000 caracteres (zona de "Información General")
    y exige patrones anclados a encabezados de duración total para evitar
    falsos positivos con duraciones de actividades puntuales.

    Devuelve string normalizado con unidad incluida, o None.
    """
    if not response_text:
        return None
    head = response_text[:3000]
    anchored_patterns = [
        r'Duraci[oó]n\s+(?:total|sugerida|del\s+curso|del\s+taller)[:\s]+([^\n]{3,80})',
        r'DURACI[OÓ]N\s+(?:TOTAL|SUGERIDA|DEL\s+CURSO|DEL\s+TALLER)[:\s]+([^\n]{3,80})',
        r'(?:^|\n)\s*[-*]?\s*\*{0,2}Duraci[oó]n\*{0,2}[:\s]+([^\n]{3,80})',
        r'(?:^|\n)\s*[-*]?\s*\*{0,2}DURACI[OÓ]N\*{0,2}[:\s]+([^\n]{3,80})',
    ]
    word_num_re = re.compile(
        r'\b(una?|dos|tres|cuatro|cinco|seis|siete|ocho|nueve|diez|once|doce|quince|veinte|treinta|cuarenta|cincuenta|sesenta|media)\b',
        re.IGNORECASE,
    )
    unit_re = re.compile(r'\b(horas?|minutos?|hrs?|mins?|h|min)\b', re.IGNORECASE)
    for pat in anchored_patterns:
        m = re.search(pat, head, re.IGNORECASE | re.MULTILINE)
        if m:
            cand = m.group(1).strip().rstrip('.,;:').strip().strip('*').strip()
            if not cand or len(cand) > 80:
                continue
            has_digit = any(c.isdigit() for c in cand)
            has_word_num = bool(word_num_re.search(cand))
            has_unit = bool(unit_re.search(cand))
            if (has_digit or has_word_num) and has_unit:
                return cand
    return None

def _is_placeholder_value(val):
    """Detecta valores sentinel/placeholder en course_info que no deben usarse como datos reales."""
    if not val:
        return True
    s = str(val).strip()
    if not s:
        return True
    if s.startswith('[') or s.startswith('__') or s.startswith('<'):
        return True
    low = s.lower()
    if low in ('por llenar', 'pendiente', 'no tengo', 'na', 'n/a', '-', '--'):
        return True
    return False

SUBNORM_MARKER_RE = re.compile(r'\[SUBNORM_NEGOCIACION:\s*minutos_solicitados\s*=\s*(\d+)\s*\]', re.IGNORECASE)

def _detect_subnorm_negotiation(response_text):
    """Detecta el marcador de negociacion sub-normativa en la respuesta de la IA.
    Devuelve dict {'minutes_requested': int} si aplica, None si no."""
    if not response_text:
        return None
    m = SUBNORM_MARKER_RE.search(response_text)
    if not m:
        return None
    try:
        minutes = int(m.group(1))
        if 0 < minutes < 120:
            return {'minutes_requested': minutes}
    except (ValueError, TypeError):
        pass
    return None

def _strip_subnorm_marker(response_text):
    """Remueve el marcador interno [SUBNORM_NEGOCIACION:...] del texto antes de mostrarlo al usuario."""
    if not response_text:
        return response_text
    return SUBNORM_MARKER_RE.sub('', response_text).rstrip()

def _classify_subnorm_user_decision(message):
    """Clasifica respuesta del usuario al protocolo de negociacion sub-normativa.
    Devuelve: 'confirmed' (usa duracion sub-120), 'declined' (acepta 120), 'unclear'."""
    if not message:
        return 'unclear'
    msg = message.lower().strip()
    has_120_signal = bool(re.search(r'\b120\b|\bdos horas\b|\bm[ií]nim[oa]\b|normativ|defecto|default', msg))
    starts_no = bool(re.match(r'^\s*no\b', msg))
    has_confirm = bool(re.search(r'\bs[ií]\b|\bconfirm|\bdale\b|\bprocede\b|\bok\b|correcto|adelante|de acuerdo|\bva\b|\bquiero\b|\bbajo mi\b|responsabilidad', msg))
    if starts_no:
        return 'declined'
    if has_120_signal and not has_confirm:
        return 'declined'
    if has_confirm and not has_120_signal:
        return 'confirmed'
    if has_confirm and has_120_signal:
        return 'declined'
    return 'unclear'

def _emit_subnorm_metric_server(metric_type, value_int, value_text=None, element_num=None, user_id=None, session_id=None):
    """Emite metrica BetaMetric server-side (sin pasar por la API HTTP)."""
    try:
        m = BetaMetric(
            user_id=user_id,
            session_id=session_id,
            metric_type=metric_type,
            element_num=element_num,
            value_int=value_int,
            value_text=(value_text or '')[:500],
        )
        db.session.add(m)
        db.session.commit()
    except Exception:
        try:
            db.session.rollback()
        except Exception:
            pass

def _build_master_doc(response_text, course_topic):
    try:
        lines = response_text.split('\n')
        master_parts = []
        if course_topic:
            master_parts.append(f"TEMA: {course_topic}")
        ci = session.get('course_info') or {}
        ci_horas_raw = (ci.get('num_horas') or '').strip()
        ci_horas = ci_horas_raw if not _is_placeholder_value(ci_horas_raw) else ''
        extracted = _extract_duracion_from_response(response_text)
        if extracted:
            duracion_master = extracted
        elif ci_horas:
            duracion_master = ci_horas
        else:
            duracion_master = "120 minutos (2 horas)"
        master_parts.append(f"DURACIÓN: {duracion_master}")
        institutional_fields = {}
        for line in lines:
            ll = line.lower().strip()
            if 'nombre del instructor' in ll and ':' in line:
                val = line.split(':', 1)[1].strip().strip('*').strip()
                if val and not val.startswith('['):
                    institutional_fields['instructor'] = val
            elif 'nombre del dise' in ll and ':' in line:
                val = line.split(':', 1)[1].strip().strip('*').strip()
                if val and not val.startswith('['):
                    institutional_fields['disenador'] = val
            elif 'lugar' in ll and 'impartici' in ll and ':' in line:
                val = line.split(':', 1)[1].strip().strip('*').strip()
                if val and not val.startswith('['):
                    institutional_fields['lugar'] = val
            elif 'fecha' in ll and 'impartici' in ll and ':' in line:
                val = line.split(':', 1)[1].strip().strip('*').strip()
                if val and not val.startswith('['):
                    institutional_fields['fecha'] = val
            elif ('número de participantes' in ll or 'numero de participantes' in ll or (ll.startswith('- participantes') and ':' in ll)) and ':' in line:
                val = line.split(':', 1)[1].strip().strip('*').strip()
                if val and not val.startswith('['):
                    institutional_fields['participantes'] = val
        if institutional_fields:
            inst_lines = []
            if 'instructor' in institutional_fields:
                inst_lines.append(f"INSTRUCTOR: {institutional_fields['instructor']}")
            if 'disenador' in institutional_fields:
                inst_lines.append(f"DISEÑADOR: {institutional_fields['disenador']}")
            if 'lugar' in institutional_fields:
                inst_lines.append(f"LUGAR: {institutional_fields['lugar']}")
            if 'fecha' in institutional_fields:
                inst_lines.append(f"FECHA: {institutional_fields['fecha']}")
            if 'participantes' in institutional_fields:
                inst_lines.append(f"PARTICIPANTES: {institutional_fields['participantes']}")
            master_parts.append("DATOS INSTITUCIONALES:\n" + '\n'.join(inst_lines))
        objetivo_lines = []
        tecnicas_lines = []
        tiempos_lines = []
        capture_obj = False
        capture_tec = False
        for line in lines:
            ll = line.lower().strip()
            if 'objetivo general' in ll or 'objetivo integrador' in ll:
                capture_obj = True
                capture_tec = False
                continue
            if 'objetivo particular' in ll or 'objetivos particular' in ll:
                capture_obj = True
                capture_tec = False
                continue
            if any(kw in ll for kw in ['técnica', 'tecnica', 'expositiva', 'demostrativa', 'diálogo', 'dialogo']):
                if len(line.strip()) > 10:
                    tecnicas_lines.append(line.strip())
                capture_tec = True
                capture_obj = False
                continue
            if any(kw in ll for kw in ['apertura', 'desarrollo', 'cierre']) and any(c.isdigit() for c in ll):
                tiempos_lines.append(line.strip())
            if capture_obj and line.strip() and not line.strip().startswith('#'):
                objetivo_lines.append(line.strip())
                if len(objetivo_lines) >= 6:
                    capture_obj = False
            if capture_tec and line.strip() and not line.strip().startswith('#'):
                tecnicas_lines.append(line.strip())
                if len(tecnicas_lines) >= 5:
                    capture_tec = False
        if objetivo_lines:
            master_parts.append("OBJETIVOS:\n" + '\n'.join(objetivo_lines[:6]))
        if tecnicas_lines:
            master_parts.append("TÉCNICAS INSTRUCCIONALES:\n" + '\n'.join(tecnicas_lines[:5]))
        if tiempos_lines:
            master_parts.append("DISTRIBUCIÓN DE TIEMPOS:\n" + '\n'.join(tiempos_lines[:4]))
        master_text = '\n\n'.join(master_parts)
        if len(master_text) > 2500:
            master_text = master_text[:2500]
        session['master_doc'] = master_text
        if course_topic:
            session['master_doc_topic'] = course_topic
        try:
            if current_user.is_authenticated:
                _persist_course_session_state(current_user)
        except Exception:
            pass
    except Exception as _e_master:
        logger.error(f"BUILD_MASTER_DOC_FAIL topic={course_topic!r} err={type(_e_master).__name__}: {_e_master}", exc_info=True)
        pass

def _extract_course_name(history):
    import re
    prefixes = [
        r"^\s*genera(?:r)?\s+(?:el\s+|un\s+|la\s+)?curso\s+(?:con\s+el\s+)?tema\s*[:\-]?\s*",
        r"^\s*genera(?:r)?\s+carta\s+descriptiva\s+(?:de|sobre|para)\s*[:\-]?\s*",
        r"^\s*carta\s+descriptiva\s+(?:de|sobre|para)\s*[:\-]?\s*",
        r"^\s*quiero\s+un\s+curso\s+(?:de|sobre|acerca\s+de)\s*[:\-]?\s*",
        r"^\s*curso\s+(?:de|sobre|acerca\s+de)\s*[:\-]?\s*",
        r"^\s*tema\s*[:\-]?\s*",
        r"^\s*disena(?:r)?\s+(?:un\s+)?curso\s+(?:de|sobre)\s*[:\-]?\s*",
    ]
    def _clean(text):
        t = text.strip()
        low = t.lower()
        for pat in prefixes:
            m = re.match(pat, low, flags=re.IGNORECASE)
            if m:
                t = t[m.end():].strip()
                break
        t = re.split(r"[,\n]|\s+(?:diseñador|disenador|instructor|participantes?|lugar|fecha|duraci[oó]n|n[uú]mero\s+de\s+horas|horario)\s*[:\-]", t, maxsplit=1, flags=re.IGNORECASE)[0]
        return t.strip(" :;-\t")
    for msg in history:
        if msg.get("role") == "user":
            text = msg.get("content", "")
            cleaned = _clean(text)
            if cleaned and 3 <= len(cleaned) <= 120 and not any(kw in cleaned.lower() for kw in ["checklist", "etapa", "manual del", "presentacion", "diagnostico"]):
                return cleaned
    return ""

_CONVERSATIONAL_PREAMBLE_PATTERNS = [
    r'^\s*(?:¡?(?:perfecto|claro|listo|excelente|de acuerdo|entendido|muy bien)[!\.,:]?\s*)+',
    r'^\s*(?:procedo a|voy a|aqu[ií] (?:tienes|te dejo|te entrego)|te (?:entrego|comparto|presento)|a continuaci[oó]n (?:te|le) (?:presento|comparto|entrego)|empiezo|comenzar[eé])[^\n]{0,200}\n+',
    r'^\s*(?:los? est[aá]ndares?|el est[aá]ndar)\s+ec0301[^\n]{0,400}(?:procedo|comenzar[eé]|empiezo|conforme a tu requerimiento)[^\n]{0,200}\n+',
    r'^\s*(?:nota|aviso|importante)\s*[:\-][^\n]{0,300}\n+',
]
_CONVERSATIONAL_PREAMBLE_RE = [re.compile(p, re.IGNORECASE) for p in _CONVERSATIONAL_PREAMBLE_PATTERNS]

def _strip_conversational_preamble(text):
    if not text:
        return text
    out = text
    for _ in range(4):
        prev = out
        for rx in _CONVERSATIONAL_PREAMBLE_RE:
            out = rx.sub('', out, count=1)
        out = out.lstrip()
        if out == prev:
            break
    return out

_COURSE_LABEL_PREFIX_RE = re.compile(
    r'^\s*\d*\s*[\.\)\-]?\s*nombre\s+del\s+curso(?:[\-\s]*taller)?\s*[:\-]?\s*',
    re.IGNORECASE)

def _strip_course_label_prefix(s):
    """Quita un prefijo de ETIQUETA de formulario tipo "1. Nombre del curso:" que
    a veces se cuela en el nombre del curso (origen de "1 Nombre del curso ...")."""
    if not s:
        return s
    return _COURSE_LABEL_PREFIX_RE.sub('', str(s)).strip()

_TOPIC_PREFIX_RE = [
    _COURSE_LABEL_PREFIX_RE,
    re.compile(r'^\s*genera(?:r|me)?\s+(?:el\s+|un\s+|la\s+|una\s+)?(?:curso|carta(?:\s+descriptiva)?|taller|capacitaci[oó]n)\s+(?:de\s+\d+\s*(?:horas?|h|min(?:utos)?)\s+)?(?:con\s+el\s+)?tema\s*[:\-]?\s*', re.IGNORECASE),
    re.compile(r'^\s*genera(?:r|me)?\s+(?:el\s+|un\s+|la\s+|una\s+)?(?:curso|carta(?:\s+descriptiva)?|taller)\s+(?:de|sobre|para|acerca\s+de)\s*[:\-]?\s*', re.IGNORECASE),
    re.compile(r'^\s*(?:dise[ñn]a(?:r|me)?|crea(?:r|me)?|hazme|haz|elabora(?:r|me)?|prepara(?:r|me)?)\s+(?:el\s+|un\s+|la\s+|una\s+)?(?:curso|carta(?:\s+descriptiva)?|taller)\s+(?:de\s+\d+\s*(?:horas?|h|min(?:utos)?)\s+)?(?:con\s+el\s+)?(?:tema|sobre|de|acerca\s+de)\s*[:\-]?\s*', re.IGNORECASE),
    re.compile(r'^\s*quiero\s+(?:un\s+|el\s+)?(?:curso|carta|taller)\s+(?:de|sobre|acerca\s+de)\s*[:\-]?\s*', re.IGNORECASE),
    re.compile(r'^\s*tema\s*[:\-]\s*', re.IGNORECASE),
]
_COMMAND_KEYWORDS_FOR_TOPIC = (
    'genera(?:r|me)?',
    'redacta(?:r|me)?',
    'dise[ñn]a(?:r|me)?',
    'crea(?:r|me)?',
    'hazme', 'haz',
    'elabora(?:r|me)?',
    'prepara(?:r|me)?',
    'arma(?:r|me)?',
    'construye',
    'quiero',
    'necesito',
)

def _sanitize_course_name_from_message(raw_message):
    if not raw_message:
        return ''
    text = raw_message.strip()
    for rx in _TOPIC_PREFIX_RE:
        m = rx.match(text)
        if m:
            text = text[m.end():].strip()
            break
    text = re.split(
        r'[,\n;]'
        r'|\.\s+(?=[A-ZÁÉÍÓÚÑ])'
        r'|\s+\d+\s+(?:horas?|h|min(?:utos)?|participantes?)\b'
        r'|\s+(?:dise[ñn]ador|instructor|participantes?|lugar|fecha|duraci[oó]n|n[uú]mero\s+de\s+horas|horario|modalidad|plataforma)\s*[:\-]',
        text, maxsplit=1, flags=re.IGNORECASE
    )[0].strip(' :;-\t.')
    if not text or len(text) < 3 or len(text) > 120:
        return ''
    low = text.lower()
    if any(re.match(rf'^{kw}\b', low) for kw in _COMMAND_KEYWORDS_FOR_TOPIC):
        return ''
    return text

def _normalize(text):
    import unicodedata
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.category(c).startswith('M'))

def _detect_element_from_filename(filename):
    fn = filename.lower()
    e1_kw = ['carta_descriptiva', 'contrato_aprendizaje', 'contrato_de_aprendizaje']
    e2_kw = ['iec', 'diagnostica', 'formativa', 'sumativa', 'satisfaccion', 'hoja_respuesta', 'instrumento']
    e3_kw = ['manual_instructor', 'manual_participante', 'manual_del_instructor', 'manual_del_participante', 'presentacion', 'slides']
    e4_kw = ['diagnostico', 'autodiagnostico', 'auto_diagnostico']
    for kw in e1_kw:
        if kw in fn:
            return 1
    for kw in e2_kw:
        if kw in fn:
            return 2
    for kw in e3_kw:
        if kw in fn:
            return 3
    for kw in e4_kw:
        if kw in fn:
            return 4
    return None

def _extract_course_name_from_docx(source):
    try:
        import docx as _docx
        import io as _io
        if isinstance(source, (bytes, bytearray)):
            doc = _docx.Document(_io.BytesIO(source))
        else:
            doc = _docx.Document(source)
        import re as _re
        for p in doc.paragraphs[:30]:
            t = (p.text or '').strip()
            m = _re.match(r'^Nombre\s+del\s+curso[\-\s]*taller\s*:\s*(.+)$', t, _re.IGNORECASE)
            if m:
                name = m.group(1).strip().strip('"').strip("'").strip()
                if name and len(name) >= 3:
                    return name
    except Exception:
        return None
    return None

def _auto_split_user_course_sessions(user_id):
    """Self-healing: separa CDs con nombres distintos en CSs separados, fusiona
    duplicados por topic, borra CSs vacíos y renumera por orden cronológico del
    primer doc/started_at. Idempotente."""
    try:
        try:
            _lock_user_row(user_id)
        except Exception:
            pass
        user = User.query.get(user_id)
        if user is None:
            return
        active_cs_id = user.active_course_session_id
        all_cs = CourseSession.query.filter_by(user_id=user_id).all()
        cs_by_id = {cs.id: cs for cs in all_cs}

        topic_to_cs_id = {}
        for cs in all_cs:
            t = (cs.topic or '').strip()
            if t and t not in topic_to_cs_id:
                topic_to_cs_id[t] = cs.id

        changed = False

        for cs in list(all_cs):
            docs = StoredFile.query.filter_by(user_id=user_id, course_session_id=cs.id, file_category='document') \
                .order_by(StoredFile.created_at.asc()).all()
            if not docs:
                continue
            doc_names = []
            for d in docs:
                nm = None
                if (d.filename or '').lower().endswith('.docx') and d.content:
                    if 'Carta_Descriptiva' in (d.filename or ''):
                        nm = _extract_course_name_from_docx(bytes(d.content))
                doc_names.append((d, nm))

            current_group_name = None
            assignments = []
            for d, nm in doc_names:
                if nm:
                    current_group_name = nm
                assignments.append((d, current_group_name))

            unique_names = []
            for _, gn in assignments:
                if gn and gn not in unique_names:
                    unique_names.append(gn)

            if not unique_names:
                continue

            primary_name = unique_names[0]
            try:
                primary_clean = _sanitize_course_name_from_message(primary_name) or primary_name
            except Exception:
                primary_clean = primary_name
            if (cs.topic or '') != primary_clean:
                cs.topic = primary_clean
                topic_to_cs_id[primary_clean] = cs.id
                changed = True

            for d, gn in assignments:
                if not gn:
                    continue
                try:
                    gn_clean = _sanitize_course_name_from_message(gn) or gn
                except Exception:
                    gn_clean = gn
                if gn_clean == primary_clean:
                    continue
                target_cs_id = topic_to_cs_id.get(gn_clean)
                if target_cs_id is None:
                    max_num = db.session.query(db.func.max(CourseSession.session_num)).filter_by(user_id=user_id).scalar() or 0
                    new_cs = CourseSession(user_id=user_id, session_num=int(max_num) + 1, is_active=False,
                                           topic=gn_clean, started_at=d.created_at or datetime.utcnow(),
                                           last_activity_at=d.created_at or datetime.utcnow())
                    db.session.add(new_cs)
                    db.session.flush()
                    cs_by_id[new_cs.id] = new_cs
                    topic_to_cs_id[gn_clean] = new_cs.id
                    target_cs_id = new_cs.id
                if d.course_session_id != target_cs_id:
                    d.course_session_id = target_cs_id
                    changed = True

        try:
            all_cs_for_chat = CourseSession.query.filter_by(user_id=user_id).all()
            chat_rows = ChatHistory.query.filter_by(user_id=user_id).all()
            cs_by_id_now = {c.id: c for c in all_cs_for_chat}
            cs_ids_set = set(cs_by_id_now.keys())
            import json as _j2, re as _re2
            def _norm_txt(s):
                s = (s or '').lower()
                s = _re2.sub(r'[^a-záéíóúñü0-9 ]+', ' ', s)
                return _re2.sub(r'\s+', ' ', s).strip()
            _STOP = {'el','la','los','las','de','del','y','en','con','un','una','unos','unas','para','sobre','por','al','que','es','su','sus','este','esta','estos','estas','como','base','tema','curso','cursos','horas','genera','crea','disena','diseña','haz','hacer','hola','dame','dare'}
            def _topic_tokens(t):
                tn = _norm_txt(t)
                if not tn: return set()
                return set(w for w in tn.split() if w not in _STOP and len(w) >= 4)
            for row in list(chat_rows):
                if row.course_session_id not in cs_ids_set:
                    continue
                try:
                    msgs = _j2.loads(row.messages_json or '[]')
                except Exception:
                    continue
                if not isinstance(msgs, list) or len(msgs) < 2:
                    continue
                origin_cs = cs_by_id_now.get(row.course_session_id)
                same_element_cs_ids = set()
                for d in StoredFile.query.filter_by(user_id=user_id, file_category='document').all():
                    if d.course_session_id in cs_ids_set:
                        same_element_cs_ids.add(d.course_session_id)
                candidates = []
                for cid in same_element_cs_ids:
                    c = cs_by_id_now.get(cid)
                    if c is None: continue
                    toks = _topic_tokens(c.topic or '')
                    if toks:
                        candidates.append((cid, toks))
                if len(candidates) <= 1:
                    continue
                anchors = []
                for i, m in enumerate(msgs):
                    if (m.get('role') if isinstance(m, dict) else None) != 'user':
                        continue
                    content_tokens = set(_norm_txt(m.get('content') if isinstance(m, dict) else '').split())
                    if not content_tokens:
                        continue
                    best_cid = None; best_score = 0.0; best_overlap = 0
                    for cid, toks in candidates:
                        overlap = len(toks & content_tokens)
                        if overlap < 2:
                            continue
                        score = overlap / max(len(toks), 1)
                        if score > best_score or (score == best_score and overlap > best_overlap):
                            best_cid = cid; best_score = score; best_overlap = overlap
                    if best_cid is not None and best_score >= 0.4:
                        anchors.append((i, best_cid))
                    elif best_cid is None:
                        for cid, toks in candidates:
                            if len(toks) == 1:
                                only_tok = next(iter(toks))
                                if len(only_tok) >= 5 and only_tok in content_tokens:
                                    anchors.append((i, cid)); break
                if not anchors:
                    continue
                blocks = {}
                first_idx = anchors[0][0]
                if first_idx > 0:
                    blocks.setdefault(row.course_session_id, []).extend(msgs[:first_idx])
                for j, (idx, cid) in enumerate(anchors):
                    end = anchors[j+1][0] if j+1 < len(anchors) else len(msgs)
                    blocks.setdefault(cid, []).extend(msgs[idx:end])
                if set(blocks.keys()) == {row.course_session_id}:
                    continue
                orig_block = blocks.get(row.course_session_id, [])
                row.messages_json = _j2.dumps(orig_block, ensure_ascii=False)
                changed = True
                for cid, blk in blocks.items():
                    if cid == row.course_session_id or not blk:
                        continue
                    existing = ChatHistory.query.filter_by(user_id=user_id, element_num=row.element_num, course_session_id=cid).first()
                    if existing is None:
                        new_row = ChatHistory(user_id=user_id, element_num=row.element_num,
                                              messages_json=_j2.dumps(blk, ensure_ascii=False),
                                              course_topic=(cs_by_id_now[cid].topic or '')[:300] or None,
                                              course_session_id=cid)
                        db.session.add(new_row)
                    else:
                        try:
                            cur = _j2.loads(existing.messages_json or '[]')
                        except Exception:
                            cur = []
                        existing.messages_json = _j2.dumps(cur + blk, ensure_ascii=False)
                logger.info(f"chat_split user={user_id} elem={row.element_num} from_cs={row.course_session_id} into={list(blocks.keys())}")
        except Exception as _e_chatsplit:
            logger.warning(f"chat_split_fail user={user_id} err={type(_e_chatsplit).__name__}: {_e_chatsplit}")

        all_cs = CourseSession.query.filter_by(user_id=user_id).all()
        for cs in list(all_cs):
            if cs.id == active_cs_id:
                continue
            _docs = StoredFile.query.filter_by(user_id=user_id, course_session_id=cs.id).count()
            _msgs = ChatHistory.query.filter_by(user_id=user_id, course_session_id=cs.id).count()
            if _docs == 0 and _msgs == 0:
                db.session.delete(cs)
                changed = True

        if changed:
            db.session.flush()

        all_cs = CourseSession.query.filter_by(user_id=user_id).all()

        def _first_event(cs):
            d = StoredFile.query.filter_by(user_id=user_id, course_session_id=cs.id) \
                .order_by(StoredFile.created_at.asc()).first()
            if d and d.created_at:
                return d.created_at
            return cs.started_at or datetime.max

        ordered = sorted(all_cs, key=lambda c: _first_event(c))
        for idx, cs in enumerate(ordered, start=1):
            if cs.session_num != idx:
                cs.session_num = idx
                changed = True

        if changed:
            db.session.commit()
            logger.info(f"auto_split_course_sessions user={user_id} normalized_total={len(ordered)}")
    except Exception as e:
        try: db.session.rollback()
        except Exception: pass
        logger.warning(f"auto_split_course_sessions_fail user={user_id} err={type(e).__name__}: {e}")

def _extract_course_name_from_cs_docs(user_id, cs_id):
    try:
        rows = StoredFile.query.filter_by(user_id=user_id, course_session_id=cs_id, file_category='document') \
            .order_by(StoredFile.created_at.asc()).all()
        cd_rows = [sf for sf in rows if 'Carta_Descriptiva' in (sf.filename or '')]
        for sf in (cd_rows + [r for r in rows if r not in cd_rows]):
            if not sf.content or not (sf.filename or '').lower().endswith('.docx'):
                continue
            name = _extract_course_name_from_docx(bytes(sf.content))
            if name:
                return name
    except Exception:
        return None
    return None

def _determine_doc_title(element_num, message, is_first_topic=False):
    if is_first_topic and element_num == 1:
        return "Carta_Descriptiva"
    msg_lower = _normalize(message.lower())
    if element_num == 1:
        if "objetivo" in msg_lower and "general" in msg_lower:
            return "Objetivo_General"
        elif "objetivo" in msg_lower and "particular" in msg_lower:
            return "Objetivos_Particulares"
        elif "carta descriptiva" in msg_lower or "completa" in msg_lower:
            return "Carta_Descriptiva"
        elif "tecnica" in msg_lower:
            return "Tecnicas_Instruccionales"
        elif "etapa" in msg_lower:
            return "Etapas_del_Curso"
        elif "contrato" in msg_lower or "aprendizaje" in msg_lower:
            return "Contrato_de_Aprendizaje"
        return "Elemento1_Carta_Descriptiva"
    elif element_num == 2:
        if "respuesta" in msg_lower:
            return "Hojas_de_Respuestas"
        elif "diagnostica" in msg_lower:
            return "Evaluacion_Diagnostica"
        elif "guia" in msg_lower or "observacion" in msg_lower:
            return "Guia_de_Observacion"
        elif "cotejo" in msg_lower:
            return "Lista_de_Cotejo"
        elif "sumativa" in msg_lower:
            return "Evaluacion_Sumativa"
        elif "satisfaccion" in msg_lower or "reaccion" in msg_lower:
            return "Evaluacion_de_Satisfaccion"
        elif "mediador" in msg_lower:
            return "Evaluacion_Mediadora"
        return "Instrumento_de_Evaluacion"
    elif element_num == 3:
        if "instructor" in msg_lower:
            return "Manual_del_Instructor"
        elif "participante" in msg_lower:
            return "Manual_del_Participante"
        elif "contenido" in msg_lower or "tematico" in msg_lower:
            return "Contenido_Tematico"
        elif "fuente" in msg_lower:
            return "Fuentes_de_Informacion"
        elif "guion" in msg_lower or "diapositiva" in msg_lower or "slide" in msg_lower or "presentacion" in msg_lower:
            return "Guion_de_Diapositivas"
        return "Manual_del_Curso"
    return "Documento_EC0301"

def _parse_slides_from_response(text):
    _SECTION_PATTERNS = [
        re.compile(r'^#{1,3}\s+.*(apertura|encuadre|desarrollo|cierre)\s*(\(.*\))?', re.IGNORECASE),
        re.compile(r'^##\s+.*(t[eé]cnica\s+(expositiva|demostrativa|di[aá]logo))', re.IGNORECASE),
    ]
    _CLOSING_KEYWORDS = ['gracias', 'cierre administrativo', 'despedida', 'fuentes de informaci', 'fuentes de información']
    _FILTER_PHRASES = ['contenido en pantalla', 'contenido de la diapositiva', 'texto en pantalla']

    def _detect_slide_type(title, content_lines):
        tl = title.lower()
        for pat in _SECTION_PATTERNS:
            if pat.search(title):
                return "section"
        if any(kw in tl for kw in _CLOSING_KEYWORDS):
            return "closing"
        has_sublists = sum(1 for l in content_lines if l.strip().startswith('-') or l.strip().startswith('*') or re.match(r'^\d+[\.\)]\s', l.strip())) > 0
        non_bullet = [l for l in content_lines if l.strip() and not l.strip().startswith('-') and not l.strip().startswith('*') and not re.match(r'^\d+[\.\)]\s', l.strip())]
        if len(non_bullet) >= 2 and has_sublists:
            return "two_column"
        return "content"

    def _filter_content(lines):
        filtered = []
        for line in lines:
            ll = line.strip().lower()
            clean_ll = re.sub(r'^[\-\*\d\.\)]+\s*', '', ll).strip()
            if any(phrase == clean_ll or phrase == clean_ll.rstrip(':') for phrase in _FILTER_PHRASES):
                continue
            if re.match(r'^#{1,3}\s+.*(apertura|encuadre|desarrollo|cierre)\s*(\(.*\))?', line.strip(), re.IGNORECASE):
                continue
            section_clean = re.sub(r'^[\-\*\d\.\)]+\s*', '', line.strip())
            if re.match(r'^#{1,3}\s+.*(t[eé]cnica|desarrollo|cierre)', section_clean, re.IGNORECASE):
                continue
            filtered.append(line)
        return filtered

    curso_datos = {}
    cd_match = re.search(r'\[CURSO_DATOS\](.*?)\[/CURSO_DATOS\]', text, re.IGNORECASE | re.DOTALL)
    if cd_match:
        def _cd_key(k):
            import unicodedata
            k = ''.join(c for c in unicodedata.normalize('NFKD', k.lower()) if not unicodedata.category(c).startswith('M')).strip()
            if 'instructor' in k or 'facilitador' in k or 'imparte' in k:
                return 'instructor'
            if 'periodo' in k or 'fecha' in k:
                return 'periodo'
            if 'horario' in k:
                return 'horario'
            if 'sede' in k or 'lugar' in k or 'ubicac' in k:
                return 'sede'
            if 'duracion' in k or 'horas' in k:
                return 'duracion'
            if 'participante' in k or 'asistente' in k:
                return 'participantes'
            if 'disenad' in k or 'diseno' in k or 'elabor' in k:
                return 'disenador'
            return ''
        for bl in cd_match.group(1).split("\n"):
            bl = re.sub(r'^[\-\*]+\s*', '', bl.strip())
            if ':' in bl:
                k, v = bl.split(":", 1)
                key = _cd_key(k)
                v = v.strip()
                if key and v and not v.startswith('['):
                    curso_datos[key] = v
        text = text.replace(cd_match.group(0), '')

    slides = []
    current_slide = None
    notes_mode = False

    for line in text.split("\n"):
        stripped = line.strip()
        section_match = False
        for pat in _SECTION_PATTERNS:
            if pat.search(stripped) and not re.match(r'^#{1,3}\s+.*[Ss]lide\s*\d+', stripped) and not re.match(r'^#{1,3}\s+.*[Dd]iapositiva\s*\d+', stripped):
                section_match = True
                break

        if section_match:
            if current_slide:
                current_slide["content"] = _filter_content(current_slide["content"])
                slides.append(current_slide)
                current_slide = None
                notes_mode = False
            sec_title = re.sub(r'^#{1,3}\s*', '', stripped).strip()
            slides.append({"title": sec_title, "content": [], "notes": "", "type": "section"})
            continue

        if re.match(r'^#{1,3}\s+.*[Ss]lide\s*\d+', stripped) or re.match(r'^#{1,3}\s+.*[Dd]iapositiva\s*\d+', stripped) or re.match(r'^[Ss]lide\s*\d+', stripped) or re.match(r'^[Dd]iapositiva\s*\d+', stripped):
            if current_slide:
                current_slide["content"] = _filter_content(current_slide["content"])
                slides.append(current_slide)
            title = re.sub(r'^#{1,3}\s*', '', stripped)
            title = re.sub(r'^[Ss]lide\s*\d+\s*[-:.\)]\s*', '', title)
            title = re.sub(r'^[Dd]iapositiva\s*\d+\s*[-:.\)]\s*', '', title)
            current_slide = {"title": title.strip(), "content": [], "notes": ""}
            notes_mode = False
        elif current_slide:
            lower = stripped.lower()
            if lower.startswith("notas para el presentador") or lower.startswith("notas del presentador") or (lower.startswith("nota") and ":" in lower):
                notes_mode = True
                after = stripped.split(":", 1)
                if len(after) > 1 and after[1].strip():
                    current_slide["notes"] += after[1].strip() + " "
                continue
            if notes_mode:
                is_content_marker = (stripped.startswith("-") or stripped.startswith("*") or
                    re.match(r'^\d+[\.\)]\s', stripped) or stripped.startswith("##") or
                    lower.startswith("contenido") or lower.startswith("texto"))
                if is_content_marker:
                    notes_mode = False
                    current_slide["content"].append(stripped)
                elif stripped and not stripped.startswith("#"):
                    current_slide["notes"] += stripped + " "
            else:
                if stripped:
                    current_slide["content"].append(stripped)

    if current_slide:
        current_slide["content"] = _filter_content(current_slide["content"])
        slides.append(current_slide)

    if not slides and text.strip():
        sections = text.split("\n\n")
        for i, section in enumerate(sections):
            lines_s = [l.strip() for l in section.split("\n") if l.strip()]
            if lines_s:
                slides.append({
                    "title": lines_s[0][:80],
                    "content": lines_s[1:],
                    "notes": "",
                    "type": "content"
                })

    for s in slides:
        if "type" not in s:
            s["type"] = _detect_slide_type(s["title"], s["content"])

    return slides, curso_datos

def _parse_response_to_sections(text):
    sections = []
    lines = text.split("\n")
    current_title = ""
    current_content = []

    for line in lines:
        if line.startswith("# ") or line.startswith("## "):
            if current_title or current_content:
                sections.append({
                    "titulo": current_title,
                    "contenido": "\n".join(current_content)
                })
            current_title = line.lstrip("#").strip()
            current_content = []
        else:
            current_content.append(line)

    if current_title or current_content:
        sections.append({
            "titulo": current_title,
            "contenido": "\n".join(current_content)
        })

    return sections if sections else [{"titulo": "", "contenido": text}]

_reengage_lock_handle = None

def _acquire_scheduler_leader_lock():
    """Adquiere un lock advisory exclusivo en /tmp. Solo UN proceso (worker) en
    toda la máquina puede mantener este lock; los demás workers de Gunicorn no
    arrancarán scheduler. El lock se libera al morir el proceso."""
    global _reengage_lock_handle
    try:
        import fcntl
        lock_path = os.environ.get('REENGAGE_LOCK_PATH', '/tmp/pertinentia_reengage.lock')
        fh = open(lock_path, 'w')
        fcntl.flock(fh.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
        fh.write(str(os.getpid()))
        fh.flush()
        _reengage_lock_handle = fh
        return True
    except (BlockingIOError, OSError) as _e:
        logger.info(f"Re-engagement scheduler: otro worker tiene el lock ({_e}); este worker NO arranca scheduler.")
        return False
    except Exception as _e:
        logger.warning(f"Re-engagement scheduler: no se pudo evaluar lock ({_e}); por seguridad NO arranca scheduler en este worker.")
        return False

def _init_reengagement_scheduler():
    """Inicializa APScheduler para correr campañas diarias a las 10:00 CDMX.
    Idempotente. Multi-worker safe vía file lock advisory en /tmp."""
    if os.environ.get('DISABLE_REENGAGE_SCHEDULER') == '1':
        logger.info("Re-engagement scheduler deshabilitado por env var.")
        return
    if app.debug and os.environ.get('WERKZEUG_RUN_MAIN') != 'true':
        return
    if getattr(app, '_reengage_scheduler_started', False):
        return
    if not _acquire_scheduler_leader_lock():
        return
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        from apscheduler.triggers.cron import CronTrigger
        sched = BackgroundScheduler(timezone='America/Mexico_City')
        def _job():
            with app.app_context():
                try:
                    logger.info("Re-engagement scheduler: iniciando corrida diaria")
                    res = reengagement.run_all_campaigns(
                        app, db, User, UserEvent, UserPurchase, EmailLog,
                        _log_email_attempt, _log_email_result, max_per_campaign=200
                    )
                    logger.info(f"Re-engagement scheduler: completado {res}")
                except Exception as _e:
                    logger.error(f"Re-engagement scheduler error: {_e}")
        sched.add_job(_job, CronTrigger(hour=10, minute=0), id='reengage_daily', replace_existing=True)
        sched.start()
        app._reengage_scheduler_started = True
        logger.info("Re-engagement scheduler iniciado: cron diario 10:00 America/Mexico_City")
    except Exception as _e:
        logger.error(f"No se pudo iniciar el scheduler de re-engagement: {_e}")

_init_reengagement_scheduler()

if __name__ == "__main__":
    is_dev = os.environ.get("REPLIT_DEV_DOMAIN") is not None
    puerto = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=puerto, debug=is_dev)
